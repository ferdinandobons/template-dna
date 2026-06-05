# SPDX-License-Identifier: MIT
from docx import Document


def text_content(path) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs] + [cell.text for t in doc.tables for row in t.rows for cell in row.cells])

