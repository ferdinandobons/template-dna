# SPDX-License-Identifier: MIT
"""Shared-spine + QA regression tests for the strengthening pass.

These cover the SHARED + QA fixes, exercised against the committed complex
fixtures (``tests/fixtures/complex/acme_complex.{docx,pptx,xlsx}``) so the checks
run on realistic OOXML, not toy inputs:

- CC-1: ``comprehend_input_bundle`` yields a real ``excerpt`` and non-empty
  ``facts.styles`` for an xlsx workbook (previously empty for every workbook,
  because only the docx/pptx catalog keys were read).
- formula_preservation: the gate now ERRORs when a generation erased or mutated a
  shell formula (the X1 silent-corruption class that previously passed clean).
- component_survival: the gate WARNs when a native component (table/chart/list/
  picture) present in the shell is missing from the output (the pptx down-render /
  docx list-flattening class the text scans are blind to).

Every input is a synthetic "Acme" fixture; no proprietary template is read.
"""

from __future__ import annotations

import os
import shutil
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from openpyxl import load_workbook

from brandkit.formats.xlsx import extract as xlsx_extract
from brandkit.profile import comprehension as comp
from brandkit.profile import schema, store
from brandkit.qa import checks_deterministic as cd
from brandkit.qa.gate import run_qa

_FIXTURES = Path(__file__).resolve().parents[0] / "fixtures" / "complex"
_XLSX = _FIXTURES / "acme_complex.xlsx"
_DOCX = _FIXTURES / "acme_complex.docx"


def _extract_xlsx(td: Path):
    """Extract the complex xlsx fixture into ``td`` and return the loaded profile."""
    old = os.getcwd()
    os.chdir(td)
    try:
        xlsx_extract.extract(_XLSX, "cx", scope="project", cwd=td)
        return store.load_profile("cx", "project")
    finally:
        os.chdir(old)


class ComprehendInputExcerptTest(unittest.TestCase):
    """CC-1: every format yields a real excerpt + facts.styles (was empty for xlsx)."""

    def test_xlsx_excerpt_and_styles_non_empty(self) -> None:
        with tempfile.TemporaryDirectory() as t:
            prof = _extract_xlsx(Path(t)).profile
            bundle = comp.comprehend_input_bundle(prof)
            # The excerpt descends into sheets[*].non_empty_cells (the xlsx nesting),
            # so the model sees real cell text rather than [].
            self.assertTrue(bundle["excerpt"], "xlsx excerpt must not be empty (CC-1)")
            self.assertTrue(
                any("report_title" in s for s in bundle["excerpt"]),
                "excerpt should carry the workbook's own cell text",
            )
            # facts.styles falls back to the xlsx 'named_styles' key.
            self.assertTrue(
                bundle["facts"]["styles"], "facts.styles must not be empty (CC-1)"
            )

    def test_excerpt_is_length_capped_and_deterministic(self) -> None:
        with tempfile.TemporaryDirectory() as t:
            prof = _extract_xlsx(Path(t)).profile
            a = comp.comprehend_input_bundle(prof, excerpt_chars=200)["excerpt"]
            b = comp.comprehend_input_bundle(prof, excerpt_chars=200)["excerpt"]
            self.assertEqual(a, b, "the same profile must yield the same excerpt")
            self.assertLessEqual(sum(len(s) for s in a), 200)


class FormulaPreservationCheckTest(unittest.TestCase):
    """formula_preservation: ERROR on any lost/mutated shell formula (X1 / QA-BLIND)."""

    PROFILE = {"kind": schema.Kind.XLSX.value, "roles": {"_index": []}}

    def test_intact_copy_has_no_findings(self) -> None:
        with tempfile.TemporaryDirectory() as t:
            out = Path(t) / "intact.xlsx"
            shutil.copyfile(_XLSX, out)
            findings = cd.check_formula_preservation(_XLSX, out, self.PROFILE)
            self.assertEqual(findings, [])

    def test_erased_formula_is_an_error(self) -> None:
        with tempfile.TemporaryDirectory() as t:
            out = Path(t) / "corrupt.xlsx"
            wb = load_workbook(_XLSX, data_only=False)
            erased = 0
            for ws in wb.worksheets:
                for cell in list(ws._cells.values()):
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        cell.value = None
                        erased += 1
            wb.save(out)
            self.assertGreater(erased, 0, "fixture must contain formulas")
            findings = cd.check_formula_preservation(_XLSX, out, self.PROFILE)
            self.assertEqual(len(findings), erased)
            self.assertTrue(
                all(f.severity == schema.Severity.ERROR.value for f in findings)
            )
            self.assertTrue(all(f.check == "formula_preservation" for f in findings))

    def test_mutated_formula_is_an_error(self) -> None:
        with tempfile.TemporaryDirectory() as t:
            out = Path(t) / "mutated.xlsx"
            wb = load_workbook(_XLSX, data_only=False)
            mutated_addr = None
            for ws in wb.worksheets:
                for cell in list(ws._cells.values()):
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        cell.value = "=BOGUS()"
                        mutated_addr = f"{ws.title}!{cell.coordinate}"
                        break
                if mutated_addr:
                    break
            wb.save(out)
            findings = cd.check_formula_preservation(_XLSX, out, self.PROFILE)
            mutated = [f for f in findings if "mutated" in f.message]
            self.assertTrue(mutated)
            self.assertEqual(mutated[0].severity, schema.Severity.ERROR.value)

    def test_missing_output_is_a_noop(self) -> None:
        # Verify time: no output to diff -> the check is silent (does not crash).
        self.assertEqual(cd.check_formula_preservation(_XLSX, None, self.PROFILE), [])
        self.assertEqual(cd.check_formula_preservation(None, _XLSX, self.PROFILE), [])

    def test_gate_fails_on_erased_formula_end_to_end(self) -> None:
        """run_qa (the real gate) now FAILS on a workbook that erased formulas."""
        with tempfile.TemporaryDirectory() as t:
            td = Path(t)
            loaded = _extract_xlsx(td)
            out = td / "out.xlsx"
            wb = load_workbook(loaded.shell_path, data_only=False)
            for ws in wb.worksheets:
                for cell in list(ws._cells.values()):
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        cell.value = None
            wb.save(out)
            report = run_qa(
                str(out),
                loaded.profile,
                qa="fast",
                shell=loaded.shell_path,
            )
            self.assertEqual(report.verdict, schema.VerificationStatus.FAILED.value)
            self.assertTrue(
                any(f.check == "formula_preservation" for f in report.findings),
                "the gate must surface a formula_preservation finding",
            )

    def test_gate_passes_on_faithful_copy(self) -> None:
        with tempfile.TemporaryDirectory() as t:
            td = Path(t)
            loaded = _extract_xlsx(td)
            out = td / "faithful.xlsx"
            shutil.copyfile(loaded.shell_path, out)
            report = run_qa(
                str(out),
                loaded.profile,
                qa="fast",
                shell=loaded.shell_path,
            )
            self.assertFalse(
                [f for f in report.findings if f.check == "formula_preservation"],
                "a verbatim copy must raise no formula_preservation finding",
            )


class ComponentSurvivalCheckTest(unittest.TestCase):
    """component_survival: WARN when a native component is lost in the output."""

    def test_xlsx_dropped_table_warns(self) -> None:
        profile = {"kind": schema.Kind.XLSX.value, "roles": {"_index": []}}
        with tempfile.TemporaryDirectory() as t:
            out = Path(t) / "notable.xlsx"
            wb = load_workbook(_XLSX, data_only=False)
            for ws in wb.worksheets:
                for tn in list(ws.tables.keys()):
                    del ws.tables[tn]
            wb.save(out)
            findings = cd.check_component_survival(_XLSX, out, profile)
            self.assertTrue(findings)
            self.assertTrue(
                all(f.severity == schema.Severity.WARNING.value for f in findings)
            )
            self.assertTrue(any("tables" in f.message for f in findings))

    def test_xlsx_intact_copy_no_warning(self) -> None:
        profile = {"kind": schema.Kind.XLSX.value, "roles": {"_index": []}}
        with tempfile.TemporaryDirectory() as t:
            out = Path(t) / "intact.xlsx"
            shutil.copyfile(_XLSX, out)
            self.assertEqual(cd.check_component_survival(_XLSX, out, profile), [])

    def test_docx_dropped_table_warns(self) -> None:
        profile = {"kind": schema.Kind.DOCX.value, "roles": {"_index": []}}
        with tempfile.TemporaryDirectory() as t:
            out = Path(t) / "notable.docx"
            doc = Document(_DOCX)
            self.assertTrue(doc.tables, "fixture must contain a table")
            for table in list(doc.tables):
                table._element.getparent().remove(table._element)
            doc.save(out)
            findings = cd.check_component_survival(_DOCX, out, profile)
            self.assertTrue(any("tables" in f.message for f in findings))
            self.assertTrue(
                all(f.severity == schema.Severity.WARNING.value for f in findings)
            )

    def test_missing_file_is_a_noop(self) -> None:
        profile = {"kind": schema.Kind.XLSX.value}
        self.assertEqual(cd.check_component_survival(_XLSX, None, profile), [])
        self.assertEqual(cd.check_component_survival(None, _XLSX, profile), [])


class NoNetStructureLossFloorTest(unittest.TestCase):
    """The destructive-floor backstop re-verifies the confidence floor when the
    caller threads ``confidence`` (additive: ``None`` skips the re-check)."""

    def _profile_with_demo(self, *, confidence: float | None) -> dict:
        comp_block = {
            "status": schema.ComprehensionStatus.PRESENT.value,
            "demo_classification": {
                "regions": [{"region_ref": "region.r1", "verdict": "demo"}]
            },
        }
        if confidence is not None:
            comp_block["confidence"] = confidence
        return {"comprehension": comp_block}

    def test_unsanctioned_removal_is_error_regardless_of_confidence(self) -> None:
        # A ref with NO corroborated verdict is an ERROR even at high confidence.
        profile = self._profile_with_demo(confidence=0.9)
        findings = cd.check_no_net_structure_loss(
            {"region.unknown"}, profile, confidence=0.9
        )
        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0].severity, schema.Severity.ERROR.value)
        self.assertIn("without a corroborated", findings[0].message)

    def test_sanctioned_removal_below_floor_is_error(self) -> None:
        # A ref WITH a corroborated demo verdict is still an ERROR when the threaded
        # confidence is below the floor (the reconcile site should have downgraded it).
        profile = self._profile_with_demo(confidence=0.3)
        findings = cd.check_no_net_structure_loss(
            {"region.r1"}, profile, confidence=0.3
        )
        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0].severity, schema.Severity.ERROR.value)
        self.assertIn("below the destructive", findings[0].message)

    def test_sanctioned_removal_above_floor_passes(self) -> None:
        profile = self._profile_with_demo(confidence=0.9)
        self.assertEqual(
            cd.check_no_net_structure_loss({"region.r1"}, profile, confidence=0.9), []
        )

    def test_confidence_none_skips_floor_recheck(self) -> None:
        # Back-compat: without a threaded confidence the floor re-check is skipped,
        # so a sanctioned removal passes even though the stored confidence is low.
        profile = self._profile_with_demo(confidence=0.1)
        self.assertEqual(cd.check_no_net_structure_loss({"region.r1"}, profile), [])


if __name__ == "__main__":
    unittest.main()
