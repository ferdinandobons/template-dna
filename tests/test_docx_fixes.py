# SPDX-License-Identifier: MIT
"""Regression tests for the CORE docx correctness + brand-guarantee fixes.

Covers the confirmed findings:
  - C1  TOC over-capture (a post-body 'Contents' heading is body, not toc;
        stacked-index front matter stays preserved).
  - C2  resolver_targets_exist (a profile pointing at a missing style FAILS).
  - M1  unhandled blocks emit no empty paragraph + record a degradation finding;
        a genuinely unknown block type raises.
  - M2  nested list items appear in output, threaded by level.
  - M3  the cover title fills a block-level w:sdt in the cover region (before TOC).
  - M4  table colspan/rowspan honored.
  - M5  intermediate w:sectPr survives clear_body_region.
  - M8/M12  a literal-markdown doc yields an ERROR finding + run_qa 'failed'.
  - refresh_toc marks only TOC fields and returns the real count.
  - arch-3  a fabricated profile is caught by validate() without shell I/O.
"""

from __future__ import annotations

import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from brandkit.formats.docx import cover as covermod
from brandkit.formats.docx import generate as docx_generate
from brandkit.formats.docx import roles as docx_roles
from brandkit.formats.docx import structure
from brandkit.formats.docx.structure import _local_name, w
from brandkit.ir import components as ir_components
from brandkit.ir import model as ir
from brandkit.profile import schema
from brandkit.profile.resolver import ProfileResolver
from brandkit.qa.gate import run_qa
from brandkit.qa.model import Finding


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def _add_toc_field(doc, instr='TOC \\o "1-3" \\h \\z \\u'):
    """Append a real TOC complex field paragraph (begin/instrText/separate/end)."""
    p = doc.add_paragraph()
    r = p.add_run()
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    r._r.append(fb)
    r2 = p.add_run()
    it = OxmlElement("w:instrText")
    it.text = instr
    r2._r.append(it)
    r3 = p.add_run()
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    r3._r.append(fs)
    # a nested PAGEREF field inside the TOC entry - must NOT be marked dirty
    rp = p.add_run()
    pb = OxmlElement("w:fldChar")
    pb.set(qn("w:fldCharType"), "begin")
    rp._r.append(pb)
    rp2 = p.add_run()
    pit = OxmlElement("w:instrText")
    pit.text = "PAGEREF _Toc1 \\h"
    rp2._r.append(pit)
    rp3 = p.add_run()
    pe = OxmlElement("w:fldChar")
    pe.set(qn("w:fldCharType"), "end")
    rp3._r.append(pe)
    p.add_run("entry .... 1")
    r6 = p.add_run()
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    r6._r.append(fe)
    return p


def _add_toc_field_split(doc, instr_parts):
    """A TOC field whose instruction is split across several ``instrText`` runs.

    Mirrors how Word emits field codes (the ``\\o`` switch and its argument can
    land in different runs), with a nested PAGEREF field inside the result.
    """
    p = doc.add_paragraph()
    r = p.add_run()
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    r._r.append(fb)
    for part in instr_parts:
        rp = p.add_run()
        it = OxmlElement("w:instrText")
        it.text = part
        rp._r.append(it)
    r3 = p.add_run()
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    r3._r.append(fs)
    # nested PAGEREF field inside a rendered entry - its instrText must not leak
    # into the outer TOC instruction.
    rb = p.add_run()
    pb = OxmlElement("w:fldChar")
    pb.set(qn("w:fldCharType"), "begin")
    rb._r.append(pb)
    rb2 = p.add_run()
    pit = OxmlElement("w:instrText")
    pit.text = "PAGEREF _Toc1 \\h"
    rb2._r.append(pit)
    rb3 = p.add_run()
    pe = OxmlElement("w:fldChar")
    pe.set(qn("w:fldCharType"), "end")
    rb3._r.append(pe)
    p.add_run("entry .... 1")
    r6 = p.add_run()
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    r6._r.append(fe)
    return p


def _append_intermediate_sectpr(doc):
    """Append a paragraph carrying an intermediate w:pPr/w:sectPr (a section break)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sect = OxmlElement("w:sectPr")
    pPr.append(sect)
    return p


def _docx_profile(roles=None):
    prof = schema.build_envelope("docx", {"name": "t"})
    prof["surface"] = {"docx": {}}
    prof["roles"] = roles or {"_index": []}
    return prof


# ---------------------------------------------------------------------------
# C1 - TOC over-capture
# ---------------------------------------------------------------------------
class TocOverCaptureTest(unittest.TestCase):
    def test_post_body_contents_heading_is_body_and_cleared(self):
        doc = Document()
        doc.add_paragraph("{{title}}", style="Title")
        _add_toc_field(doc)  # real TOC
        doc.add_paragraph("Real Section", style="Heading 1")  # body
        doc.add_paragraph("real body text", style="Normal")  # body
        doc.add_paragraph("Contents", style="Heading 1")  # body heading named Contents
        doc.add_paragraph("more real body text", style="Normal")

        cls = structure.classify_body_children(doc)
        # exactly one toc child (the field); the 'Contents' heading is body.
        regions = [c["region"] for c in cls if not c["is_sectpr"]]
        self.assertEqual(regions.count("toc"), 1)
        contents_idx = next(
            i
            for i, c in enumerate(cls)
            if _local_name(list(doc.element.body)[c["index"]].tag) == "p"
            and "Contents"
            in "".join(
                t.text or "" for t in list(doc.element.body)[c["index"]].iter(w("t"))
            )
        )
        self.assertEqual(cls[contents_idx]["region"], "body")

        # clearing the body removes the stale 'Contents' + real body, keeps TOC.
        structure.clear_body_region(doc, preserve_cover=True, preserve_toc=True)
        body_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("Real Section", body_text)
        self.assertNotIn("more real body text", body_text)
        self.assertNotIn("Contents", body_text)

    def test_stacked_index_front_matter_is_preserved(self):
        doc = Document()
        doc.add_paragraph("{{title}}", style="Title")
        doc.add_paragraph("Sommario", style="Heading 1")  # TOC heading (folded in)
        _add_toc_field(doc, 'TOC \\o "1-3" \\h')  # main TOC
        doc.add_paragraph("Indice delle Tabelle", style="Heading 1")  # index separator
        _add_toc_field(doc, 'TOC \\h \\c "Tabella"')  # table-of-tables
        doc.add_paragraph("Indice delle figure", style="Heading 1")  # index separator
        _add_toc_field(doc, 'TOC \\h \\c "Figura"')  # table-of-figures
        doc.add_paragraph("1. Introduction", style="Heading 1")  # body starts
        doc.add_paragraph("body para", style="Normal")

        cls = structure.classify_body_children(doc)
        regions = [c["region"] for c in cls if not c["is_sectpr"]]
        # cover (title) + 6 toc (Sommario heading + 3 fields + 2 index headings) + 2 body
        self.assertEqual(regions.count("toc"), 6)
        self.assertEqual(regions.count("body"), 2)

        structure.clear_body_region(doc, preserve_cover=True, preserve_toc=True)
        kept = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Sommario", kept)
        self.assertIn("Indice delle Tabelle", kept)
        self.assertIn("Indice delle figure", kept)
        self.assertNotIn("1. Introduction", kept)
        self.assertNotIn("body para", kept)


# ---------------------------------------------------------------------------
# M5 - multi-section sectPr preserved
# ---------------------------------------------------------------------------
class SectPrPreservationTest(unittest.TestCase):
    def _count_sect(self, doc):
        body = doc.element.body
        top = sum(1 for c in body if _local_name(c.tag) == "sectPr")
        inter = sum(
            1
            for c in body
            if _local_name(c.tag) == "p"
            and c.find(w("pPr")) is not None
            and c.find(w("pPr")).find(w("sectPr")) is not None
        )
        return top + inter

    def test_intermediate_sectpr_survives_clear(self):
        doc = Document()
        doc.add_paragraph("Body Section A", style="Heading 1")
        doc.add_paragraph("text a", style="Normal")
        _append_intermediate_sectpr(doc)  # section break -> body region, holds_sectpr
        doc.add_paragraph("Body Section B", style="Heading 1")
        doc.add_paragraph("text b", style="Normal")

        before = self._count_sect(doc)
        self.assertEqual(before, 2)  # 1 intermediate + 1 final top-level
        cls = structure.classify_body_children(doc)
        self.assertTrue(any(c.get("holds_sectpr") for c in cls))
        structure.clear_body_region(doc, preserve_cover=True, preserve_toc=True)
        self.assertEqual(self._count_sect(doc), before)


# ---------------------------------------------------------------------------
# refresh_toc
# ---------------------------------------------------------------------------
class RefreshTocTest(unittest.TestCase):
    def test_marks_only_toc_fields_and_returns_count(self):
        doc = Document()
        doc.add_paragraph("{{title}}", style="Title")
        _add_toc_field(doc)  # 1 TOC field (with a nested PAGEREF inside)
        _add_toc_field(doc, 'TOC \\h \\c "Tabella"')  # 2nd TOC field
        doc.add_paragraph("Body", style="Heading 1")

        n = structure.refresh_toc(doc)
        self.assertEqual(n, 2)  # only the two TOC fields, not the nested PAGEREFs
        body = doc.element.body
        dirty = [
            f
            for f in body.iter(w("fldChar"))
            if f.get(w("fldCharType")) == "begin" and f.get(w("dirty")) == "true"
        ]
        self.assertEqual(len(dirty), 2)
        # updateFields present in settings, before any element that must follow it.
        names = [_local_name(c.tag) for c in doc.settings.element]
        self.assertIn("updateFields", names)

    def test_no_toc_returns_zero_and_writes_nothing(self):
        doc = Document()
        doc.add_paragraph("Body", style="Heading 1")
        self.assertEqual(structure.refresh_toc(doc), 0)

    def test_visible_outline_toc_cache_is_rewritten_from_generated_headings(self):
        doc = Document()
        doc.add_paragraph("{{title}}", style="Title")
        _add_toc_field(doc)
        doc.add_paragraph("Old Template Body", style="Heading 1")

        rewritten = structure.refresh_visible_outline_toc_cache(
            doc,
            [(1, "New Real Section"), (2, "Nested Real Topic")],
        )

        self.assertEqual(rewritten, 1)
        toc_text = "\n".join(
            "".join(
                t.text or "" for t in list(doc.element.body)[c["index"]].iter(w("t"))
            )
            for c in structure.classify_body_children(doc)
            if c["region"] == "toc"
        )
        self.assertIn("New Real Section", toc_text)
        self.assertIn("Nested Real Topic", toc_text)
        self.assertNotIn("entry .... 1", toc_text)

    def test_split_instrtext_outline_toc_preserves_full_field_code(self):
        # Word splits the field instruction across consecutive instrText runs;
        # the rewritten field must carry the COMPLETE code, not just 'TOC \o',
        # and must not absorb the nested PAGEREF instruction.
        doc = Document()
        doc.add_paragraph("{{title}}", style="Title")
        _add_toc_field_split(doc, ["TOC \\o ", '"1-3" \\h \\z \\u'])
        doc.add_paragraph("Old Template Body", style="Heading 1")

        rewritten = structure.refresh_visible_outline_toc_cache(
            doc,
            [(1, "New Real Section")],
        )
        self.assertEqual(rewritten, 1)

        toc_instr = "".join(
            it.text or ""
            for it in doc.element.body.iter(w("instrText"))
            if (it.text or "").strip().startswith("TOC")
        )
        self.assertIn('\\o "1-3"', toc_instr)  # full range preserved
        self.assertIn("\\h", toc_instr)
        self.assertIn("\\z", toc_instr)
        self.assertNotIn("PAGEREF", toc_instr)  # nested field not over-captured
        toc_text = "\n".join(
            "".join(
                t.text or "" for t in list(doc.element.body)[c["index"]].iter(w("t"))
            )
            for c in structure.classify_body_children(doc)
            if c["region"] == "toc"
        )
        self.assertIn("New Real Section", toc_text)

    def test_split_caption_index_is_not_rewritten_as_outline_toc(self):
        # A caption index whose \c switch lands in a later instrText run must be
        # left for Word to recompute, not overwritten with outline headings.
        doc = Document()
        doc.add_paragraph("{{title}}", style="Title")
        _add_toc_field_split(doc, ["TOC \\h ", '\\c "Figura"'])
        doc.add_paragraph("Old Template Body", style="Heading 1")

        rewritten = structure.refresh_visible_outline_toc_cache(
            doc,
            [(1, "New Real Section")],
        )
        self.assertEqual(rewritten, 0)


# ---------------------------------------------------------------------------
# C2 - resolver_targets_exist
# ---------------------------------------------------------------------------
class ResolverTargetsExistTest(unittest.TestCase):
    def _shell(self, tmp_path):
        shell = tmp_path / "shell.docx"
        d = Document()
        d.add_paragraph("hi", style="Heading 1")
        d.save(shell)
        return shell

    def test_missing_style_fails_verification(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            prof = _docx_profile(
                {
                    "_index": ["heading.1"],
                    "heading.1": {
                        "resolver": {
                            "type": "named_style",
                            "style_id": "DoesNotExistXYZ",
                            "style_name": "Nope",
                        }
                    },
                }
            )
            report = run_qa(None, prof, shell=shell)
            self.assertEqual(report.verdict, schema.VerificationStatus.FAILED.value)
            self.assertTrue(
                any(
                    f.check == "resolver_targets_exist"
                    and f.severity == schema.Severity.ERROR.value
                    for f in report.findings
                )
            )

    def test_present_style_passes(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            prof = _docx_profile(
                {
                    "_index": ["heading.1"],
                    "heading.1": {
                        "resolver": {
                            "type": "named_style",
                            "style_id": "Heading1",
                            "style_name": "Heading 1",
                        }
                    },
                }
            )
            report = run_qa(None, prof, shell=shell)
            self.assertTrue(report.passed)


# ---------------------------------------------------------------------------
# M8 / M12 - literal markdown is an ERROR + run_qa failed
# ---------------------------------------------------------------------------
class LiteralMarkdownGateTest(unittest.TestCase):
    def test_literal_markdown_doc_fails(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "leaky.docx"
            d = Document()
            d.add_paragraph("This has **bold** literal markdown that leaked.")
            d.save(out)
            prof = _docx_profile({"_index": []})
            report = run_qa(out, prof, qa="fast")
            self.assertEqual(report.verdict, schema.VerificationStatus.FAILED.value)
            self.assertTrue(
                any(
                    f.check == "no_literal_markdown"
                    and f.severity == schema.Severity.ERROR.value
                    for f in report.findings
                )
            )

    def test_lists_use_named_numbering_not_advertised(self):
        # M8: profiles must not claim an unenforced invariant.
        self.assertNotIn("lists_use_named_numbering", schema.DEFAULT_L0_INVARIANTS)


# ---------------------------------------------------------------------------
# M1 - unhandled blocks
# ---------------------------------------------------------------------------
class UnhandledBlockTest(unittest.TestCase):
    def _shell(self, tmp_path):
        shell = tmp_path / "shell.docx"
        d = Document()
        d.add_paragraph("x", style="Heading 1")
        d.save(shell)
        return shell

    def test_unhandled_block_emits_no_empty_paragraph_and_records_finding(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _docx_profile({"_index": []})
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(runs=[{"t": "real text"}]),
                    ir.Image(
                        src="/nope.png"
                    ),  # missing source -> degrades, never crashes
                    ir.Chart(),  # empty chart (no data) -> degrades loudly, no blank para
                ]
            )
            before = len(Document(shell).paragraphs)
            findings: list[Finding] = []
            docx_generate.generate(prof, shell, idoc, out, findings=findings)
            gen = Document(out)
            texts = [p.text for p in gen.paragraphs]
            # No empty paragraph injected for the degraded blocks.
            self.assertNotIn(
                "", [t for t in texts if t == "" and texts.index(t) >= before]
            )
            self.assertIn("real text", "\n".join(texts))
            # Exactly the two degradation findings recorded (image source missing + chart).
            degraded = [f for f in findings if f.check == "block_degraded"]
            self.assertEqual(len(degraded), 2)
            self.assertTrue(
                all(f.severity == schema.Severity.WARNING.value for f in degraded)
            )

    def test_unknown_block_type_raises(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _docx_profile({"_index": []})

            class Mystery(ir.Block):
                TYPE = "mystery"

            idoc = ir.IntermediateDocument(blocks=[Mystery()])
            with self.assertRaises(docx_generate.GenerationError):
                docx_generate.generate(prof, shell, idoc, out)


# ---------------------------------------------------------------------------
# Native docx TOC (ir.Toc -> a real, updateable outline TOC field, or deferral
# to a preserved structural TOC so the document never carries two)
# ---------------------------------------------------------------------------
def _toc_instr_count(path) -> int:
    """Count outer TOC field instructions in the output document.xml."""
    import re
    import zipfile

    xml = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8", "replace")
    return len(re.findall(r"<w:instrText[^>]*>\s*TOC\b", xml))


def _settings_has_update_fields(path) -> bool:
    import zipfile

    with zipfile.ZipFile(path) as z:
        return "updateFields" in z.read("word/settings.xml").decode("utf-8", "replace")


class NativeTocTest(unittest.TestCase):
    def _bare_shell(self, tmp_path):
        shell = tmp_path / "shell.docx"
        d = Document()
        d.add_paragraph("x", style="Heading 1")
        d.save(shell)
        return shell

    def _toc_shell(self, tmp_path):
        shell = tmp_path / "shell.docx"
        d = Document()
        _add_toc_field(d)  # a real structural TOC field
        d.add_paragraph("Body heading", style="Heading 1")
        d.save(shell)
        return shell

    def _idoc(self):
        return ir.IntermediateDocument(
            blocks=[
                ir.Heading(level=1, runs=[{"t": "Alpha"}]),
                ir.Heading(level=2, runs=[{"t": "Beta"}]),
                ir.Toc(title="Contents", max_level=2),
                ir.Paragraph(runs=[{"t": "Body."}]),
            ]
        )

    def test_native_toc_emitted_when_no_structural_toc(self):
        import tempfile
        import zipfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._bare_shell(Path(td))
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(
                _docx_profile({"_index": []}),
                shell,
                self._idoc(),
                out,
                findings=findings,
            )
            # A real, updateable TOC field is authored (not a degradation).
            self.assertEqual(_toc_instr_count(out), 1)
            self.assertTrue(_settings_has_update_fields(out))
            xml = (
                zipfile.ZipFile(out)
                .read("word/document.xml")
                .decode("utf-8", "replace")
            )
            self.assertIn('w:dirty="true"', xml)  # begin fldChar marked dirty
            # Visible cache carries the generated headings, not template samples.
            self.assertIn("Alpha", xml)
            self.assertIn("Beta", xml)
            # The toc block is NOT degraded.
            self.assertFalse(any(f.check == "block_degraded" for f in findings))

    def test_toc_block_defers_to_preserved_structural_toc(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._toc_shell(Path(td))
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(
                _docx_profile({"_index": []}),
                shell,
                self._idoc(),
                out,
                findings=findings,
            )
            # Exactly ONE TOC field survives - the preserved structural one; the
            # block deferred to it instead of emitting a duplicate.
            self.assertEqual(_toc_instr_count(out), 1)
            self.assertTrue(any(f.check == "toc_deferred" for f in findings))
            self.assertFalse(any(f.check == "block_degraded" for f in findings))

    def test_toc_block_is_not_suppressed_by_a_caption_index_only_shell(self):
        # A shell whose only strong TOC is a table-of-figures (TOC \c) must NOT
        # suppress the requested outline TOC: the block defers only to a real
        # OUTLINE TOC, so here it authors its own.
        import tempfile

        from docx import Document as _Doc

        with tempfile.TemporaryDirectory() as td:
            shell = Path(td) / "shell.docx"
            d = Document()
            _add_toc_field(d, instr='TOC \\h \\c "Figure"')  # caption index only
            d.add_paragraph("Body heading", style="Heading 1")
            d.save(shell)
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(
                _docx_profile({"_index": []}),
                shell,
                self._idoc(),
                out,
                findings=findings,
            )
            self.assertFalse(any(f.check == "toc_deferred" for f in findings))
            # An outline TOC now exists in the output (alongside the caption index).
            self.assertTrue(structure.is_outline_toc_present(_Doc(out)))

    def test_native_toc_generation_is_idempotent(self):
        import hashlib
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._bare_shell(Path(td))
            a, b = Path(td) / "a.docx", Path(td) / "b.docx"
            docx_generate.generate(
                _docx_profile({"_index": []}), shell, self._idoc(), a
            )
            docx_generate.generate(
                _docx_profile({"_index": []}), shell, self._idoc(), b
            )
            self.assertEqual(
                hashlib.sha256(a.read_bytes()).hexdigest(),
                hashlib.sha256(b.read_bytes()).hexdigest(),
            )


# ---------------------------------------------------------------------------
# Native docx charts (ir.Chart -> a DrawingML c:chartSpace part, inline data)
# ---------------------------------------------------------------------------
class NativeDocxChartTest(unittest.TestCase):
    """A Chart block is authored as a REAL Word chart (a c:chartSpace part an inline
    w:drawing references), not flattened to text: valid part, inline cached data (NO
    embedded workbook), theme-colored, byte-idempotent, unknown-type fallback, empty
    degrades, pie truncation surfaced."""

    def _shell(self, tmp_path):
        shell = tmp_path / "shell.docx"
        d = Document()
        d.add_paragraph("x", style="Heading 1")
        d.save(shell)
        return shell

    def _chart_parts(self, path):
        import zipfile

        with zipfile.ZipFile(path) as z:
            return {
                n: z.read(n).decode("utf-8")
                for n in z.namelist()
                if n.startswith("word/charts/chart") and n.endswith(".xml")
            }

    def _gen(self, td, idoc, *, sink=None):
        shell = self._shell(td)
        out = td / "out.docx"
        docx_generate.generate(
            _docx_profile({"_index": []}), shell, idoc, out, findings=sink
        )
        return out

    def test_chart_blocks_become_native_chart_parts(self):
        import tempfile
        import zipfile

        idoc = ir.IntermediateDocument(
            blocks=[
                ir.Heading(level=1, runs=[{"t": "Ricavi"}]),
                ir.Chart(
                    chart_type="bar",
                    title="Ricavi",
                    categories=["Q1", "Q2", "Q3"],
                    series=[
                        {"name": "A", "values": [1, 2, 3]},
                        {"name": "B", "values": [3, 2, 1]},
                    ],
                ),
                ir.Chart(
                    chart_type="pie",
                    title="Quota",
                    categories=["X", "Y"],
                    series=[{"name": "S", "values": [60, 40]}],
                ),
            ]
        )
        with tempfile.TemporaryDirectory() as t:
            td = Path(t)
            sink: list[Finding] = []
            out = self._gen(td, idoc, sink=sink)
            parts = self._chart_parts(out)
            self.assertEqual(len(parts), 2, "two native chart parts expected")
            joined = "".join(parts.values())
            for xml in parts.values():
                self.assertIn("<c:chartSpace", xml)
            self.assertIn("<c:barChart", joined)
            self.assertIn("<c:pieChart", joined)
            self.assertIn("<c:numCache", joined)  # data cached INLINE
            self.assertIn("<c:v>3</c:v>", joined)  # a series value round-tripped
            # The inline-data chart embeds NO workbook (that is what keeps it
            # deterministic) and the body references the chart part via a drawing.
            with zipfile.ZipFile(out) as z:
                self.assertFalse(
                    [n for n in z.namelist() if "embeddings" in n],
                    "inline-data chart must not embed a workbook",
                )
                body = z.read("word/document.xml").decode("utf-8")
            self.assertIn("c:chart", body)
            self.assertFalse(
                any(f.check == "block_degraded" and "chart" in f.message for f in sink)
            )

    def test_chart_is_byte_idempotent(self):
        import tempfile

        idoc = ir.IntermediateDocument(
            blocks=[
                ir.Chart(
                    chart_type="bar",
                    categories=["Q1", "Q2"],
                    series=[{"name": "A", "values": [1, 2]}],
                )
            ]
        )
        with tempfile.TemporaryDirectory() as t:
            td = Path(t)
            shell = self._shell(td)
            a, b = td / "a.docx", td / "b.docx"
            docx_generate.generate(_docx_profile({"_index": []}), shell, idoc, a)
            docx_generate.generate(_docx_profile({"_index": []}), shell, idoc, b)
            self.assertEqual(
                a.read_bytes(), b.read_bytes(), "native-chart docx not byte-idempotent"
            )

    def test_unknown_type_falls_back_and_empty_degrades(self):
        import tempfile

        idoc = ir.IntermediateDocument(
            blocks=[
                ir.Chart(
                    chart_type="nonsense",
                    categories=["A", "B"],
                    series=[{"name": "S", "values": [1, 2]}],
                ),
                ir.Chart(chart_type="bar", categories=[], series=[]),  # empty
            ]
        )
        with tempfile.TemporaryDirectory() as t:
            td = Path(t)
            sink: list[Finding] = []
            out = self._gen(td, idoc, sink=sink)
            # the unknown type still produces ONE native chart part (fallback)...
            self.assertEqual(len(self._chart_parts(out)), 1)
            self.assertTrue(any(f.check == "chart_type_fallback" for f in sink))
            # ...and the empty chart degrades loudly (never a silent drop).
            self.assertTrue(
                any(f.check == "block_degraded" and "chart" in f.message for f in sink)
            )

    def test_multi_series_pie_warns(self):
        import tempfile

        idoc = ir.IntermediateDocument(
            blocks=[
                ir.Chart(
                    chart_type="pie",
                    categories=["A", "B"],
                    series=[
                        {"name": "First", "values": [1, 2]},
                        {"name": "Second", "values": [3, 4]},
                    ],
                )
            ]
        )
        with tempfile.TemporaryDirectory() as t:
            td = Path(t)
            sink: list[Finding] = []
            self._gen(td, idoc, sink=sink)
            self.assertTrue(any(f.check == "chart_series_truncated" for f in sink))

    def test_docpr_ids_unique_across_image_and_chart(self):
        # An image and a chart in the same document must get DISTINCT wp:docPr ids
        # (a chart-only counter used to collide with the image's id -> Word rejects).
        import tempfile

        from PIL import Image as PILImage

        with tempfile.TemporaryDirectory() as t:
            td = Path(t)
            png = td / "fig.png"
            PILImage.new("RGB", (32, 16), (20, 40, 80)).save(png)
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Image(src=str(png)),
                    ir.Chart(
                        chart_type="bar",
                        categories=["A", "B"],
                        series=[{"name": "S", "values": [1, 2]}],
                    ),
                ]
            )
            out = self._gen(td, idoc)
            doc = Document(out)
            ids = [
                el.get("id")
                for el in doc.element.iter(qn("wp:docPr"))
                if el.get("id") is not None
            ]
            self.assertGreaterEqual(len(ids), 2, "expected an image + a chart docPr")
            self.assertEqual(len(ids), len(set(ids)), f"duplicate wp:docPr ids: {ids}")

    def test_non_finite_and_ragged_values_are_safe(self):
        # inf/nan must not crash (int(inf) would); a series shorter than the category
        # axis must be padded so the value ptCount matches the categories (else Word
        # drops the trailing ones). The chart still generates with no crash.
        import tempfile

        idoc = ir.IntermediateDocument(
            blocks=[
                ir.Chart(
                    chart_type="bar",
                    categories=["A", "B", "C", "D"],
                    series=[{"name": "S", "values": [10, float("inf"), 12]}],  # ragged
                )
            ]
        )
        with tempfile.TemporaryDirectory() as t:
            td = Path(t)
            sink: list[Finding] = []
            out = self._gen(td, idoc, sink=sink)
            parts = self._chart_parts(out)
            self.assertEqual(len(parts), 1)
            xml = next(iter(parts.values()))
            # value ptCount matches the 4 categories (short series padded to 4)...
            self.assertIn(
                '<c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="4"/>',
                xml,
            )
            # ...inf was dropped (rendered as a gap, not "inf"/crash).
            self.assertNotIn("inf", xml.lower())
            self.assertFalse(
                any(f.check == "block_degraded" and "chart" in f.message for f in sink)
            )


# ---------------------------------------------------------------------------
# Native docx SmartArt (ir.SmartArt -> a brand-styled table, not flattened)
# ---------------------------------------------------------------------------
class NativeDocxSmartArtTest(unittest.TestCase):
    """SmartArt is authored as a native brand table: a process is a single ROW (one
    cell per step), a list a single COLUMN (one row per node); a node's children are
    inlined so nothing is lost. An empty diagram degrades loudly."""

    def _shell(self, tp):
        shell = tp / "shell.docx"
        d = Document()
        d.add_paragraph("x", style="Heading 1")
        d.save(shell)
        return shell

    def test_process_and_list_become_brand_tables(self):
        import tempfile

        idoc = ir.IntermediateDocument(
            blocks=[
                ir.SmartArt(
                    diagram="process",
                    nodes=[{"text": "Plan"}, {"text": "Build"}, {"text": "Ship"}],
                ),
                ir.SmartArt(
                    diagram="list",
                    nodes=[
                        {"text": "Quality", "children": [{"text": "tests"}]},
                        {"text": "Speed"},
                    ],
                ),
            ]
        )
        with tempfile.TemporaryDirectory() as t:
            tp = Path(t)
            out = tp / "out.docx"
            sink: list[Finding] = []
            docx_generate.generate(
                _docx_profile({"_index": []}), self._shell(tp), idoc, out, findings=sink
            )
            tables = Document(out).tables
            self.assertEqual(len(tables), 2)
            proc, lst = tables[0], tables[1]
            # process -> one row, one cell per step
            self.assertEqual(len(proc.rows), 1)
            self.assertEqual(len(proc.columns), 3)
            self.assertEqual(proc.rows[0].cells[0].text, "Plan")
            # list -> one column, one row per node; the child is inlined (not lost)
            self.assertEqual(len(lst.columns), 1)
            self.assertEqual(len(lst.rows), 2)
            self.assertIn("tests", lst.rows[0].cells[0].text)
            self.assertFalse(
                any(
                    f.check == "block_degraded" and "smartart" in f.message
                    for f in sink
                )
            )

    def test_empty_smartart_degrades(self):
        import tempfile

        idoc = ir.IntermediateDocument(
            blocks=[ir.SmartArt(diagram="process", nodes=[])]
        )
        with tempfile.TemporaryDirectory() as t:
            tp = Path(t)
            out = tp / "out.docx"
            sink: list[Finding] = []
            docx_generate.generate(
                _docx_profile({"_index": []}), self._shell(tp), idoc, out, findings=sink
            )
            self.assertEqual(len(Document(out).tables), 0)
            self.assertTrue(
                any(
                    f.check == "block_degraded" and "smartart" in f.message
                    for f in sink
                )
            )


# ---------------------------------------------------------------------------
# M2 - nested lists
# ---------------------------------------------------------------------------
class NestedListTest(unittest.TestCase):
    def _shell(self, tmp_path):
        shell = tmp_path / "shell.docx"
        Document().save(shell)
        return shell

    def test_nested_items_appear_in_output(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _docx_profile({"_index": []})
            block = ir.ListBlock(
                ordered=False,
                items=[
                    ir.ListItem(
                        runs=[{"t": "Top1"}],
                        level=0,
                        items=[
                            ir.ListItem(runs=[{"t": "Nested1a"}], level=1),
                            ir.ListItem(runs=[{"t": "Nested1b"}], level=1),
                        ],
                    ),
                    ir.ListItem(runs=[{"t": "Top2"}], level=0),
                ],
            )
            docx_generate.generate(
                prof, shell, ir.IntermediateDocument(blocks=[block]), out
            )
            text = "\n".join(p.text for p in Document(out).paragraphs)
            for expected in ("Top1", "Nested1a", "Nested1b", "Top2"):
                self.assertIn(expected, text)

    def test_resolver_threads_item_level(self):
        prof = _docx_profile(
            {
                "_index": ["list.bullet.1", "list.bullet.2"],
                "list.bullet.1": {
                    "resolver": {"type": "named_style", "style_name": "List Bullet"}
                },
                "list.bullet.2": {
                    "resolver": {"type": "named_style", "style_name": "List Bullet 2"}
                },
            }
        )
        r = ProfileResolver(prof)
        block = ir.ListBlock(ordered=False, items=[])
        lvl0 = ir.ListItem(level=0)
        lvl1 = ir.ListItem(level=1)
        self.assertEqual(r.resolve_list_item(block, lvl0).role_id, "list.bullet.1")
        self.assertEqual(r.resolve_list_item(block, lvl1).role_id, "list.bullet.2")


# ---------------------------------------------------------------------------
# M4 - table spans
# ---------------------------------------------------------------------------
class TableSpanTest(unittest.TestCase):
    def test_colspan_honored(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = Path(td) / "shell.docx"
            Document().save(shell)
            out = Path(td) / "out.docx"
            prof = _docx_profile({"_index": []})
            block = ir.Table(
                columns=[{"t": "A"}, {"t": "B"}, {"t": "C"}],
                rows=[
                    [
                        ir.TableCell(runs=[{"t": "X"}], colspan=2),
                        ir.TableCell(runs=[{"t": "Y"}]),
                    ],
                    [
                        ir.TableCell(runs=[{"t": "1"}]),
                        ir.TableCell(runs=[{"t": "2"}]),
                        ir.TableCell(runs=[{"t": "3"}]),
                    ],
                ],
            )
            docx_generate.generate(
                prof, shell, ir.IntermediateDocument(blocks=[block]), out
            )
            t = Document(out).tables[-1]
            self.assertEqual(len(t.columns), 3)  # span-expanded width, not cell count
            # The spanned 'X' cell occupies columns 0 and 1; 'Y' lands in column 2,
            # not shifted left into a phantom cell.
            self.assertEqual(t.cell(1, 0).text, "X")
            self.assertEqual(t.cell(1, 2).text, "Y")


# ---------------------------------------------------------------------------
# M3 - cover SDT
# ---------------------------------------------------------------------------
class CoverSdtTest(unittest.TestCase):
    def _shell_with_sdt_cover(self, tmp_path):
        """A shell whose cover title is a block-level w:sdt, then a TOC, then body."""
        shell = tmp_path / "shell.docx"
        doc = Document()
        body = doc.element.body
        sectpr = body.find(w("sectPr"))
        # Build a block-level sdt with an alias 'Titolo' and placeholder text.
        sdt = OxmlElement("w:sdt")
        sdtPr = OxmlElement("w:sdtPr")
        alias = OxmlElement("w:alias")
        alias.set(w("val"), "Titolo")
        sdtPr.append(alias)
        sdt.append(sdtPr)
        sdtContent = OxmlElement("w:sdtContent")
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "Insert title here"
        r.append(t)
        p.append(r)
        sdtContent.append(p)
        sdt.append(sdtContent)
        body.insert(list(body).index(sectpr), sdt)
        # a TOC field after the cover
        _add_toc_field(doc)
        doc.add_paragraph("Body Heading", style="Heading 1")
        doc.save(shell)
        return shell

    def test_title_fills_cover_sdt_before_toc(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_sdt_cover(Path(td))
            out = Path(td) / "out.docx"
            prof = _docx_profile({"_index": []})
            idoc = ir.IntermediateDocument(
                blocks=[ir.Heading(level=1, runs=[{"t": "Section"}])],
                cover=ir.Cover(title=[{"t": "My Brand Title"}]),
            )
            docx_generate.generate(prof, shell, idoc, out)
            gen = Document(out)
            # title is in the SDT (a text node), not appended after the TOC.
            all_text = "".join(t.text or "" for t in gen.element.body.iter(w("t")))
            self.assertIn("My Brand Title", all_text)
            self.assertNotIn("Insert title here", all_text)
            # The SDT carrying the title sits in the cover region (before the TOC).
            cls = structure.classify_body_children(gen)
            children = list(gen.element.body)
            sdt_region = None
            for c in cls:
                el = children[c["index"]]
                if _local_name(el.tag) == "sdt":
                    txt = "".join(t.text or "" for t in el.iter(w("t")))
                    if "My Brand Title" in txt:
                        sdt_region = c["region"]
            self.assertEqual(sdt_region, "cover")

    def test_no_cover_anchor_appends_before_toc_with_finding(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = Path(td) / "shell.docx"
            doc = Document()
            _add_toc_field(doc)
            doc.add_paragraph("Body Heading", style="Heading 1")
            doc.save(shell)
            out = Path(td) / "out.docx"
            prof = _docx_profile({"_index": []})
            idoc = ir.IntermediateDocument(
                blocks=[], cover=ir.Cover(title=[{"t": "Fallback Title"}])
            )
            findings: list[Finding] = []
            docx_generate.generate(prof, shell, idoc, out, findings=findings)
            gen = Document(out)
            cls = structure.classify_body_children(gen)
            children = list(gen.element.body)
            # The appended title paragraph must NOT be in the toc/after-toc region.
            title_region = None
            for c in cls:
                el = children[c["index"]]
                if _local_name(el.tag) == "p":
                    txt = "".join(t.text or "" for t in el.iter(w("t")))
                    if "Fallback Title" in txt:
                        title_region = c["region"]
            self.assertEqual(title_region, "cover")
            self.assertTrue(any(f.check == "cover_degraded" for f in findings))


# ---------------------------------------------------------------------------
# arch-3 - intra-profile consistency (no shell I/O)
# ---------------------------------------------------------------------------
class IntraProfileConsistencyTest(unittest.TestCase):
    def test_fabricated_pptx_layout_is_caught(self):
        prof = schema.build_envelope("pptx", {"name": "deck"})
        prof["surface"] = {"pptx": {"layouts": {"Cover": {}, "Full Image Only": {}}}}
        prof["roles"] = {
            "_index": ["cover.title"],
            "cover.title": {
                "resolver": {
                    "type": "placeholder",
                    "layout": "Title Slide",
                    "ph_idx": 0,
                }
            },
        }
        problems = schema.validate(prof)
        self.assertTrue(any("Title Slide" in p and "layouts" in p for p in problems))

    def test_fabricated_xlsx_named_range_is_caught(self):
        prof = schema.build_envelope("xlsx", {"name": "model"})
        prof["surface"] = {"xlsx": {"named_regions": {"data_region": {}}}}
        prof["roles"] = {
            "_index": ["title"],
            "title": {"resolver": {"type": "named_range", "name": "ghost_range"}},
        }
        problems = schema.validate(prof)
        self.assertTrue(any("ghost_range" in p for p in problems))

    def test_consistent_profile_has_no_consistency_problem(self):
        prof = schema.build_envelope("pptx", {"name": "deck"})
        prof["surface"] = {"pptx": {"layouts": {"Cover": {}}}}
        prof["roles"] = {
            "_index": ["cover.title"],
            "cover.title": {
                "resolver": {"type": "placeholder", "layout": "Cover", "ph_idx": 0}
            },
        }
        problems = schema.validate(prof)
        self.assertFalse(any("layouts" in p for p in problems))


# ---------------------------------------------------------------------------
# arch-5 - component/section expansion
# ---------------------------------------------------------------------------
class ComponentExpansionTest(unittest.TestCase):
    def test_undefined_component_is_rejected(self):
        prof = _docx_profile({"_index": []})
        doc = ir.IntermediateDocument(blocks=[ir.Component(ref="ghost")])
        with self.assertRaises(ir_components.ComponentExpansionError):
            ir_components.expand_components(doc, prof)

    def test_defined_component_expands_to_primitives(self):
        prof = _docx_profile({"_index": []})
        prof["components"] = {
            "intro": {
                "blocks": [
                    {"type": "heading", "level": 1, "text": "Intro"},
                    {"type": "paragraph", "text": "body"},
                ]
            }
        }
        doc = ir.IntermediateDocument(blocks=[ir.Component(ref="intro")])
        expanded = ir_components.expand_components(doc, prof)
        self.assertEqual([b.TYPE for b in expanded.blocks], ["heading", "paragraph"])


# ---------------------------------------------------------------------------
# comprehension-driven DOCX reconciliation (model-free: a hand-authored
# comprehension exercises the same code path the model would drive). Proves the
# Napoleon fixes: multi-slot cover FILL in place, no duplicate title, stale
# caption-index REMOVE, and the destructive-action floor.
# ---------------------------------------------------------------------------
def _add_multi_paragraph_index(doc, instr):
    """Append a caption index whose field span covers SEVERAL paragraphs.

    The ``begin``/``instrText`` sit in the first paragraph, a styled entry in the
    middle, and the closing ``end`` fldChar in a final (otherwise empty) paragraph
    - the real-template shape that the TOC-span field-extension must cover so a
    body clear cannot sever the field.
    """
    p1 = doc.add_paragraph()
    r = p1.add_run()
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    r._r.append(fb)
    r2 = p1.add_run()
    it = OxmlElement("w:instrText")
    it.text = instr
    r2._r.append(it)
    r3 = p1.add_run()
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    r3._r.append(fs)
    p1.add_run("entry one .... 1")
    doc.add_paragraph("entry two .... 2")  # a styled index entry
    p3 = doc.add_paragraph()  # closing end fldChar paragraph
    re = p3.add_run()
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    re._r.append(fe)
    return p1, p3


def _present_comp(prof, comp, *, confidence=0.9):
    """Stamp a PRESENT, sha-current comprehension onto ``prof`` (bypassing merge)."""
    prof.setdefault("provenance", {}).setdefault("shell", {})["sha256"] = "shellsha"
    block = schema.empty_comprehension()
    block["status"] = schema.ComprehensionStatus.PRESENT.value
    block["source_shell_sha256"] = "shellsha"
    block["confidence"] = confidence
    block.update(comp)
    prof["comprehension"] = block
    return prof


class ComprehensionReconcileTest(unittest.TestCase):
    def _build_shell(self, td):
        """A shell with a multi-slot cover (2 placeholder paras + 1 SDT), a real
        outline TOC, and a stacked caption index with a multi-paragraph span."""
        shell = Path(td) / "shell.docx"
        doc = Document()
        doc.add_paragraph("Brand Header")  # para slot (leave)
        title_p = doc.add_paragraph("Insert title here")  # para slot (title)
        # a subtitle SDT (alias) inserted into the cover region, right AFTER the
        # title paragraph (a raw body.append would land after the final sectPr).
        sdt = OxmlElement("w:sdt")
        sdtPr = OxmlElement("w:sdtPr")
        alias = OxmlElement("w:alias")
        alias.set(qn("w:val"), "Subtitle")
        sdtPr.append(alias)
        sdt.append(sdtPr)
        sdtContent = OxmlElement("w:sdtContent")
        sp = OxmlElement("w:p")
        sr = OxmlElement("w:r")
        st = OxmlElement("w:t")
        st.text = "Subtitle prompt"
        sr.append(st)
        sp.append(sr)
        sdtContent.append(sp)
        sdt.append(sdtContent)
        title_p._p.addnext(sdt)
        _add_toc_field(doc, 'TOC \\o "1-3" \\h')  # outline TOC
        doc.add_paragraph("Index of Tables")  # index heading
        _add_multi_paragraph_index(doc, 'TOC \\h \\c "Tabella"')  # caption index
        doc.add_paragraph("Body Heading", style="Heading 1")
        doc.save(shell)
        return shell

    def _profile_for(self, shell):
        from brandkit.formats.docx import extract as docx_extract
        import tempfile
        import os
        import json

        # Extract a real profile so the surfaced inventory ids are authentic.
        with tempfile.TemporaryDirectory() as ed:
            old = Path.cwd()
            os.chdir(ed)
            try:
                pj = docx_extract.extract(shell, "recon", scope="project")
                return json.loads(Path(pj).read_text())
            finally:
                os.chdir(old)

    def test_multi_slot_cover_fill_no_duplicate_and_index_removed(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            prof = self._profile_for(shell)
            s = prof["surface"]["docx"]
            anchor_ids = [a["id"] for a in s["cover_anchors"]]
            [f["id"] for f in s["fields"]]
            # The multi-slot cover surfaced 3 slots; the caption index surfaced.
            self.assertEqual(len(anchor_ids), 3, anchor_ids)
            tabella = next(f["id"] for f in s["fields"] if f["seq_id"] == "Tabella")
            title_anchor = next(
                a["id"]
                for a in s["cover_anchors"]
                if "Insert title here" in (a.get("placeholder") or "")
            )
            sub_anchor = next(
                a["id"]
                for a in s["cover_anchors"]
                if "Subtitle" in (a.get("placeholder") or "")
            )
            comp = {
                "cover_slots": {
                    title_anchor: {
                        "binds_to": "title",
                        "fill_rule": "in_place",
                        "demo_value": "Insert title here",
                    },
                    sub_anchor: {
                        "binds_to": "subtitle",
                        "fill_rule": "in_place",
                        "demo_value": "Subtitle prompt",
                    },
                },
                "conventions": {
                    "indexes": [
                        {
                            "index_ref": tabella,
                            "seq_id": "Tabella",
                            "reconcile": "clear",
                        }
                    ],
                    "sections": [],
                },
                "demo_classification": {
                    "regions": [{"region_ref": f"region.{tabella}", "verdict": "demo"}]
                },
            }
            _present_comp(prof, comp)
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(
                prof,
                shell,
                ir.IntermediateDocument(
                    blocks=[ir.Heading(level=1, runs=[{"t": "Real Section"}])],
                    cover=ir.Cover(
                        title=[{"t": "Real Title"}], subtitle=[{"t": "Real Subtitle"}]
                    ),
                ),
                out,
                findings=findings,
            )
            gen = Document(out)
            text = "".join(t.text or "" for t in gen.element.body.iter(w("t")))
            # Cover filled in place (both slots), no leftover prompt.
            self.assertIn("Real Title", text)
            self.assertIn("Real Subtitle", text)
            self.assertNotIn("Insert title here", text)
            self.assertNotIn("Subtitle prompt", text)
            # Cover fields bound to Word core properties must render the new
            # values when LibreOffice/Word refreshes document fields.
            self.assertEqual(gen.core_properties.title, "Real Title")
            self.assertEqual(gen.core_properties.subject, "Real Subtitle")
            # No duplicate title appended (there WAS a title-bearing slot).
            self.assertEqual(text.count("Real Title"), 1)
            self.assertFalse(any(f.check == "cover_degraded" for f in findings))
            # Stale caption index removed (entries + heading gone).
            self.assertNotIn("entry one", text)
            self.assertNotIn("entry two", text)
            self.assertNotIn("Index of Tables", text)
            # The destructive floor did not flag a net loss (verdict corroborated).
            self.assertFalse(any(f.check == "no_net_structure_loss" for f in findings))

    def test_index_clear_downgraded_when_not_corroborated(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            prof = self._profile_for(shell)
            s = prof["surface"]["docx"]
            tabella = next(f["id"] for f in s["fields"] if f["seq_id"] == "Tabella")
            # reconcile=clear but NO demo verdict and the content HAS a captionable
            # item -> the destructive floor downgrades to KEEP + WARNING.
            comp = {
                "conventions": {
                    "indexes": [
                        {
                            "index_ref": tabella,
                            "seq_id": "Tabella",
                            "reconcile": "clear",
                        }
                    ],
                    "sections": [],
                },
                "demo_classification": {"regions": []},
            }
            _present_comp(prof, comp)
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(
                prof,
                shell,
                ir.IntermediateDocument(
                    blocks=[
                        ir.Caption(
                            runs=[{"t": "Table 1. A real caption"}], target="table"
                        )
                    ],
                    cover=None,
                ),
                out,
                findings=findings,
            )
            gen = Document(out)
            text = "".join(t.text or "" for t in gen.element.body.iter(w("t")))
            # The index was KEPT (not corroborated) and a WARNING was recorded.
            self.assertIn("entry one", text)
            self.assertTrue(any(f.check == "index_clear_downgraded" for f in findings))

    def test_index_clear_downgraded_when_below_confidence_floor(self):
        # A CLEAR the model is NOT confident about (< the destructive floor) is
        # downgraded to KEEP + WARNING EVEN WHEN the verdict is otherwise corroborated
        # (here the index's region is tagged demo). This proves the confidence floor
        # gates the docx index clear independently, mirroring the docx/pptx cover
        # reconcilers - the SAME confidence now yields the SAME behavior per format.
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            prof = self._profile_for(shell)
            s = prof["surface"]["docx"]
            tabella = next(f["id"] for f in s["fields"] if f["seq_id"] == "Tabella")
            comp = {
                "conventions": {
                    "indexes": [
                        {
                            "index_ref": tabella,
                            "seq_id": "Tabella",
                            "reconcile": "clear",
                        }
                    ],
                    "sections": [],
                },
                # The verdict WOULD corroborate (region tagged demo); only the low
                # confidence blocks the clear.
                "demo_classification": {
                    "regions": [{"region_ref": f"region.{tabella}", "verdict": "demo"}]
                },
            }
            _present_comp(prof, comp, confidence=0.3)  # below the floor (0.5)
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(
                prof,
                shell,
                ir.IntermediateDocument(
                    blocks=[ir.Heading(level=1, runs=[{"t": "Real Section"}])],
                    cover=None,
                ),
                out,
                findings=findings,
            )
            gen = Document(out)
            text = "".join(t.text or "" for t in gen.element.body.iter(w("t")))
            # The index was KEPT (low confidence) and a WARNING was recorded.
            self.assertIn("entry one", text)
            downgrades = [f for f in findings if f.check == "index_clear_downgraded"]
            self.assertTrue(downgrades)
            self.assertIn("confidence 0.30", downgrades[0].message)
            # Nothing was removed -> no net-loss ERROR.
            self.assertFalse(
                any(
                    f.check == "no_net_structure_loss"
                    and f.severity == schema.Severity.ERROR.value
                    for f in findings
                )
            )

    def test_high_confidence_corroborated_index_clear_still_fires(self):
        # The positive control for the floor: a corroborated CLEAR whose confidence
        # clears the floor is still honored (the index is removed), so the floor only
        # blocks the LOW-confidence case, never the confident one.
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            prof = self._profile_for(shell)
            s = prof["surface"]["docx"]
            tabella = next(f["id"] for f in s["fields"] if f["seq_id"] == "Tabella")
            comp = {
                "conventions": {
                    "indexes": [
                        {
                            "index_ref": tabella,
                            "seq_id": "Tabella",
                            "reconcile": "clear",
                        }
                    ],
                    "sections": [],
                },
                "demo_classification": {
                    "regions": [{"region_ref": f"region.{tabella}", "verdict": "demo"}]
                },
            }
            _present_comp(prof, comp, confidence=0.9)  # clears the floor
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(
                prof,
                shell,
                ir.IntermediateDocument(
                    blocks=[ir.Heading(level=1, runs=[{"t": "Real Section"}])],
                    cover=None,
                ),
                out,
                findings=findings,
            )
            gen = Document(out)
            text = "".join(t.text or "" for t in gen.element.body.iter(w("t")))
            # The stale caption index was removed (entries + heading gone).
            self.assertNotIn("entry one", text)
            self.assertNotIn("Index of Tables", text)
            self.assertFalse(any(f.check == "index_clear_downgraded" for f in findings))
            self.assertFalse(any(f.check == "no_net_structure_loss" for f in findings))

    def test_index_clear_corroborated_by_absent_captionables_only(self):
        # The OTHER corroboration path: a CLEAR with NO demo verdict still fires when
        # the new content has no captionables to feed the index (absence corroborates
        # removal), distinct from the demo-verdict path above.
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            prof = self._profile_for(shell)
            s = prof["surface"]["docx"]
            tabella = next(f["id"] for f in s["fields"] if f["seq_id"] == "Tabella")
            comp = {
                "conventions": {
                    "indexes": [
                        {
                            "index_ref": tabella,
                            "seq_id": "Tabella",
                            "reconcile": "clear",
                        }
                    ],
                    "sections": [],
                },
                "demo_classification": {"regions": []},  # NO demo verdict
            }
            _present_comp(prof, comp, confidence=0.9)
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(
                prof,
                shell,
                # Content has NO captionable (no caption/table/image) -> nothing to
                # feed the table index, which corroborates the clear by absence.
                ir.IntermediateDocument(
                    blocks=[ir.Heading(level=1, runs=[{"t": "Real Section"}])],
                    cover=None,
                ),
                out,
                findings=findings,
            )
            gen = Document(out)
            text = "".join(t.text or "" for t in gen.element.body.iter(w("t")))
            self.assertNotIn("entry one", text)
            self.assertNotIn("Index of Tables", text)
            self.assertFalse(any(f.check == "index_clear_downgraded" for f in findings))

    def test_absent_comprehension_uses_deterministic_path(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            prof = self._profile_for(shell)
            # comprehension absent (default) -> today's single-title behaviour.
            self.assertEqual(prof["comprehension"]["status"], "absent")
            Path(td) / "out.docx"
            findings: list[Finding] = []
            cleared = covermod.compose_cover(
                Document(shell), ir.Cover(title=[{"t": "X"}]), prof, findings=findings
            )
            self.assertEqual(
                cleared, set()
            )  # deterministic path returns no cleared refs


# ---------------------------------------------------------------------------
# ORPHAN-INDEX-HEADING fix - when a caption index is removed, its introducing
# heading (a non-field paragraph immediately above the span, identified
# STRUCTURALLY by toc-region membership) is removed too, leaving no orphan
# heading. Reproduces the Napoleon 'Indice delle Tabelle'/'figure' case and the
# index-shift regression (the cover fill inserts a paragraph, so the index refs
# must be resolved before the cover shifts the positions).
# ---------------------------------------------------------------------------
class OrphanIndexHeadingTest(unittest.TestCase):
    def _build_shell(self, td):
        """Cover (a fillable SDT title so compose_cover INSERTS and shifts indices)
        + outline TOC + an index heading + a caption index + a post-index body
        heading. Mirrors the real template_chat front matter."""
        shell = Path(td) / "shell.docx"
        doc = Document()
        # A fillable cover title SDT in the cover region (before the TOC). Filling
        # it appends a paragraph -> shifts every subsequent top-level child index.
        title_p = doc.add_paragraph("Insert title here")
        sdt = OxmlElement("w:sdt")
        sdtPr = OxmlElement("w:sdtPr")
        alias = OxmlElement("w:alias")
        alias.set(qn("w:val"), "Titolo")
        sdtPr.append(alias)
        sdt.append(sdtPr)
        sc = OxmlElement("w:sdtContent")
        sp = OxmlElement("w:p")
        sr = OxmlElement("w:r")
        st = OxmlElement("w:t")
        st.text = "Insert title here"
        sr.append(st)
        sp.append(sr)
        sc.append(sp)
        sdt.append(sc)
        title_p._p.addnext(sdt)
        _add_toc_field(doc, 'TOC \\o "1-3" \\h')  # outline TOC
        doc.add_paragraph("Indice delle Tabelle")  # introducing heading
        _add_multi_paragraph_index(doc, 'TOC \\h \\c "Tabella"')  # caption index
        doc.add_paragraph("Real Body Heading", style="Heading 1")  # body (post-index)
        doc.save(shell)
        return shell

    def _build_shell_with_empty_body_artifacts(self, td):
        """Like the real front matter, but with empty TOC/body break artifacts.

        The blank paragraphs between the outline TOC and caption index are
        initially folded into the TOC span and preserved. Once the caption index is
        reconciled away they become leading body blanks; the empty section-break
        paragraphs are demo-body artifacts that used to force blank rendered pages
        before generated content.
        """
        shell = Path(td) / "shell-empty-artifacts.docx"
        doc = Document()
        doc.add_paragraph("Insert title here")
        _add_toc_field(doc, 'TOC \\o "1-3" \\h')
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("Indice delle Tabelle")
        _add_multi_paragraph_index(doc, 'TOC \\h \\c "Tabella"')
        _append_intermediate_sectpr(doc)
        _append_intermediate_sectpr(doc)
        doc.add_paragraph("Template Body Heading", style="Heading 1")
        doc.save(shell)
        return shell

    def _profile_for(self, shell):
        from brandkit.formats.docx import extract as docx_extract
        import tempfile
        import os
        import json

        with tempfile.TemporaryDirectory() as ed:
            old = Path.cwd()
            os.chdir(ed)
            try:
                pj = docx_extract.extract(shell, "orphan", scope="project")
                return json.loads(Path(pj).read_text())
            finally:
                os.chdir(old)

    def test_clearing_caption_index_removes_its_introducing_heading_in_place(self):
        # Unit-level: the structural remove pulls in the toc-region heading above
        # the field span, and the post-index body heading is NEVER touched.
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            doc = Document(shell)
            tabella = next(
                f["id"]
                for f in structure.inventory_fields(doc)
                if f["seq_id"] == "Tabella"
            )
            idxs = structure._index_field_remove_indices(doc, tabella)
            children = list(doc.element.body)
            removed_texts = {
                "".join(t.text or "" for t in children[i].iter(w("t"))).strip()
                for i in idxs
            }
            self.assertIn("Indice delle Tabelle", removed_texts)  # orphan heading in
            self.assertNotIn("Real Body Heading", removed_texts)  # body heading out
            structure.remove_index_fields(doc, [tabella])
            kept = "\n".join(p.text for p in doc.paragraphs)
            self.assertNotIn("Indice delle Tabelle", kept)
            self.assertIn("Real Body Heading", kept)

    def test_orphan_heading_gone_through_full_generate_after_cover_shift(self):
        # End-to-end regression: the cover fill inserts a paragraph (index shift);
        # the index reconcile must resolve refs BEFORE that shift, so the orphan
        # 'Indice delle Tabelle' heading is removed and idempotency holds.
        import tempfile
        import hashlib

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            prof = self._profile_for(shell)
            s = prof["surface"]["docx"]
            tabella = next(f["id"] for f in s["fields"] if f["seq_id"] == "Tabella")
            comp = {
                "conventions": {
                    "indexes": [
                        {
                            "index_ref": tabella,
                            "seq_id": "Tabella",
                            "reconcile": "clear",
                        }
                    ],
                    "sections": [],
                },
                "demo_classification": {
                    "regions": [{"region_ref": f"region.{tabella}", "verdict": "demo"}]
                },
            }
            _present_comp(prof, comp)
            idoc = ir.IntermediateDocument(
                blocks=[ir.Heading(level=1, runs=[{"t": "New Real Section"}])],
                cover=ir.Cover(title=[{"t": "Filled Title"}]),
            )
            out1 = Path(td) / "o1.docx"
            findings: list[Finding] = []
            docx_generate.generate(prof, shell, idoc, out1, findings=findings)
            gen = Document(out1)
            body_text = "\n".join(p.text for p in gen.paragraphs)
            alltext = "".join(t.text or "" for t in gen.element.body.iter(w("t")))
            # The orphan introducing heading is gone, the stale entries are gone.
            self.assertNotIn("Indice delle Tabelle", body_text)
            self.assertNotIn("entry one", body_text)
            self.assertNotIn("entry two", body_text)
            # The outline TOC + cover fill + new content survived; no false net-loss.
            self.assertIn("Filled Title", alltext)
            self.assertIn("New Real Section", body_text)
            self.assertFalse(
                any(
                    f.check == "no_net_structure_loss"
                    and f.severity == schema.Severity.ERROR.value
                    for f in findings
                )
            )
            # Idempotent: generate twice -> byte identical.
            out2 = Path(td) / "o2.docx"
            docx_generate.generate(prof, shell, idoc, out2)
            self.assertEqual(
                hashlib.sha256(out1.read_bytes()).hexdigest(),
                hashlib.sha256(out2.read_bytes()).hexdigest(),
            )

    def test_full_generate_prunes_leading_empty_body_artifacts_after_index_clear(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell_with_empty_body_artifacts(td)
            prof = self._profile_for(shell)
            s = prof["surface"]["docx"]
            tabella = next(f["id"] for f in s["fields"] if f["seq_id"] == "Tabella")
            comp = {
                "conventions": {
                    "indexes": [
                        {
                            "index_ref": tabella,
                            "seq_id": "Tabella",
                            "reconcile": "clear",
                        }
                    ],
                    "sections": [],
                },
                "demo_classification": {
                    "regions": [{"region_ref": f"region.{tabella}", "verdict": "demo"}]
                },
            }
            _present_comp(prof, comp)
            idoc = ir.IntermediateDocument(
                blocks=[ir.Heading(level=1, runs=[{"t": "New Real Section"}])],
                cover=ir.Cover(title=[{"t": "Filled Title"}]),
            )
            out = Path(td) / "out.docx"
            docx_generate.generate(prof, shell, idoc, out)

            gen = Document(out)
            children = list(gen.element.body)
            leading_body = []
            for c in structure.classify_body_children(gen):
                if c["region"] != "body" or c["is_sectpr"]:
                    continue
                el = children[c["index"]]
                text = "".join(t.text or "" for t in el.iter(w("t"))).strip()
                leading_body.append((c, text))
                if text:
                    break

            self.assertTrue(leading_body, "generated body disappeared")
            self.assertEqual(leading_body[0][1], "New Real Section")
            self.assertFalse(
                any(c.get("holds_sectpr") for c, _ in leading_body[:-1]),
                "empty body section breaks survived before generated content",
            )

    def test_demo_marker_literals_are_gone(self):
        # The DEMO_MARKERS literal list is removed; detect_demo_region no longer
        # matches any global phrase (instruction_markers is always empty) and works
        # purely structurally off the first body Heading-1.
        self.assertFalse(hasattr(structure, "DEMO_MARKERS"))
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._build_shell(td)
            demo = structure.detect_demo_region(Document(shell))
            self.assertEqual(demo["instruction_markers"], [])
            self.assertTrue(demo["present"])
            # start_text is THIS template's own captured body heading, not a literal.
            self.assertEqual(demo["start_text"], "Real Body Heading")


# ---------------------------------------------------------------------------
# Deterministic cover subtitle fill (style-identified, comprehension absent)
# ---------------------------------------------------------------------------
class DeterministicCoverSubtitleFillTest(unittest.TestCase):
    """Without comprehension, the cover fill now places the authored subtitle into
    a slot identified by its resolved ``cover.subtitle`` style (correct-by-style,
    not by guessing the template's placeholder text), so the output no longer shows
    the template's stale demo subtitle. Extra fields stay the comprehension path's
    job and are still surfaced as unplaced."""

    def test_subtitle_slot_filled_in_place_by_style(self):
        doc = Document()
        doc.add_paragraph("{{title}}", style="Title")  # cover title slot
        doc.add_paragraph("Stale template subtitle", style="Subtitle")  # cover subtitle
        doc.add_paragraph("Body", style="Heading 1")  # body -> closes the cover region

        roles = docx_roles.infer_roles(doc)
        self.assertIn("cover.subtitle", roles)  # resolved from the builtin Subtitle
        prof = _docx_profile(roles)
        cover = ir.Cover(
            fields={"title": "Napoleon", "subtitle": "Full history", "doc_id": "X-1"}
        )
        sink: list[Finding] = []
        covermod.compose_cover(doc, cover, prof, findings=sink)

        texts = [p.text for p in doc.paragraphs]
        # the authored subtitle replaced the template's stale subtitle IN PLACE
        self.assertIn("Full history", texts)
        self.assertNotIn("Stale template subtitle", "\n".join(texts))
        filled = next(p for p in doc.paragraphs if p.text == "Full history")
        self.assertEqual(filled.style.name, "Subtitle")  # cover.subtitle style kept

        # subtitle is no longer reported unplaced; the extra field still is.
        notes = "\n".join(f.message for f in sink if f.check == "cover_degraded")
        self.assertNotIn("subtitle", notes)
        self.assertIn("doc_id", notes)

    def test_no_subtitle_style_leaves_subtitle_unplaced(self):
        # A template with no subtitle style at all (no builtin Subtitle, no custom
        # subtitle-named style) cannot place the subtitle deterministically: it is
        # surfaced as unplaced (the comprehension path fills the full cover).
        doc = Document()
        title = doc.add_paragraph("{{title}}", style="Title")  # noqa: F841
        doc.add_paragraph("Body", style="Heading 1")
        roles = {
            "_index": ["cover.title"],
            "cover.title": {
                "resolver": {"type": "named_style", "style_id": "Title"},
            },
        }
        prof = _docx_profile(roles)
        cover = ir.Cover(fields={"title": "Napoleon", "subtitle": "Full history"})
        sink: list[Finding] = []
        covermod.compose_cover(doc, cover, prof, findings=sink)
        notes = "\n".join(f.message for f in sink if f.check == "cover_degraded")
        self.assertIn("subtitle", notes)  # honestly surfaced, never silently dropped


if __name__ == "__main__":
    unittest.main()
