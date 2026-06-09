# SPDX-License-Identifier: MIT
"""Adversarial integration tests proving the brand guarantee holds end-to-end.

Two layers:

1. A synthetic unit regression (`TocSdtCollisionTest`) for a defect surfaced while
   integrating the parallel-group fixes: on a shell whose Table-of-Contents is a
   block-level ``w:sdt`` (``docPartGallery='Table of Contents'``) whose cached
   entry text contains title-like words, the M3 cover-fill used to misidentify
   the TOC content control as the cover title, dump the user's title into it, and
   blank every TOC entry. `_sdt_is_title` now excludes any SDT carrying a strong
   TOC marker, so the cover fill never touches the TOC.

2. A real-template end-to-end test (`RealTemplateIntegrationTest`) that runs
   extract -> generate on a real structured ``.docx`` (kept OUT of the repo - it
   lives under ``/tmp/agents_extract`` and the brand-kit is written to a temp dir,
   never into the source tree). It asserts, from the generated ``document.xml``
   region order and content, that every reviewed critical/major leak is fixed:
     - C1: no stale shell *body* text leaks into the generated body;
     - M3: the user's cover title lands on the cover (before the real TOC field),
           and the preserved TOC content control is NOT wiped;
     - M5: every intermediate ``w:sectPr`` (multi-section geometry) survives;
     - M1: an unhandled block emits no empty Normal paragraph + a degraded finding;
     - M2: nested list items are all present.
   It is skipped automatically when the proprietary template is absent, so the
   suite stays green on any machine.
"""

from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from brandkit.formats.docx import cover as covermod
from brandkit.formats.docx import extract as docx_extract
from brandkit.formats.docx import generate as docx_generate
from brandkit.formats.docx import structure
from brandkit.formats.docx.structure import _local_name, w
from brandkit.ir import model as ir
from brandkit.profile import store
from brandkit.qa.model import Finding

REAL_TEMPLATE = Path(
    "/tmp/agents_extract/agents/word-exporter/assets/template_extraction.docx"
)


def _ptext(el) -> str:
    return "".join(t.text for t in el.iter(w("t")) if t.text)


def _docx_profile(roles=None):
    from brandkit.profile import schema

    prof = schema.build_envelope("docx", {"name": "t"})
    prof["surface"] = {"docx": {}}
    prof["roles"] = roles or {"_index": []}
    return prof


# ---------------------------------------------------------------------------
# Synthetic unit regression for the TOC-SDT / cover-title collision.
# ---------------------------------------------------------------------------
class TocSdtCollisionTest(unittest.TestCase):
    def _toc_sdt(self):
        """A block-level w:sdt that IS the Table of Contents (docPartGallery),
        whose cached inner text contains a title-like word ('Titolo'/'Sommario')."""
        sdt = OxmlElement("w:sdt")
        sdtPr = OxmlElement("w:sdtPr")
        gallery_holder = OxmlElement("w:docPartObj")
        gallery = OxmlElement("w:docPartGallery")
        gallery.set(w("val"), "Table of Contents")
        gallery_holder.append(gallery)
        sdtPr.append(gallery_holder)
        sdt.append(sdtPr)
        content = OxmlElement("w:sdtContent")
        # A TOC heading paragraph whose text contains a title token, then a TOC field.
        ph = OxmlElement("w:p")
        pr = OxmlElement("w:r")
        pt = OxmlElement("w:t")
        pt.text = "Sommario / Titolo del documento"
        pr.append(pt)
        ph.append(pr)
        content.append(ph)
        fp = OxmlElement("w:p")
        r1 = OxmlElement("w:r")
        fb = OxmlElement("w:fldChar")
        fb.set(qn("w:fldCharType"), "begin")
        r1.append(fb)
        fp.append(r1)
        r2 = OxmlElement("w:r")
        it = OxmlElement("w:instrText")
        it.text = 'TOC \\o "1-3" \\h \\z \\u'
        r2.append(it)
        fp.append(r2)
        r3 = OxmlElement("w:r")
        fs = OxmlElement("w:fldChar")
        fs.set(qn("w:fldCharType"), "separate")
        r3.append(fs)
        fp.append(r3)
        r4 = OxmlElement("w:r")
        et = OxmlElement("w:t")
        et.text = "1. Esempio di titolo cached entry ....... 3"
        r4.append(et)
        fp.append(r4)
        r5 = OxmlElement("w:r")
        fe = OxmlElement("w:fldChar")
        fe.set(qn("w:fldCharType"), "end")
        r5.append(fe)
        fp.append(r5)
        content.append(fp)
        sdt.append(content)
        return sdt

    def test_toc_sdt_is_not_classified_as_title(self):
        """A Table-of-Contents SDT must never be flagged as the cover title."""
        sdt = self._toc_sdt()
        self.assertTrue(structure._element_holds_strong_toc(sdt))
        self.assertFalse(covermod._sdt_is_title(sdt))

    def test_cover_fill_does_not_wipe_the_toc_sdt(self):
        """Regression: filling the cover title must leave the TOC SDT intact and
        must NOT write the title into the TOC content control."""
        with tempfile.TemporaryDirectory() as td:
            shell = Path(td) / "shell.docx"
            doc = Document()
            body = doc.element.body
            sectpr = body.find(w("sectPr"))
            body.insert(list(body).index(sectpr), self._toc_sdt())
            doc.add_paragraph("Body Heading", style="Heading 1")
            doc.save(shell)

            # Count the TOC SDT's non-empty text runs before generation.
            pre = Document(shell)
            toc_sdt_pre = next(
                el
                for el in pre.element.body
                if _local_name(el.tag) == "sdt"
                and structure._element_holds_strong_toc(el)
            )
            pre_runs = len(
                [t for t in toc_sdt_pre.iter(w("t")) if (t.text or "").strip()]
            )
            self.assertGreater(pre_runs, 0)

            out = Path(td) / "out.docx"
            prof = _docx_profile({"_index": []})
            findings: list[Finding] = []
            idoc = ir.IntermediateDocument(
                blocks=[ir.Heading(level=1, runs=[{"t": "Section"}])],
                cover=ir.Cover(title=[{"t": "USER COVER TITLE 9Z"}]),
            )
            docx_generate.generate(prof, shell, idoc, out, findings=findings)

            gen = Document(out)
            children = list(gen.element.body)
            toc_sdt = next(
                el
                for el in children
                if _local_name(el.tag) == "sdt"
                and structure._element_holds_strong_toc(el)
            )
            # The user's title was NOT dumped into the TOC SDT.
            self.assertNotIn("USER COVER TITLE 9Z", _ptext(toc_sdt))
            # The TOC SDT's field survived, and its visible cache now reflects
            # generated headings instead of stale template entries.
            self.assertIn("Section", _ptext(toc_sdt))
            self.assertNotIn("Esempio di titolo cached entry", _ptext(toc_sdt))
            self.assertTrue(
                any(
                    (d.text or "").strip().startswith("TOC")
                    for d in toc_sdt.iter(w("instrText"))
                )
            )
            # The title still landed somewhere in the document (on the cover),
            # never inside the TOC, with a degraded finding (no real cover anchor).
            title_p = None
            for c in structure.classify_body_children(gen):
                el = children[c["index"]]
                if _local_name(el.tag) == "p" and "USER COVER TITLE 9Z" in _ptext(el):
                    title_p = c["region"]
            self.assertEqual(title_p, "cover")
            self.assertTrue(any(f.check == "cover_degraded" for f in findings))


# ---------------------------------------------------------------------------
# Real-template end-to-end integration (skipped when the template is absent).
# ---------------------------------------------------------------------------
@unittest.skipUnless(REAL_TEMPLATE.is_file(), "real structured template not available")
class RealTemplateIntegrationTest(unittest.TestCase):
    def _extract_and_generate(self, td: Path):
        # Extract into a temp project store (NEVER the repo).
        docx_extract.extract(REAL_TEMPLATE, "rt", scope="project", cwd=td)
        loaded = store.load_profile("rt", "project", cwd=td)
        iid = {
            "cover": {"title": [{"t": "INTEGRATION COVER XZ9"}]},
            "blocks": [
                {"type": "heading", "level": 1, "runs": [{"t": "First Heading QQ1"}]},
                {"type": "paragraph", "runs": [{"t": "Body alpha UNIQUEBODY42."}]},
                {
                    "type": "list",
                    "ordered": False,
                    "items": [
                        {
                            "text": "Top NL_TOP",
                            "items": [
                                {"text": "Child NL_CHILD_A"},
                                {"text": "Child NL_CHILD_B"},
                            ],
                        },
                        {"text": "Second NL_TOP2"},
                    ],
                },
                {"type": "image", "alt": "unhandled image block"},
                {"type": "heading", "level": 2, "runs": [{"t": "Second Heading QQ2"}]},
                {"type": "paragraph", "runs": [{"t": "Closing ENDBODY77."}]},
            ],
        }
        out = td / "out.docx"
        findings: list[Finding] = []
        docx_generate.generate(
            loaded.profile,
            loaded.shell_path,
            ir.parse_idoc(iid),
            out,
            findings=findings,
        )
        return out, findings

    def test_end_to_end_region_order_and_leaks(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            out, findings = self._extract_and_generate(td)
            self.assertIsNone(
                zipfile.ZipFile(out).testzip(), "generated docx is corrupt"
            )

            gen = Document(out)
            children = list(gen.element.body)
            classes = structure.classify_body_children(gen)
            region_of = {c["index"]: c["region"] for c in classes}

            def region_text(region):
                return "\n".join(
                    _ptext(children[c["index"]])
                    for c in classes
                    if c["region"] == region
                )

            body_text = region_text("body")
            toc_text = region_text("toc")

            # --- C1: no stale shell BODY text leaks (TOC field cache is allowed) ---
            self.assertNotIn("Esempio di titolo di Primo livello", body_text)
            self.assertNotIn("perspiciatis unde omnis iste natus", body_text)
            # our authored body IS present
            self.assertIn("UNIQUEBODY42", body_text)
            self.assertIn("ENDBODY77", body_text)

            # --- M3: cover title on the cover, before the real TOC field ---
            title_idx = next(
                (
                    i
                    for i, el in enumerate(children)
                    if _local_name(el.tag) == "p"
                    and "INTEGRATION COVER XZ9" in _ptext(el)
                ),
                None,
            )
            self.assertIsNotNone(title_idx, "cover title not found as a paragraph")
            self.assertEqual(region_of.get(title_idx), "cover")
            # the title must NOT sit inside any TOC content control
            self.assertNotIn("INTEGRATION COVER XZ9", toc_text)
            first_toc_field = next(
                (
                    i
                    for i, el in enumerate(children)
                    if any(
                        (d.text or "").strip().startswith("TOC")
                        for d in el.iter(w("instrText"))
                    )
                ),
                None,
            )
            self.assertIsNotNone(
                first_toc_field, "TOC field disappeared (M3 regression)"
            )
            self.assertLess(title_idx, first_toc_field, "title landed after the TOC")
            # the preserved TOC content survived (not blanked by the cover fill)
            self.assertIn("Indice delle figure", toc_text)

            # --- M5: every intermediate sectPr (4 sections) preserved ---
            n_intermediate = sum(1 for c in classes if c.get("holds_sectpr"))
            n_final = sum(1 for c in classes if c["is_sectpr"])
            self.assertEqual(n_intermediate + n_final, 4, "a section break was dropped")

            # --- M1: unhandled image block -> no empty body para + degraded finding ---
            empty_body = 0
            for c in classes:
                if c["region"] != "body" or c.get("holds_sectpr"):
                    continue
                el = children[c["index"]]
                if _local_name(el.tag) == "p" and _ptext(el).strip() == "":
                    empty_body += 1
            self.assertEqual(
                empty_body, 0, "unhandled block injected an empty paragraph"
            )
            self.assertTrue(
                any(
                    f.check == "block_degraded" and "image" in f.message
                    for f in findings
                )
            )

            # --- M2: every nested list item present, in order ---
            order = []
            for tok in ["NL_TOP", "NL_CHILD_A", "NL_CHILD_B", "NL_TOP2"]:
                idx = next(
                    (i for i, el in enumerate(children) if tok in _ptext(el)), None
                )
                self.assertIsNotNone(idx, f"nested list item {tok} dropped")
                order.append(idx)
            self.assertEqual(order, sorted(order), "nested list items out of order")

            # sanity: regions appear in cover -> toc -> body order
            seen = [r for r in (region_of[i] for i in sorted(region_of)) if r]
            # collapse consecutive duplicates
            collapsed = [r for j, r in enumerate(seen) if j == 0 or r != seen[j - 1]]
            self.assertEqual(
                collapsed, ["cover", "toc", "body"], f"region order wrong: {collapsed}"
            )

    def test_verify_passes_on_real_profile_with_shell(self):
        """The real extracted profile must VERIFY clean (resolver_targets_exist
        passes because every role's style genuinely exists in the shell)."""
        from brandkit.qa.gate import run_qa

        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            docx_extract.extract(REAL_TEMPLATE, "rt", scope="project", cwd=td)
            loaded = store.load_profile("rt", "project", cwd=td)
            report = run_qa(None, loaded.profile, qa="fast", shell=loaded.shell_path)
            self.assertTrue(report.passed, [f.message for f in report.findings])
            # no resolver_targets_exist ERROR on a real, faithful profile
            self.assertFalse(
                any(
                    f.check == "resolver_targets_exist" and f.severity == "ERROR"
                    for f in report.findings
                )
            )


class TableD2VerifyTest(unittest.TestCase):
    """Cluster D2: a profile that captures the template's OWN table facts (tblLook +
    style + cell margins) must VERIFY clean - the table-target check surfaces no ERROR
    when the captured facts are shell-backed."""

    def _template(self, td):
        from docx.oxml import OxmlElement as _Ox

        template = td / "tbl_template.docx"
        d = Document()
        styles = d.styles.element
        st = _Ox("w:style")
        st.set(qn("w:type"), "table")
        st.set(qn("w:styleId"), "AcmeTable")
        st.set(qn("w:customStyle"), "1")
        nm = _Ox("w:name")
        nm.set(qn("w:val"), "Acme Table")
        st.append(nm)
        styles.append(st)
        for _ in range(3):
            t = d.add_table(rows=2, cols=2)
            t.style = "Acme Table"
            tblpr = t._tbl.tblPr
            tblpr.find(qn("w:tblLook")).set(qn("w:val"), "01E0")
            cm = _Ox("w:tblCellMar")
            for side in ("left", "right"):
                el = _Ox(f"w:{side}")
                el.set(qn("w:w"), "80")
                el.set(qn("w:type"), "dxa")
                cm.append(el)
            tblpr.append(cm)
        d.save(template)
        return template

    def test_verify_does_not_error_on_observed_table_facts(self):
        from brandkit.qa.gate import run_qa

        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            template = self._template(td)
            docx_extract.extract(template, "tbl", scope="project", cwd=td)
            loaded = store.load_profile("tbl", "project", cwd=td)
            # the captured table facts are present (the template declared them)
            self.assertIn("table", loaded.profile.get("theme", {}))
            report = run_qa(None, loaded.profile, qa="fast", shell=loaded.shell_path)
            self.assertFalse(
                any(
                    f.check == "appearance_table_targets" and f.severity == "ERROR"
                    for f in report.findings
                ),
                [f.message for f in report.findings],
            )


if __name__ == "__main__":
    unittest.main()
