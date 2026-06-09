# SPDX-License-Identifier: MIT
"""B2 tests: cross-run regression findings (a pure side artifact on top of B1).

``compute_regression_findings`` diffs THIS run's findings against prior
same-shell ``generation_report.json`` digests and emits advisory
``regression.recurred`` / ``regression.reintroduced`` findings, keyed STRICTLY on
the ``(check, location)`` multiset (never the ``message`` body), partitioned by
``shell_sha256``, ignoring the volatile ``generated_at``. They are folded into
the QAReport (and so the persisted report) of the run that detects them, which
makes recurrence self-recording for the next run.

These prove:
  * a finding seen in run1 + run2 -> ``regression.recurred`` with recurred_runs=2;
  * present, gone, present -> ``regression.reintroduced``;
  * a DIFFERENT-shell prior report does not contribute;
  * recurrence is advisory only -- never an ERROR, never in DEFAULT_L0_INVARIANTS,
    never flips a verdict to failed;
  * a differing ``generated_at`` between otherwise-identical findings still
    counts as recurred (the timestamp is ignored);
  * the CLI self-records the regression finding end-to-end across two generates.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

from brandkit.cli import main
from brandkit.profile import schema
from brandkit.qa import report as vreport
from brandkit.qa.model import Finding, QAReport


def _report_dict(
    *,
    shell_sha256: str,
    findings: list[tuple[str, str | None]],
    generated_at: str = "2026-06-09T00:00:00Z",
) -> dict:
    """A minimal serialized generation_report.json digest (B1 schema subset)."""
    return {
        "schema_version": vreport.REPORT_SCHEMA_VERSION,
        "kind": "docx",
        "shell_sha256": shell_sha256,
        "findings": [
            {
                "check": c,
                "severity": "WARNING",
                "message": "irrelevant body",
                "location": loc,
            }
            for (c, loc) in findings
        ],
        "generated_at": generated_at,
    }


def _qareport(findings: list[tuple[str, str | None]]) -> QAReport:
    return QAReport(
        verdict=schema.VerificationStatus.PASSED.value,
        findings=[
            Finding(c, schema.Severity.WARNING.value, "live body", location=loc)
            for (c, loc) in findings
        ],
    )


class RegressionFindingsUnitTest(unittest.TestCase):
    # -- test_recurred_across_two_runs -------------------------------------
    def test_recurred_across_two_runs(self) -> None:
        prior = [
            _report_dict(shell_sha256="abc", findings=[("style_fallback", "Heading 1")])
        ]
        report = _qareport([("style_fallback", "Heading 1")])
        out = vreport.compute_regression_findings({}, report, prior)
        recurred = [f for f in out if f.check == vreport.REGRESSION_RECURRED]
        self.assertEqual(len(recurred), 1)
        self.assertEqual(recurred[0].location, "Heading 1")
        self.assertEqual(recurred[0].severity, schema.Severity.INFO.value)
        # recurred_runs counts this run + the priors that show it -> 2.
        self.assertIn("recurred_runs=2", recurred[0].message)

    def test_recurred_runs_counts_all_priors(self) -> None:
        priors = [
            _report_dict(
                shell_sha256="abc",
                findings=[("style_fallback", "H1")],
                generated_at="2026-06-01T00:00:00Z",
            ),
            _report_dict(
                shell_sha256="abc",
                findings=[("style_fallback", "H1")],
                generated_at="2026-06-02T00:00:00Z",
            ),
        ]
        report = _qareport([("style_fallback", "H1")])
        out = vreport.compute_regression_findings({}, report, priors)
        recurred = [f for f in out if f.check == vreport.REGRESSION_RECURRED]
        self.assertEqual(len(recurred), 1)
        self.assertIn("recurred_runs=3", recurred[0].message)

    def test_brand_new_finding_is_not_recurred(self) -> None:
        prior = [_report_dict(shell_sha256="abc", findings=[("style_fallback", "H1")])]
        report = _qareport([("no_residual_template_text", "Body")])
        out = vreport.compute_regression_findings({}, report, prior)
        self.assertEqual(out, [])

    # -- test_reintroduced -------------------------------------------------
    def test_reintroduced(self) -> None:
        # present (oldest), gone (immediately-prior), present (this run).
        priors = [
            _report_dict(
                shell_sha256="abc",
                findings=[("style_fallback", "H1")],
                generated_at="2026-06-01T00:00:00Z",
            ),
            _report_dict(
                shell_sha256="abc",
                findings=[("other_check", "X")],
                generated_at="2026-06-02T00:00:00Z",
            ),
        ]
        report = _qareport([("style_fallback", "H1")])
        out = vreport.compute_regression_findings({}, report, priors)
        reintroduced = [f for f in out if f.check == vreport.REGRESSION_REINTRODUCED]
        self.assertEqual(len(reintroduced), 1)
        self.assertEqual(reintroduced[0].location, "H1")
        self.assertEqual(reintroduced[0].severity, schema.Severity.WARNING.value)
        # A reintroduced pair is ALSO recurred (it appears in an earlier run).
        recurred = [f for f in out if f.check == vreport.REGRESSION_RECURRED]
        self.assertEqual(len(recurred), 1)

    def test_recurred_in_immediately_prior_is_not_reintroduced(self) -> None:
        # present in the immediately-prior run -> recurred but NOT reintroduced.
        priors = [
            _report_dict(
                shell_sha256="abc",
                findings=[("style_fallback", "H1")],
                generated_at="2026-06-02T00:00:00Z",
            )
        ]
        report = _qareport([("style_fallback", "H1")])
        out = vreport.compute_regression_findings({}, report, priors)
        self.assertEqual(
            [f for f in out if f.check == vreport.REGRESSION_REINTRODUCED], []
        )
        self.assertEqual(
            len([f for f in out if f.check == vreport.REGRESSION_RECURRED]), 1
        )

    # -- test_regression_partitioned_by_shell ------------------------------
    def test_regression_partitioned_by_shell(self) -> None:
        # compute_regression_findings trusts the caller's partition: a prior list
        # that has already been filtered to the live shell yields a recurrence;
        # discover_prior_reports is what enforces the filter (tested below).
        same = [_report_dict(shell_sha256="abc", findings=[("style_fallback", "H1")])]
        report = _qareport([("style_fallback", "H1")])
        self.assertEqual(
            len(
                [
                    f
                    for f in vreport.compute_regression_findings({}, report, same)
                    if f.check == vreport.REGRESSION_RECURRED
                ]
            ),
            1,
        )

    def test_discover_partitions_by_shell(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            # Two sibling .visual dirs: one same-shell, one different-shell.
            same_dir = base / "a.docx.visual"
            diff_dir = base / "b.docx.visual"
            same_dir.mkdir()
            diff_dir.mkdir()
            (same_dir / vreport.REPORT_FILENAME).write_text(
                json.dumps(
                    _report_dict(
                        shell_sha256="LIVE", findings=[("style_fallback", "H1")]
                    )
                ),
                encoding="utf-8",
            )
            (diff_dir / vreport.REPORT_FILENAME).write_text(
                json.dumps(
                    _report_dict(
                        shell_sha256="OTHER", findings=[("style_fallback", "H1")]
                    )
                ),
                encoding="utf-8",
            )
            # New run writes into c.docx.visual; discover priors for shell LIVE.
            new_dir = base / "c.docx.visual"
            priors = vreport.discover_prior_reports(
                new_dir,
                shell_sha256="LIVE",
                exclude=new_dir / vreport.REPORT_FILENAME,
            )
            self.assertEqual(len(priors), 1)
            self.assertEqual(priors[0]["shell_sha256"], "LIVE")

    def test_discover_excludes_own_report(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            own = base / "out.docx.visual"
            own.mkdir()
            (own / vreport.REPORT_FILENAME).write_text(
                json.dumps(
                    _report_dict(
                        shell_sha256="LIVE", findings=[("style_fallback", "H1")]
                    )
                ),
                encoding="utf-8",
            )
            # Re-generate to the SAME output: its own stale report must not count.
            priors = vreport.discover_prior_reports(
                own, shell_sha256="LIVE", exclude=own / vreport.REPORT_FILENAME
            )
            self.assertEqual(priors, [])

    def test_discover_none_shell_matches_nothing(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            sib = base / "a.docx.visual"
            sib.mkdir()
            (sib / vreport.REPORT_FILENAME).write_text(
                json.dumps(_report_dict(shell_sha256="LIVE", findings=[("x", "y")])),
                encoding="utf-8",
            )
            self.assertEqual(
                vreport.discover_prior_reports(
                    base / "c.docx.visual", shell_sha256=None
                ),
                [],
            )

    # -- test_regression_advisory_only -------------------------------------
    def test_regression_advisory_only(self) -> None:
        # Regression ids are never in the L0 invariant list and are never ERROR.
        self.assertNotIn(vreport.REGRESSION_RECURRED, schema.DEFAULT_L0_INVARIANTS)
        self.assertNotIn(vreport.REGRESSION_REINTRODUCED, schema.DEFAULT_L0_INVARIANTS)
        priors = [
            _report_dict(
                shell_sha256="abc",
                findings=[("style_fallback", "H1")],
                generated_at="2026-06-01T00:00:00Z",
            ),
            _report_dict(
                shell_sha256="abc",
                findings=[("other", "Z")],
                generated_at="2026-06-02T00:00:00Z",
            ),
        ]
        report = _qareport([("style_fallback", "H1")])
        out = vreport.compute_regression_findings({}, report, priors)
        self.assertTrue(out)  # both recurred + reintroduced fired
        for f in out:
            self.assertNotEqual(f.severity, schema.Severity.ERROR.value)
        # A QAReport carrying only these advisory findings still passes (no ERROR).
        self.assertTrue(QAReport(verdict="x", findings=out).passed)

    # -- test_regression_ignores_timestamp ---------------------------------
    def test_regression_ignores_timestamp(self) -> None:
        # Identical (check, location) findings, DIFFERENT generated_at -> recurred.
        prior = [
            _report_dict(
                shell_sha256="abc",
                findings=[("style_fallback", "H1")],
                generated_at="1999-01-01T00:00:00Z",
            )
        ]
        report = _qareport([("style_fallback", "H1")])
        out = vreport.compute_regression_findings({}, report, prior)
        self.assertEqual(
            len([f for f in out if f.check == vreport.REGRESSION_RECURRED]), 1
        )

    def test_key_is_check_location_not_message(self) -> None:
        # Same (check, location), but the prior's message body differs entirely.
        # Keying on message would miss this recurrence; keying on (check, location)
        # catches it. Different LOCATION must NOT recur.
        prior = [_report_dict(shell_sha256="abc", findings=[("style_fallback", "H1")])]
        same_loc = _qareport([("style_fallback", "H1")])
        diff_loc = _qareport([("style_fallback", "H2")])
        self.assertEqual(
            len(
                [
                    f
                    for f in vreport.compute_regression_findings({}, same_loc, prior)
                    if f.check == vreport.REGRESSION_RECURRED
                ]
            ),
            1,
        )
        self.assertEqual(
            vreport.compute_regression_findings({}, diff_loc, prior),
            [],
        )

    def test_no_priors_no_findings(self) -> None:
        self.assertEqual(
            vreport.compute_regression_findings({}, _qareport([("x", "y")]), []), []
        )


# ---------------------------------------------------------------------------
# CLI-level: the regression finding self-records end-to-end across two generates.
# ---------------------------------------------------------------------------
def _synthetic_template(path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    h1 = styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x5A, 0xAB)
    callout = styles.add_style("ACME Callout Info", WD_STYLE_TYPE.PARAGRAPH)
    callout.base_style = styles["Normal"]
    doc.add_paragraph("{{title}}", style="Title")
    doc.add_paragraph("Example first-level title", style="Heading 1")
    doc.add_paragraph("General instructions: replace this demo text.", style="Normal")
    doc.save(path)


# An idoc that re-emits the template's captured demo text ("Example first-level
# title"), so the deterministic gate raises the SAME ``no_residual_template_text``
# finding (check='no_residual_template_text', location=None) on EVERY run -- a
# stable recurrence signal across runs. (It is an ERROR, so both runs return 1;
# that is exactly what proves regression folding is purely ADDITIVE: it records
# the recurrence without changing the verdict the base findings already dictate.)
_IDOC = {
    "cover": {"title": "Quarterly Review"},
    "blocks": [
        {"type": "heading", "level": 1, "text": "Highlights"},
        {"type": "paragraph", "text": "Example first-level title"},
    ],
}


class RegressionFindingsCliTest(unittest.TestCase):
    def _run(self, args: list[str]) -> int:
        return main(args)

    def _extract(self, tmp_path: Path) -> None:
        template = tmp_path / "synthetic-template.docx"
        _synthetic_template(template)
        self.assertEqual(
            self._run(
                [
                    "extract",
                    "--name",
                    "acme",
                    "--template",
                    str(template),
                    "--scope",
                    "project",
                ]
            ),
            0,
        )

    def _generate(self, tmp_path: Path, out: Path) -> int:
        idoc = tmp_path / "idoc.json"
        idoc.write_text(json.dumps(_IDOC), encoding="utf-8")
        return self._run(
            [
                "generate",
                "--name",
                "acme",
                "--input",
                str(idoc),
                "--output",
                str(out),
                "--scope",
                "project",
                "--qa",
                "fast",
            ]
        )

    def _read_report(self, out: Path) -> dict:
        path = out.parent / (out.name + ".visual") / vreport.REPORT_FILENAME
        return json.loads(path.read_text(encoding="utf-8"))

    def test_regression_self_records_across_runs(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            old_cwd = Path.cwd()
            os.chdir(tmp_path)
            try:
                self._extract(tmp_path)

                # Run 1 -> distinct output (so its report is a sibling prior). The
                # base finding (no_residual_template_text) is an ERROR -> rc 1.
                out1 = tmp_path / "run1.docx"
                self.assertEqual(self._generate(tmp_path, out1), 1)
                r1 = self._read_report(out1)
                # No prior reports yet: run1 has no regression findings.
                self.assertEqual(
                    [f for f in r1["findings"] if f["check"].startswith("regression.")],
                    [],
                    "run1 emitted a regression finding with no prior",
                )
                base_checks = [f["check"] for f in r1["findings"]]
                self.assertIn("no_residual_template_text", base_checks)

                # Run 2 -> a different output so run1's report is discoverable.
                out2 = tmp_path / "run2.docx"
                self.assertEqual(self._generate(tmp_path, out2), 1)
                r2 = self._read_report(out2)
                recurred = [
                    f
                    for f in r2["findings"]
                    if f["check"] == vreport.REGRESSION_RECURRED
                ]
                self.assertTrue(
                    recurred,
                    "run2 did not self-record a regression.recurred finding",
                )
                # The recurred finding points at the SAME base check, carries the
                # count, and is advisory INFO -- never an ERROR.
                self.assertIn("recurred_runs=2", recurred[0]["message"])
                self.assertTrue(all(f["severity"] == "INFO" for f in recurred))
            finally:
                os.chdir(old_cwd)

    def test_regression_is_purely_additive_to_return_code(self) -> None:
        # The return code is governed ENTIRELY by the base findings (the ERROR
        # here); the folded-in advisory regression finding never changes it.
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            old_cwd = Path.cwd()
            os.chdir(tmp_path)
            try:
                self._extract(tmp_path)
                out1 = tmp_path / "run1.docx"
                out2 = tmp_path / "run2.docx"
                self.assertEqual(self._generate(tmp_path, out1), 1)
                # run2 folds in regression.recurred yet the rc is still the base 1.
                self.assertEqual(self._generate(tmp_path, out2), 1)
            finally:
                os.chdir(old_cwd)


if __name__ == "__main__":
    unittest.main()
