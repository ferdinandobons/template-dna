# SPDX-License-Identifier: MIT
"""Two-stage visual audit tests (no soffice required).

The L1 proxy tests feed SYNTHETIC PIL images (blank page, edge bleed, centered
content, landscape) so they assert the deterministic pixel proxies without ever
invoking ``soffice``. The wiring/degrade tests monkeypatch the renderer so the
gate's ``--qa fast|auto|deep|strict`` semantics and the clean CI degrade are
proven without external tools. One gated end-to-end test runs the real render
when the binaries are present (skipped in CI).
"""

from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
import unittest
from contextlib import redirect_stdout
from io import StringIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from PIL import Image, ImageDraw

from brandkit import doctor
from brandkit.profile import schema
from brandkit.qa import checks_deterministic
from brandkit.qa import gate
from brandkit.qa import visual as vqa
from brandkit.qa.model import Finding


def _real_docx(path: Path) -> Path:
    """Write a minimal but VALID .docx so the L0 docx check can open it as a zip.

    The wiring tests only care about the visual path; the file just has to be a
    real OOXML package the L0 layer can read without errors.
    """
    doc = Document()
    doc.add_paragraph("placeholder body")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Synthetic-image helpers (US-Letter-ish portrait at 100 DPI = 850x1100)
# ---------------------------------------------------------------------------
def _blank(width: int = 850, height: int = 1100) -> Image.Image:
    return Image.new("L", (width, height), 255)


def _centered_content(width: int = 850, height: int = 1100) -> Image.Image:
    """A page with a dark block well inside the margins (clean edges)."""
    img = _blank(width, height)
    draw = ImageDraw.Draw(img)
    # A large central block (luma 0 = ink) covering > the blank ink threshold.
    draw.rectangle([width // 4, height // 4, 3 * width // 4, 3 * height // 4], fill=0)
    return img


def _bottom_bleed(width: int = 850, height: int = 1100) -> Image.Image:
    """A near-blank page with a solid ink bar along the very bottom edge."""
    img = _blank(width, height)
    draw = ImageDraw.Draw(img)
    # A bar inside the bottom 1% of the page -> falls in the bottom margin band.
    draw.rectangle([0, height - max(1, height // 100), width, height], fill=0)
    return img


# ---------------------------------------------------------------------------
# §7.1 L1 proxies on synthetic PNGs
# ---------------------------------------------------------------------------
class L1ProxyTest(unittest.TestCase):
    def test_default_out_dir_keeps_extension_to_avoid_format_collisions(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            self.assertEqual(
                vqa.default_out_dir(base / "out.docx"), base / "out.docx.visual"
            )
            self.assertEqual(
                vqa.default_out_dir(base / "out.pptx"), base / "out.pptx.visual"
            )
            self.assertEqual(
                vqa.default_out_dir(base / "out.xlsx"), base / "out.xlsx.visual"
            )

    def test_blank_page_flagged(self) -> None:
        findings = vqa.check_blank_page(_blank(), page_index=0)
        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0].check, "visual.blank_page")
        self.assertEqual(findings[0].severity, schema.Severity.WARNING.value)
        self.assertEqual(findings[0].location, "page:1")

    def test_content_page_not_flagged(self) -> None:
        self.assertEqual(vqa.check_blank_page(_centered_content(), page_index=0), [])

    def test_edge_bleed_flagged(self) -> None:
        findings = vqa.check_edge_bleed(_bottom_bleed(), page_index=1)
        self.assertTrue(findings)
        self.assertTrue(all(f.check == "visual.edge_bleed" for f in findings))
        self.assertTrue(
            all(f.severity == schema.Severity.WARNING.value for f in findings)
        )
        self.assertTrue(any("bottom" in (f.location or "") for f in findings))
        self.assertTrue(any(f.location == "page:2:bottom" for f in findings))

    def test_centered_content_no_bleed(self) -> None:
        self.assertEqual(vqa.check_edge_bleed(_centered_content(), page_index=0), [])

    def test_landscape_dimensions_handled(self) -> None:
        # 1100x850 (landscape): a centered block must still read as clean edges,
        # proving the band is computed off the correct (per-side) dimension.
        landscape = _centered_content(width=1100, height=850)
        self.assertEqual(vqa.check_edge_bleed(landscape, page_index=0), [])
        self.assertEqual(vqa.check_blank_page(landscape, page_index=0), [])
        # And a landscape bottom bleed is still caught.
        bleed = _bottom_bleed(width=1100, height=850)
        findings = vqa.check_edge_bleed(bleed, page_index=0)
        self.assertTrue(any("bottom" in (f.location or "") for f in findings))

    def test_proxies_accept_path_and_image(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            png = Path(td) / "page-1.png"
            _blank().save(png)
            from_image = vqa.check_blank_page(_blank(), page_index=0)
            from_path = vqa.check_blank_page(png, page_index=0)
            self.assertEqual(len(from_image), len(from_path))
            self.assertEqual(from_image[0].check, from_path[0].check)
            self.assertEqual(from_image[0].location, from_path[0].location)

    def test_proxies_never_raise_on_bad_image(self) -> None:
        missing = Path("/nonexistent/page-1.png")
        self.assertEqual(vqa.check_blank_page(missing, page_index=0), [])
        self.assertEqual(vqa.check_edge_bleed(missing, page_index=0), [])
        with tempfile.TemporaryDirectory() as td:
            bogus = Path(td) / "page-1.png"
            bogus.write_text("not a png", encoding="utf-8")
            self.assertEqual(vqa.check_blank_page(bogus, page_index=0), [])
            self.assertEqual(vqa.check_edge_bleed(bogus, page_index=0), [])

    def test_run_visual_l1_concatenates(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            p1 = Path(td) / "page-1.png"
            p2 = Path(td) / "page-2.png"
            _blank().save(p1)  # -> blank_page on page 1
            _bottom_bleed().save(p2)  # -> edge_bleed on page 2
            findings = vqa.run_visual_l1([p1, p2])
            checks = {f.check for f in findings}
            self.assertIn("visual.blank_page", checks)
            self.assertIn("visual.edge_bleed", checks)

    def test_page_count_sane_flags_empty(self) -> None:
        findings = vqa.check_page_count_sane([])
        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0].check, "visual.no_pages")
        self.assertEqual(vqa.check_page_count_sane([_blank()]), [])


class RendererAvailabilityTest(unittest.TestCase):
    def setUp(self) -> None:
        # The renderer cache is a process-global; reset it so a leaked probe
        # result from another test (or collection order) cannot leak in.
        self._reset_renderer_cache()
        self.addCleanup(self._reset_renderer_cache)

    @staticmethod
    def _reset_renderer_cache() -> None:
        vqa._LAST_RENDERER_STATUS = None

    def test_renderers_available_caches_doctor_status(self) -> None:
        def fake_probe():
            return {
                "python_deps": {},
                "binaries": {"soffice": True, "pdftoppm": True},
                "binary_paths": {
                    "soffice": "/fake/soffice",
                    "pdftoppm": "/fake/pdftoppm",
                },
                "binary_errors": {},
                "visual_qa": True,
            }

        with patch.object(doctor, "probe", fake_probe):
            self.assertTrue(vqa.renderers_available())
            self.assertEqual(
                vqa.last_renderer_status()["binary_paths"]["soffice"], "/fake/soffice"
            )

    def test_doctor_prints_install_hints_for_missing_dependencies(self) -> None:
        def fake_probe():
            return {
                "python_deps": {
                    "docx": False,
                    "pptx": True,
                    "openpyxl": True,
                    "lxml": True,
                    "PIL": True,
                },
                "optional_python_deps": {"fitz": False},
                "binaries": {"soffice": False, "pdftoppm": False},
                "binary_paths": {"soffice": None, "pdftoppm": None},
                "binary_errors": {},
                "visual_qa": False,
                "ocr_binaries": {"tesseract": False},
                "ocr_binary_paths": {"tesseract": None},
                "ocr_binary_errors": {},
                "ocr_qa": False,
            }

        buf = StringIO()
        with patch.object(doctor, "probe", fake_probe), redirect_stdout(buf):
            doctor.print_report()

        out = buf.getvalue()
        self.assertIn("install:python:", out)
        self.assertIn("pip install -r requirements.txt", out)
        self.assertIn("install:soffice:", out)
        self.assertIn("libreoffice", out)
        self.assertIn("install:pdftoppm:", out)
        self.assertIn("poppler", out)
        self.assertIn("install:fitz:", out)
        self.assertIn("PyMuPDF", out)
        self.assertIn("ocr:tesseract:", out)
        self.assertIn("install:tesseract:", out)

    def test_probe_marks_visual_unavailable_when_binary_probe_fails(self) -> None:
        """PATH presence alone is not enough; broken renderers must degrade."""

        def fake_run(*args, **kwargs):
            return SimpleNamespace(returncode=134, stdout=b"", stderr=b"abort")

        with (
            patch.object(doctor.shutil, "which", lambda name: f"/fake/{name}"),
            patch.object(doctor, "_soffice_app_signature_error", lambda path: None),
            patch.object(subprocess, "run", fake_run),
        ):
            status = doctor.probe()

        self.assertFalse(status["binaries"]["soffice"])
        self.assertFalse(status["visual_qa"])

    def test_signature_invalid_but_functional_soffice_is_usable(self) -> None:
        """A macOS signature quibble is NON-fatal: if soffice actually RUNS, it is
        usable. The functional probe is authoritative, not codesign - a LibreOffice
        whose seal broke after an update/quarantine removal still renders."""
        calls = []

        def fake_signature_error(path):
            return (
                "LibreOffice.app signature invalid: bad signature"
                if path.endswith("soffice")
                else None
            )

        def fake_run(args, *unused_args, **unused_kwargs):
            calls.append(list(args))
            return SimpleNamespace(returncode=0, stdout=b"version ok", stderr=b"")

        with (
            patch.object(doctor.shutil, "which", lambda name: f"/fake/{name}"),
            patch.object(doctor, "_soffice_app_signature_error", fake_signature_error),
            patch.object(subprocess, "run", fake_run),
        ):
            status = doctor.probe(skip_visual_pipeline=True)

        # Usable despite the signature quibble, and it WAS launched (functional wins).
        self.assertTrue(status["binaries"]["soffice"])
        self.assertNotIn("soffice", status["binary_errors"])
        self.assertTrue(any(call and call[0].endswith("soffice") for call in calls))

    def test_signature_invalid_and_crashing_soffice_is_unusable_with_reason(
        self,
    ) -> None:
        """When soffice BOTH has a bad seal AND fails to run, it is unusable and the
        signature problem is surfaced as the reason."""

        def fake_signature_error(path):
            return (
                "LibreOffice.app signature invalid: bad signature"
                if path.endswith("soffice")
                else None
            )

        def fake_run(args, *unused_args, **unused_kwargs):
            if args[0].endswith("soffice"):
                return SimpleNamespace(returncode=134, stdout=b"", stderr=b"abort")
            return SimpleNamespace(returncode=0, stdout=b"version ok", stderr=b"")

        with (
            patch.object(doctor.shutil, "which", lambda name: f"/fake/{name}"),
            patch.object(doctor, "_soffice_app_signature_error", fake_signature_error),
            patch.object(subprocess, "run", fake_run),
        ):
            status = doctor.probe(skip_visual_pipeline=True)

        self.assertFalse(status["binaries"]["soffice"])
        self.assertIn("signature invalid", status["binary_errors"]["soffice"])

    def test_probe_marks_visual_unavailable_when_conversion_probe_fails(self) -> None:
        """Version commands can pass while headless conversion is unusable."""
        calls = []

        def fake_run(args, *unused_args, **unused_kwargs):
            calls.append(list(args))
            if "--convert-to" in args:
                return SimpleNamespace(returncode=134, stdout=b"", stderr=b"abort")
            return SimpleNamespace(returncode=0, stdout=b"version ok", stderr=b"")

        with (
            patch.object(doctor.shutil, "which", lambda name: f"/fake/{name}"),
            patch.object(doctor, "_soffice_app_signature_error", lambda path: None),
            patch.object(subprocess, "run", fake_run),
        ):
            status = doctor.probe()

        self.assertTrue(status["binaries"]["soffice"])
        self.assertTrue(status["binaries"]["pdftoppm"])
        self.assertFalse(status["visual_qa"])
        self.assertIn("soffice convert failed", status["binary_errors"]["visual_qa"])
        self.assertTrue(any("--convert-to" in call for call in calls))

    def test_probe_marks_visual_unavailable_when_conversion_times_out(self) -> None:
        def fake_run(args, *unused_args, **unused_kwargs):
            if "--convert-to" in args:
                raise subprocess.TimeoutExpired(args, 1)
            return SimpleNamespace(returncode=0, stdout=b"version ok", stderr=b"")

        with (
            patch.object(doctor.shutil, "which", lambda name: f"/fake/{name}"),
            patch.object(doctor, "_soffice_app_signature_error", lambda path: None),
            patch.object(subprocess, "run", fake_run),
        ):
            status = doctor.probe()

        self.assertFalse(status["visual_qa"])
        self.assertIn("timed out", status["binary_errors"]["visual_qa"])

    def test_conversion_probe_smoke_tests_docx_pptx_and_xlsx(self) -> None:
        convert_launches: list[list[str]] = []
        rasterized_pdfs: list[str] = []

        def fake_run(args, *unused_args, **unused_kwargs):
            if "--convert-to" in args:
                convert_launches.append(list(args))
                outdir = Path(args[args.index("--outdir") + 1])
                outdir.mkdir(parents=True, exist_ok=True)
                for raw in args[args.index("--outdir") + 2 :]:
                    document = Path(raw)
                    (outdir / f"{document.stem}.pdf").write_bytes(b"%PDF-1.4\n%%EOF\n")
                return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")
            if args and Path(args[0]).name == "pdftoppm":
                pdf = Path(args[-2])
                rasterized_pdfs.append(pdf.stem)
                prefix = Path(args[-1])
                prefix.parent.mkdir(parents=True, exist_ok=True)
                _centered_content().save(prefix.with_name(prefix.name + "-1.png"))
                return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")
            return SimpleNamespace(returncode=0, stdout=b"version ok", stderr=b"")

        with patch.object(subprocess, "run", fake_run):
            ok, error = doctor._probe_visual_pipeline(
                {"soffice": "/fake/soffice", "pdftoppm": "/fake/pdftoppm"}
            )

        self.assertTrue(ok, error)
        # ONE soffice launch converts all three probe documents (startup +
        # fresh-UserInstallation cost paid once)...
        self.assertEqual(1, len(convert_launches))
        launch = convert_launches[0]
        converted_suffixes = {
            Path(raw).suffix for raw in launch[launch.index("--outdir") + 2 :]
        }
        self.assertEqual({".docx", ".pptx", ".xlsx"}, converted_suffixes)
        # ...and every format's PDF is still individually rasterized.
        self.assertEqual(3, len(rasterized_pdfs))

    def test_conversion_probe_uses_pymupdf_when_pdftoppm_missing(self) -> None:
        converted: list[str] = []
        pymupdf_renders: list[str] = []

        def fake_run(args, *unused_args, **unused_kwargs):
            if "--convert-to" in args:
                outdir = Path(args[args.index("--outdir") + 1])
                outdir.mkdir(parents=True, exist_ok=True)
                for raw in args[args.index("--outdir") + 2 :]:
                    document = Path(raw)
                    converted.append(document.suffix)
                    (outdir / f"{document.stem}.pdf").write_bytes(b"%PDF-1.4\n%%EOF\n")
                return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")
            if args and Path(args[0]).name == "pdftoppm":
                raise AssertionError(
                    "pdftoppm must not be used when rasterizer is disabled"
                )
            return SimpleNamespace(returncode=0, stdout=b"version ok", stderr=b"")

        def fake_pymupdf(pdf: Path, png_dir: Path, *, dpi: int):
            pymupdf_renders.append(pdf.stem)
            _centered_content(width=100, height=100).save(png_dir / "page-1.png")
            return True, None

        with (
            patch.object(subprocess, "run", fake_run),
            patch.object(doctor, "_rasterize_pdf_with_pymupdf", fake_pymupdf),
        ):
            ok, error = doctor._probe_visual_pipeline(
                {"soffice": "/fake/soffice", "pdftoppm": None},
                {"pdftoppm": False, "fitz": True},
            )

        self.assertTrue(ok, error)
        self.assertEqual({".docx", ".pptx", ".xlsx"}, set(converted))
        self.assertEqual(3, len(pymupdf_renders))


# ---------------------------------------------------------------------------
# §7.2 Manifest + checklist (model-free)
# ---------------------------------------------------------------------------
def _extract_real_profile(td: Path) -> dict:
    """Extract the repo template into a temp store and return its profile dict."""
    from brandkit.formats.docx import extract as docx_extract
    from brandkit.profile import store

    template = (
        Path(__file__).resolve().parents[1]
        / "examples"
        / "templates"
        / "branddocs_template.docx"
    )
    docx_extract.extract(template, "vqa", scope="project", cwd=td)
    return store.load_profile("vqa", "project", cwd=td).profile


class ManifestTest(unittest.TestCase):
    def test_captured_template_texts_include_surface_demo_values(self) -> None:
        profile = {
            "kind": "pptx",
            "surface": {
                "pptx": {
                    "cover_anchors": [
                        {"placeholder": "Click to edit Master title style"},
                        {"demo_value": "Template demo value"},
                    ]
                }
            },
        }
        self.assertEqual(checks_deterministic.captured_template_texts(profile), [])
        texts = checks_deterministic.captured_template_texts(
            profile,
            include_surface_prompts=True,
        )
        self.assertIn("Click to edit Master title style", texts)
        self.assertIn("Template demo value", texts)

    def test_build_manifest_shape(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            profile = _extract_real_profile(td)
            out_dir = td / "out.visual"
            out_dir.mkdir(parents=True, exist_ok=True)
            png_paths = []
            for i in (1, 2):
                p = out_dir / f"page-{i}.png"
                _centered_content().save(p)
                png_paths.append(p)
            l1 = vqa.run_visual_l1(png_paths)
            manifest_path = vqa.build_visual_manifest(
                profile=profile,
                document=td / "out.docx",
                png_paths=png_paths,
                l1_findings=l1,
                renderers_ok=True,
                out_dir=out_dir,
                environment_status={
                    "visual_qa": True,
                    "binary_paths": {
                        "soffice": "/usr/bin/soffice",
                        "pdftoppm": "/usr/bin/pdftoppm",
                    },
                    "binaries": {"soffice": True, "pdftoppm": True},
                    "binary_errors": {},
                    "ocr_binary_paths": {"tesseract": "/usr/bin/tesseract"},
                    "ocr_binaries": {"tesseract": True},
                    "ocr_binary_errors": {},
                    "ocr_qa": True,
                },
                shell_sha256="deadbeef",
                content_sha256="cafef00d",
            )
            self.assertTrue(manifest_path.is_file())
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(data["schema_version"], "visual-manifest-2")
            # C1: the manifest carries the exact-artifact shas the L2 model scopes a
            # verdict to (the short-circuit later requires both to match).
            self.assertEqual(data["shell_sha256"], "deadbeef")
            self.assertEqual(data["content_sha256"], "cafef00d")
            self.assertEqual(data["kind"], "docx")
            self.assertEqual(len(data["pages"]), 2)
            self.assertIn("checklist", data)
            self.assertIn("l1_findings", data)
            self.assertIn("instructions", data)
            self.assertTrue(data["renderers_available"])
            # Page records carry dimensions + orientation.
            self.assertEqual(data["pages"][0]["width"], 850)
            self.assertEqual(data["pages"][0]["orientation"], "portrait")
            self.assertEqual(data["environment"]["visual_qa"], True)
            self.assertEqual(
                data["environment"]["renderers"]["soffice"]["path"], "/usr/bin/soffice"
            )
            self.assertEqual(
                data["environment"]["renderers"]["pdftoppm"]["available"], True
            )
            self.assertIn("fitz", data["environment"]["optional_python"])
            self.assertEqual(
                data["environment"]["ocr"]["tesseract"]["path"], "/usr/bin/tesseract"
            )
            self.assertIn("ocr", data)
            self.assertEqual(data["ocr"]["status"], "not_run")
            self.assertIn("platform", data["environment"])

    def test_checklist_derives_from_profile(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            profile = _extract_real_profile(Path(td))
            checklist = vqa.derive_visual_checklist(profile)
            ids = {item["id"] for item in checklist}
            self.assertIn("regions_present", ids)  # from structure.skeleton
            self.assertIn("cover_correct", ids)  # from anchors.cover
            self.assertIn("palette_on_brand", ids)  # from theme.colors
            self.assertIn("roles_styled", ids)  # from roles._index
            self.assertIn("overflow_clean", ids)  # from qa.overflow_capability=render
            # No chart in this template -> charts_rendered must be ABSENT.
            self.assertNotIn("charts_rendered", ids)
            # Every item is traceable.
            for item in checklist:
                self.assertIn("derived_from", item)
                self.assertIn("what", item)

    def test_checklist_includes_charts_when_present(self) -> None:
        profile = {"kind": "docx", "roles": {"_index": ["chart.bar"]}}
        ids = {item["id"] for item in vqa.derive_visual_checklist(profile)}
        self.assertIn("charts_rendered", ids)

    def test_manifest_paths_relative(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            profile = _extract_real_profile(td)
            out_dir = td / "out.visual"
            out_dir.mkdir(parents=True, exist_ok=True)
            png = out_dir / "page-1.png"
            _centered_content().save(png)
            manifest_path = vqa.build_visual_manifest(
                profile=profile,
                document=td / "out.docx",
                png_paths=[png],
                l1_findings=[],
                renderers_ok=True,
                out_dir=out_dir,
            )
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(data["pages"][0]["png"], "page-1.png")  # relative, no dir

    def test_manifest_degraded_keeps_checklist(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            profile = _extract_real_profile(td)
            out_dir = td / "out.visual"
            manifest_path = vqa.build_visual_manifest(
                profile=profile,
                document=td / "out.docx",
                png_paths=[],
                l1_findings=[],
                renderers_ok=False,
                out_dir=out_dir,
            )
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertTrue(data["degraded"])
            self.assertEqual(data["pages"], [])
            self.assertEqual(data["l1_findings"], [])
            self.assertTrue(data["checklist"])  # still populated
            self.assertEqual(data["ocr"]["status"], "not_run")

    def test_manifest_degraded_can_still_include_fallback_pages(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            profile = _extract_real_profile(td)
            out_dir = td / "out.visual"
            out_dir.mkdir(parents=True, exist_ok=True)
            png = out_dir / "page-1.png"
            _blank().save(png)
            l1 = vqa.run_visual_l1([png])
            manifest_path = vqa.build_visual_manifest(
                profile=profile,
                document=td / "out.docx",
                png_paths=[png],
                l1_findings=l1,
                renderers_ok=False,
                out_dir=out_dir,
                degraded=True,
            )
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertTrue(data["degraded"])
            self.assertFalse(data["renderers_available"])
            self.assertEqual(data["pages"][0]["png"], "page-1.png")
            self.assertTrue(
                any(f["check"] == "visual.blank_page" for f in data["l1_findings"])
            )

    def test_manifest_degraded_keeps_visual_findings_without_pages(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            profile = _extract_real_profile(td)
            out_dir = td / "out.visual"
            manifest_path = vqa.build_visual_manifest(
                profile=profile,
                document=td / "out.docx",
                png_paths=[],
                l1_findings=[
                    Finding(
                        "visual.no_pages",
                        schema.Severity.WARNING.value,
                        "output rendered zero pages",
                    )
                ],
                renderers_ok=False,
                out_dir=out_dir,
                degraded=True,
            )
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(data["pages"], [])
            self.assertEqual(data["l1_findings"][0]["check"], "visual.no_pages")

    def test_visual_ocr_flags_rendered_residual_text(self) -> None:
        orig_which = vqa.shutil.which
        orig_run = subprocess.run
        profile = {
            "kind": "docx",
            "surface": {
                "docx": {"cover_anchors": [{"placeholder": "Old template subtitle"}]}
            },
        }

        def fake_which(name):
            if name == "tesseract":
                return "/fake/tesseract"
            return orig_which(name)

        def fake_run(args, *unused_args, **unused_kwargs):
            if args and Path(args[0]).name == "tesseract":
                # tesseract output is captured as BYTES (no text=True), so the fake
                # mirrors that; _ocr_png decodes it tolerantly.
                return SimpleNamespace(
                    returncode=0,
                    stdout=b"Generated page\nOld template subtitle\n",
                    stderr=b"",
                )
            return orig_run(args, *unused_args, **unused_kwargs)

        with (
            patch.object(vqa.shutil, "which", fake_which),
            patch.object(subprocess, "run", fake_run),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                png = td / "page-1.png"
                _centered_content(width=100, height=100).save(png)
                report = vqa.run_visual_ocr([png], profile)

        self.assertEqual(report["status"], "ok")
        self.assertEqual(report["hits"][0]["term"], "Old template subtitle")
        findings = vqa.ocr_findings(report)
        self.assertEqual(findings[0].check, "visual.ocr_residual_text")

    def test_rasterize_pdf_uses_pymupdf_fallback_when_pdftoppm_missing(self) -> None:
        orig_which = vqa.shutil.which

        def fake_which(name):
            return None if name == "pdftoppm" else orig_which(name)

        def fake_pymupdf(pdf: Path, out_dir: Path, *, dpi: int):
            png = out_dir / "page-1.png"
            _centered_content(width=100, height=100).save(png)
            return [png], None

        with (
            patch.object(vqa.shutil, "which", fake_which),
            patch.object(vqa, "_rasterize_pdf_with_pymupdf", fake_pymupdf),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                pdf = td / "out.pdf"
                pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
                errors: list[str] = []
                warnings: list[str] = []
                pngs = vqa._rasterize_pdf_to_pngs(
                    pdf,
                    td,
                    dpi=100,
                    timeout_s=1,
                    render_errors=errors,
                    render_warnings=warnings,
                )

        self.assertEqual([p.name for p in pngs], ["page-1.png"])
        self.assertTrue(any("pdftoppm unavailable" in e for e in errors))
        self.assertTrue(any("PyMuPDF PDF raster fallback used" in w for w in warnings))


# ---------------------------------------------------------------------------
# §7.3 Degrade and gate wiring (no renderer)
# ---------------------------------------------------------------------------
def _minimal_profile() -> dict:
    prof = schema.build_envelope("docx", {"name": "wire"})
    prof["surface"] = {"docx": {}}
    prof["roles"] = {"_index": []}
    return prof


class GateWiringTest(unittest.TestCase):
    def setUp(self) -> None:
        # ``renderers_available`` caches into this process-global; reset it so a
        # patched probe result cannot leak across tests or collection order.
        vqa._LAST_RENDERER_STATUS = None
        self.addCleanup(setattr, vqa, "_LAST_RENDERER_STATUS", None)

    def test_run_qa_fast_never_renders(self) -> None:
        called = {"n": 0}

        def boom(*a, **k):  # pragma: no cover - must never run
            called["n"] += 1
            raise AssertionError("render_to_pngs must not be called on --qa fast")

        with patch.object(vqa, "render_to_pngs", boom):
            with tempfile.TemporaryDirectory() as td:
                target = Path(td) / "out.docx"
                _real_docx(target)  # presence only; never rendered
                report = gate.run_qa(target, _minimal_profile(), qa="fast")
        self.assertEqual(called["n"], 0)
        self.assertFalse(any(f.check.startswith("visual.") for f in report.findings))

    def test_run_qa_auto_degrades_when_no_renderers(self) -> None:
        with patch.object(vqa, "renderers_available", lambda: False):
            with tempfile.TemporaryDirectory() as td:
                target = Path(td) / "out.docx"
                _real_docx(target)
                report = gate.run_qa(target, _minimal_profile(), qa="auto")
        unavailable = [f for f in report.findings if f.check == "visual.unavailable"]
        self.assertEqual(len(unavailable), 1)
        self.assertEqual(unavailable[0].severity, schema.Severity.INFO.value)
        self.assertTrue(report.passed)
        self.assertFalse(
            any(f.severity == schema.Severity.ERROR.value for f in report.findings)
        )

    def test_run_qa_deep_degraded_writes_manifest(self) -> None:
        with (
            patch.object(vqa, "renderers_available", lambda: False),
            patch.object(vqa, "render_to_pngs", lambda *args, **kwargs: []),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                report = gate.run_qa(
                    target,
                    _minimal_profile(),
                    qa="deep",
                    out_dir=out_dir,
                )
                unavailable = [
                    f for f in report.findings if f.check == "visual.unavailable"
                ]
                self.assertEqual(len(unavailable), 1)
                manifest = [f for f in report.findings if f.check == "visual.manifest"]
                self.assertEqual(len(manifest), 1)
                self.assertEqual(manifest[0].severity, schema.Severity.INFO.value)
                data = json.loads(
                    Path(manifest[0].location).read_text(encoding="utf-8")
                )
        self.assertTrue(data["degraded"])
        self.assertEqual(data["pages"], [])
        self.assertTrue(data["checklist"])
        self.assertTrue(report.passed)

    def test_run_qa_deep_uses_quicklook_when_primary_unavailable(self) -> None:
        def fake_render(*args, **kwargs):
            self.assertTrue(kwargs.get("quicklook_only"))
            warnings = kwargs.get("render_warnings")
            if warnings is not None:
                warnings.append(
                    "Quick Look thumbnail fallback used because LibreOffice is unavailable"
                )
            out_dir = Path(args[1])
            out_dir.mkdir(parents=True, exist_ok=True)
            png = out_dir / "page-1.png"
            _centered_content().save(png)
            return [png]

        with (
            patch.object(vqa, "renderers_available", lambda: False),
            patch.object(vqa, "render_to_pngs", fake_render),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                report = gate.run_qa(
                    target,
                    _minimal_profile(),
                    qa="deep",
                    out_dir=out_dir,
                )
                manifest = [f for f in report.findings if f.check == "visual.manifest"]
                self.assertEqual(len(manifest), 1)
                data = json.loads(
                    Path(manifest[0].location).read_text(encoding="utf-8")
                )
        self.assertTrue(any(f.check == "visual.unavailable" for f in report.findings))
        degraded = [f for f in report.findings if f.check == "visual.render_degraded"]
        self.assertEqual(len(degraded), 1)
        self.assertEqual(len(data["pages"]), 1)
        self.assertTrue(data["degraded"])
        self.assertFalse(data["renderers_available"])

    def test_run_qa_deep_reports_render_failure_after_probe(self) -> None:
        def fake_render(*args, **kwargs):
            errors = kwargs.get("render_errors")
            if errors is not None:
                errors.append("soffice convert failed: abort trap")
            return []

        with (
            patch.object(vqa, "renderers_available", lambda: True),
            patch.object(vqa, "render_to_pngs", fake_render),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                report = gate.run_qa(
                    target,
                    _minimal_profile(),
                    qa="deep",
                    out_dir=out_dir,
                )
                failed = [
                    f for f in report.findings if f.check == "visual.render_failed"
                ]
                self.assertEqual(len(failed), 1)
                self.assertIn("abort trap", failed[0].message)
                manifest = [f for f in report.findings if f.check == "visual.manifest"]
                self.assertEqual(len(manifest), 1)
                data = json.loads(
                    Path(manifest[0].location).read_text(encoding="utf-8")
                )
        self.assertTrue(data["degraded"])
        self.assertTrue(report.passed)

    def test_run_qa_deep_reports_degraded_render_when_fallback_used(self) -> None:
        def fake_render(*args, **kwargs):
            warnings = kwargs.get("render_warnings")
            if warnings is not None:
                warnings.append(
                    "Quick Look thumbnail fallback used after soffice failure"
                )
            out_dir = Path(args[1])
            out_dir.mkdir(parents=True, exist_ok=True)
            png = out_dir / "page-1.png"
            _centered_content().save(png)
            return [png]

        with (
            patch.object(vqa, "renderers_available", lambda: True),
            patch.object(vqa, "render_to_pngs", fake_render),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                report = gate.run_qa(
                    target,
                    _minimal_profile(),
                    qa="deep",
                    out_dir=out_dir,
                )
                manifest = [f for f in report.findings if f.check == "visual.manifest"]
                data = json.loads(
                    Path(manifest[0].location).read_text(encoding="utf-8")
                )
        degraded = [f for f in report.findings if f.check == "visual.render_degraded"]
        self.assertEqual(len(degraded), 1)
        self.assertIn("Quick Look", degraded[0].message)
        self.assertFalse(
            any(f.check == "visual.render_failed" for f in report.findings)
        )
        self.assertTrue(data["degraded"])
        self.assertFalse(data["renderers_available"])

    def test_run_qa_strict_fails_when_renderers_unavailable(self) -> None:
        with patch.object(vqa, "renderers_available", lambda: False):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                report = gate.run_qa(
                    target,
                    _minimal_profile(),
                    qa="strict",
                    out_dir=out_dir,
                )
                manifest = [f for f in report.findings if f.check == "visual.manifest"]
                self.assertEqual(len(manifest), 1)
                data = json.loads(
                    Path(manifest[0].location).read_text(encoding="utf-8")
                )

        self.assertFalse(report.passed)
        self.assertTrue(
            any(f.check == "visual.strict_unavailable" for f in report.findings)
        )
        self.assertTrue(data["degraded"])
        self.assertEqual(data["qa_mode"], "strict")

    def test_run_qa_strict_promotes_l1_findings_to_errors(self) -> None:
        def fake_ocr(*args, **kwargs):
            return {
                "engine": "tesseract",
                "available": False,
                "status": "unavailable",
                "terms_checked": [],
                "pages": [],
                "hits": [],
                "errors": [],
                "reason": "test",
            }

        with patch.object(vqa, "run_visual_ocr", fake_ocr):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                out_dir.mkdir(parents=True, exist_ok=True)
                png = out_dir / "page-1.png"
                _blank().save(png)
                report = gate.run_qa(
                    target,
                    _minimal_profile(),
                    qa="strict",
                    out_dir=out_dir,
                    visual=(True, [png]),
                )
                manifest = [f for f in report.findings if f.check == "visual.manifest"]
                data = json.loads(
                    Path(manifest[0].location).read_text(encoding="utf-8")
                )

        self.assertFalse(report.passed)
        self.assertTrue(any(f.check == "visual.blank_page" for f in report.findings))
        strict = [f for f in report.findings if f.check == "visual.strict"]
        self.assertTrue(strict)
        self.assertTrue(any("visual.blank_page" in f.message for f in strict))
        self.assertEqual(data["qa_mode"], "strict")

    def test_run_qa_strict_promotes_ocr_degraded_to_error(self) -> None:
        # A flaky/failed OCR pass (status='failed' with errors) surfaces a
        # visual.ocr_degraded INFO, which strict mode must promote to a blocking
        # visual.strict ERROR. Use non-blank content so blank_page does not also fire.
        def fake_ocr(*args, **kwargs):
            return {
                "engine": "tesseract",
                "available": True,
                "status": "failed",
                "terms_checked": [],
                "pages": [],
                "hits": [],
                "errors": [{"page": 1, "error": "tesseract crashed"}],
            }

        with patch.object(vqa, "run_visual_ocr", fake_ocr):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                out_dir.mkdir(parents=True, exist_ok=True)
                png = out_dir / "page-1.png"
                _centered_content().save(png)  # not blank -> isolates ocr_degraded
                report = gate.run_qa(
                    target,
                    _minimal_profile(),
                    qa="strict",
                    out_dir=out_dir,
                    visual=(True, [png]),
                )

        self.assertFalse(report.passed)
        strict = [f for f in report.findings if f.check == "visual.strict"]
        self.assertTrue(any("visual.ocr_degraded" in f.message for f in strict), strict)

    def test_run_qa_deep_surfaces_ocr_hits_in_manifest(self) -> None:
        def fake_ocr(*args, **kwargs):
            return {
                "engine": "tesseract",
                "available": True,
                "status": "ok",
                "terms_checked": ["Old template subtitle"],
                "pages": [
                    {
                        "index": 1,
                        "text": "Old template subtitle",
                        "hits": [{"term": "Old template subtitle"}],
                    }
                ],
                "hits": [{"page": 1, "term": "Old template subtitle"}],
                "errors": [],
            }

        with patch.object(vqa, "run_visual_ocr", fake_ocr):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                png = out_dir / "page-1.png"
                out_dir.mkdir(parents=True, exist_ok=True)
                _centered_content().save(png)
                report = gate.run_qa(
                    target,
                    _minimal_profile(),
                    qa="deep",
                    out_dir=out_dir,
                    visual=(True, [png]),
                )
                manifest = [f for f in report.findings if f.check == "visual.manifest"]
                data = json.loads(
                    Path(manifest[0].location).read_text(encoding="utf-8")
                )

        self.assertTrue(
            any(f.check == "visual.ocr_residual_text" for f in report.findings)
        )
        self.assertEqual(data["ocr"]["hits"][0]["term"], "Old template subtitle")

    def test_render_to_pngs_can_skip_availability_recheck(self) -> None:
        orig_which = vqa.shutil.which

        def boom_available():  # pragma: no cover - must not be called
            raise AssertionError("availability already checked by gate")

        # Force the pdftoppm rasterization branch regardless of host: on CI
        # (Ubuntu) the binary is absent, so without this stub shutil.which
        # returns None and render_to_pngs falls through to the quicklook path.
        def fake_which(name):
            if name == "pdftoppm":
                return "/fake/pdftoppm"
            return orig_which(name)

        def fake_run(args, *unused_args, **unused_kwargs):
            if "--convert-to" in args:
                outdir = Path(args[args.index("--outdir") + 1])
                outdir.mkdir(parents=True, exist_ok=True)
                (outdir / "out.pdf").write_bytes(b"%PDF-1.4\n%%EOF\n")
                return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")
            if args and Path(args[0]).name == "pdftoppm":
                prefix = Path(args[-1])
                _centered_content().save(prefix.with_name(prefix.name + "-1.png"))
                return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")
            return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")

        with (
            patch.object(vqa, "renderers_available", boom_available),
            patch.object(vqa.shutil, "which", fake_which),
            patch.object(subprocess, "run", fake_run),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                pngs = vqa.render_to_pngs(
                    target,
                    out_dir,
                    check_available=False,
                )
            self.assertEqual([p.name for p in pngs], ["page-1.png"])

    def test_render_to_pngs_falls_back_to_quicklook_thumbnail(self) -> None:
        def fake_which(name):
            if name in {"qlmanage", "soffice", "pdftoppm"}:
                return f"/fake/{name}"
            return None

        def fake_run(args, *unused_args, **unused_kwargs):
            exe = Path(args[0]).name
            if "--convert-to" in args:
                return SimpleNamespace(returncode=134, stdout=b"", stderr=b"abort trap")
            if exe == "qlmanage":
                outdir = Path(args[args.index("-o") + 1])
                doc = Path(args[-1])
                _centered_content().save(outdir / f"{doc.name}.png")
                return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")
            if exe == "pdftoppm":
                raise AssertionError(
                    "pdftoppm should not run after soffice failure fallback"
                )
            return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")

        with (
            patch.object(subprocess, "run", fake_run),
            patch.object(vqa.shutil, "which", fake_which),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                warnings: list[str] = []
                errors: list[str] = []
                pngs = vqa.render_to_pngs(
                    target,
                    out_dir,
                    check_available=False,
                    quicklook_only=False,
                    render_errors=errors,
                    render_warnings=warnings,
                )
            self.assertEqual([p.name for p in pngs], ["page-1.png"])
            self.assertTrue(any("Quick Look" in w for w in warnings))
            self.assertTrue(any("soffice convert failed" in e for e in errors))

    def test_render_to_pngs_degrades_when_soffice_fails_and_quicklook_absent(
        self,
    ) -> None:
        # soffice convert fails AND qlmanage is absent -> no pages, with both reasons
        # surfaced as render errors (never a crash, never a stale/bogus frame).
        def fake_which(name):
            return f"/fake/{name}" if name == "soffice" else None  # no qlmanage

        def fake_run(args, *unused_args, **unused_kwargs):
            return SimpleNamespace(returncode=134, stdout=b"", stderr=b"abort trap")

        with (
            patch.object(subprocess, "run", fake_run),
            patch.object(vqa.shutil, "which", fake_which),
        ):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                out_dir = td / "out.visual"
                _real_docx(target)
                errors: list[str] = []
                pngs = vqa.render_to_pngs(
                    target,
                    out_dir,
                    check_available=False,
                    render_errors=errors,
                    render_warnings=[],
                )
            self.assertEqual(pngs, [])
            self.assertTrue(any("qlmanage not found" in e for e in errors), errors)

    def test_run_qa_deep_injected_visual(self) -> None:
        """Drive ``deep`` without soffice via the ``visual=`` injection hook."""
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            profile = _extract_real_profile(td)
            out_dir = td / "out.visual"
            out_dir.mkdir(parents=True, exist_ok=True)
            p1 = out_dir / "page-1.png"
            p2 = out_dir / "page-2.png"
            _blank().save(p1)  # blank page -> L1 finding
            _centered_content().save(p2)
            target = td / "out.docx"
            _real_docx(target)

            report = gate.run_qa(
                target,
                profile,
                qa="deep",
                out_dir=out_dir,
                visual=(True, [p1, p2]),
            )
            # L1 blank-page finding present.
            self.assertTrue(
                any(f.check == "visual.blank_page" for f in report.findings)
            )
            # Manifest finding present, pointing at a written file.
            manifest = [f for f in report.findings if f.check == "visual.manifest"]
            self.assertEqual(len(manifest), 1)
            self.assertEqual(manifest[0].severity, schema.Severity.INFO.value)
            self.assertTrue(Path(manifest[0].location).is_file())
            # The manifest records both pages and the L1 finding.
            data = json.loads(Path(manifest[0].location).read_text(encoding="utf-8"))
            self.assertEqual(len(data["pages"]), 2)
            self.assertTrue(
                any(lf["check"] == "visual.blank_page" for lf in data["l1_findings"])
            )
            # Still passes (only INFO/WARNING).
            self.assertTrue(report.passed)

    def test_run_qa_auto_injected_visual_runs_l1_no_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            out_dir = td / "out.visual"
            out_dir.mkdir(parents=True, exist_ok=True)
            p1 = out_dir / "page-1.png"
            _blank().save(p1)
            target = td / "out.docx"
            _real_docx(target)
            report = gate.run_qa(
                target,
                _minimal_profile(),
                qa="auto",
                out_dir=out_dir,
                visual=(True, [p1]),
            )
            self.assertTrue(
                any(f.check == "visual.blank_page" for f in report.findings)
            )
            # auto does NOT emit a manifest.
            self.assertFalse(any(f.check == "visual.manifest" for f in report.findings))

    def test_visual_findings_never_change_verdict(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            out_dir = td / "out.visual"
            out_dir.mkdir(parents=True, exist_ok=True)
            p1 = out_dir / "page-1.png"
            _blank().save(p1)  # WARNING-producing
            target = td / "out.docx"
            _real_docx(target)
            report = gate.run_qa(
                target,
                _minimal_profile(),
                qa="auto",
                out_dir=out_dir,
                visual=(True, [p1]),
            )
            self.assertTrue(
                any(
                    f.severity == schema.Severity.WARNING.value for f in report.findings
                )
            )
            self.assertEqual(
                report.verdict, schema.VerificationStatus.PASSED_WITH_WARNINGS.value
            )
            self.assertNotEqual(report.verdict, schema.VerificationStatus.FAILED.value)
            self.assertTrue(report.passed)

    def test_verify_time_target_none_never_renders(self) -> None:
        called = {"n": 0}

        def boom(*a, **k):  # pragma: no cover
            called["n"] += 1
            raise AssertionError("no render at verify time (target is None)")

        with patch.object(vqa, "render_to_pngs", boom):
            report = gate.run_qa(None, _minimal_profile(), qa="deep")
        self.assertEqual(called["n"], 0)
        self.assertFalse(any(f.check.startswith("visual.") for f in report.findings))


# ---------------------------------------------------------------------------
# §7.4 Backward-compat guard: CI (no renderer) == fast modulo one INFO
# ---------------------------------------------------------------------------
class BackwardCompatTest(unittest.TestCase):
    def setUp(self) -> None:
        vqa._LAST_RENDERER_STATUS = None
        self.addCleanup(setattr, vqa, "_LAST_RENDERER_STATUS", None)

    def test_ci_degrade_matches_fast_modulo_info(self) -> None:
        """CI (no renderer) on --qa auto == --qa fast, modulo one INFO finding.

        The load-bearing invariant is the EXIT CODE: ``report.passed`` (no ERROR)
        is identical, so a green gate stays green and a red one stays red. The
        only delta is a single INFO ``visual.unavailable`` finding (which flips
        the human-readable verdict label from ``passed`` to
        ``passed_with_warnings`` but never the exit code). Crucially, no existing
        test exercises this path: every smoke test passes ``--qa fast``.
        """
        with patch.object(vqa, "renderers_available", lambda: False):
            with tempfile.TemporaryDirectory() as td:
                target = Path(td) / "out.docx"
                _real_docx(target)
                profile = _minimal_profile()
                fast = gate.run_qa(target, profile, qa="fast")
                auto = gate.run_qa(target, profile, qa="auto")
        # Exit code (passed) is invariant -- the real backward-compat promise.
        self.assertEqual(fast.passed, auto.passed)
        self.assertTrue(auto.passed)
        # No new ERROR is ever introduced by the degrade path.
        self.assertFalse(
            any(f.severity == schema.Severity.ERROR.value for f in auto.findings)
        )
        # The only delta is exactly one INFO visual.unavailable finding.
        extra = [f for f in auto.findings if f not in fast.findings]
        self.assertEqual(len(extra), 1)
        self.assertEqual(extra[0].check, "visual.unavailable")
        self.assertEqual(extra[0].severity, schema.Severity.INFO.value)


# ---------------------------------------------------------------------------
# §7.5 Gated real end-to-end render (skipped in CI)
# ---------------------------------------------------------------------------
def _real_render_requested() -> bool:
    return (
        os.environ.get("BRANDDOCS_RUN_REAL_RENDER") == "1" and vqa.renderers_available()
    )


@unittest.skipUnless(
    _real_render_requested(),
    "set BRANDDOCS_RUN_REAL_RENDER=1 to run real soffice/pdftoppm render tests",
)
class RealRenderE2ETest(unittest.TestCase):
    def test_render_produces_ordered_pngs(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            template = (
                Path(__file__).resolve().parents[1]
                / "examples"
                / "templates"
                / "branddocs_template.docx"
            )
            out_dir = td / "render"
            pngs = vqa.render_to_pngs(template, out_dir)
            self.assertTrue(pngs, "real render produced no PNGs")
            self.assertTrue(all(p.is_file() for p in pngs))
            # Ordered numerically (page-1 before page-2 ... page-10).
            indices = [vqa._page_sort_key(p) for p in pngs]
            self.assertEqual(indices, sorted(indices))

    def test_deep_generate_writes_manifest_and_pngs(self) -> None:
        from brandkit.cli import main

        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            old = Path.cwd()
            import os

            os.chdir(td)
            try:
                template = (
                    Path(__file__).resolve().parents[1]
                    / "examples"
                    / "templates"
                    / "branddocs_template.docx"
                )
                self.assertEqual(
                    main(
                        [
                            "extract",
                            "--name",
                            "e2e",
                            "--template",
                            str(template),
                            "--scope",
                            "project",
                        ]
                    ),
                    0,
                )
                idoc = td / "idoc.json"
                idoc.write_text(
                    json.dumps(
                        {
                            "cover": {"title": "Visual Audit Demo"},
                            "blocks": [
                                {"type": "heading", "level": 1, "text": "Section"},
                                {
                                    "type": "paragraph",
                                    "text": "Rendered for the visual audit.",
                                },
                            ],
                        }
                    ),
                    encoding="utf-8",
                )
                out = td / "out.docx"
                rc = main(
                    [
                        "generate",
                        "--name",
                        "e2e",
                        "--input",
                        str(idoc),
                        "--output",
                        str(out),
                        "--scope",
                        "project",
                        "--qa",
                        "deep",
                    ]
                )
                self.assertEqual(rc, 0)
                visual_dir = out.parent / "out.docx.visual"
                self.assertTrue((visual_dir / "visual_manifest.json").is_file())
                self.assertTrue((visual_dir / "page-1.png").is_file())
            finally:
                os.chdir(old)

    def test_deep_generate_pptx_writes_manifest_and_pngs(self) -> None:
        from brandkit.cli import main

        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            old = Path.cwd()
            import os

            os.chdir(td)
            try:
                template = (
                    Path(__file__).resolve().parents[1]
                    / "examples"
                    / "templates"
                    / "branddocs_template.pptx"
                )
                self.assertEqual(
                    main(
                        [
                            "extract",
                            "--name",
                            "deck",
                            "--template",
                            str(template),
                            "--scope",
                            "project",
                        ]
                    ),
                    0,
                )
                idoc = td / "deck.json"
                idoc.write_text(
                    json.dumps(
                        {
                            "cover": {"title": "Visual Audit Deck"},
                            "blocks": [
                                {
                                    "type": "heading",
                                    "level": 1,
                                    "text": "Market Context",
                                },
                                {
                                    "type": "paragraph",
                                    "text": "Rendered for the PPTX visual audit.",
                                },
                            ],
                        }
                    ),
                    encoding="utf-8",
                )
                out = td / "out.pptx"
                rc = main(
                    [
                        "generate",
                        "--name",
                        "deck",
                        "--input",
                        str(idoc),
                        "--output",
                        str(out),
                        "--scope",
                        "project",
                        "--qa",
                        "deep",
                    ]
                )
                self.assertEqual(rc, 0)
                visual_dir = out.parent / "out.pptx.visual"
                manifest = visual_dir / "visual_manifest.json"
                self.assertTrue(manifest.is_file())
                self.assertTrue((visual_dir / "page-1.png").is_file())
                data = json.loads(manifest.read_text(encoding="utf-8"))
                self.assertEqual(data["kind"], "pptx")
                self.assertTrue(data["pages"])
            finally:
                os.chdir(old)

    def test_deep_generate_xlsx_writes_manifest_and_pngs(self) -> None:
        from brandkit.cli import main

        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            old = Path.cwd()
            import os

            os.chdir(td)
            try:
                template = (
                    Path(__file__).resolve().parents[1]
                    / "examples"
                    / "templates"
                    / "branddocs_template.xlsx"
                )
                self.assertEqual(
                    main(
                        [
                            "extract",
                            "--name",
                            "book",
                            "--template",
                            str(template),
                            "--scope",
                            "project",
                        ]
                    ),
                    0,
                )
                grid = td / "grid.json"
                grid.write_text(
                    json.dumps(
                        {
                            "cells": {
                                "report_title": "Visual Audit Workbook",
                                "report_subtitle": "Excel render proof",
                                "client_name": "BrandDocs",
                                "period": "FY 2026",
                                "headline_kpi": "On-brand",
                            },
                            "regions": {
                                "data_block": [
                                    ["Metric", "Q1", "Q2", "Status"],
                                    ["Pipeline", 42, 48, "Healthy"],
                                    ["Delivery", 91, 94, "Green"],
                                ],
                            },
                        }
                    ),
                    encoding="utf-8",
                )
                out = td / "out.xlsx"
                rc = main(
                    [
                        "generate",
                        "--name",
                        "book",
                        "--input",
                        str(grid),
                        "--output",
                        str(out),
                        "--scope",
                        "project",
                        "--qa",
                        "deep",
                    ]
                )
                self.assertEqual(rc, 0)
                visual_dir = out.parent / "out.xlsx.visual"
                manifest = visual_dir / "visual_manifest.json"
                self.assertTrue(manifest.is_file())
                self.assertTrue((visual_dir / "page-1.png").is_file())
                data = json.loads(manifest.read_text(encoding="utf-8"))
                self.assertEqual(data["kind"], "xlsx")
                self.assertTrue(data["pages"])
            finally:
                os.chdir(old)


class OcrRobustnessTest(unittest.TestCase):
    """``_ocr_png`` must tolerate non-UTF-8 bytes from tesseract instead of crashing
    the whole QA run (tesseract can emit stray bytes - e.g. a 0x89 PNG byte in a
    fallback error path - which ``text=True`` would hard-fail on)."""

    def test_non_utf8_stdout_is_decoded_not_crashed(self) -> None:
        def fake_run(args, *unused_args, **unused_kwargs):
            return SimpleNamespace(returncode=0, stdout=b"ok \x89PNG text", stderr=b"")

        with patch.object(subprocess, "run", fake_run):
            text, err = vqa._ocr_png("tesseract", Path("/x/page-1.png"), timeout_s=5)
        self.assertIsNone(err)
        self.assertIn("ok", text)  # decoded tolerantly, no UnicodeDecodeError

    def test_non_utf8_error_output_degrades_gracefully(self) -> None:
        def fake_run(args, *unused_args, **unused_kwargs):
            return SimpleNamespace(returncode=1, stdout=b"", stderr=b"boom \x89PNG")

        with patch.object(subprocess, "run", fake_run):
            text, err = vqa._ocr_png("tesseract", Path("/x/page-1.png"), timeout_s=5)
        self.assertEqual(text, "")
        self.assertIsNotNone(err)  # error surfaced (decoded), run not crashed

    def test_timeout_degrades_gracefully(self) -> None:
        def fake_run(args, *unused_args, **unused_kwargs):
            raise subprocess.TimeoutExpired(args, 5)

        with patch.object(subprocess, "run", fake_run):
            text, err = vqa._ocr_png("tesseract", Path("/x/page-1.png"), timeout_s=5)
        self.assertEqual(text, "")
        self.assertIn("timed out", err)

    def test_oserror_degrades_gracefully(self) -> None:
        def fake_run(args, *unused_args, **unused_kwargs):
            raise OSError("no such binary")

        with patch.object(subprocess, "run", fake_run):
            text, err = vqa._ocr_png("tesseract", Path("/x/page-1.png"), timeout_s=5)
        self.assertEqual(text, "")
        self.assertIn("failed", err)


class L2ShortCircuitTest(unittest.TestCase):
    """C1: the generate-time L2 short-circuit can NEVER mask a regression."""

    SHELL_SHA = "abc123shell"
    CONTENT_SHA = "def456content"

    def setUp(self) -> None:
        vqa._LAST_RENDERER_STATUS = None
        self.addCleanup(setattr, vqa, "_LAST_RENDERER_STATUS", None)

    def _present_profile(
        self, *, verdicts=None, content_sha=None, shell_sha=None
    ) -> dict:
        """A docx profile with a PRESENT comprehension carrying an audit map that
        covers every current checklist id (default all PASS at the matching shas)."""
        prof = _minimal_profile()
        prof["provenance"]["shell"]["sha256"] = self.SHELL_SHA
        ids = vqa.visual_checklist_ids(prof)
        self.assertTrue(ids)
        audit = {}
        for cid in ids:
            audit[cid] = {
                "verdict": (verdicts or {}).get(cid, "PASS"),
                "shell_sha256": shell_sha if shell_sha is not None else self.SHELL_SHA,
                "content_sha256": content_sha
                if content_sha is not None
                else self.CONTENT_SHA,
            }
        prof["comprehension"] = {
            "schema_version": "comprehension-1",
            "status": "present",
            "source_shell_sha256": self.SHELL_SHA,
            "confidence": 1.0,
            "cover_slots": {},
            "conventions": {"indexes": [], "sections": []},
            "role_annotations": {},
            "demo_classification": {"regions": []},
            "palette_annotations": {},
            "audit": audit,
        }
        return prof

    def _run(self, prof, *, qa="deep", content_hash=CONTENT_SHA):
        boom = {"n": 0}

        def no_render(*a, **k):  # render must NOT run when short-circuiting
            boom["n"] += 1
            return []

        with patch.object(vqa, "render_to_pngs", no_render):
            with tempfile.TemporaryDirectory() as td:
                td = Path(td)
                target = td / "out.docx"
                _real_docx(target)
                report = gate.run_qa(
                    target,
                    prof,
                    qa=qa,
                    shell=None,  # falls back to provenance.shell.sha256
                    out_dir=td / "out.visual",
                    visual=(True, [td / "out.visual" / "page-1.png"]),
                    content_hash=content_hash,
                )
        return report, boom["n"]

    def _checks(self, report) -> set:
        return {f.check for f in report.findings}

    def test_short_circuit_only_on_all_pass_same_sha(self) -> None:
        prof = self._present_profile()
        report, renders = self._run(prof)
        self.assertIn("visual.l2_short_circuit", self._checks(report))
        self.assertNotIn("visual.manifest", self._checks(report))
        self.assertEqual(renders, 0, "render must be skipped when short-circuiting")
        self.assertTrue(report.passed)
        sc = [f for f in report.findings if f.check == "visual.l2_short_circuit"]
        self.assertEqual(sc[0].severity, schema.Severity.INFO.value)

    def test_single_FAIL_or_NA_reruns_l2(self) -> None:
        ids = vqa.visual_checklist_ids(self._present_profile())
        for bad in ("FAIL", "NA"):
            prof = self._present_profile(verdicts={ids[0]: bad})
            report, renders = self._run(prof)
            self.assertNotIn(
                "visual.l2_short_circuit",
                self._checks(report),
                f"{bad} short-circuited",
            )
            self.assertIn("visual.manifest", self._checks(report))

    def test_content_sha_mismatch_reruns_l2(self) -> None:
        prof = self._present_profile(content_sha="STALE")
        report, _ = self._run(prof, content_hash=self.CONTENT_SHA)
        self.assertNotIn("visual.l2_short_circuit", self._checks(report))
        self.assertIn("visual.manifest", self._checks(report))

    def test_shell_sha_mismatch_reruns_l2(self) -> None:
        prof = self._present_profile(shell_sha="OLDSHELL")
        report, _ = self._run(prof)
        self.assertNotIn("visual.l2_short_circuit", self._checks(report))
        self.assertIn("visual.manifest", self._checks(report))

    def test_added_checklist_id_defeats_stale_audit(self) -> None:
        # A present audit covering only SOME current ids must NOT short-circuit:
        # a newly-derived id with no audit row forces a full L2.
        prof = self._present_profile()
        # Drop one row so coverage of the current derived set is incomplete.
        some_id = next(iter(prof["comprehension"]["audit"]))
        del prof["comprehension"]["audit"][some_id]
        report, _ = self._run(prof)
        self.assertNotIn("visual.l2_short_circuit", self._checks(report))
        self.assertIn("visual.manifest", self._checks(report))

    def test_short_circuit_disabled_under_strict(self) -> None:
        prof = self._present_profile()
        report, renders = self._run(prof, qa="strict")
        self.assertNotIn("visual.l2_short_circuit", self._checks(report))
        self.assertIn("visual.manifest", self._checks(report))

    def test_verify_never_short_circuits(self) -> None:
        # verify passes no content_hash => stays None => never fires.
        prof = self._present_profile()
        report, renders = self._run(prof, content_hash=None)
        self.assertNotIn("visual.l2_short_circuit", self._checks(report))
        self.assertIn("visual.manifest", self._checks(report))

    def test_absent_audit_runs_full_l2(self) -> None:
        # An empty audit map (the byte-identical default) never short-circuits.
        prof = self._present_profile()
        prof["comprehension"]["audit"] = {}
        report, _ = self._run(prof)
        self.assertNotIn("visual.l2_short_circuit", self._checks(report))
        self.assertIn("visual.manifest", self._checks(report))


class TriageGateTest(unittest.TestCase):
    """C2: model-assisted QA triage demotes an AMBIGUOUS WARNING to INFO, and can
    NEVER demote an ERROR (the load-bearing safety property)."""

    def setUp(self) -> None:
        vqa._LAST_RENDERER_STATUS = None
        self.addCleanup(setattr, vqa, "_LAST_RENDERER_STATUS", None)

    def _profile_with_triage(self, triage: list) -> dict:
        prof = _minimal_profile()
        prof["comprehension"] = {
            "schema_version": "comprehension-1",
            "status": "present",
            "source_shell_sha256": "abc",
            "confidence": 1.0,
            "cover_slots": {},
            "conventions": {"indexes": [], "sections": []},
            "role_annotations": {},
            "demo_classification": {"regions": []},
            "audit": {},
            "triage": triage,
        }
        return prof

    def _run(self, prof, extra_findings, *, qa="auto"):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            target = td / "out.docx"
            _real_docx(target)
            return gate.run_qa(
                target,
                prof,
                qa=qa,
                shell=None,
                out_dir=td / "out.visual",
                visual=(True, []),  # no PNGs -> only the injected findings drive L1
                extra_findings=extra_findings,
            )

    def _find(self, report, check):
        return [f for f in report.findings if f.check == check]

    def test_expected_demotes_warning_to_info(self) -> None:
        prof = self._profile_with_triage(
            [
                {
                    "check": "component_survival",
                    "location": "tables",
                    "disposition": "expected",
                    "evidence": "fewer tables by design",
                }
            ]
        )
        warn = Finding(
            "component_survival",
            schema.Severity.WARNING.value,
            "native tables count dropped 2 -> 1",
            location="tables",
        )
        report = self._run(prof, [warn])
        cs = self._find(report, "component_survival")
        self.assertEqual(len(cs), 1)
        self.assertEqual(cs[0].severity, schema.Severity.INFO.value)
        self.assertIn("triaged EXPECTED: fewer tables by design", cs[0].message)
        self.assertTrue(report.passed)

    def test_defect_keeps_warning(self) -> None:
        prof = self._profile_with_triage(
            [
                {
                    "check": "component_survival",
                    "location": "tables",
                    "disposition": "defect",
                }
            ]
        )
        warn = Finding(
            "component_survival",
            schema.Severity.WARNING.value,
            "native tables count dropped 2 -> 1",
            location="tables",
        )
        report = self._run(prof, [warn])
        cs = self._find(report, "component_survival")
        self.assertEqual(len(cs), 1)
        self.assertEqual(cs[0].severity, schema.Severity.WARNING.value)
        self.assertNotIn("triaged", cs[0].message)

    def test_triage_only_matches_its_own_location(self) -> None:
        # An EXPECTED entry for one family must NOT demote a WARNING on another.
        prof = self._profile_with_triage(
            [
                {
                    "check": "component_survival",
                    "location": "tables",
                    "disposition": "expected",
                }
            ]
        )
        other = Finding(
            "component_survival",
            schema.Severity.WARNING.value,
            "native charts count dropped 1 -> 0",
            location="charts",
        )
        report = self._run(prof, [other])
        cs = self._find(report, "component_survival")
        self.assertEqual(cs[0].severity, schema.Severity.WARNING.value)

    def test_triage_never_demotes_error(self) -> None:
        # An ERROR sharing a (check, location) with an EXPECTED triage entry stays an
        # ERROR and FAILS the gate. This is the load-bearing proof: _apply_triage
        # guards on severity == WARNING, so no triage value can lower an ERROR.
        prof = self._profile_with_triage(
            [
                {
                    "check": "component_survival",
                    "location": "tables",
                    "disposition": "expected",
                    "evidence": "should-not-matter",
                }
            ]
        )
        err = Finding(
            "component_survival",
            schema.Severity.ERROR.value,  # deliberately an ERROR
            "native tables erased entirely",
            location="tables",
        )
        report = self._run(prof, [err])
        cs = self._find(report, "component_survival")
        self.assertEqual(len(cs), 1)
        self.assertEqual(cs[0].severity, schema.Severity.ERROR.value)
        self.assertNotIn("triaged", cs[0].message)
        self.assertFalse(report.passed)

    def test_empty_triage_is_byte_identical_noop(self) -> None:
        # With no triage entries, the WARNING is left exactly as emitted.
        prof = self._profile_with_triage([])
        warn = Finding(
            "component_survival",
            schema.Severity.WARNING.value,
            "native tables count dropped 2 -> 1",
            location="tables",
        )
        report = self._run(prof, [warn])
        cs = self._find(report, "component_survival")
        self.assertEqual(cs[0].severity, schema.Severity.WARNING.value)
        self.assertNotIn("triaged", cs[0].message)

    def test_component_survival_location_addressable(self) -> None:
        # The check now stamps the dropped FAMILY as location, so two dropped
        # families no longer collide on (check, None) and triage can name one.
        from openpyxl import Workbook
        from openpyxl.worksheet.table import Table

        with tempfile.TemporaryDirectory() as t:
            t = Path(t)
            shell = t / "shell.xlsx"
            wb = Workbook()
            ws = wb.active
            ws["A1"] = "h"
            ws["A2"] = "v"
            ws.add_table(Table(displayName="T1", ref="A1:A2"))
            wb.save(shell)
            out = t / "out.xlsx"
            wb2 = Workbook()
            wb2.active["A1"] = "h"
            wb2.save(out)
            findings = checks_deterministic.check_component_survival(
                shell, out, {"kind": "xlsx", "roles": {"_index": []}}
            )
            self.assertTrue(findings)
            self.assertTrue(
                all(f.location is not None for f in findings),
                "component_survival must carry an addressable location",
            )
            self.assertTrue(any(f.location == "tables" for f in findings))

    def test_strict_promotion_reads_triaged_list(self) -> None:
        # An EXPECTED full-bleed cover (visual.edge_bleed) must NOT hard-fail strict:
        # triage demotes the WARNING to INFO BEFORE the strict promoter runs, so it is
        # never promoted to a visual.strict ERROR. The single bottom-band bar below is
        # inset from the left/right edges so it produces ONLY the page:1:bottom finding.
        from PIL import ImageDraw

        prof = self._profile_with_triage(
            [
                {
                    "check": "visual.edge_bleed",
                    "location": "page:1:bottom",
                    "disposition": "expected",
                    "evidence": "cover is intentionally full-bleed",
                }
            ]
        )
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            target = td / "out.docx"
            _real_docx(target)
            png = td / "page-1.png"
            w, h = 850, 1100
            img = _blank(w, h)
            draw = ImageDraw.Draw(img)
            # Bottom-band bar inset from left/right -> ONLY page:1:bottom bleeds.
            draw.rectangle([w // 4, h - max(1, h // 100), 3 * w // 4, h], fill=0)
            img.save(png)
            l1 = vqa.run_visual_l1([png])
            self.assertEqual(
                sorted(f.location for f in l1 if f.check == "visual.edge_bleed"),
                ["page:1:bottom"],
                "fixture must bleed ONLY the bottom band",
            )
            report = gate.run_qa(
                target,
                prof,
                qa="strict",
                shell=None,
                out_dir=td / "out.visual",
                visual=(True, [png]),
            )
        # The edge_bleed WARNING was demoted to INFO and NEVER promoted to strict.
        eb = self._find(report, "visual.edge_bleed")
        self.assertTrue(eb)
        self.assertTrue(all(f.severity == schema.Severity.INFO.value for f in eb))
        strict_on_edge = [
            f
            for f in report.findings
            if f.check == "visual.strict" and "edge_bleed" in f.message
        ]
        self.assertEqual(
            strict_on_edge, [], "EXPECTED edge_bleed must not be promoted to strict"
        )


if __name__ == "__main__":
    unittest.main()
