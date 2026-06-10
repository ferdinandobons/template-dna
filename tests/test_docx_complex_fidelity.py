# SPDX-License-Identifier: MIT
"""DOCX fidelity regression tests on the committed COMPLEX fixture.

These reproduce, on ``tests/fixtures/complex/acme_complex.docx`` (a 100% synthetic
Acme template), the confirmed DOCX faithfulness fixes and assert the brand
artifacts now win over the generic/builtin floor:

  - D1  list ROLE NOMINATION from ``word/numbering.xml`` (a paragraph style whose
        definition carries a ``w:numPr`` binds ``list.<family>.<level>`` to the
        custom brand style, family read from the abstractNum ``w:numFmt``), AND
        the generator writes a real ``w:numPr`` (numId + per-item ilvl) so the
        list renders bulleted/numbered - not flat ``Normal`` paragraphs.
  - D2  ``table.default`` NOMINATION from the custom ``w:type="table"`` brand
        style (``Acme Table``), not the builtin ``Normal Table`` / ``Table Grid``.
  - D6  ``callout.info`` nominated from the brand boxed style (``Acme Callout``,
        ``w:shd`` + ``w:pBdr``), NOT the builtin ``Footnote Text`` (the ``note``
        lexicon token must not let a builtin win).
  - D4  a filled cover PARAGRAPH slot (not an SDT) re-applies its bound role
        style (``cover.title``), mirroring the SDT branch.

All assertions inspect the produced role registry and the real output OOXML.
"""

from __future__ import annotations

import copy
import hashlib
import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from brandkit.formats.docx import extract as docx_extract
from brandkit.formats.docx import generate as docx_generate
from brandkit.formats.docx import structure as docx_structure
from brandkit.formats.docx.structure import w
from brandkit.ir import model as ir
from brandkit.profile import schema

_COMPLEX_DOCX = (
    Path(__file__).resolve().parents[0] / "fixtures" / "complex" / "acme_complex.docx"
)


def _extract_profile(td: str) -> dict:
    """Extract the complex docx fixture into ``td`` and return the loaded profile."""
    old = Path.cwd()
    os.chdir(td)
    try:
        pj = docx_extract.extract(_COMPLEX_DOCX, "acme", scope="project")
        return json.loads(Path(pj).read_text())
    finally:
        os.chdir(old)


def _num_pr(p):
    """Return ``(numId, ilvl)`` of a paragraph's ``w:numPr``, or ``(None, None)``."""
    pPr = p._p.find(w("pPr"))
    if pPr is None:
        return None, None
    numPr = pPr.find(w("numPr"))
    if numPr is None:
        return None, None
    n = numPr.find(w("numId"))
    il = numPr.find(w("ilvl"))
    return (
        n.get(w("val")) if n is not None else None,
        il.get(w("val")) if il is not None else None,
    )


class _ClassScopedExtractTest(unittest.TestCase):
    """Extract the committed complex fixture ONCE per class and hand each test a
    deep copy (the same setUpClass pattern test_xlsx_complex_fidelity.py and
    test_examples_kitchen_sink.py already use).

    Safe by construction: the fixture is committed and immutable, extraction is
    deterministic (guarded by tests/test_fixture_determinism.py and the
    canonical JSON writer), and the per-test ``copy.deepcopy`` preserves
    isolation against any in-test profile mutation."""

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls._extract_td = tempfile.TemporaryDirectory()
        cls.addClassCleanup(cls._extract_td.cleanup)
        cls._class_profile = _extract_profile(cls._extract_td.name)

    def _profile_copy(self) -> dict:
        return copy.deepcopy(self._class_profile)


@unittest.skipUnless(_COMPLEX_DOCX.exists(), "complex docx fixture missing")
class DocxRoleNominationTest(_ClassScopedExtractTest):
    """The role registry promotes the PRESENT brand styles into their roles."""

    def setUp(self):
        self.prof = self._profile_copy()
        self.roles = self.prof["roles"]

    def test_callout_is_brand_style_not_builtin_footnote(self):
        # D6: 'Acme Callout' (boxed custom style) wins over builtin 'Footnote Text'.
        callout = self.roles.get("callout.info")
        self.assertIsNotNone(callout, "callout.info not nominated")
        r = callout["resolver"]
        self.assertEqual(r.get("style_id"), "AcmeCallout")
        self.assertEqual(r.get("style_name"), "Acme Callout")
        self.assertNotIn(r.get("style_id"), {"FootnoteText", "EndnoteText"})

    def test_table_default_is_custom_brand_style(self):
        # D2: the custom w:type='table' 'Acme Table' wins over Normal Table/Grid.
        table = self.roles.get("table.default")
        self.assertIsNotNone(table, "table.default not nominated")
        r = table["resolver"]
        self.assertEqual(r.get("style_id"), "AcmeTable")
        self.assertNotIn(r.get("style_id"), {"TableNormal", "TableGrid"})

    def test_list_roles_bind_custom_styles_with_numbering(self):
        # D1: list.bullet.1 / list.number.1 bind the CUSTOM brand styles and carry
        # the verbatim num_id from the numbering part (family from numFmt, not name).
        bullet = self.roles.get("list.bullet.1")
        number = self.roles.get("list.number.1")
        self.assertIsNotNone(bullet, "list.bullet.1 not nominated")
        self.assertIsNotNone(number, "list.number.1 not nominated")
        self.assertEqual(bullet["resolver"].get("style_id"), "AcmeBulletL1")
        self.assertEqual(number["resolver"].get("style_id"), "AcmeNumberL1")
        # The resolver carries the verbatim numbering id the generator re-asserts.
        self.assertEqual(str(bullet["resolver"].get("num_id")), "1")
        self.assertEqual(str(number["resolver"].get("num_id")), "3")

    def test_d3_numbering_facts_captured_per_list_role(self):
        # D3: each list role carries appearance.numbering with the referenced num_id /
        # abstract_num_id and the shell's OWN per-level numFmt / lvlText / indent facts.
        bullet = (self.roles["list.bullet.1"].get("appearance") or {}).get("numbering")
        self.assertIsNotNone(bullet, "list.bullet.1 appearance.numbering not captured")
        self.assertEqual(str(bullet["num_id"]), "1")
        self.assertEqual(str(bullet["abstract_num_id"]), "0")
        # The profile is loaded from JSON, so per_level_facts keys are STRINGS (the apply
        # and check layers coerce with int()).
        per = bullet["per_level_facts"]
        # abstractNum 0 is a TWO-level bullet: both declared levels are captured.
        self.assertEqual(per["0"]["numFmt"], "bullet")
        self.assertEqual(per["1"]["numFmt"], "bullet")
        # lvlText is captured VERBATIM (a bullet glyph, kept byte-for-byte).
        self.assertIn("lvlText", per["0"])
        self.assertEqual(per["0"]["indent"]["left"], 720)
        self.assertEqual(per["1"]["indent"]["left"], 1440)

        number = (self.roles["list.number.1"].get("appearance") or {}).get("numbering")
        self.assertIsNotNone(number, "list.number.1 appearance.numbering not captured")
        self.assertEqual(str(number["abstract_num_id"]), "1")
        self.assertEqual(number["per_level_facts"]["0"]["numFmt"], "decimal")
        self.assertEqual(number["per_level_facts"]["0"]["lvlText"], "%1.")

    def test_d3_numbering_axis_absent_when_no_list_numbering(self):
        # A non-list role (paragraph) never carries appearance.numbering.
        para = self.roles.get("paragraph") or {}
        self.assertNotIn("numbering", para.get("appearance") or {})

    def test_list_role_is_not_the_normal_floor(self):
        # The old 'Normal as list' floor must NOT be used when a real list style
        # exists - list.bullet.1 must not resolve to the body 'Normal' style.
        bullet = self.roles["list.bullet.1"]["resolver"]
        self.assertNotEqual(bullet.get("style_id"), "Normal")

    def test_builtin_list_style_loses_to_custom(self):
        # CC-2: Word's builtin latent 'List Bullet' (id ListBullet) references the
        # same numId, but the custom 'Acme Bullet L1' is nominated instead.
        self.assertEqual(
            self.roles["list.bullet.1"]["resolver"].get("style_id"), "AcmeBulletL1"
        )


@unittest.skipUnless(_COMPLEX_DOCX.exists(), "complex docx fixture missing")
class DocxListGenerationTest(_ClassScopedExtractTest):
    """The generator writes real numbering onto list items keyed by item.level."""

    def _generate(self, idoc, td) -> Document:
        prof = self._profile_copy()
        out = Path(td) / "out.docx"
        docx_generate.generate(prof, _COMPLEX_DOCX, idoc, out)
        return Document(out)

    def test_bulleted_and_numbered_lists_get_numPr_and_brand_style(self):
        with tempfile.TemporaryDirectory() as td:
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.ListBlock(
                        ordered=False,
                        items=[
                            ir.ListItem(
                                runs=[{"t": "Top bullet"}],
                                level=0,
                                items=[
                                    ir.ListItem(runs=[{"t": "Nested bullet"}], level=1)
                                ],
                            ),
                            ir.ListItem(runs=[{"t": "Second top"}], level=0),
                        ],
                    ),
                    ir.ListBlock(
                        ordered=True,
                        items=[
                            ir.ListItem(runs=[{"t": "Step one"}], level=0),
                            ir.ListItem(runs=[{"t": "Step two"}], level=0),
                        ],
                    ),
                ]
            )
            gen = self._generate(idoc, td)
            by_text = {p.text: p for p in gen.paragraphs if p.text}

            # Bullet items: brand style + numId 1; nested item gets ilvl 1.
            top = by_text["Top bullet"]
            self.assertEqual(top.style.name, "Acme Bullet L1")
            nid, ilvl = _num_pr(top)
            self.assertEqual(nid, "1")
            self.assertEqual(ilvl, "0")
            nested = by_text["Nested bullet"]
            nid_n, ilvl_n = _num_pr(nested)
            self.assertEqual(nid_n, "1")
            self.assertEqual(ilvl_n, "1")  # keyed by item.level

            # Numbered items: brand number style + numId 3.
            step = by_text["Step one"]
            self.assertEqual(step.style.name, "Acme Number L1")
            nid_s, _ = _num_pr(step)
            self.assertEqual(nid_s, "3")

    def test_list_items_are_not_flat_normal_paragraphs(self):
        # Regression on D1's failure mode: a list must not render as 'Normal' with
        # no numbering.
        with tempfile.TemporaryDirectory() as td:
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.ListBlock(
                        ordered=False,
                        items=[ir.ListItem(runs=[{"t": "Has bullets"}], level=0)],
                    )
                ]
            )
            gen = self._generate(idoc, td)
            p = next(p for p in gen.paragraphs if p.text == "Has bullets")
            self.assertNotEqual(p.style.name, "Normal")
            nid, _ = _num_pr(p)
            self.assertIsNotNone(nid, "list paragraph has no numPr")

    def test_table_uses_brand_table_style(self):
        # D2 through full generation: the produced table carries 'Acme Table'.
        with tempfile.TemporaryDirectory() as td:
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Table(
                        columns=[{"t": "A"}, {"t": "B"}],
                        rows=[
                            [
                                ir.TableCell(runs=[{"t": "1"}]),
                                ir.TableCell(runs=[{"t": "2"}]),
                            ]
                        ],
                    )
                ]
            )
            gen = self._generate(idoc, td)
            self.assertEqual(gen.tables[-1].style.name, "Acme Table")

    def test_callout_renders_with_brand_callout_style(self):
        with tempfile.TemporaryDirectory() as td:
            idoc = ir.IntermediateDocument(
                blocks=[ir.Callout(intent="info", runs=[{"t": "Important note"}])]
            )
            gen = self._generate(idoc, td)
            p = next(p for p in gen.paragraphs if "Important note" in p.text)
            self.assertEqual(p.style.name, "Acme Callout")

    def test_list_generation_is_idempotent(self):
        # The numPr writes must be deterministic (idempotent generation).
        with tempfile.TemporaryDirectory() as td:
            prof = self._profile_copy()
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.ListBlock(
                        ordered=False,
                        items=[
                            ir.ListItem(
                                runs=[{"t": "A"}],
                                level=0,
                                items=[ir.ListItem(runs=[{"t": "B"}], level=1)],
                            )
                        ],
                    )
                ]
            )
            o1 = Path(td) / "o1.docx"
            o2 = Path(td) / "o2.docx"
            docx_generate.generate(prof, _COMPLEX_DOCX, idoc, o1)
            docx_generate.generate(prof, _COMPLEX_DOCX, idoc, o2)
            self.assertEqual(
                hashlib.sha256(o1.read_bytes()).hexdigest(),
                hashlib.sha256(o2.read_bytes()).hexdigest(),
            )


@unittest.skipUnless(_COMPLEX_DOCX.exists(), "complex docx fixture missing")
class DocxCoverParagraphStyleTest(unittest.TestCase):
    """D4: a filled cover PARAGRAPH slot re-applies its bound role style."""

    def _build_paragraph_cover_shell(self, td) -> Path:
        """A shell whose cover title is a plain paragraph slot (no SDT), with a
        custom 'Brand Cover Title' style and a TOC after, so the slot is in the
        cover region."""
        shell = Path(td) / "shell.docx"
        doc = Document()
        styles = doc.styles.element
        st = OxmlElement("w:style")
        st.set(qn("w:type"), "paragraph")
        st.set(qn("w:styleId"), "BrandCoverTitle")
        st.set(qn("w:customStyle"), "1")
        nm = OxmlElement("w:name")
        nm.set(qn("w:val"), "Brand Cover Title")
        st.append(nm)
        styles.append(st)
        doc.add_paragraph("{{title}}")  # plain-paragraph cover slot
        # a real TOC field after the slot
        p = doc.add_paragraph()
        r = p.add_run()
        fb = OxmlElement("w:fldChar")
        fb.set(qn("w:fldCharType"), "begin")
        r._r.append(fb)
        r2 = p.add_run()
        it = OxmlElement("w:instrText")
        it.text = 'TOC \\o "1-3" \\h'
        r2._r.append(it)
        r3 = p.add_run()
        fs = OxmlElement("w:fldChar")
        fs.set(qn("w:fldCharType"), "separate")
        r3._r.append(fs)
        p.add_run("entry 1")
        r4 = p.add_run()
        fe = OxmlElement("w:fldChar")
        fe.set(qn("w:fldCharType"), "end")
        r4._r.append(fe)
        doc.add_paragraph("Body Heading", style="Heading 1")
        doc.save(shell)
        return shell

    def test_filled_paragraph_slot_carries_cover_title_style(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._build_paragraph_cover_shell(td)
            # extract a real profile, then bind cover.title to the brand style.
            old = Path.cwd()
            os.chdir(td)
            try:
                pj = docx_extract.extract(shell, "cv", scope="project")
                prof = json.loads(Path(pj).read_text())
            finally:
                os.chdir(old)
            if "cover.title" not in prof["roles"]:
                prof["roles"].setdefault("_index", []).append("cover.title")
            prof["roles"]["cover.title"] = {
                "resolver": {
                    "type": "named_style",
                    "style_id": "BrandCoverTitle",
                    "style_name": "Brand Cover Title",
                },
                "status": "robust",
                "confidence": 1.0,
            }
            para_anchor = next(
                a["id"]
                for a in prof["surface"]["docx"]["cover_anchors"]
                if a["id"].startswith("para.")
            )
            prof.setdefault("provenance", {}).setdefault("shell", {})["sha256"] = "sh"
            block = schema.empty_comprehension()
            block["status"] = schema.ComprehensionStatus.PRESENT.value
            block["source_shell_sha256"] = "sh"
            block["confidence"] = 0.9
            block["cover_slots"] = {
                para_anchor: {
                    "binds_to": "title",
                    "fill_rule": "in_place",
                    "demo_value": "{{title}}",
                }
            }
            prof["comprehension"] = block

            out = Path(td) / "out.docx"
            docx_generate.generate(
                prof,
                shell,
                ir.IntermediateDocument(
                    blocks=[ir.Heading(level=1, runs=[{"t": "Sec"}])],
                    cover=ir.Cover(title=[{"t": "My Real Title"}]),
                ),
                out,
            )
            gen = Document(out)
            filled = next(p for p in gen.paragraphs if "My Real Title" in p.text)
            self.assertEqual(filled.style.name, "Brand Cover Title")


@unittest.skipUnless(_COMPLEX_DOCX.exists(), "complex docx fixture missing")
class DocxNumberingHelpersTest(unittest.TestCase):
    """Unit coverage for the structural numbering helpers behind D1."""

    def test_num_family_for_reads_numFmt_not_name(self):
        doc = Document(_COMPLEX_DOCX)
        # numId 1 -> abstractNum 0 (bullet); numId 3 -> abstractNum 1 (decimal).
        self.assertEqual(docx_structure.num_family_for(doc, "1", 0), "bullet")
        self.assertEqual(docx_structure.num_family_for(doc, "3", 0), "number")
        # an unknown id yields None (skip rather than guess).
        self.assertIsNone(docx_structure.num_family_for(doc, "9999", 0))

    def test_style_num_binding_reads_style_definition_numPr(self):
        doc = Document(_COMPLEX_DOCX)
        from docx.enum.style import WD_STYLE_TYPE

        acme_bullet = next(
            s
            for s in doc.styles
            if s.type == WD_STYLE_TYPE.PARAGRAPH
            and getattr(s, "style_id", None) == "AcmeBulletL1"
        )
        binding = docx_structure.style_num_binding(acme_bullet)
        self.assertEqual(binding, ("1", 0))

    def test_no_numbering_part_is_not_an_error(self):
        # A document with no list numbering must not crash the helpers.
        doc = Document()
        # default template has a numbering part, but no nominated list style with
        # a custom numPr binding maps to a family resolvable via num_family_for for
        # an absent id; the helper must simply return None.
        self.assertIsNone(docx_structure.num_family_for(doc, "424242", 0))


def _define_table_style(doc, style_id="AcmeTable", name="Acme Table"):
    """Add a custom ``w:type='table'`` style so a captured/applied style id is a member
    of the shell's table-style inventory."""
    styles = doc.styles.element
    st = OxmlElement("w:style")
    st.set(qn("w:type"), "table")
    st.set(qn("w:styleId"), style_id)
    st.set(qn("w:customStyle"), "1")
    nm = OxmlElement("w:name")
    nm.set(qn("w:val"), name)
    st.append(nm)
    bo = OxmlElement("w:basedOn")
    bo.set(qn("w:val"), "TableNormal")
    st.append(bo)
    styles.append(st)


def _brand_table(doc, *, tbllook="01E0", margins=None, style_name="Acme Table"):
    """Add a 2x2 table with an explicit ``w:tblLook@w:val`` + ``w:tblStyle`` + optional
    ``w:tblCellMar`` margins."""
    t = doc.add_table(rows=2, cols=2)
    if style_name is not None:
        t.style = style_name
    tblpr = t._tbl.tblPr
    if tbllook is not None:
        look = tblpr.find(qn("w:tblLook"))
        if look is None:
            look = OxmlElement("w:tblLook")
            tblpr.append(look)
        look.set(qn("w:val"), tbllook)
    if margins:
        cm = OxmlElement("w:tblCellMar")
        for side, w in margins.items():
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(w))
            el.set(qn("w:type"), "dxa")
            cm.append(el)
        tblpr.append(cm)
    return t


class DocxTableD2FidelityTest(unittest.TestCase):
    """D2: the template's own table conditional-format facts (tblLook bitmask, table
    style reference, cell margins) are captured and re-applied set-only-when-unset; the
    band FILLS stay in the shell's style part (the engine only toggles)."""

    def _template(self, td, *, tbllook="01E0", margins=None):
        template = Path(td) / "template.docx"
        d = Document()
        _define_table_style(d)
        for _ in range(3):
            _brand_table(d, tbllook=tbllook, margins=margins)
        d.save(template)
        return template

    def _profile(self, td, template):
        cwd = Path.cwd()
        os.chdir(td)
        try:
            pj = docx_extract.extract(template, "tbl", scope="project")
            return json.loads(Path(pj).read_text())
        finally:
            os.chdir(cwd)

    def _tblpr_xml(self, out):
        from lxml import etree

        return "\n".join(
            etree.tostring(t._tbl.tblPr, encoding="unicode")
            for t in Document(out).tables
        )

    def _idoc(self):
        return ir.IntermediateDocument(
            blocks=[
                ir.Table(
                    columns=[{"t": "A"}, {"t": "B"}],
                    rows=[
                        [
                            ir.TableCell(runs=[{"t": "1"}]),
                            ir.TableCell(runs=[{"t": "2"}]),
                        ]
                    ],
                )
            ]
        )

    def test_table_tblLook_applied_set_only_when_unset(self):
        with tempfile.TemporaryDirectory() as td:
            template = self._template(td, tbllook="01E0")
            prof = self._profile(td, template)
            self.assertEqual(prof["theme"]["table"]["body"]["tblLook"], 0x01E0)
            out = Path(td) / "out.docx"
            docx_generate.generate(prof, template, self._idoc(), out)
            xml = self._tblpr_xml(out)
            self.assertIn('w:val="01E0"', xml)
            # python-docx's synthetic default 04A0 is replaced (the captured look wins).
            self.assertNotIn('w:val="04A0"', xml)

    def test_table_style_applied(self):
        with tempfile.TemporaryDirectory() as td:
            template = self._template(td)
            prof = self._profile(td, template)
            self.assertEqual(prof["theme"]["table"]["body"]["style_id"], "AcmeTable")
            out = Path(td) / "out.docx"
            docx_generate.generate(prof, template, self._idoc(), out)
            xml = self._tblpr_xml(out)
            # the table style reference is present (only once per table).
            self.assertEqual(xml.count('w:tblStyle w:val="AcmeTable"'), 1)

    def test_table_cell_margins_applied_set_only_when_unset(self):
        with tempfile.TemporaryDirectory() as td:
            template = self._template(
                td, margins={"top": 120, "bottom": 120, "left": 80}
            )
            prof = self._profile(td, template)
            margins = prof["theme"]["table"]["body"]["cell_margins"]
            self.assertEqual(margins["top_twips"], 120)
            self.assertEqual(margins["left_twips"], 80)
            out = Path(td) / "out.docx"
            docx_generate.generate(prof, template, self._idoc(), out)
            xml = self._tblpr_xml(out)
            self.assertIn('w:w="120"', xml)
            self.assertIn('w:w="80"', xml)

    def test_table_end_to_end(self):
        # Extract a template with brand table style + tblLook + margins, generate on the
        # same shell, verify the output table carries the same facts and the QA gate
        # surfaces NO table-target ERROR.
        from brandkit.qa import gate

        with tempfile.TemporaryDirectory() as td:
            template = self._template(td, tbllook="01E0", margins={"left": 80})
            prof = self._profile(td, template)
            out = Path(td) / "out.docx"
            docx_generate.generate(prof, template, self._idoc(), out)
            xml = self._tblpr_xml(out)
            self.assertIn('w:val="01E0"', xml)
            self.assertIn("AcmeTable", xml)
            self.assertIn('w:w="80"', xml)
            # The injected visual seam keeps the visual-audit branch executing in
            # degraded mode without spawning soffice; the asserted finding is a
            # deterministic L0 fact assembled BEFORE the visual path.
            report = gate.run_qa(out, prof, shell=template, visual=(False, []))
            self.assertFalse(
                any(
                    f.check == "appearance_table_targets"
                    and f.severity == schema.Severity.ERROR.value
                    for f in report.findings
                )
            )


@unittest.skipUnless(_COMPLEX_DOCX.exists(), "complex docx fixture missing")
class DocxNumberingD3FidelityTest(_ClassScopedExtractTest):
    """D3: the template's own per-level numbering facts (numFmt / lvlText / indent) are
    captured and re-asserted onto the output's cloned w:abstractNum set-only-when-unset;
    the numbering DEFINITION stays the shell's (referenced/cloned by id, never minted)."""

    def _generate(self, idoc, td) -> Document:
        prof = self._profile_copy()
        out = Path(td) / "out.docx"
        docx_generate.generate(prof, _COMPLEX_DOCX, idoc, out)
        return Document(out)

    def _list_idoc(self):
        return ir.IntermediateDocument(
            blocks=[
                ir.ListBlock(
                    ordered=False,
                    items=[
                        ir.ListItem(
                            runs=[{"t": "Top bullet"}],
                            level=0,
                            items=[ir.ListItem(runs=[{"t": "Nested bullet"}], level=1)],
                        )
                    ],
                ),
                ir.ListBlock(
                    ordered=True,
                    items=[ir.ListItem(runs=[{"t": "Step one"}], level=0)],
                ),
            ]
        )

    def _abstract_level(self, doc, abstract_num_id, ilvl):
        """The ``w:lvl`` element of ``abstract_num_id`` at ``ilvl`` in ``doc``'s
        numbering part, or None."""
        root = docx_structure._numbering_root(doc)
        for an in root.findall(w("abstractNum")):
            if an.get(w("abstractNumId")) == str(abstract_num_id):
                for lvl in an.findall(w("lvl")):
                    if (lvl.get(w("ilvl")) or "0") == str(ilvl):
                        return lvl
        return None

    def test_referenced_abstractnum_is_present_in_output(self):
        # The generated list references the shell's num_id; the shell's own w:abstractNum
        # (0 for bullets, 1 for decimals) is present in the output's numbering part.
        with tempfile.TemporaryDirectory() as td:
            gen = self._generate(self._list_idoc(), td)
            root = docx_structure._numbering_root(gen)
            ids = {an.get(w("abstractNumId")) for an in root.findall(w("abstractNum"))}
            self.assertIn("0", ids)
            self.assertIn("1", ids)

    def test_per_level_facts_match_template_on_output(self):
        # The output's w:abstractNum carries the template's OWN per-level numFmt /
        # lvlText / indent (re-asserted set-only-when-unset onto the cloned def).
        with tempfile.TemporaryDirectory() as td:
            shell = Document(_COMPLEX_DOCX)
            gen = self._generate(self._list_idoc(), td)
            # bullet abstractNum 0, level 0: numFmt/lvlText/indent equal the shell's.
            for aid, ilvl in (("0", 0), ("0", 1), ("1", 0)):
                shell_lvl = self._abstract_level(shell, aid, ilvl)
                out_lvl = self._abstract_level(gen, aid, ilvl)
                self.assertIsNotNone(out_lvl, f"abstract {aid} level {ilvl} missing")
                for tag in ("numFmt", "lvlText"):
                    s = shell_lvl.find(w(tag))
                    o = out_lvl.find(w(tag))
                    if s is not None:
                        self.assertEqual(
                            o.get(w("val")), s.get(w("val")), f"{aid}.{ilvl} {tag}"
                        )

    def test_lists_render_with_template_numbering(self):
        # The generated list paragraphs reference the template's numId, and the abstractNum
        # those ids resolve to carries the template's numFmt (bullet vs decimal).
        with tempfile.TemporaryDirectory() as td:
            gen = self._generate(self._list_idoc(), td)
            by_text = {p.text: p for p in gen.paragraphs if p.text}
            nid_b, _ = _num_pr(by_text["Top bullet"])
            self.assertEqual(nid_b, "1")
            self.assertEqual(docx_structure.num_family_for(gen, nid_b, 0), "bullet")
            nid_n, _ = _num_pr(by_text["Step one"])
            self.assertEqual(nid_n, "3")
            self.assertEqual(docx_structure.num_family_for(gen, nid_n, 0), "number")

    def test_numbering_generation_is_idempotent(self):
        # Re-generating the same content twice yields byte-identical output (the per-level
        # re-assert is set-only-when-unset, so a second run never re-touches the def).
        with tempfile.TemporaryDirectory() as td:
            prof = self._profile_copy()
            idoc = self._list_idoc()
            o1 = Path(td) / "o1.docx"
            o2 = Path(td) / "o2.docx"
            docx_generate.generate(prof, _COMPLEX_DOCX, idoc, o1)
            docx_generate.generate(prof, _COMPLEX_DOCX, idoc, o2)
            self.assertEqual(
                hashlib.sha256(o1.read_bytes()).hexdigest(),
                hashlib.sha256(o2.read_bytes()).hexdigest(),
            )

    def test_end_to_end_gate_has_no_numbering_error(self):
        from brandkit.qa import gate

        with tempfile.TemporaryDirectory() as td:
            prof = self._profile_copy()
            out = Path(td) / "out.docx"
            docx_generate.generate(prof, _COMPLEX_DOCX, self._list_idoc(), out)
            # The injected visual seam keeps the visual-audit branch executing in
            # degraded mode without spawning soffice; the asserted ERROR-absence is
            # a deterministic L0 fact assembled BEFORE the visual path.
            report = gate.run_qa(out, prof, shell=_COMPLEX_DOCX, visual=(False, []))
            self.assertFalse(
                any(
                    f.check == "appearance_numbering_targets"
                    and f.severity == schema.Severity.ERROR.value
                    for f in report.findings
                ),
                [
                    f.message
                    for f in report.findings
                    if f.check == "appearance_numbering_targets"
                ],
            )

    def test_set_only_when_unset_does_not_clobber_authored_level(self):
        # Manually edit the output's abstractNum to carry a DIFFERENT (authored) lvlText on
        # bullet level 0, then re-generate the same content onto that edited package: the
        # authored value must survive (the per-level re-assert is set-only-when-unset).
        with tempfile.TemporaryDirectory() as td:
            prof = self._profile_copy()
            stage = Path(td) / "stage.docx"
            docx_generate.generate(prof, _COMPLEX_DOCX, self._list_idoc(), stage)
            d = Document(stage)
            lvl = self._abstract_level(d, "0", 0)
            lt = lvl.find(w("lvlText"))
            lt.set(w("val"), "AUTHORED")
            d.save(stage)
            # Re-generate onto the EDITED package (use it as the shell).
            out = Path(td) / "out.docx"
            docx_generate.generate(prof, stage, self._list_idoc(), out)
            out_lvl = self._abstract_level(Document(out), "0", 0)
            self.assertEqual(out_lvl.find(w("lvlText")).get(w("val")), "AUTHORED")


if __name__ == "__main__":
    unittest.main()
