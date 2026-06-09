# SPDX-License-Identifier: MIT
"""Deterministic ``learn`` verb + single ``merge_overrides`` sink (Cluster B, B3c).

These exercise the model-free learn writer and its canonical sink:

  - ``learn`` THRESHOLD: a ``(check, location)`` pair seen ONCE is not distilled;
    one that recurred across ``>= N`` same-shell runs is;
  - ``learn`` ACYCLIC guard: an A->B->A reroute proposal is rejected and NOTHING is
    written (the colored-DFS backstop);
  - ``merge_overrides`` ALL-OR-NOTHING: one unbound pointer rejects the WHOLE
    proposal and writes nothing load-bearing (``status='rejected'``);
  - ``merge_overrides`` CLEAN: a fully-bound, acyclic proposal becomes a canonical
    sorted ``present`` block with ``source_shell_sha256`` stamped + provenance;
  - number_format SWAP membership: a swap to a shell-backed mask binds; a swap to a
    non-shell mask is rejected by membership AND ERRORed by ``check_override_targets``;
  - SHELL-FROZEN: the sink stamps the LIVE shell sha; a re-extract drift empties the
    consumer (``overrides_are_present`` False);
  - the single-sink invariant: ``learn`` writes the IDENTICAL canonical shape
    ``merge_overrides`` does (no parallel writer/sort order).
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.shared import Pt, RGBColor

from brandkit.cli import main
from brandkit.formats.docx import generate as docx_generate
from brandkit.ir import model as ir
from brandkit.profile import overrides as ov
from brandkit.profile import schema, store
from brandkit.qa import checks_deterministic as cd
from brandkit.qa import report as vreport
from brandkit.qa.model import QAReport


# ---------------------------------------------------------------------------
# Profile builders (mirror test_overrides_resolver.py, but schema-valid so the
# sink's shape gate passes)
# ---------------------------------------------------------------------------
def _docx_profile() -> dict:
    """A schema-valid docx profile: body + heading.1 (both concretely resolved)."""
    prof = schema.build_envelope("docx", {"name": "t"})
    prof["provenance"]["shell"]["sha256"] = "shell-sha-1"
    prof["roles"] = {
        "_index": ["paragraph", "heading.1"],
        "paragraph": {
            "resolver": {
                "type": "named_style",
                "style_id": "Normal",
                "style_name": "Normal",
            },
            "status": "robust",
            "confidence": 1.0,
        },
        "heading.1": {
            "resolver": {
                "type": "named_style",
                "style_id": "Heading1",
                "style_name": "Heading 1",
            },
            "status": "robust",
            "confidence": 1.0,
        },
    }
    return prof


def _with_stub_heading(prof: dict, role: str = "heading.9") -> dict:
    """Add a STUB heading role (empty resolver) so a reroute lesson can target a
    healthy same-family sibling (heading.1)."""
    prof["roles"]["_index"].append(role)
    prof["roles"][role] = {"resolver": {}, "status": "stub"}
    return prof


def _xlsx_profile() -> dict:
    prof = schema.build_envelope("xlsx", {"name": "t"})
    prof["provenance"]["shell"]["sha256"] = "xlsx-sha-1"
    prof["roles"] = {
        "_index": ["metric.value"],
        "metric.value": {
            "resolver": {"type": "number_format", "number_format": "0.00"},
            "status": "robust",
            "confidence": 1.0,
        },
    }
    # The surface inventory the swap mask binds against (shell uses #,##0 + 0.00).
    prof["surface"] = {
        "xlsx": {
            "number_formats": [
                {"format": "#,##0", "count": 3},
                {"format": "0.00", "count": 1},
            ]
        }
    }
    return prof


def _report(
    loc: str, *, check: str = "resolver_targets_exist", sha: str = "shell-sha-1"
) -> dict:
    """A minimal same-shell generation_report.json digest carrying one finding."""
    return {
        "schema_version": "generation-report-1",
        "kind": "docx",
        "shell_sha256": sha,
        "findings": [
            {
                "check": check,
                "severity": "ERROR",
                "message": "brand body",
                "location": loc,
            }
        ],
        "generated_at": "2026-06-09T00:00:00Z",
    }


# ---------------------------------------------------------------------------
# learn threshold
# ---------------------------------------------------------------------------
class LearnThresholdTest(unittest.TestCase):
    def test_seen_once_is_not_distilled(self):
        """A finding present in a SINGLE run does not cross min_runs (default 2)."""
        prof = _with_stub_heading(_docx_profile())
        res = ov.learn(prof, [_report("heading.9")])
        self.assertFalse(res.ok)
        self.assertEqual(res.distilled, 0)
        self.assertEqual(res.status, schema.ComprehensionStatus.ABSENT.value)
        # Nothing live-eligible was written.
        self.assertFalse(store.overrides_are_present(prof))
        self.assertEqual(prof["rules"]["overrides"]["reroute_roles"], {})

    def test_recurred_twice_is_distilled(self):
        """A finding recurred across >= 2 same-shell runs IS distilled into a lesson."""
        prof = _with_stub_heading(_docx_profile())
        res = ov.learn(prof, [_report("heading.9"), _report("heading.9")])
        self.assertTrue(res.ok)
        self.assertEqual(res.distilled, 1)
        self.assertEqual(res.status, schema.ComprehensionStatus.PRESENT.value)
        block = prof["rules"]["overrides"]
        self.assertEqual(block["reroute_roles"], {"heading.9": "heading.1"})
        # Provenance records which (check, location, recurred_runs) produced it.
        prov = block["provenance"]["reroute_roles.heading.9"]
        self.assertEqual(prov["check"], "resolver_targets_exist")
        self.assertEqual(prov["location"], "heading.9")
        self.assertEqual(prov["recurred_runs"], 2)
        # The lesson is live-eligible (present + sha-bound + non-empty).
        self.assertTrue(store.overrides_are_present(prof))

    def test_custom_min_runs_threshold(self):
        """At min_runs=3, two runs is below threshold; three runs crosses it."""
        below = _with_stub_heading(_docx_profile())
        r1 = ov.learn(below, [_report("heading.9"), _report("heading.9")], min_runs=3)
        self.assertFalse(r1.ok)
        self.assertEqual(r1.distilled, 0)

        at = _with_stub_heading(_docx_profile())
        r2 = ov.learn(
            at,
            [_report("heading.9"), _report("heading.9"), _report("heading.9")],
            min_runs=3,
        )
        self.assertTrue(r2.ok)
        self.assertEqual(r2.distilled, 1)

    def test_non_learnable_check_is_ignored(self):
        """A recurring finding whose check is NOT in LEARNABLE_CHECKS is not distilled."""
        prof = _with_stub_heading(_docx_profile())
        res = ov.learn(
            prof,
            [
                _report("heading.9", check="overflow"),
                _report("heading.9", check="overflow"),
            ],
        )
        self.assertFalse(res.ok)
        self.assertEqual(res.distilled, 0)

    def test_different_shell_history_does_not_contribute(self):
        """A prior report for a DIFFERENT shell sha is partitioned out (SHELL-FROZEN)."""
        prof = _with_stub_heading(_docx_profile())
        res = ov.learn(
            prof,
            [
                _report("heading.9", sha="other-sha"),
                _report("heading.9", sha="other-sha"),
            ],
        )
        self.assertFalse(res.ok)
        self.assertEqual(res.distilled, 0)


# ---------------------------------------------------------------------------
# learn acyclic guard
# ---------------------------------------------------------------------------
class LearnAcyclicGuardTest(unittest.TestCase):
    def test_cycle_rejected_nothing_written(self):
        """An A->B->A reroute proposal is rejected by the sink; nothing is written."""
        prof = _docx_profile()
        # Direct merge_overrides proposal (learn maps single-source->target, which is
        # structurally acyclic; this exercises the backstop the sink enforces for any
        # multi-entry proposal, incl. a future B4 model one).
        result = ov.merge_overrides(
            prof,
            {"reroute_roles": {"paragraph": "heading.1", "heading.1": "paragraph"}},
        )
        self.assertFalse(result.ok)
        self.assertEqual(result.status, schema.ComprehensionStatus.REJECTED.value)
        self.assertTrue(any("cycle" in p.lower() for p in result.problems))
        # Nothing load-bearing written.
        block = prof["rules"]["overrides"]
        self.assertEqual(block["status"], schema.ComprehensionStatus.REJECTED.value)
        self.assertEqual(block["reroute_roles"], {})
        self.assertFalse(store.overrides_are_present(prof))


# ---------------------------------------------------------------------------
# merge_overrides all-or-nothing + clean
# ---------------------------------------------------------------------------
class MergeOverridesTest(unittest.TestCase):
    def test_all_or_nothing_one_unbound_pointer_rejects(self):
        """One unbound reroute target rejects the WHOLE proposal; the bound entry is
        NOT partially written."""
        prof = _docx_profile()
        result = ov.merge_overrides(
            prof,
            {"reroute_roles": {"paragraph": "heading.1", "x": "role.not.declared"}},
        )
        self.assertFalse(result.ok)
        self.assertEqual(result.status, schema.ComprehensionStatus.REJECTED.value)
        block = prof["rules"]["overrides"]
        self.assertEqual(block["reroute_roles"], {})  # nothing partial
        self.assertFalse(store.overrides_are_present(prof))

    def test_clean_proposal_writes_present_sorted_stamped(self):
        """A fully-bound acyclic proposal becomes a canonical present block."""
        prof = _docx_profile()
        result = ov.merge_overrides(
            prof,
            {"reroute_roles": {"paragraph": "heading.1"}, "confidence": 0.9},
            generated_by={"model": "deterministic", "prompt_version": "n/a"},
        )
        self.assertTrue(result.ok)
        block = prof["rules"]["overrides"]
        self.assertEqual(block["status"], schema.ComprehensionStatus.PRESENT.value)
        self.assertEqual(block["source_shell_sha256"], "shell-sha-1")
        self.assertEqual(block["reroute_roles"], {"paragraph": "heading.1"})
        self.assertEqual(block["confidence"], 0.9)
        self.assertEqual(block["generated_by"]["model"], "deterministic")
        self.assertTrue(store.overrides_are_present(prof))

    def test_supplied_status_is_disposed(self):
        """A proposal-supplied status='absent' cannot short-circuit the binding gate
        (the sink forces present for the trial, then writes present on a clean bind)."""
        prof = _docx_profile()
        result = ov.merge_overrides(
            prof, {"status": "absent", "reroute_roles": {"paragraph": "heading.1"}}
        )
        self.assertTrue(result.ok)
        self.assertEqual(
            prof["rules"]["overrides"]["status"],
            schema.ComprehensionStatus.PRESENT.value,
        )

    def test_idempotent_canonical_shape(self):
        """Re-merging the same proposal yields the identical canonical block."""
        p1 = _docx_profile()
        p2 = _docx_profile()
        proposal = {"reroute_roles": {"paragraph": "heading.1"}, "confidence": 0.7}
        ov.merge_overrides(p1, dict(proposal))
        ov.merge_overrides(p2, dict(proposal))
        self.assertEqual(p1["rules"]["overrides"], p2["rules"]["overrides"])


# ---------------------------------------------------------------------------
# number_format swap membership (learn + merge + check_override_targets)
# ---------------------------------------------------------------------------
class NumberFormatSwapMembershipTest(unittest.TestCase):
    def test_swap_to_shell_mask_binds(self):
        prof = _xlsx_profile()
        result = ov.merge_overrides(
            prof, {"number_format_swaps": {"metric.value": "#,##0"}}
        )
        self.assertTrue(result.ok)
        self.assertEqual(
            prof["rules"]["overrides"]["number_format_swaps"], {"metric.value": "#,##0"}
        )

    def test_swap_to_non_shell_mask_rejected_by_merge(self):
        prof = _xlsx_profile()
        result = ov.merge_overrides(
            prof, {"number_format_swaps": {"metric.value": "0%"}}
        )
        self.assertFalse(result.ok)
        self.assertEqual(prof["rules"]["overrides"]["number_format_swaps"], {})

    def test_non_shell_swap_errored_by_check_override_targets(self):
        """A non-shell mask that somehow reaches a present block is ERRORed at verify."""
        prof = _xlsx_profile()
        ovb = prof["rules"]["overrides"]
        ovb["status"] = "present"
        ovb["source_shell_sha256"] = prof["provenance"]["shell"]["sha256"]
        ovb["number_format_swaps"] = {"metric.value": "0%"}

        class _FakeShell:
            pass

        orig_masks = cd._xlsx_number_format_masks
        orig_names = cd._xlsx_defined_names
        cd._xlsx_number_format_masks = lambda shell: {"#,##0", "0.00"}
        cd._xlsx_defined_names = lambda shell: set()
        try:
            errs = cd.check_override_targets(_FakeShell(), prof)
            self.assertTrue(errs)
            self.assertEqual(errs[0].check, "override_targets_exist")
            self.assertEqual(errs[0].severity, schema.Severity.ERROR.value)
        finally:
            cd._xlsx_number_format_masks = orig_masks
            cd._xlsx_defined_names = orig_names


# ---------------------------------------------------------------------------
# demo-clear distillation
# ---------------------------------------------------------------------------
class LearnDemoClearTest(unittest.TestCase):
    def test_recurring_residual_distills_demo_clear(self):
        prof = _docx_profile()
        # Surface a captured demo string so the residual value binds.
        prof["surface"] = {
            "docx": {"demo_region": {"instruction_markers": ["Lorem ipsum"]}}
        }
        reports = [
            _report("Lorem ipsum", check="no_residual_template_text"),
            _report("Lorem ipsum", check="no_residual_template_text"),
        ]
        res = ov.learn(prof, reports)
        self.assertTrue(res.ok)
        self.assertEqual(prof["rules"]["overrides"]["demo_clears"], ["Lorem ipsum"])

    def test_uncaptured_residual_is_dropped(self):
        prof = _docx_profile()  # no captured demo set
        reports = [
            _report("Lorem ipsum", check="no_residual_template_text"),
            _report("Lorem ipsum", check="no_residual_template_text"),
        ]
        res = ov.learn(prof, reports)
        self.assertFalse(res.ok)
        self.assertEqual(res.distilled, 0)


# ---------------------------------------------------------------------------
# single-sink invariant
# ---------------------------------------------------------------------------
class SingleSinkTest(unittest.TestCase):
    def test_learn_writes_identical_shape_to_direct_merge(self):
        """learn's output is byte-identical to a direct merge_overrides of the same
        derived proposal (one sink, one sort order)."""
        learned = _with_stub_heading(_docx_profile())
        ov.learn(learned, [_report("heading.9"), _report("heading.9")])

        # The proposal learn derives for this history: reroute heading.9 -> heading.1.
        direct = _with_stub_heading(_docx_profile())
        ov.merge_overrides(
            direct,
            {
                "reroute_roles": {"heading.9": "heading.1"},
                "confidence": learned["rules"]["overrides"]["confidence"],
                "provenance": learned["rules"]["overrides"]["provenance"],
            },
        )
        self.assertEqual(learned["rules"]["overrides"], direct["rules"]["overrides"])


# ---------------------------------------------------------------------------
# shell-frozen
# ---------------------------------------------------------------------------
class LearnShellFrozenTest(unittest.TestCase):
    def test_reextract_drift_empties_consumer(self):
        prof = _with_stub_heading(_docx_profile())
        ov.learn(prof, [_report("heading.9"), _report("heading.9")])
        self.assertTrue(store.overrides_are_present(prof))
        # A re-extract re-stamps provenance.shell.sha256; the recorded sha drifts.
        prof["provenance"]["shell"]["sha256"] = "re-extracted-sha"
        self.assertFalse(store.overrides_are_present(prof))


# ---------------------------------------------------------------------------
# REGRESSION: learn distills from findings the REAL producers emit
# ---------------------------------------------------------------------------
class LearnFromRealProducerFindingsTest(unittest.TestCase):
    """The learn loop must work on findings as PRODUCTION code emits them.

    Regression for the review HIGH: every other learn test hand-populates
    ``location`` in its fixture reports - values the producers historically never
    emitted (they carried the pointer only in the brand-bearing ``message``), so
    the loop was inert on real history. These tests drive the REAL producers
    (``check_residual_template_text``, ``check_resolver_targets``, and the docx
    ``generate`` writer's ``style_fallback``), persist their findings through the
    REAL ``build_generation_report`` writer, re-discover them through the REAL
    ``discover_prior_reports``, and assert ``learn`` distills a lesson - proving
    the structured ``location`` is populated at the source, end-to-end.
    """

    def _persist_and_discover(self, prof: dict, findings: list, tmp: Path) -> list:
        """Write two same-shell runs via the REAL report writer, then re-discover
        them via the REAL sibling-dir discovery (the exact CLI read path)."""
        qa = QAReport(verdict="failed", findings=list(findings))
        for i in range(2):
            path = vreport.build_generation_report(
                profile=prof,
                document=tmp / f"run{i}.docx",
                report=qa,
                shell_path=None,  # falls back to provenance.shell.sha256
                out_dir=tmp / f"run{i}.docx.visual",
                content_hash="c" * 64,
                generated_at=f"2026-06-0{i + 1}T00:00:00Z",
            )
            self.assertIsNotNone(path)
        sha = prof["provenance"]["shell"]["sha256"]
        return vreport.discover_prior_reports(tmp / "new.docx.visual", shell_sha256=sha)

    def test_real_residual_findings_distill_demo_clear(self):
        """check_residual_template_text -> report writer -> discovery -> learn."""
        prof = _docx_profile()
        prof["surface"] = {
            "docx": {"demo_region": {"instruction_markers": ["Lorem ipsum"]}}
        }
        # The REAL producer: location is populated by the check itself, never by
        # this test.
        findings = cd.check_residual_template_text(
            "body text still carrying Lorem ipsum here", prof
        )
        self.assertTrue(findings)
        self.assertEqual(findings[0].check, "no_residual_template_text")
        self.assertEqual(findings[0].location, "Lorem ipsum")

        with tempfile.TemporaryDirectory() as td:
            reports = self._persist_and_discover(prof, findings, Path(td))
            self.assertEqual(len(reports), 2)
            res = ov.learn(prof, reports)
        self.assertTrue(res.ok)
        self.assertEqual(prof["rules"]["overrides"]["demo_clears"], ["Lorem ipsum"])

    def test_real_resolver_targets_finding_carries_role_location(self):
        """check_resolver_targets (against a real shell) keys ERRORs on the role id."""
        prof = _docx_profile()
        prof["roles"]["_index"].append("callout.info")
        prof["roles"]["callout.info"] = {
            "resolver": {
                "type": "named_style",
                "style_id": "FancyCallout",
                "style_name": "Fancy Callout",
            },
            "status": "robust",
            "confidence": 1.0,
        }
        with tempfile.TemporaryDirectory() as td:
            shell = Path(td) / "shell.docx"
            Document().save(shell)  # no FancyCallout style in the shell
            findings = cd.check_resolver_targets(shell, prof)
        hits = [f for f in findings if f.check == "resolver_targets_exist"]
        self.assertTrue(hits)
        self.assertEqual(hits[0].location, "callout.info")

    def test_real_generate_style_fallback_distills_reroute(self):
        """The docx writer's own style_fallback finding (a table with no table role)
        recurs across two persisted runs and learn reroutes it to the healthy
        same-family sibling - the full producer -> report -> learn chain."""
        prof = schema.build_envelope("docx", {"name": "t"})
        prof["provenance"]["shell"]["sha256"] = "shell-sha-1"
        prof["surface"] = {"docx": {}}
        prof["roles"] = {
            "_index": ["paragraph", "table.fancy"],
            "paragraph": {
                "resolver": {
                    "type": "named_style",
                    "style_id": "Normal",
                    "style_name": "Normal",
                },
                "status": "robust",
                "confidence": 1.0,
            },
            # The healthy same-family sibling the lesson should re-point at.
            "table.fancy": {
                "resolver": {
                    "type": "named_style",
                    "style_id": "TableGrid",
                    "style_name": "Table Grid",
                },
                "status": "robust",
                "confidence": 1.0,
            },
        }
        idoc = ir.IntermediateDocument(
            blocks=[ir.Table.from_dict({"columns": ["A"], "rows": [["1"]]})]
        )
        with tempfile.TemporaryDirectory() as td:
            tmp = Path(td)
            shell = tmp / "shell.docx"
            Document().save(shell)
            sink: list = []
            docx_generate.generate(prof, shell, idoc, tmp / "out.docx", findings=sink)
            # The REAL producer emitted the stub-role pointer in `location`.
            fallbacks = [f for f in sink if f.check == "style_fallback"]
            self.assertTrue(fallbacks)
            self.assertEqual(fallbacks[0].location, "table.default")

            reports = self._persist_and_discover(prof, fallbacks, tmp)
            self.assertEqual(len(reports), 2)
            res = ov.learn(prof, reports)
        self.assertTrue(res.ok)
        self.assertEqual(
            prof["rules"]["overrides"]["reroute_roles"],
            {"table.default": "table.fancy"},
        )


# ---------------------------------------------------------------------------
# CLI: the learn verb + ADVISORY accept gate
# ---------------------------------------------------------------------------
def _synthetic_template(path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    h1 = styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0x00, 0x5A, 0xAB)
    doc.add_paragraph("Example first-level title", style="Heading 1")
    doc.add_paragraph("Body.", style="Normal")
    doc.save(path)


class LearnCliAcceptGateTest(unittest.TestCase):
    """Drive the ``learn`` verb end-to-end and prove the advisory accept gate:
    without --accept the distilled lesson is written but kept OUT of the live
    resolver (status 'absent', so byte-identical); --accept promotes it to 'present'.
    """

    def _extract(self) -> None:
        template = Path("synthetic-template.docx")
        _synthetic_template(template)
        self.assertEqual(
            main(
                [
                    "extract",
                    "--name",
                    "acme",
                    "--template",
                    str(template),
                    "--scope",
                    "project",
                ]
            ),
            0,
        )

    def _stage_prior_reports(self, n: int) -> str:
        """Stage ``n`` sibling generation_report.json files (same live shell) that
        each carry a recurring, structured-location learnable finding for a stub
        role (heading.9), so learn can bind a brand-safe reroute to heading.1."""
        loaded = store.load_profile("acme", "project")
        shell_sha = vreport.report_shell_sha256(loaded.profile, loaded.shell_path)
        self.assertTrue(shell_sha)
        for i in range(n):
            vdir = Path(f"run{i}.docx.visual")
            vdir.mkdir(parents=True, exist_ok=True)
            (vdir / vreport.REPORT_FILENAME).write_text(
                json.dumps(
                    {
                        "schema_version": vreport.REPORT_SCHEMA_VERSION,
                        "kind": "docx",
                        "shell_sha256": shell_sha,
                        "findings": [
                            {
                                "check": "resolver_targets_exist",
                                "severity": "ERROR",
                                "message": "brand body text",
                                "location": "heading.9",
                            }
                        ],
                        "generated_at": f"2026-06-0{i + 1}T00:00:00Z",
                    }
                ),
                encoding="utf-8",
            )
        return shell_sha

    def test_learn_advisory_then_accept(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            old_cwd = Path.cwd()
            os.chdir(td)
            try:
                self._extract()
                self._stage_prior_reports(2)  # recurred across 2 runs -> distilled

                # Advisory run: writes the lesson but keeps it OUT of the live
                # resolver (status forced 'absent').
                self.assertEqual(
                    main(["learn", "--name", "acme", "--scope", "project"]), 0
                )
                loaded = store.load_profile("acme", "project")
                block = loaded.profile["rules"]["overrides"]
                self.assertEqual(
                    block["status"], schema.ComprehensionStatus.ABSENT.value
                )
                self.assertEqual(block["reroute_roles"], {"heading.9": "heading.1"})
                # Not live: the resolver takes zero new branches (byte-identical).
                self.assertFalse(store.overrides_are_present(loaded.profile))

                # Accept run: promotes the SAME distilled lesson to LIVE.
                self.assertEqual(
                    main(["learn", "--name", "acme", "--scope", "project", "--accept"]),
                    0,
                )
                loaded2 = store.load_profile("acme", "project")
                block2 = loaded2.profile["rules"]["overrides"]
                self.assertEqual(
                    block2["status"], schema.ComprehensionStatus.PRESENT.value
                )
                self.assertEqual(block2["reroute_roles"], {"heading.9": "heading.1"})
                self.assertTrue(store.overrides_are_present(loaded2.profile))
            finally:
                os.chdir(old_cwd)

    def test_learn_no_history_distills_nothing(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            old_cwd = Path.cwd()
            os.chdir(td)
            try:
                self._extract()
                # No prior reports -> nothing recurs -> nothing distilled, rc 0.
                self.assertEqual(
                    main(["learn", "--name", "acme", "--scope", "project"]), 0
                )
                loaded = store.load_profile("acme", "project")
                self.assertFalse(store.overrides_are_present(loaded.profile))
            finally:
                os.chdir(old_cwd)

    def test_learn_single_run_below_threshold(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            old_cwd = Path.cwd()
            os.chdir(td)
            try:
                self._extract()
                self._stage_prior_reports(1)  # seen once -> below default threshold
                self.assertEqual(
                    main(["learn", "--name", "acme", "--scope", "project"]), 0
                )
                loaded = store.load_profile("acme", "project")
                self.assertEqual(
                    loaded.profile["rules"]["overrides"]["reroute_roles"], {}
                )
                self.assertFalse(store.overrides_are_present(loaded.profile))
            finally:
                os.chdir(old_cwd)


if __name__ == "__main__":
    unittest.main()
