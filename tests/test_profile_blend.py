# SPDX-License-Identifier: MIT
"""Multi-template VALUE-FACT blending (REFLECTIONS P3): unit + e2e proof.

Covers the non-negotiable blend constraints:

  - same-format only (kind mismatch rejected, profile bytes untouched);
  - value-facts only (pointer classes NEVER cross - adversarial donor);
  - primary-wins precedence (conflicts kept-primary);
  - bounded, deterministic, order-independent corroboration math;
  - primary-shell proof (unprovable fills rejected, verify stays green);
  - fail-closed all-or-nothing atomicity (a trial-validation failure leaves
    profile.json byte-identical, no binary written);
  - sha-dedupe idempotence (re-blend is a structural no-op);
  - single-template profiles gain NOT ONE new key;
  - the ``blend_shell_provenance`` QA check has an honest fail path;
  - e2e on REAL example-derived shells: a sparse-primary variant (derived at
    test time, never committed) picks up a blended size and a generated
    document applies it set-only-when-unset.

No brand literal is asserted anywhere: e2e expectations are read from the
extracted profiles; unit tests use synthetic shells whose facts the tests
author themselves (Arial is the WordprocessingML docDefaults baseline the
allow-set always carries, not a brand value).
"""

from __future__ import annotations

import contextlib
import copy
import io
import json
import os
import sys
import tempfile
import unittest
import zipfile
from contextlib import redirect_stdout
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.shared import Pt, RGBColor
from lxml import etree

from brandkit.profile import blend as blend_mod
from brandkit.profile import schema, store
from brandkit.qa.gate import run_qa

ROOT = Path(__file__).resolve().parents[1]
DOCX_TEMPLATE = ROOT / "examples" / "templates" / "branddocs_template.docx"
PPTX_TEMPLATE = ROOT / "examples" / "templates" / "branddocs_template.pptx"

_WML = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Synthetic shell facts the unit tests author themselves (half-points / hex).
_PROVABLE_SIZE = 28
_OTHER_SIZE = 36
_UNPROVABLE_SIZE = 990
_PROVABLE_HEX = "112233"
_OTHER_HEX = "445566"
# The docDefaults baseline face every shell allow-set carries (not a brand value).
_BASELINE_FONT = "Arial"


@contextlib.contextmanager
def _chdir(path):
    old = Path.cwd()
    os.chdir(path)
    try:
        yield
    finally:
        os.chdir(old)


def _shell_bytes(
    *, sizes_hp=(_PROVABLE_SIZE, _OTHER_SIZE), hexes=(_PROVABLE_HEX, _OTHER_HEX)
) -> bytes:
    """A synthetic docx shell carrying exactly the run sizes/hexes given (the
    provable fact sets the blend pre-proof and QA validate against)."""
    doc = Document()
    for i, hp in enumerate(sizes_hp):
        para = doc.add_paragraph(f"size run {i}")
        para.runs[0].font.size = Pt(hp / 2)
    for i, hexval in enumerate(hexes):
        para = doc.add_paragraph(f"color run {i}")
        para.runs[0].font.color.rgb = RGBColor.from_string(hexval)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _role(
    style_id="Normal",
    style_name="Normal",
    *,
    appearance=None,
    confidence=0.8,
) -> dict:
    entry = {
        "resolver": {
            "type": "named_style",
            "style_id": style_id,
            "style_name": style_name,
            "style_type": "paragraph",
        },
        "status": "robust",
        "confidence": confidence,
    }
    if appearance is not None:
        entry["appearance"] = appearance
    return entry


def _roles(entries: dict) -> dict:
    out: dict = {"_index": list(entries)}
    out.update(entries)
    return out


def _primary(
    tmp_root: Path,
    *,
    name="prim",
    roles=None,
    theme_extra=None,
    profile_extra=None,
    shell=None,
) -> store.LoadedProfile:
    """Build + save a synthetic docx primary profile and load it back."""
    profile = schema.build_envelope("docx", {"name": name})
    if roles is not None:
        profile["roles"] = roles
    for key, val in (theme_extra or {}).items():
        profile["theme"][key] = val
    for key, val in (profile_extra or {}).items():
        profile[key] = val
    shell_bytes = shell if shell is not None else _shell_bytes()
    target = store.target_dir_for_save(name, "project", cwd=tmp_root)
    store.save_profile(target, profile, shell_bytes)
    return store.load_profile(name, "project", cwd=tmp_root)


def _secondary(*, roles=None, theme_extra=None, kind="docx") -> dict:
    """An in-memory secondary profile dict (the temp-extraction stand-in)."""
    profile = schema.build_envelope(kind, {"name": "donor"})
    if roles is not None:
        profile["roles"] = roles
    for key, val in (theme_extra or {}).items():
        profile["theme"][key] = val
    return profile


def _profile_bytes(loaded: store.LoadedProfile) -> bytes:
    return (loaded.directory / store.PROFILE_JSON).read_bytes()


def _blend_binaries(loaded: store.LoadedProfile) -> list[str]:
    return sorted(p.name for p in (loaded.directory / "template").glob("blend-*"))


class BlendGuardsTest(unittest.TestCase):
    def setUp(self):
        td = tempfile.TemporaryDirectory()
        self.addCleanup(td.cleanup)
        self.tmp = Path(td.name)

    def test_blend_rejects_kind_mismatch(self):
        primary = _primary(self.tmp, roles=_roles({"paragraph": _role()}))
        before = _profile_bytes(primary)
        result = blend_mod.blend(
            primary, _secondary(kind="pptx"), b"donor-bytes-pptx", "donor.pptx"
        )
        self.assertFalse(result.ok)
        self.assertIsNone(result.report)
        self.assertTrue(any("kind mismatch" in p for p in result.problems))
        self.assertEqual(before, _profile_bytes(primary))
        self.assertEqual([], _blend_binaries(primary))
        # CLI level: the same rejection with the cross-format message. The fake
        # donor is never parsed (the suffix gate fires first).
        fake = self.tmp / "donor.pptx"
        fake.write_bytes(b"not parsed")
        from brandkit import cli

        with _chdir(self.tmp):
            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = cli.main(
                    [
                        "extract",
                        "--name",
                        "prim",
                        "--template",
                        str(fake),
                        "--scope",
                        "project",
                        "--blend",
                    ]
                )
        self.assertEqual(rc, 1)
        self.assertIn("cross-format comparison is the read-only", buf.getvalue())
        self.assertEqual(before, _profile_bytes(primary))

    def test_blend_requires_existing_profile(self):
        from brandkit import cli

        fake = self.tmp / "anything.docx"
        fake.write_bytes(b"never read")
        with _chdir(self.tmp):
            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = cli.main(
                    [
                        "extract",
                        "--name",
                        "no-such-profile",
                        "--template",
                        str(fake),
                        "--scope",
                        "project",
                        "--blend",
                    ]
                )
        self.assertEqual(rc, 1)
        self.assertIn("no existing profile", buf.getvalue())
        # WITHOUT --blend an existing name takes the ordinary full re-extract
        # path (here: it actually tries to parse the template - the blend branch
        # never would have touched it).
        _primary(self.tmp, roles=_roles({"paragraph": _role()}))
        with _chdir(self.tmp):
            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = cli.main(
                    [
                        "extract",
                        "--name",
                        "prim",
                        "--template",
                        str(fake),
                        "--scope",
                        "project",
                    ]
                )
        self.assertEqual(rc, 1)
        self.assertIn("ERROR extract:", buf.getvalue())
        self.assertNotIn("blend", buf.getvalue())


class BlendMergeTest(unittest.TestCase):
    def setUp(self):
        td = tempfile.TemporaryDirectory()
        self.addCleanup(td.cleanup)
        self.tmp = Path(td.name)

    def test_blend_fills_unset_axis(self):
        primary = _primary(
            self.tmp,
            roles=_roles(
                {
                    # appearance key missing entirely: all axes unset.
                    "paragraph": _role(),
                    # appearance present but {}: equally all axes unset.
                    "heading.1": _role("Heading1", "Heading 1", appearance={}),
                }
            ),
        )
        donor_font = {"latin": _BASELINE_FONT}
        donor_color = {"kind": "hex", "hex": _PROVABLE_HEX}
        secondary = _secondary(
            roles=_roles(
                {
                    "paragraph": _role(
                        appearance={
                            "font": dict(donor_font),
                            "confidence": 0.7,
                            "size_hp": _PROVABLE_SIZE,
                            "size_confidence": 0.66,
                            "color": dict(donor_color),
                            "color_confidence": 0.91,
                        }
                    ),
                    "heading.1": _role(
                        "Heading1",
                        "Heading 1",
                        appearance={"size_hp": _OTHER_SIZE, "size_confidence": 0.62},
                    ),
                }
            ),
            theme_extra={
                "fonts": {
                    "major": {"latin": None, "fallback": None},
                    "minor": {"latin": None, "fallback": None},
                    "body": {
                        "latin": _BASELINE_FONT,
                        "confidence": 0.75,
                        "size_hp": _PROVABLE_SIZE,
                        "size_confidence": 0.71,
                    },
                }
            },
        )
        result = blend_mod.blend(primary, secondary, b"donor-bytes-1", "donor.docx")
        self.assertTrue(result.ok, result.problems)
        written = json.loads(_profile_bytes(primary))
        sha = store.sha256_bytes(b"donor-bytes-1")
        para = written["roles"]["paragraph"]["appearance"]
        self.assertEqual(para["font"], donor_font)
        self.assertEqual(para["confidence"], 0.7)
        self.assertEqual(para["size_hp"], _PROVABLE_SIZE)
        self.assertEqual(para["size_confidence"], 0.66)
        self.assertEqual(para["color"], donor_color)
        self.assertEqual(para["color_confidence"], 0.91)
        head = written["roles"]["heading.1"]["appearance"]
        self.assertEqual(head["size_hp"], _OTHER_SIZE)
        self.assertEqual(head["size_confidence"], 0.62)
        body = written["theme"]["fonts"]["body"]
        self.assertEqual(body["latin"], _BASELINE_FONT)
        self.assertEqual(body["size_hp"], _PROVABLE_SIZE)
        ledger = written["blend"]["ledger"]["filled"]
        for path in (
            "roles.paragraph.appearance.font",
            "roles.paragraph.appearance.size_hp",
            "roles.paragraph.appearance.color",
            "roles.heading.1.appearance.size_hp",
            "theme.fonts.body.latin",
            "theme.fonts.body.size_hp",
        ):
            self.assertEqual(ledger.get(path), {"from": sha}, path)
        shells = written["provenance"]["blended_shells"]
        self.assertEqual(
            shells,
            [
                {
                    "filename": "donor.docx",
                    "path": f"template/blend-{sha[:12]}.docx",
                    "sha256": sha,
                }
            ],
        )
        self.assertEqual([], schema.validate(written))
        self.assertEqual([f"blend-{sha[:12]}.docx"], _blend_binaries(primary))

    def test_blend_primary_wins_conflict(self):
        appearance = {"size_hp": _PROVABLE_SIZE, "size_confidence": 0.8}
        primary = _primary(
            self.tmp,
            roles=_roles({"paragraph": _role(appearance=dict(appearance))}),
        )
        before_entry = copy.deepcopy(primary.profile["roles"]["paragraph"])
        # A foreign resolver isolates the conflict behavior: with resolver
        # inequality not even the role-detection confidence may move, so the
        # whole primary subtree must be deep-equal pre/post.
        secondary = _secondary(
            roles=_roles(
                {
                    "paragraph": _role(
                        "OtherStyle",
                        "Other Style",
                        appearance={"size_hp": _OTHER_SIZE, "size_confidence": 0.9},
                    )
                }
            )
        )
        result = blend_mod.blend(primary, secondary, b"donor-bytes-1", "donor.docx")
        self.assertTrue(result.ok, result.problems)
        written = json.loads(_profile_bytes(primary))
        self.assertEqual(written["roles"]["paragraph"], before_entry)
        self.assertEqual(
            result.report.conflicts,
            [
                {
                    "path": "roles.paragraph.appearance.size_hp",
                    "primary": _PROVABLE_SIZE,
                    "secondary": _OTHER_SIZE,
                }
            ],
        )
        self.assertEqual(written["blend"]["ledger"]["filled"], {})
        self.assertEqual(written["blend"]["ledger"]["corroborated"], {})

    def test_blend_corroboration_math(self):
        donor_roles = _roles(
            {
                "paragraph": _role(
                    appearance={"font": {"latin": _BASELINE_FONT}, "confidence": 0.5},
                    confidence=0.97,
                )
            }
        )

        def fresh(name: str) -> store.LoadedProfile:
            return _primary(
                self.tmp,
                name=name,
                roles=copy.deepcopy(donor_roles),
            )

        def donor() -> dict:
            return _secondary(roles=copy.deepcopy(donor_roles))

        # Two DISTINCT donors: +0.05 each on the axis confidence, the role
        # confidence saturates at the 1.0 cap from 0.97.
        primary = fresh("math")
        result = blend_mod.blend(primary, donor(), b"donor-bytes-a", "a.docx")
        self.assertTrue(result.ok, result.problems)
        one = json.loads(_profile_bytes(primary))
        self.assertEqual(one["roles"]["paragraph"]["appearance"]["confidence"], 0.55)
        self.assertEqual(one["roles"]["paragraph"]["confidence"], 1.0)
        reloaded = store.load_profile("math", "project", cwd=self.tmp)
        result = blend_mod.blend(reloaded, donor(), b"donor-bytes-b", "b.docx")
        self.assertTrue(result.ok, result.problems)
        two = json.loads(_profile_bytes(reloaded))
        self.assertEqual(two["roles"]["paragraph"]["appearance"]["confidence"], 0.6)
        self.assertEqual(two["roles"]["paragraph"]["confidence"], 1.0)
        sha_a = store.sha256_bytes(b"donor-bytes-a")
        sha_b = store.sha256_bytes(b"donor-bytes-b")
        self.assertEqual(
            two["blend"]["ledger"]["corroborated"]["roles.paragraph.appearance.font"][
                "by"
            ],
            sorted([sha_a, sha_b]),
        )

        # Order independence: blending the same two donors in either order
        # produces byte-identical profile.json (commutative corroboration).
        # Each order runs in its own root with the SAME profile name and the
        # donors keep their own stable filenames (both are recorded facts).
        donor_a = (b"donor-bytes-a", "a.docx")
        donor_b = (b"donor-bytes-b", "b.docx")
        outputs = []
        for sub, order in (("ab", (donor_a, donor_b)), ("ba", (donor_b, donor_a))):
            root = self.tmp / f"order-{sub}"
            root.mkdir()
            _primary(root, name="order", roles=copy.deepcopy(donor_roles))
            for payload, filename in order:
                loaded = store.load_profile("order", "project", cwd=root)
                result = blend_mod.blend(loaded, donor(), payload, filename)
                self.assertTrue(result.ok, result.problems)
            outputs.append(
                (root / "brand-kit" / "order" / store.PROFILE_JSON).read_bytes()
            )
        self.assertEqual(outputs[0], outputs[1])

        # Pointer-equality rule: an agreeing axis VALUE with a DIFFERENT resolver
        # bumps the axis confidence but NEVER the role-detection confidence.
        primary = fresh("resolver-neq")
        other = donor()
        other["roles"]["paragraph"]["resolver"]["style_id"] = "OtherStyle"
        result = blend_mod.blend(primary, other, b"donor-bytes-c", "c.docx")
        self.assertTrue(result.ok, result.problems)
        written = json.loads(_profile_bytes(primary))
        self.assertEqual(
            written["roles"]["paragraph"]["appearance"]["confidence"], 0.55
        )
        self.assertEqual(written["roles"]["paragraph"]["confidence"], 0.97)
        # The donor's foreign resolver pointer never crossed.
        self.assertEqual(
            written["roles"]["paragraph"]["resolver"]["style_id"], "Normal"
        )

    def test_blend_never_imports_pointers(self):
        primary = _primary(
            self.tmp,
            roles=_roles({"paragraph": _role(confidence=0.8)}),
        )
        before = json.loads(_profile_bytes(primary))
        secondary = _secondary(
            roles=_roles(
                {
                    "paragraph": _role(
                        "EvilStyle",
                        "Evil Style",
                        appearance={
                            "font": {"latin": _BASELINE_FONT},
                            "confidence": 0.9,
                            "geometry": {"spacing": {"after_twips": 240}},
                            "table": {"style_id": "EvilTable"},
                            "numbering": {"num_id": "7", "abstract_num_id": "7"},
                        },
                    ),
                    "extra.role": _role("EvilStyle", "Evil Style"),
                }
            ),
            theme_extra={
                "palette": {
                    "accent1": {
                        "ref": {"kind": "theme", "theme": "accent1"},
                        "provenance": [],
                        "frequency": "rare",
                        "name": None,
                        "purpose": None,
                        "use_when": None,
                    }
                }
            },
        )
        result = blend_mod.blend(primary, secondary, b"donor-bytes-1", "donor.docx")
        self.assertTrue(result.ok, result.problems)
        written = json.loads(_profile_bytes(primary))
        # The value-fact DID land (this is a blend, not a wholesale rejection)...
        self.assertEqual(
            written["roles"]["paragraph"]["appearance"]["font"],
            {"latin": _BASELINE_FONT},
        )
        # ...but nothing pointer-shaped crossed.
        self.assertEqual(
            written["roles"]["paragraph"]["resolver"],
            before["roles"]["paragraph"]["resolver"],
        )
        self.assertEqual(written["roles"]["_index"], before["roles"]["_index"])
        self.assertNotIn("extra.role", written["roles"])
        for frozen in ("geometry", "table", "numbering"):
            self.assertNotIn(frozen, written["roles"]["paragraph"]["appearance"])
        for top in ("anchors", "structure", "surface", "components", "sections"):
            self.assertEqual(written[top], before[top], top)
        self.assertNotIn("accent1", written["theme"]["palette"])
        self.assertEqual(
            [k for k in written["theme"]["palette"] if not k.startswith("hex:")], []
        )
        # The role-detection confidence did not move (resolver inequality).
        self.assertEqual(written["roles"]["paragraph"]["confidence"], 0.8)
        rejected_classes = {entry["class"] for entry in result.report.rejected}
        self.assertEqual(
            rejected_classes,
            {
                "roles",
                "appearance.geometry",
                "appearance.table",
                "appearance.numbering",
                "theme.palette.slot",
            },
        )

    def test_blend_preserves_caches(self):
        shell_bytes = _shell_bytes()
        shell_sha = store.sha256_bytes(shell_bytes)
        comp = schema.empty_comprehension()
        comp["status"] = "present"
        comp["source_shell_sha256"] = shell_sha
        overrides = schema.empty_overrides()
        overrides["status"] = "present"
        overrides["source_shell_sha256"] = shell_sha
        overrides["reroute_roles"] = {"callout.warn": "paragraph"}
        primary = _primary(
            self.tmp,
            roles=_roles({"paragraph": _role()}),
            profile_extra={"comprehension": comp},
            shell=shell_bytes,
        )
        # rules.overrides is nested; seed it after envelope construction.
        primary.profile["rules"]["overrides"] = overrides
        store.write_profile_json(primary.directory, primary.profile)
        primary = store.load_profile("prim", "project", cwd=self.tmp)
        before = json.loads(_profile_bytes(primary))
        self.assertTrue(store.comprehension_is_present(primary.profile))
        self.assertTrue(store.overrides_are_present(primary.profile))

        secondary = _secondary(
            roles=_roles({"paragraph": _role(appearance={"size_hp": _PROVABLE_SIZE})})
        )
        result = blend_mod.blend(primary, secondary, b"donor-bytes-1", "donor.docx")
        self.assertTrue(result.ok, result.problems)
        written = json.loads(_profile_bytes(primary))
        self.assertEqual(written["comprehension"], before["comprehension"])
        self.assertEqual(written["rules"]["overrides"], before["rules"]["overrides"])
        self.assertEqual(written["provenance"]["shell"], before["provenance"]["shell"])
        self.assertTrue(store.comprehension_is_present(written))
        self.assertTrue(store.overrides_are_present(written))

    def test_blend_unprovable_fill_rejected_class(self):
        primary = _primary(self.tmp, roles=_roles({"paragraph": _role()}))
        secondary = _secondary(
            roles=_roles(
                {
                    "paragraph": _role(
                        appearance={
                            "size_hp": _UNPROVABLE_SIZE,
                            "size_confidence": 0.9,
                        }
                    )
                }
            )
        )
        result = blend_mod.blend(primary, secondary, b"donor-bytes-1", "donor.docx")
        self.assertTrue(result.ok, result.problems)
        written = json.loads(_profile_bytes(primary))
        self.assertNotIn("appearance", written["roles"]["paragraph"])
        unprovable = [
            entry for entry in result.report.rejected if entry["class"] == "unprovable"
        ]
        self.assertEqual(len(unprovable), 1)
        self.assertIn("roles.paragraph.appearance.size_hp", unprovable[0]["detail"])
        self.assertIn(str(_UNPROVABLE_SIZE), unprovable[0]["detail"])
        # The blended profile still verifies green (fast gate, shell-backed).
        report = run_qa(None, written, qa="fast", shell=primary.shell_path)
        errors = [f for f in report.findings if f.severity == "ERROR"]
        self.assertEqual(errors, [])

    def test_blend_idempotent_and_dedupe(self):
        primary = _primary(self.tmp, roles=_roles({"paragraph": _role()}))
        secondary = _secondary(
            roles=_roles({"paragraph": _role(appearance={"size_hp": _PROVABLE_SIZE})})
        )
        result = blend_mod.blend(primary, secondary, b"donor-bytes-1", "donor.docx")
        self.assertTrue(result.ok, result.problems)
        self.assertFalse(result.report.noop)
        after = _profile_bytes(primary)
        binaries = _blend_binaries(primary)

        reloaded = store.load_profile("prim", "project", cwd=self.tmp)
        again = blend_mod.blend(reloaded, secondary, b"donor-bytes-1", "donor.docx")
        self.assertTrue(again.ok)
        self.assertTrue(again.report.noop)
        self.assertEqual(after, _profile_bytes(reloaded))
        self.assertEqual(binaries, _blend_binaries(reloaded))

        # Blending the primary template itself is equally a no-op (sha match).
        own_shell = reloaded.shell_path.read_bytes()
        own = blend_mod.blend(reloaded, secondary, own_shell, "self.docx")
        self.assertTrue(own.ok)
        self.assertTrue(own.report.noop)
        self.assertEqual(after, _profile_bytes(reloaded))

    def test_blend_fail_closed_atomicity(self):
        primary = _primary(self.tmp, roles=_roles({"paragraph": _role()}))
        before = _profile_bytes(primary)
        # A donor palette entry whose KEY is provable (so the plan imports it
        # verbatim) but whose ref.kind is illegal: the trial-wide schema
        # validation must reject the WHOLE blend, leaving bytes untouched.
        secondary = _secondary(
            theme_extra={
                "palette": {
                    f"hex:{_PROVABLE_HEX}": {
                        "ref": {"kind": "rgb", "hex": _PROVABLE_HEX},
                        "provenance": [],
                        "frequency": "rare",
                        "name": None,
                        "purpose": None,
                        "use_when": None,
                    }
                }
            }
        )
        result = blend_mod.blend(primary, secondary, b"donor-bytes-1", "donor.docx")
        self.assertFalse(result.ok)
        self.assertIsNone(result.report)
        self.assertIsNone(result.profile)
        self.assertTrue(any("ref.kind" in p for p in result.problems))
        self.assertEqual(before, _profile_bytes(primary))
        self.assertEqual([], _blend_binaries(primary))
        written = json.loads(_profile_bytes(primary))
        self.assertNotIn("blend", written)
        self.assertNotIn("blended_shells", written["provenance"])

    def test_blend_ledger_and_provenance_shapes(self):
        primary = _primary(self.tmp, roles=_roles({"paragraph": _role()}))
        live_sha = primary.profile["provenance"]["shell"]["sha256"]
        fill_donor = _secondary(
            roles=_roles(
                {
                    "paragraph": _role(
                        appearance={"size_hp": _PROVABLE_SIZE, "size_confidence": 0.7}
                    )
                }
            )
        )
        result = blend_mod.blend(primary, fill_donor, b"donor-bytes-a", "a.docx")
        self.assertTrue(result.ok, result.problems)
        reloaded = store.load_profile("prim", "project", cwd=self.tmp)
        result = blend_mod.blend(reloaded, fill_donor, b"donor-bytes-b", "b.docx")
        self.assertTrue(result.ok, result.problems)

        written = json.loads(_profile_bytes(reloaded))
        shells = written["provenance"]["blended_shells"]
        shas = [entry["sha256"] for entry in shells]
        self.assertEqual(shas, sorted(shas))
        self.assertEqual(len(shas), len(set(shas)))
        for entry in shells:
            self.assertEqual(
                entry["path"], f"template/blend-{entry['sha256'][:12]}.docx"
            )
        block = written["blend"]
        self.assertEqual(block["schema_version"], schema.BLEND_SCHEMA_VERSION)
        self.assertEqual(block["status"], "present")
        self.assertEqual(block["source_shell_sha256"], live_sha)
        ledger_shas = {m["from"] for m in block["ledger"]["filled"].values()}
        for mark in block["ledger"]["corroborated"].values():
            ledger_shas.update(mark["by"])
        self.assertTrue(ledger_shas <= set(shas))
        # First-writer-wins: the fill stays attributed to the FIRST donor; the
        # second agreeing donor corroborates it instead.
        sha_a = store.sha256_bytes(b"donor-bytes-a")
        sha_b = store.sha256_bytes(b"donor-bytes-b")
        self.assertEqual(
            block["ledger"]["filled"]["roles.paragraph.appearance.size_hp"],
            {"from": sha_a},
        )
        self.assertEqual(
            block["ledger"]["corroborated"]["roles.paragraph.appearance.size_hp"]["by"],
            [sha_b],
        )
        self.assertTrue(store.blend_is_present(written))
        self.assertEqual([], schema.validate(written))
        self.assertEqual([], reloaded.blended_shell_drift())


class BlendShellDriftCheckTest(unittest.TestCase):
    def setUp(self):
        td = tempfile.TemporaryDirectory()
        self.addCleanup(td.cleanup)
        self.tmp = Path(td.name)

    def _findings(self, profile: dict, shell) -> list:
        report = run_qa(None, profile, qa="fast", shell=shell)
        return [f for f in report.findings if f.check == "blend_shell_provenance"]

    def test_blend_shell_drift_check(self):
        # (d) never-blended profile: zero findings, single-template QA unchanged.
        pristine = _primary(
            self.tmp, name="pristine", roles=_roles({"paragraph": _role()})
        )
        self.assertEqual([], self._findings(pristine.profile, pristine.shell_path))

        primary = _primary(self.tmp, roles=_roles({"paragraph": _role()}))
        secondary = _secondary(
            roles=_roles({"paragraph": _role(appearance={"size_hp": _PROVABLE_SIZE})})
        )
        result = blend_mod.blend(primary, secondary, b"donor-bytes-1", "donor.docx")
        self.assertTrue(result.ok, result.problems)
        written = result.profile
        rel = written["provenance"]["blended_shells"][0]["path"]
        binary = primary.directory / rel

        # (c) clean blended profile: zero findings from the check.
        self.assertEqual([], self._findings(written, primary.shell_path))
        reloaded = store.load_profile("prim", "project", cwd=self.tmp)
        self.assertEqual([], reloaded.blended_shell_drift())

        # (a) overwrite the donor binary: ERROR with the drifted-hash message,
        # surfaced through run_qa, the drift mirror, and the verify verb.
        binary.write_bytes(b"tampered donor bytes")
        findings = self._findings(written, primary.shell_path)
        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0].severity, "ERROR")
        self.assertIn("blend shell hash drifted", findings[0].message)
        drift = store.load_profile("prim", "project", cwd=self.tmp)
        self.assertTrue(
            any("drifted" in problem for problem in drift.blended_shell_drift())
        )
        from brandkit import cli

        with _chdir(self.tmp):
            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = cli.main(
                    ["verify", "--name", "prim", "--scope", "project", "--qa", "fast"]
                )
        self.assertEqual(rc, 1)
        self.assertIn("blend_shell_provenance", buf.getvalue())

        # (b) delete it: the missing-file ERROR path.
        binary.unlink()
        findings = self._findings(written, primary.shell_path)
        self.assertEqual(len(findings), 1)
        self.assertIn("missing", findings[0].message)


class SingleTemplateUnchangedTest(unittest.TestCase):
    @unittest.skipUnless(DOCX_TEMPLATE.exists(), "example docx template missing")
    def test_single_template_profiles_write_no_new_keys(self):
        from brandkit import cli

        with tempfile.TemporaryDirectory() as td:
            with _chdir(td):
                snapshots = []
                for _ in range(2):
                    rc = cli.main(
                        [
                            "extract",
                            "--name",
                            "solo",
                            "--template",
                            str(DOCX_TEMPLATE),
                            "--scope",
                            "project",
                        ]
                    )
                    self.assertEqual(rc, 0)
                    profile = json.loads(
                        (Path(td) / "brand-kit" / "solo" / "profile.json").read_text(
                            encoding="utf-8"
                        )
                    )
                    self.assertNotIn("blend", profile)
                    self.assertNotIn("blended_shells", profile["provenance"])
                    # Normalize the one pre-existing wall-clock field so the
                    # comparison proves "no new keys, no changed values".
                    profile["provenance"]["extracted_at"] = None
                    snapshots.append(
                        json.dumps(profile, sort_keys=True, ensure_ascii=False)
                    )
                self.assertEqual(snapshots[0], snapshots[1])


def _para_with_marker(docx_path: Path, marker: str):
    xml = zipfile.ZipFile(docx_path).read("word/document.xml")
    root = etree.fromstring(xml)
    for para in root.iter(f"{_WML}p"):
        text = "".join(t.text or "" for t in para.iter(f"{_WML}t"))
        if marker in text:
            return para
    raise AssertionError(f"no paragraph carrying marker {marker!r} in {docx_path}")


def _run_sizes(para) -> list[str]:
    return [sz.get(f"{_WML}val") for sz in para.iter(f"{_WML}sz")]


@unittest.skipUnless(DOCX_TEMPLATE.exists(), "example docx template missing")
class BlendEndToEndTest(unittest.TestCase):
    """Blend two REAL example-derived docx shells: the sparse-primary variant is
    derived in memory at test time (never committed) and a generated document
    provably picks up the blended size set-only-when-unset."""

    _MAX_PADS = 40

    def _extract_probe(self, template: Path) -> dict:
        from brandkit.formats.docx import extract as docx_extract

        with tempfile.TemporaryDirectory() as tmp:
            docx_extract.extract(template, "probe", scope="project", cwd=tmp)
            return json.loads(
                (Path(tmp) / "brand-kit" / "probe" / "profile.json").read_text(
                    encoding="utf-8"
                )
            )

    def _pick_size_role(self, control: dict) -> str:
        """A non-paragraph-family role whose donor appearance carries size_hp.

        The design named ``heading.1``, but the live example template captures
        no heading size (its headings inherit size from the style); the proof
        needs a role with a captured size AND no body-size fallback at apply
        time, picked from the profile itself so the test never hardcodes a
        brand fact.
        """
        for rid in control["roles"]["_index"]:
            if rid == "_index" or rid.split(".")[0] == "paragraph":
                continue
            entry = control["roles"][rid]
            appearance = entry.get("appearance") or {}
            resolver = entry.get("resolver") or {}
            if "size_hp" in appearance and resolver.get("style_name"):
                return rid
        self.fail(
            "the example docx template no longer captures a role-level size_hp "
            "on any non-paragraph role; adapt the e2e to a role the example "
            "does capture (see roles[*].appearance in the extracted profile)"
        )

    def _derive_sparse(
        self, control: dict, rid: str, out_path: Path
    ) -> tuple[int, int]:
        """Pad the donor template with same-style runs of DISTINCT explicit
        sizes until a re-extract drops the role's ``size_hp`` (the dominance
        floor). Returns ``(captured_size, pads_used)``."""
        size_hp = int(control["roles"][rid]["appearance"]["size_hp"])
        style_name = control["roles"][rid]["resolver"]["style_name"]
        doc = Document(str(DOCX_TEMPLATE))
        pads = 0
        while True:
            for _ in range(4):
                pads += 1
                para = doc.add_paragraph(f"sparse pad {pads}", style=style_name)
                # Distinct per pad and never the captured winner, so no new
                # dominant emerges while the old one falls under the floor.
                para.runs[0].font.size = Pt((size_hp + 2 * pads) / 2)
            doc.save(str(out_path))
            probe = self._extract_probe(out_path)
            if "size_hp" not in (probe["roles"][rid].get("appearance") or {}):
                return size_hp, pads
            self.assertLessEqual(
                pads,
                self._MAX_PADS,
                f"could not derive a sparse variant for role {rid!r} within "
                f"{self._MAX_PADS} pads",
            )

    def test_blend_e2e_sparse_primary_docx(self):
        from brandkit import cli

        marker = "blend e2e marker text"
        control = self._extract_probe(DOCX_TEMPLATE)
        rid = self._pick_size_role(control)
        size_hp = int(control["roles"][rid]["appearance"]["size_hp"])

        with tempfile.TemporaryDirectory() as td:
            tdp = Path(td)
            sparse_template = tdp / "sparse_variant.docx"
            captured, pads = self._derive_sparse(control, rid, sparse_template)
            self.assertEqual(captured, size_hp)
            self.assertLessEqual(pads, self._MAX_PADS)

            with _chdir(td):
                rc = cli.main(
                    [
                        "extract",
                        "--name",
                        "sparse",
                        "--template",
                        str(sparse_template),
                        "--scope",
                        "project",
                    ]
                )
                self.assertEqual(rc, 0)
                pj = tdp / "brand-kit" / "sparse" / "profile.json"
                sparse_profile = json.loads(pj.read_text(encoding="utf-8"))
                self.assertNotIn(
                    "size_hp", sparse_profile["roles"][rid].get("appearance") or {}
                )

                # The idoc exercises the size-bearing role plus a level-1 heading.
                block_type = {
                    "caption": "caption",
                }.get(rid.split(".")[0])
                self.assertIsNotNone(
                    block_type,
                    f"e2e picked role {rid!r} with no simple idoc block mapping; "
                    "extend the mapping for the evolved example template",
                )
                idoc = {
                    "blocks": [
                        {"type": "heading", "level": 1, "runs": [{"t": "E2E Title"}]},
                        {"type": block_type, "runs": [{"t": marker}]},
                    ]
                }
                idoc_path = tdp / "input.json"
                idoc_path.write_text(
                    json.dumps(idoc, ensure_ascii=False), encoding="utf-8"
                )

                # (a) Un-blended sparse profile: the marker run inherits its size
                # from the style (no direct w:sz).
                rc = cli.main(
                    [
                        "generate",
                        "--name",
                        "sparse",
                        "--input",
                        str(idoc_path),
                        "--output",
                        str(tdp / "before.docx"),
                        "--scope",
                        "project",
                        "--qa",
                        "fast",
                    ]
                )
                self.assertEqual(rc, 0)
                para = _para_with_marker(tdp / "before.docx", marker)
                self.assertEqual(_run_sizes(para), [])

                # Blend the ORIGINAL example template into the sparse profile.
                before_blend = pj.read_bytes()
                buf = io.StringIO()
                with redirect_stdout(buf):
                    rc = cli.main(
                        [
                            "extract",
                            "--name",
                            "sparse",
                            "--template",
                            str(DOCX_TEMPLATE),
                            "--scope",
                            "project",
                            "--blend",
                        ]
                    )
                self.assertEqual(rc, 0, buf.getvalue())
                self.assertNotEqual(before_blend, pj.read_bytes())
                blended = json.loads(pj.read_text(encoding="utf-8"))
                self.assertEqual(
                    blended["roles"][rid]["appearance"]["size_hp"], size_hp
                )
                donor_sha = store.sha256_bytes(DOCX_TEMPLATE.read_bytes())
                self.assertEqual(
                    blended["blend"]["ledger"]["filled"][
                        f"roles.{rid}.appearance.size_hp"
                    ],
                    {"from": donor_sha},
                )
                self.assertIn(f"roles.{rid}.appearance.size_hp", buf.getvalue())
                self.assertEqual([], schema.validate(blended))

                # (b) Blended profile: the marker run now carries the blended
                # size as direct formatting (set-only-when-unset).
                rc = cli.main(
                    [
                        "generate",
                        "--name",
                        "sparse",
                        "--input",
                        str(idoc_path),
                        "--output",
                        str(tdp / "after.docx"),
                        "--scope",
                        "project",
                        "--qa",
                        "fast",
                    ]
                )
                self.assertEqual(rc, 0)
                para = _para_with_marker(tdp / "after.docx", marker)
                self.assertEqual(set(_run_sizes(para)), {str(size_hp)})

                # Determinism: a second generation is byte-identical.
                rc = cli.main(
                    [
                        "generate",
                        "--name",
                        "sparse",
                        "--input",
                        str(idoc_path),
                        "--output",
                        str(tdp / "after2.docx"),
                        "--scope",
                        "project",
                        "--qa",
                        "fast",
                    ]
                )
                self.assertEqual(rc, 0)
                self.assertEqual(
                    (tdp / "after.docx").read_bytes(),
                    (tdp / "after2.docx").read_bytes(),
                )

                # Idempotence: re-blending the same donor is a no-op.
                after_blend = pj.read_bytes()
                buf = io.StringIO()
                with redirect_stdout(buf):
                    rc = cli.main(
                        [
                            "extract",
                            "--name",
                            "sparse",
                            "--template",
                            str(DOCX_TEMPLATE),
                            "--scope",
                            "project",
                            "--blend",
                        ]
                    )
                self.assertEqual(rc, 0)
                self.assertIn("no-op", buf.getvalue())
                self.assertEqual(after_blend, pj.read_bytes())


if __name__ == "__main__":
    unittest.main()
