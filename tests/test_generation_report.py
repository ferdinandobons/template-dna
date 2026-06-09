# SPDX-License-Identifier: MIT
"""B1 tests: the persisted ``generation_report.json`` side artifact.

These drive the real CLI ``extract`` -> ``generate`` flow on a synthetic docx
template (mirroring ``test_smoke``) and assert the durable report substrate that
Cluster B's cross-run "learn-from-errors" reads:

  * the report lands next to the output, with all keys, ``findings`` 1:1 in order;
  * the generated ``.docx`` bytes are identical across runs even though the
    report's ``generated_at`` differs (the timestamp lives only in the JSON);
  * a writer failure degrades to a no-op (no exception, verdict unchanged);
  * ``verify`` writes NO report (generate-only);
  * the ``content_sha256`` is stable under a cosmetic-only input change.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

from brandkit.cli import main
from brandkit.qa import report as vreport


def _synthetic_template(path: Path) -> None:
    """A minimal branded docx template (same shape test_smoke extracts)."""
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    h1 = styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x5A, 0xAB)
    callout = styles.add_style("ACME Callout Info", WD_STYLE_TYPE.PARAGRAPH)
    callout.base_style = styles["Normal"]
    doc.add_paragraph("{{title}}", style="Title")
    doc.add_paragraph("Example first-level title", style="Heading 1")
    doc.add_paragraph("General instructions: replace this demo text.", style="Normal")
    doc.save(path)


_IDOC = {
    "cover": {"title": "Quarterly Review"},
    "blocks": [
        {"type": "heading", "level": 1, "text": "Highlights"},
        {"type": "paragraph", "text": "Revenue grew without markdown literals."},
    ],
}


class GenerationReportB1Test(unittest.TestCase):
    def _extract(self, tmp_path: Path) -> None:
        template = tmp_path / "synthetic-template.docx"
        _synthetic_template(template)
        self.assertEqual(
            main(
                [
                    "extract",
                    "--name",
                    "acme",
                    "--template",
                    str(template),
                    "--scope",
                    "project",
                ]
            ),
            0,
        )

    def _write_idoc(self, tmp_path: Path, data: dict, name: str = "idoc.json") -> Path:
        p = tmp_path / name
        p.write_text(json.dumps(data), encoding="utf-8")
        return p

    def _generate(self, tmp_path: Path, idoc: Path, out: Path) -> int:
        return main(
            [
                "generate",
                "--name",
                "acme",
                "--input",
                str(idoc),
                "--output",
                str(out),
                "--scope",
                "project",
                "--qa",
                "fast",
            ]
        )

    # -- test_report_written_next_to_output --------------------------------
    def test_report_written_next_to_output(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            old_cwd = Path.cwd()
            os.chdir(tmp_path)
            try:
                self._extract(tmp_path)
                idoc = self._write_idoc(tmp_path, _IDOC)
                out = tmp_path / "out.docx"
                self.assertEqual(self._generate(tmp_path, idoc, out), 0)

                report_path = (
                    tmp_path / (out.name + ".visual") / vreport.REPORT_FILENAME
                )
                self.assertTrue(report_path.is_file(), "report not written next to out")
                doc = json.loads(report_path.read_text(encoding="utf-8"))

                # All keys present, with the locked schema version.
                for key in (
                    "schema_version",
                    "kind",
                    "profile_name",
                    "document",
                    "verdict",
                    "shell_sha256",
                    "content_sha256",
                    "output_sha256",
                    "findings",
                    "generated_at",
                ):
                    self.assertIn(key, doc, f"missing report key: {key}")
                self.assertEqual(doc["schema_version"], vreport.REPORT_SCHEMA_VERSION)
                self.assertEqual(doc["kind"], "docx")
                self.assertEqual(doc["document"], "out.docx")
                # content/output/shell hashes are real 64-hex digests.
                for key in ("shell_sha256", "content_sha256", "output_sha256"):
                    self.assertIsInstance(doc[key], str)
                    self.assertEqual(len(doc[key]), 64)
                # output_sha256 matches the on-disk doc bytes.
                import hashlib

                self.assertEqual(
                    doc["output_sha256"],
                    hashlib.sha256(out.read_bytes()).hexdigest(),
                )

                # findings 1:1 in order with the report the gate would print.
                from brandkit.profile import store
                from brandkit.qa.gate import run_qa

                loaded = store.load_profile("acme", "project")
                live = run_qa(out, loaded.profile, qa="fast", shell=loaded.shell_path)
                self.assertEqual(
                    [f["check"] for f in doc["findings"]],
                    [f.check for f in live.findings],
                )
            finally:
                os.chdir(old_cwd)

    # -- test_report_byte_identical_output ---------------------------------
    def test_report_byte_identical_output(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            old_cwd = Path.cwd()
            os.chdir(tmp_path)
            try:
                self._extract(tmp_path)
                idoc = self._write_idoc(tmp_path, _IDOC)

                # Generate to the SAME output path twice so EVERYTHING but the
                # volatile timestamp must match -- the strongest byte-identity proof.
                out = tmp_path / "out.docx"
                report_path = (
                    tmp_path / (out.name + ".visual") / vreport.REPORT_FILENAME
                )

                self.assertEqual(self._generate(tmp_path, idoc, out), 0)
                bytes1 = out.read_bytes()
                r1 = json.loads(report_path.read_text(encoding="utf-8"))

                self.assertEqual(self._generate(tmp_path, idoc, out), 0)
                bytes2 = out.read_bytes()
                r2 = json.loads(report_path.read_text(encoding="utf-8"))

                # The generated .docx bytes are identical across runs.
                self.assertEqual(
                    bytes1, bytes2, "generate output bytes drifted between runs"
                )
                # Same content hash (same input), same output hash (same bytes)...
                self.assertEqual(r1["content_sha256"], r2["content_sha256"])
                self.assertEqual(r1["output_sha256"], r2["output_sha256"])
                # ...and dropping the ONLY volatile field, the reports are equal
                # (the timestamp lives only in the JSON, never in the doc bytes).
                self.assertIn("generated_at", r1)
                self.assertIn("generated_at", r2)
                r1.pop("generated_at")
                r2.pop("generated_at")
                self.assertEqual(r1, r2)
            finally:
                os.chdir(old_cwd)

    # -- test_report_writer_failure_no_op ----------------------------------
    def test_report_writer_failure_no_op(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            old_cwd = Path.cwd()
            os.chdir(tmp_path)
            try:
                self._extract(tmp_path)
                idoc = self._write_idoc(tmp_path, _IDOC)
                out = tmp_path / "out.docx"

                def _boom(*a, **k):
                    raise OSError("synthetic sha failure")

                # Make the report writer's hashing raise; the run must still
                # return 0 (passed), and no report file is written. Patch the
                # report module's LOCAL binding so the unrelated load_profile
                # shell-hash call is unaffected.
                with patch.object(vreport, "sha256_file", _boom):
                    rc = self._generate(tmp_path, idoc, out)
                self.assertEqual(rc, 0, "report-write failure flipped the verdict")
                self.assertTrue(out.is_file())
                report_path = (
                    tmp_path / (out.name + ".visual") / vreport.REPORT_FILENAME
                )
                self.assertFalse(
                    report_path.is_file(),
                    "a degraded report writer must not leave a partial file",
                )
            finally:
                os.chdir(old_cwd)

    # -- test_no_report_on_verify ------------------------------------------
    def test_no_report_on_verify(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            old_cwd = Path.cwd()
            os.chdir(tmp_path)
            try:
                self._extract(tmp_path)
                self.assertEqual(
                    main(
                        [
                            "verify",
                            "--name",
                            "acme",
                            "--scope",
                            "project",
                            "--qa",
                            "fast",
                        ]
                    ),
                    0,
                )
                # No generation_report.json anywhere under the temp tree.
                hits = list(tmp_path.rglob(vreport.REPORT_FILENAME))
                self.assertEqual(hits, [], f"verify wrote a report: {hits}")
            finally:
                os.chdir(old_cwd)

    # -- test_content_hash_stable_under_cosmetic_input ---------------------
    def test_content_hash_stable_under_cosmetic_input(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            old_cwd = Path.cwd()
            os.chdir(tmp_path)
            try:
                self._extract(tmp_path)

                # Same logical idoc, written with different key order / whitespace.
                canonical = self._write_idoc(tmp_path, _IDOC, "canonical.json")
                reordered_blocks = [
                    {"text": "Highlights", "level": 1, "type": "heading"},
                    {
                        "text": "Revenue grew without markdown literals.",
                        "type": "paragraph",
                    },
                ]
                cosmetic_path = tmp_path / "cosmetic.json"
                # Hand-roll a JSON string with extra whitespace + reordered keys.
                cosmetic_path.write_text(
                    json.dumps(
                        {
                            "blocks": reordered_blocks,
                            "cover": {"title": "Quarterly Review"},
                        },
                        indent=4,
                        sort_keys=False,
                    ),
                    encoding="utf-8",
                )

                out1 = tmp_path / "c1.docx"
                out2 = tmp_path / "c2.docx"
                self.assertEqual(self._generate(tmp_path, canonical, out1), 0)
                self.assertEqual(self._generate(tmp_path, cosmetic_path, out2), 0)

                r1 = json.loads(
                    (
                        tmp_path / (out1.name + ".visual") / vreport.REPORT_FILENAME
                    ).read_text(encoding="utf-8")
                )
                r2 = json.loads(
                    (
                        tmp_path / (out2.name + ".visual") / vreport.REPORT_FILENAME
                    ).read_text(encoding="utf-8")
                )
                self.assertEqual(
                    r1["content_sha256"],
                    r2["content_sha256"],
                    "cosmetic-only input difference changed the content hash",
                )
            finally:
                os.chdir(old_cwd)


if __name__ == "__main__":
    unittest.main()
