# SPDX-License-Identifier: MIT
"""Word-faithful outline-TOC cache rewrite (bookmarks + hyperlinks + PAGEREF).

``refresh_visible_outline_toc_cache`` rewrites the visible cache of every outline
TOC field in the shape Word itself writes for a TOC result: bookmarks authored
around the generated heading paragraphs, one cache entry per heading whose
``pPr`` is deep-copied from the template's own old cached entry of the SAME
level, a ``w:hyperlink`` to the heading's bookmark and a nested dirty PAGEREF
field with no cached result. Covered here:

  - per-level pPr mapping (style trailing digit + indentation fallback);
  - hyperlink anchors match the authored bookmarks (names/ids deterministic,
    unique against pre-existing bookmarks);
  - PAGEREF instructions match the anchors, are dirty, cache no result;
  - the full bare-paragraph span is replaced (no stale template entry text);
  - the SDT-wrapped variant converges to the same shape;
  - fail-closed: no outline TOC -> zero bookmarks and ZERO body mutation;
    a malformed cache (or legacy 2-tuple headings) -> the exact simple plain
    rewrite, zero bookmarks; mixed docs degrade per field, all-or-nothing;
  - end-to-end through docx ``generate()``: rich cache + byte idempotency, and
    a no-TOC generation authors zero bookmarks.
"""

from __future__ import annotations

import re
import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from brandkit.formats.docx import generate as docx_generate
from brandkit.formats.docx import structure
from brandkit.formats.docx.structure import _local_name, w
from brandkit.ir import model as ir
from brandkit.profile import schema
from brandkit.qa.model import Finding

_BOOKMARK_NAME_RE = re.compile(r"^_TocBD\d{6}$")


# ---------------------------------------------------------------------------
# fixture builders
# ---------------------------------------------------------------------------
def _field_char(kind, dirty=False):
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), kind)
    if dirty:
        fc.set(qn("w:dirty"), "true")
    return fc


def _run_with(el):
    r = OxmlElement("w:r")
    r.append(el)
    return r


def _instr(text):
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = text
    return it


def _text_el(text):
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    return t


def _begin_para(doc, instr='TOC \\o "1-3" \\h \\z \\u '):
    """The field-code paragraph: begin / instrText / separate (template shape)."""
    p = doc.add_paragraph()._p
    p.append(_run_with(_field_char("begin", dirty=True)))
    p.append(_run_with(_instr(instr)))
    p.append(_run_with(_field_char("separate")))
    return p


def _entry_para(doc, style, text, indent=None):
    """An OLD cached entry: styled pPr, text, tab, nested PAGEREF with a result."""
    para = doc.add_paragraph()
    p = para._p
    pPr = p.get_or_add_pPr()
    if style:
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), style)
        pPr.append(ps)
    if indent is not None:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent))
        pPr.append(ind)
    p.append(_run_with(_text_el(text)))
    p.append(_run_with(OxmlElement("w:tab")))
    p.append(_run_with(_field_char("begin")))
    p.append(_run_with(_instr(" PAGEREF _Toc0001 \\h ")))
    p.append(_run_with(_field_char("separate")))
    p.append(_run_with(_text_el("3")))
    p.append(_run_with(_field_char("end")))
    return p


def _end_para(doc):
    p = doc.add_paragraph()._p
    p.append(_run_with(_field_char("end")))
    return p


_DEFAULT_ENTRIES = (
    ("TOC1", "1  Overview", None),
    ("TOC2", "1.1  Scope", None),
)


def _multi_para_bare_toc(doc, entries=_DEFAULT_ENTRIES):
    """A real bare-paragraph outline TOC spanning several top-level children."""
    begin = _begin_para(doc)
    entry_ps = [_entry_para(doc, s, t, ind) for s, t, ind in entries]
    end = _end_para(doc)
    return begin, entry_ps, end


def _sdt_toc(doc, entries=_DEFAULT_ENTRIES):
    """The same TOC wrapped in a block-level w:sdt with a docPartGallery."""
    begin, entry_ps, end = _multi_para_bare_toc(doc, entries)
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    obj = OxmlElement("w:docPartObj")
    gal = OxmlElement("w:docPartGallery")
    gal.set(qn("w:val"), "Table of Contents")
    obj.append(gal)
    sdt_pr.append(obj)
    sdt.append(sdt_pr)
    content = OxmlElement("w:sdtContent")
    for el in [begin, *entry_ps, end]:
        content.append(el)  # moves out of the body into sdtContent
    sdt.append(content)
    body = doc.element.body
    sectpr = body.find(w("sectPr"))
    if sectpr is not None:
        sectpr.addprevious(sdt)
    else:
        body.append(sdt)
    return sdt


def _headings3(doc, specs):
    """Real generated heading paragraphs; returns [(level, text, w:p), ...]."""
    out = []
    for level, text in specs:
        para = doc.add_paragraph(text, style="Heading 1")
        out.append((level, text, para._p))
    return out


# ---------------------------------------------------------------------------
# assertion helpers
# ---------------------------------------------------------------------------
def _entry_paragraphs(root):
    """Rewritten rich cache entries = paragraphs carrying a w:hyperlink."""
    return [p for p in root.iter(w("p")) if p.find(w("hyperlink")) is not None]


def _bookmark_starts(doc):
    return list(doc.element.iter(w("bookmarkStart")))


def _body_text(doc):
    return "".join(t.text or "" for t in doc.element.body.iter(w("t")))


def _pageref_parts(entry_p):
    """(instr_el, begin_fldchar, between_separate_and_end_w_t_count) of the
    entry's nested PAGEREF field."""
    instr_el = None
    begin_fld = None
    cached = 0
    state = "before"
    for el in entry_p.iter():
        ln = _local_name(el.tag)
        if ln == "fldChar":
            kind = el.get(w("fldCharType"))
            if kind == "begin":
                begin_fld = el
                state = "code"
            elif kind == "separate":
                state = "result"
            elif kind == "end":
                state = "after"
        elif ln == "instrText" and state == "code":
            instr_el = el
        elif ln == "t" and state == "result":
            cached += 1
    return instr_el, begin_fld, cached


# ---------------------------------------------------------------------------
# rich rewrite
# ---------------------------------------------------------------------------
class OutlineTocRichRewriteTest(unittest.TestCase):
    def test_per_level_ppr_mapping_from_template_entries(self):
        doc = Document()
        _, entry_ps, _ = _multi_para_bare_toc(
            doc, (("TOC1", "1  Overview", 100), ("TOC2", "1.1  Scope", 400))
        )
        old_pprs = [p.find(w("pPr")) for p in entry_ps]
        old_xml = [etree.tostring(ppr) for ppr in old_pprs]
        headings = _headings3(doc, [(1, "One"), (2, "Two"), (3, "Three")])

        self.assertEqual(structure.refresh_visible_outline_toc_cache(doc, headings), 1)

        entries = _entry_paragraphs(doc.element.body)
        self.assertEqual(len(entries), 3)
        styles = [structure._p_style_val(p) for p in entries]
        # level 3 has no harvested pPr -> nearest harvested LOWER level (2).
        self.assertEqual(styles, ["TOC1", "TOC2", "TOC2"])
        new_pprs = [p.find(w("pPr")) for p in entries]
        # deepcopy semantics: equal serialization, distinct element objects.
        self.assertEqual(etree.tostring(new_pprs[0]), old_xml[0])
        self.assertEqual(etree.tostring(new_pprs[1]), old_xml[1])
        self.assertEqual(etree.tostring(new_pprs[2]), old_xml[1])
        for new in new_pprs:
            for old in old_pprs:
                self.assertIsNot(new, old)

    def test_indentation_fallback_when_styles_carry_no_digits(self):
        doc = Document()
        _multi_para_bare_toc(
            doc,
            (
                ("SommarioEntry", "Vecchio uno", 0),
                ("SommarioEntry", "Vecchio due", 400),
            ),
        )
        headings = _headings3(doc, [(1, "Uno"), (2, "Due")])

        self.assertEqual(structure.refresh_visible_outline_toc_cache(doc, headings), 1)

        entries = _entry_paragraphs(doc.element.body)
        self.assertEqual(len(entries), 2)
        # level follows the indent rank: 0 -> level 1, 400 -> level 2.
        self.assertEqual(structure._entry_indent_left(entries[0]), 0)
        self.assertEqual(structure._entry_indent_left(entries[1]), 400)

    def test_hyperlink_anchors_match_authored_bookmarks(self):
        doc = Document()
        _multi_para_bare_toc(doc)
        headings = _headings3(doc, [(1, "One"), (2, "Two")])

        structure.refresh_visible_outline_toc_cache(doc, headings)

        starts = _bookmark_starts(doc)
        names = [bs.get(w("name")) for bs in starts]
        self.assertEqual(names, ["_TocBD000001", "_TocBD000002"])
        for name in names:
            self.assertRegex(name, _BOOKMARK_NAME_RE)
        anchors = [h.get(w("anchor")) for h in doc.element.body.iter(w("hyperlink"))]
        self.assertEqual(anchors, names)
        # Each bookmarkStart sits on its heading paragraph, directly after pPr,
        # with a same-id bookmarkEnd in the same paragraph.
        for (_, _, hp), bs in zip(headings, starts):
            self.assertIs(bs.getparent(), hp)
            children = list(hp)
            self.assertEqual(_local_name(children[0].tag), "pPr")
            self.assertIs(children[1], bs)
            ends = [
                be
                for be in hp.findall(w("bookmarkEnd"))
                if be.get(w("id")) == bs.get(w("id"))
            ]
            self.assertEqual(len(ends), 1)
            self.assertIs(list(hp)[-1], ends[0])

    def test_bookmark_ids_and_names_unique_and_deterministic(self):
        def build():
            doc = Document()
            _multi_para_bare_toc(doc)
            # Pre-seed a colliding bookmark: id 5, the first candidate name.
            seed_p = doc.add_paragraph("seed")._p
            bs = OxmlElement("w:bookmarkStart")
            bs.set(qn("w:id"), "5")
            bs.set(qn("w:name"), "_TocBD000001")
            seed_p.append(bs)
            be = OxmlElement("w:bookmarkEnd")
            be.set(qn("w:id"), "5")
            seed_p.append(be)
            headings = _headings3(doc, [(1, "One"), (2, "Two")])
            structure.refresh_visible_outline_toc_cache(doc, headings)
            return doc

        doc = build()
        new_starts = [
            bs for bs in _bookmark_starts(doc) if bs.get(w("name")) != "_TocBD000001"
        ]
        # Names skip the taken candidate; ids continue past the existing max.
        self.assertEqual(
            [bs.get(w("name")) for bs in new_starts],
            ["_TocBD000002", "_TocBD000003"],
        )
        self.assertEqual([bs.get(w("id")) for bs in new_starts], ["6", "7"])
        all_names = [bs.get(w("name")) for bs in _bookmark_starts(doc)]
        self.assertEqual(len(all_names), len(set(all_names)))
        # Determinism: the identical input doc yields byte-identical XML.
        self.assertEqual(etree.tostring(build().element), etree.tostring(doc.element))

    def test_pageref_instr_matches_anchor_no_cached_result_and_dirty(self):
        doc = Document()
        _multi_para_bare_toc(doc)
        headings = _headings3(doc, [(1, "One"), (2, "Two")])

        structure.refresh_visible_outline_toc_cache(doc, headings)

        names = [bs.get(w("name")) for bs in _bookmark_starts(doc)]
        entries = _entry_paragraphs(doc.element.body)
        self.assertEqual(len(entries), len(names))
        for entry, name in zip(entries, names):
            instr_el, begin_fld, cached = _pageref_parts(entry)
            self.assertIsNotNone(instr_el)
            self.assertEqual(instr_el.text, f" PAGEREF {name} \\h ")
            self.assertEqual(instr_el.get(qn("xml:space")), "preserve")
            self.assertEqual(begin_fld.get(w("dirty")), "true")
            self.assertEqual(cached, 0)  # no cached result between separate/end
        # The OUTER field begin is also dirty.
        outer_begin = next(
            fc
            for fc in doc.element.body.iter(w("fldChar"))
            if fc.get(w("fldCharType")) == "begin"
        )
        self.assertEqual(outer_begin.get(w("dirty")), "true")

    def test_no_stale_template_entry_text_survives_bare_span(self):
        doc = Document()
        begin, _, _ = _multi_para_bare_toc(doc)
        begin_index = list(doc.element.body).index(begin)
        headings = _headings3(doc, [(1, "One"), (2, "Two")])

        structure.refresh_visible_outline_toc_cache(doc, headings)

        self.assertNotIn("Overview", _body_text(doc))
        self.assertNotIn("Scope", _body_text(doc))
        # The new field-start paragraph sits at the old begin index.
        children = list(doc.element.body)
        start_p = children[begin_index]
        instrs = [
            it.text or ""
            for it in start_p.iter(w("instrText"))
            if (it.text or "").strip().startswith("TOC")
        ]
        self.assertEqual(len(instrs), 1)
        # Exactly one OUTER TOC begin in the body, and begins/ends balance.
        toc_instrs = [
            it
            for it in doc.element.body.iter(w("instrText"))
            if (it.text or "").strip().startswith("TOC")
        ]
        self.assertEqual(len(toc_instrs), 1)
        kinds = [fc.get(w("fldCharType")) for fc in doc.element.body.iter(w("fldChar"))]
        self.assertEqual(kinds.count("begin"), kinds.count("end"))

    def test_sdt_wrapped_toc_converges_to_same_shape(self):
        doc = Document()
        sdt = _sdt_toc(doc)
        top_sdt_count = sum(1 for c in doc.element.body if _local_name(c.tag) == "sdt")
        headings = _headings3(doc, [(1, "One"), (2, "Two")])

        self.assertEqual(structure.refresh_visible_outline_toc_cache(doc, headings), 1)

        # SDT survives at top level (the rewrite happens inside sdtContent).
        self.assertEqual(
            sum(1 for c in doc.element.body if _local_name(c.tag) == "sdt"),
            top_sdt_count,
        )
        self.assertIsNotNone(sdt.getparent())
        # Anchors match bookmarks (assertion 3).
        names = [bs.get(w("name")) for bs in _bookmark_starts(doc)]
        anchors = [h.get(w("anchor")) for h in sdt.iter(w("hyperlink"))]
        self.assertEqual(anchors, names)
        # PAGEREF shape (assertion 5).
        entries = _entry_paragraphs(sdt)
        self.assertEqual(len(entries), 2)
        for entry, name in zip(entries, names):
            instr_el, begin_fld, cached = _pageref_parts(entry)
            self.assertEqual(instr_el.text, f" PAGEREF {name} \\h ")
            self.assertEqual(begin_fld.get(w("dirty")), "true")
            self.assertEqual(cached, 0)
        # No stale entry text survives (assertion 6).
        self.assertNotIn("Overview", _body_text(doc))
        self.assertNotIn("Scope", _body_text(doc))


# ---------------------------------------------------------------------------
# fail-closed
# ---------------------------------------------------------------------------
class OutlineTocFailClosedTest(unittest.TestCase):
    def test_no_outline_toc_authors_zero_bookmarks_and_zero_mutation(self):
        doc = Document()
        # The only field is a CAPTION index (TOC \c) - not an outline TOC.
        begin = _begin_para(doc, 'TOC \\h \\c "Figura" ')
        _entry_para(doc, "TableOfFigures", "Figura 1. Vecchia")
        _end_para(doc)
        self.assertIsNotNone(begin)
        headings = _headings3(doc, [(1, "One"), (2, "Two")])
        before = etree.tostring(doc.element.body)

        self.assertEqual(structure.refresh_visible_outline_toc_cache(doc, headings), 0)

        self.assertEqual(len(_bookmark_starts(doc)), 0)
        self.assertEqual(etree.tostring(doc.element.body), before)

    def test_malformed_cache_falls_back_to_simple_plain_shape(self):
        doc = Document()
        # Unterminated field: begin/instr/separate but NO end fldChar anywhere,
        # so the top-level span is unresolvable.
        field_p = _begin_para(doc)
        field_p.append(_run_with(_text_el("stale entry .... 1")))
        headings = _headings3(doc, [(1, "One"), (2, "Two")])

        self.assertEqual(structure.refresh_visible_outline_toc_cache(doc, headings), 1)

        # All-or-nothing: the only field fell back, so rich mode never engaged.
        self.assertEqual(len(_bookmark_starts(doc)), 0)
        self.assertEqual(len(list(doc.element.body.iter(w("hyperlink")))), 0)
        # The CURRENT single-paragraph plain shape: one text run, " | "-joined
        # entries with two-space level indents, no PAGEREF.
        texts = [t.text for t in field_p.iter(w("t"))]
        self.assertEqual(texts, ["One |   Two"])
        self.assertNotIn(
            "PAGEREF",
            "".join(it.text or "" for it in field_p.iter(w("instrText"))),
        )

    def test_legacy_two_tuple_headings_keep_plain_shape(self):
        doc = Document()
        begin, _, _ = _multi_para_bare_toc(doc)

        rewritten = structure.refresh_visible_outline_toc_cache(
            doc, [(1, "One"), (2, "Two")]
        )

        self.assertEqual(rewritten, 1)
        self.assertEqual(len(_bookmark_starts(doc)), 0)
        self.assertEqual(len(list(doc.element.body.iter(w("hyperlink")))), 0)
        texts = [t.text for t in begin.iter(w("t"))]
        self.assertEqual(texts, ["One |   Two"])

    def test_mixed_fields_one_malformed_one_rich(self):
        doc = Document()
        # Field 1: unterminated (no end fldChar) -> plain.
        bad_p = _begin_para(doc)
        bad_p.append(_run_with(_text_el("bad stale .... 1")))
        # Field 2: well-formed multi-paragraph -> rich.
        _multi_para_bare_toc(doc)
        headings = _headings3(doc, [(1, "One"), (2, "Two")])

        self.assertEqual(structure.refresh_visible_outline_toc_cache(doc, headings), 2)

        # The well-formed field is rich: bookmarks + hyperlinks exist.
        names = [bs.get(w("name")) for bs in _bookmark_starts(doc)]
        self.assertEqual(len(names), 2)
        anchors = [h.get(w("anchor")) for h in doc.element.body.iter(w("hyperlink"))]
        self.assertEqual(anchors, names)
        # The malformed field kept the byte-plain single-paragraph shape.
        self.assertEqual(len(list(bad_p.iter(w("hyperlink")))), 0)
        self.assertEqual([t.text for t in bad_p.iter(w("t"))], ["One |   Two"])
        self.assertNotIn("Overview", _body_text(doc))


# ---------------------------------------------------------------------------
# end-to-end through docx generate()
# ---------------------------------------------------------------------------
def _profile():
    prof = schema.build_envelope("docx", {"name": "t"})
    prof["surface"] = {"docx": {}}
    prof["roles"] = {"_index": []}
    return prof


def _idoc():
    return ir.IntermediateDocument(
        blocks=[
            ir.Heading(level=1, runs=[{"t": "Alpha"}]),
            ir.Heading(level=2, runs=[{"t": "Beta"}]),
            ir.Paragraph(runs=[{"t": "Body."}]),
        ]
    )


class OutlineTocGenerateE2ETest(unittest.TestCase):
    def _toc_shell(self, tmp_path):
        shell = tmp_path / "shell.docx"
        d = Document()
        _multi_para_bare_toc(d)
        d.add_paragraph("Body heading", style="Heading 1")
        d.save(shell)
        return shell

    def _bare_shell(self, tmp_path):
        shell = tmp_path / "shell.docx"
        d = Document()
        d.add_paragraph("x", style="Heading 1")
        d.save(shell)
        return shell

    def _doc_xml(self, path):
        import zipfile

        return (
            zipfile.ZipFile(path).read("word/document.xml").decode("utf-8", "replace")
        )

    def test_generate_authors_bookmarks_and_rich_cache(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._toc_shell(Path(td))
            out = Path(td) / "out.docx"
            findings: list[Finding] = []
            docx_generate.generate(_profile(), shell, _idoc(), out, findings=findings)
            xml = self._doc_xml(out)
            names = re.findall(r'w:name="(_TocBD\d{6})"', xml)
            self.assertEqual(len(names), 2)
            for name in names:
                self.assertIn(f'w:anchor="{name}"', xml)
                self.assertIn(f" PAGEREF {name} \\h ", xml)
            # Stale shell entry text is gone.
            self.assertNotIn("Overview", xml)
            self.assertNotIn("Scope", xml)
            # The field still classifies as an outline TOC.
            self.assertTrue(structure.is_outline_toc_present(Document(out)))

    def test_generate_with_outline_toc_is_byte_idempotent(self):
        import hashlib
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._toc_shell(Path(td))
            a, b = Path(td) / "a.docx", Path(td) / "b.docx"
            docx_generate.generate(_profile(), shell, _idoc(), a)
            docx_generate.generate(_profile(), shell, _idoc(), b)
            self.assertEqual(
                hashlib.sha256(a.read_bytes()).hexdigest(),
                hashlib.sha256(b.read_bytes()).hexdigest(),
            )

    def test_generate_without_outline_toc_authors_zero_bookmarks(self):
        import tempfile

        with tempfile.TemporaryDirectory() as td:
            shell = self._bare_shell(Path(td))
            out = Path(td) / "out.docx"
            docx_generate.generate(_profile(), shell, _idoc(), out)
            self.assertNotIn("bookmarkStart", self._doc_xml(out))


class EmptyHeadingsRebuildTest(unittest.TestCase):
    """Heading-less generation rebuilds every resolvable TOC cache EMPTY.

    Regression guard for the empty-headings bare multi-paragraph span: the
    plain begin-paragraph-only writer used to leave the template's demo
    entries in place AND orphan the original end fldChar (malformed field
    structure). The full span must collapse to [field-start, field-end] with
    balanced fldChars, zero entries, zero bookmarks.
    """

    def _assert_empty_balanced(self, doc):
        body = doc.element.body
        text = "".join(t.text or "" for t in body.iter(w("t")))
        self.assertNotIn("Overview", text)
        self.assertNotIn("Scope", text)
        kinds = [fc.get(qn("w:fldCharType")) for fc in body.iter(w("fldChar"))]
        self.assertEqual(kinds.count("begin"), kinds.count("end"), kinds)
        self.assertEqual(kinds, ["begin", "separate", "end"], kinds)
        self.assertEqual(len(list(body.iter(w("bookmarkStart")))), 0)

    def test_bare_multi_paragraph_span_rebuilds_empty_and_balanced(self):
        doc = Document()
        _multi_para_bare_toc(doc)
        self.assertEqual(structure.refresh_visible_outline_toc_cache(doc, []), 1)
        self._assert_empty_balanced(doc)

    def test_sdt_span_rebuilds_empty_and_balanced(self):
        doc = Document()
        _sdt_toc(doc)
        self.assertEqual(structure.refresh_visible_outline_toc_cache(doc, []), 1)
        self._assert_empty_balanced(doc)

    def test_empty_rebuild_is_idempotent(self):
        doc = Document()
        _multi_para_bare_toc(doc)
        structure.refresh_visible_outline_toc_cache(doc, [])
        first = etree.tostring(doc.element.body)
        structure.refresh_visible_outline_toc_cache(doc, [])
        self.assertEqual(first, etree.tostring(doc.element.body))


if __name__ == "__main__":
    unittest.main()
