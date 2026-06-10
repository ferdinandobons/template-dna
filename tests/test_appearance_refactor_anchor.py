# SPDX-License-Identifier: MIT
"""Frozen-hash anchor for the cross-format appearance refactor (Cluster A, PR-0).

The docx capture/apply appearance logic is being lifted into format-neutral shared
modules (``common.typography`` / ``common.appearance``) that pptx/xlsx will also
drive. The HARD invariant is that the docx adapter still produces BYTE-IDENTICAL
output: it delegates to the shared engine but emits its current WordprocessingML
tokens verbatim, with the set-only-when-unset guards unchanged.

This module pins a ``sha256`` of a docx ``generate()`` output for a FIXED
``(profile, shell, idoc)`` triple into a frozen constant and asserts it stays
identical. The shell is the committed, byte-stable synthetic ``acme_complex.docx``
fixture (NOT a freshly built ``Document()`` whose ``core.xml`` would carry a
wall-clock timestamp); the profile is an inline, self-contained dict that exercises
every refactored apply path:

  - body font + body size (``theme.fonts.body.latin`` / ``size_hp``);
  - body color (``theme.text.body.color`` as a hex);
  - a role-level THEME-token color (``heading.1`` -> ``accent1``), enriched to a hex
    by the resolver from ``theme.colors`` and applied through the WML theme-color map;
  - a per-run palette color TOKEN (``color: "accent1"``) on a hyperlink run, the
    raw-XML hyperlink path that stays inside docx.

Keeping the profile inline (instead of re-extracting it each run) isolates the proof
to ``generate()``'s output bytes, so this anchor breaks ONLY if the refactor changed
what the docx writer emits - which is exactly what PR-0 must never do.

If a LATER, intentional change to the docx writer alters these bytes, recompute the
constant deliberately (the test prints the actual hash on failure) - never silence
the assertion.

NOTE: the ``acme_complex.docx`` shell itself CONTAINS a bare-paragraph outline TOC
field, so ``_FROZEN_SHA256`` moves whenever the outline-TOC cache writer changes
shape (see the recompute log at the constant). The byte-identity guarantee for
documents WITHOUT an outline TOC is pinned separately - and independently of TOC
writer changes - by ``NoOutlineTocAnchorTest`` below, against the same fixture with
the outline TOC field span stripped.
"""

from __future__ import annotations

import hashlib
import io
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from lxml import etree

from brandkit.formats.docx import generate as docx_generate
from brandkit.formats.docx import structure as docx_structure
from brandkit.ir import model as ir
from brandkit.profile import schema

# The committed, byte-stable synthetic shell (also used by the complex-fidelity
# suite). Its on-disk bytes do not change between runs, so the generated output hash
# is reproducible across processes.
_SHELL = (
    Path(__file__).resolve().parents[0] / "fixtures" / "complex" / "acme_complex.docx"
)

# The frozen output hash. Computed from the CURRENT docx writer; the refactor must
# keep it identical. (Recompute deliberately only on an intentional writer change.)
#
# DELIBERATE RECOMPUTE (2026-06-10), per the protocol in the module docstring: the
# Word-faithful outline-TOC cache rewrite intentionally changed these bytes. This
# shell contains a bare-paragraph outline TOC field (TOC \o "1-3" \h \z \u), so the
# triple's output now carries authored heading bookmarks and hyperlink + nested
# dirty-PAGEREF cache entries, and the FULL old field span is replaced (which also
# removes the stale template entries the previous single-paragraph rewrite leaked
# after the field). No implementation of that feature can keep a TOC-bearing
# shell's bytes frozen. Previous value (plain-text cache shape):
#   c96548539684d65df6e91f5ee52009df191ad09670b1e1498672e2add16fa878
# The no-outline-TOC byte-identity guarantee that this anchor used to stand in for
# is now pinned genuinely by NoOutlineTocAnchorTest / _FROZEN_NO_TOC_SHA256 below.
#
# DELIBERATE RECOMPUTE (2026-06-10, second of the day): a kept caption index whose
# sequence received no captions is now rebuilt EMPTY instead of keeping the
# template's demo entries (the "stale derived index" contract). This shell carries
# two caption indexes (TOC \c) and the anchor IDoc emits no captions, so their
# demo caches are now cleared. Previous value (rich outline TOC, demo caption
# caches kept):
#   d6f261d75bcbe3319298d24e249f414856361078c4f860fd0ea19065aceb75b9
_FROZEN_SHA256 = "c20599c329e7c9f3a7b696c8992b20c06b5d9d3f56b7d67f0c60c234ee063c24"


def _anchor_profile() -> dict:
    """A fixed, self-contained profile exercising every refactored apply path."""
    prof = schema.build_envelope("docx", {"name": "anchor"})
    prof["surface"] = {"docx": {}}
    prof["roles"] = {
        "_index": ["paragraph", "heading.1"],
        "paragraph": {
            "resolver": {
                "type": "named_style",
                "style_id": "Normal",
                "style_name": "Normal",
            },
        },
        "heading.1": {
            "resolver": {
                "type": "named_style",
                "style_id": "Heading1",
                "style_name": "Heading 1",
            },
            # A theme-TOKEN color: the resolver enriches it with the concrete hex from
            # theme.colors and the writer applies it via the WML theme-color map.
            "appearance": {"color": {"kind": "theme", "theme": "accent1"}},
        },
    }
    prof["theme"] = {
        "colors": {"accent1": {"hex": "4F81BD"}},
        "fonts": {"body": {"latin": "Roboto", "size_hp": 22}},
        "text": {"body": {"color": {"kind": "hex", "hex": "1F4E79"}}},
        "palette": {
            "accent1": {
                "ref": {"kind": "theme", "theme": "accent1"},
                "provenance": [],
                "frequency": "rare",
                "name": None,
                "purpose": None,
                "use_when": None,
            }
        },
    }
    return prof


def _anchor_idoc() -> ir.IntermediateDocument:
    """A fixed IR exercising headings, a body paragraph with mixed runs, and a
    hyperlink run carrying a per-run palette color TOKEN (the raw-XML link path)."""
    return ir.IntermediateDocument(
        blocks=[
            ir.Heading(level=1, runs=[{"t": "Anchor Title"}]),
            ir.Paragraph(
                runs=[
                    {"t": "Body paragraph with "},
                    {"t": "bold", "b": True},
                    {"t": " text."},
                ]
            ),
            ir.Paragraph(
                runs=[
                    {"t": "A link: "},
                    {"t": "site", "link": "https://example.com", "color": "accent1"},
                ]
            ),
        ]
    )


def _generate_hash() -> str:
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "out.docx"
        docx_generate.generate(_anchor_profile(), _SHELL, _anchor_idoc(), out)
        return hashlib.sha256(out.read_bytes()).hexdigest()


@unittest.skipUnless(_SHELL.exists(), "complex docx fixture missing")
class AppearanceRefactorAnchorTest(unittest.TestCase):
    def test_docx_generate_output_matches_frozen_hash(self):
        """The docx writer's output for the fixed triple is byte-for-byte unchanged."""
        actual = _generate_hash()
        self.assertEqual(
            actual,
            _FROZEN_SHA256,
            "docx generate() output bytes changed: the appearance refactor must "
            "keep the docx adapter byte-identical. If this is an INTENTIONAL writer "
            f"change, update _FROZEN_SHA256 to {actual!r} deliberately.",
        )

    def test_docx_generate_is_byte_idempotent(self):
        """Two generations of the fixed triple hash identically (determinism guard)."""
        self.assertEqual(_generate_hash(), _generate_hash())


# ---------------------------------------------------------------------------
# The TOC-free anchor: documents WITHOUT an outline TOC stay byte-identical.
# ---------------------------------------------------------------------------
# The frozen output hash for the SAME (profile, idoc) pair over a TOC-free shell:
# the committed acme_complex.docx fixture with its outline TOC field span stripped
# (derived deterministically at test time - no second committed binary, and the
# provenance stays auditable). This anchor pins the hard guarantee that writer
# changes scoped to outline-TOC handling (cache shape, bookmark authoring) leave
# documents WITHOUT an outline TOC byte-for-byte unchanged: bookmarks and rich
# cache entries are authored ONLY when an outline TOC field is present. Unlike
# _FROZEN_SHA256 above, an outline-TOC writer change must NEVER move this value.
# Verified identical between pre-change HEAD and the rich-TOC writer (2026-06-10):
# both produce exactly these bytes for the TOC-free triple.
#
# DELIBERATE RECOMPUTE (2026-06-10): moved by the caption-index empty-rebuild (the
# derived TOC-free shell deliberately KEEPS the fixture's two caption indexes, and
# their demo caches are now cleared when no captions are emitted), NOT by an
# outline-TOC writer change - the contract above still holds: zero bookmarks, zero
# rich entries on a TOC-free document. Previous value (demo caption caches kept):
#   64d8f0963e0a67cdb012db922d9467aadc1c625ffd10f27b49d9d4c5d5504da6
_FROZEN_NO_TOC_SHA256 = (
    "7520dfa3fc00c3f02c66bb8fcebf44025068892670ed0194f19bb563cf4fb84e"
)

_WML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _strip_outline_toc(src: Path, dst: Path) -> None:
    """Write ``dst`` = ``src`` minus the outline TOC field's full body span.

    Independent of the engine under test: scans ``word/document.xml`` for the
    top-level child whose field code is an outline TOC instruction (starts with
    ``TOC``, no ``\\c`` caption switch), tracks ``fldChar`` begin/end depth to the
    span's last child, removes the span, and repacks every part with fixed zip
    metadata so the derived shell is itself byte-stable across runs. The two
    caption indexes (``TOC \\c``) in the fixture are deliberately kept.
    """
    with zipfile.ZipFile(src) as zin:
        parts = [(info.filename, zin.read(info.filename)) for info in zin.infolist()]
    parts_map = dict(parts)
    root = etree.fromstring(parts_map["word/document.xml"])
    body = root.find(f"{_WML_NS}body")
    children = list(body)

    def _is_outline_field_code(el) -> bool:
        return any(
            (it.text or "").lstrip().startswith("TOC") and "\\c" not in (it.text or "")
            for it in el.iter(f"{_WML_NS}instrText")
        )

    begin_i = next(i for i, ch in enumerate(children) if _is_outline_field_code(ch))
    depth = 0
    end_i = begin_i
    for i in range(begin_i, len(children)):
        for fc in children[i].iter(f"{_WML_NS}fldChar"):
            kind = fc.get(f"{_WML_NS}fldCharType")
            if kind == "begin":
                depth += 1
            elif kind == "end":
                depth -= 1
        if depth <= 0:
            end_i = i
            break
    for ch in children[begin_i : end_i + 1]:
        body.remove(ch)
    parts_map["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    with zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, _ in parts:
            zi = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            zi.compress_type = zipfile.ZIP_DEFLATED
            zout.writestr(zi, parts_map[name])


def _generate_no_toc_output() -> bytes:
    with tempfile.TemporaryDirectory() as td:
        shell = Path(td) / "no_toc_shell.docx"
        _strip_outline_toc(_SHELL, shell)
        # The derived shell must really be TOC-free, or this anchor pins nothing.
        assert not docx_structure.is_outline_toc_present(Document(str(shell)))
        out = Path(td) / "out.docx"
        docx_generate.generate(_anchor_profile(), shell, _anchor_idoc(), out)
        return out.read_bytes()


@unittest.skipUnless(_SHELL.exists(), "complex docx fixture missing")
class NoOutlineTocAnchorTest(unittest.TestCase):
    def test_no_toc_generate_output_matches_frozen_hash(self):
        """Without an outline TOC the writer's output is byte-for-byte unchanged.

        This is the operative byte-identity guarantee: outline-TOC writer changes
        (cache shape, bookmark authoring) must never reach a TOC-free document.
        """
        data = _generate_no_toc_output()
        actual = hashlib.sha256(data).hexdigest()
        self.assertEqual(
            actual,
            _FROZEN_NO_TOC_SHA256,
            "docx generate() output for a TOC-FREE shell changed: outline-TOC "
            "writer work must leave documents without an outline TOC "
            "byte-identical. If an UNRELATED intentional writer change moved "
            f"these bytes, update _FROZEN_NO_TOC_SHA256 to {actual!r} "
            "deliberately.",
        )

    def test_no_toc_output_authors_zero_bookmarks(self):
        """Bookmarks are authored ONLY when an outline TOC field is present."""
        data = _generate_no_toc_output()
        xml = zipfile.ZipFile(io.BytesIO(data)).read("word/document.xml")
        self.assertNotIn(b"bookmarkStart", xml)

    def test_no_toc_generate_is_byte_idempotent(self):
        """Shell derivation + generation is deterministic across runs."""
        self.assertEqual(_generate_no_toc_output(), _generate_no_toc_output())


if __name__ == "__main__":
    unittest.main()
