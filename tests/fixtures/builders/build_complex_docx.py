# SPDX-License-Identifier: MIT
"""Deterministic builder for the COMPLEX synthetic DOCX fixture.

Produces ``tests/fixtures/complex/acme_complex.docx``: a 100% synthetic
(``Acme Corp``, never proprietary) Word template that stresses the brand-docx
extractor / generator across as many Word component types as can be authored
without Microsoft Word - python-docx for the bulk, and raw lxml for the many
parts python-docx cannot reach (block-level ``w:sdt`` content controls, real
TOC / SEQ complex fields, ``word/numbering.xml`` abstractNum/num, a custom
``w:type="table"`` style, header logo drawing, ``PAGE`` field, footnotes, and a
landscape ``w:sectPr``).

Components authored (each is a real OOXML structure, not a text approximation):

  COVER (multi-slot front matter, all in the cover region before the TOC):
    * a block-level ``w:sdt`` TITLE content control with ``w:alias='Title'`` and
      a placeholder prompt - the shape ``cover.discover_cover`` keys on as an SDT
      anchor (python-docx cannot author this, so it is lxml);
    * a SUBTITLE / description placeholder paragraph;
    * a DOCUMENT-ID placeholder paragraph ("Document ID: {{doc_id}}");
    * a DATE placeholder paragraph ("{{date}}").

  INDEX FRONT MATTER (three real complex fields, each with cached demo entries):
    * an outline Table of Contents  ``TOC \\o "1-3" \\h \\z \\u``;
    * a Table of Tables             ``TOC \\h \\z \\c "Table"``;
    * a Table of Figures            ``TOC \\h \\z \\c "Figure"``.
    Each carries a cached result (styled entry paragraphs with PAGEREF) so the
    field renders before Word recomputes it, and each ``\\c`` switch is the
    opaque seq id ``structure.inventory_fields`` surfaces.

  NUMBERING (real ``word/numbering.xml``, referenced via ``w:numPr`` from named
    paragraph styles, not direct formatting):
    * a 2-level BULLET list  -> styles "Acme Bullet L1" / "Acme Bullet L2";
    * a 1-level NUMBERED list -> style "Acme Number L1".

  TABLE: a custom ``w:type="table"`` style "Acme Table" (header-row shading +
    row banding via ``w:tblStylePr`` conditional formatting) applied to a sample
    table, with a real ``SEQ Table`` caption ("Table 1. ...").

  FIGURE: an inline PNG logo (synthetic, generated in-process) with a real
    ``SEQ Figure`` caption ("Figure 1. ...").

  CALLOUT: a paragraph style "Acme Callout" with shading + a box border.

  HEADER / FOOTER: a synthetic Acme logo image in the default header, and a
    ``PAGE`` field in the default footer.

  SECTIONS: a PORTRAIT first section and a LANDSCAPE second section (distinct
    ``w:sectPr`` page size + orientation).

  FOOTNOTE: one real footnote (``word/footnotes.xml`` + a referencing run).

  DEMO BODY: instruction / lorem-ipsum body content (an "Example heading"
    Heading-1 and following paragraphs) a generation is expected to clear.

The output is content-reproducible within a fixed library set: every id / image
byte / part is fixed, with no randomness or wall-clock, so two rebuilds in the
SAME environment are identical. The committed binary is the source of truth;
rebuilds may differ byte-for-byte across python-docx / lxml versions (benign
serialization noise), so equality is asserted STRUCTURALLY, not by raw bytes
(see tests/test_fixture_determinism.py).

Run:
    PYTHONPATH=scripts .venv/bin/python tests/fixtures/builders/build_complex_docx.py
"""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Twips
from lxml import etree

OUT = Path(__file__).resolve().parents[1] / "complex" / "acme_complex.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"

# Synthetic Acme Corp brand palette (made-up; never proprietary).
ACME_NAVY = "1F3864"
ACME_TEAL = "2E8B8B"
ACME_AMBER = "E8A33D"
ACME_LIGHT = "EAF0F6"
ACME_BAND = "DCE6F1"
WHITE = "FFFFFF"


# ---------------------------------------------------------------------------
# lxml element helpers
# ---------------------------------------------------------------------------
def _w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def _el(tag: str, **attrs) -> etree._Element:
    """Make a ``w:``-namespaced element with ``w:``-namespaced attributes."""
    e = etree.SubElement(etree.Element(_w("_root")), _w(tag))
    e.getparent().remove(e)
    for k, v in attrs.items():
        e.set(_w(k), v)
    return e


def _sub(parent: etree._Element, tag: str, **attrs) -> etree._Element:
    e = etree.SubElement(parent, _w(tag))
    for k, v in attrs.items():
        e.set(_w(k), v)
    return e


def _run(text: str, *, preserve: bool = True, instr: bool = False) -> etree._Element:
    """A ``w:r`` carrying either a ``w:t`` or a ``w:instrText``."""
    r = _el("r")
    leaf = _sub(r, "instrText" if instr else "t")
    leaf.text = text
    if preserve:
        leaf.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _fldchar(kind: str, *, dirty: bool = False) -> etree._Element:
    r = _el("r")
    fc = _sub(r, "fldChar", fldCharType=kind)
    if dirty and kind == "begin":
        fc.set(_w("dirty"), "true")
    return r


def _set_pstyle(paragraph, style_id: str) -> None:
    """Stamp ``w:pPr/w:pStyle@w:val`` on a paragraph by STYLE ID directly.

    Bypasses python-docx's ``style=`` name lookup (which is deprecated for ids and
    cannot reference a custom style not registered under its display name), and
    matches how the extractor keys on style ids.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    existing = pPr.find(_w("pStyle"))
    if existing is not None:
        existing.set(_w("val"), style_id)
        return
    pStyle = etree.SubElement(pPr, _w("pStyle"))
    pStyle.set(_w("val"), style_id)
    pPr.insert(0, pStyle)


def _p(doc, text: str = "", style_id: str | None = None):
    """Add a body paragraph, optionally stamping a style id directly."""
    paragraph = doc.add_paragraph(text)
    if style_id:
        _set_pstyle(paragraph, style_id)
    return paragraph


# ---------------------------------------------------------------------------
# A tiny synthetic PNG logo generated in-process (no external asset on disk).
# ---------------------------------------------------------------------------
def _synthetic_logo_png() -> bytes:
    """Return bytes of a deterministic 96x32 RGBA PNG (an 'Acme' navy block)."""
    w, h = 96, 32
    navy = (31, 56, 100, 255)
    amber = (232, 163, 61, 255)
    raw = bytearray()
    for y in range(h):
        raw.append(0)  # PNG filter type 0 (None) per scanline
        for x in range(w):
            px = amber if (12 <= y < 20 and 6 <= x < 90) else navy
            raw.extend(px)

    def _chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", w, h, 8, 6, 0, 0, 0)  # 8-bit RGBA
    idat = zlib.compress(bytes(raw), 9)
    return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


# ---------------------------------------------------------------------------
# Custom STYLES (paragraph styles + the branded table style). Authored straight
# into ``word/styles.xml`` via lxml so we control header shading / banding /
# borders python-docx cannot express.
# ---------------------------------------------------------------------------
def _add_paragraph_style(
    styles,
    style_id,
    name,
    *,
    based_on="Normal",
    color=None,
    bold=False,
    size_pt=None,
    shading=None,
    box_border=None,
):
    st = _sub(styles, "style", type="paragraph", styleId=style_id)
    st.set(_w("customStyle"), "1")
    _sub(st, "name", val=name)
    _sub(st, "basedOn", val=based_on)
    _sub(st, "qFormat")
    pPr = _sub(st, "pPr")
    if shading:
        _sub(pPr, "shd", val="clear", color="auto", fill=shading)
    if box_border:
        pbdr = _sub(pPr, "pBdr")
        for side in ("top", "left", "bottom", "right"):
            _sub(pbdr, side, val="single", sz="12", space="6", color=box_border)
        _sub(pPr, "spacing", before="120", after="120")
    rPr = _sub(st, "rPr")
    if bold:
        _sub(rPr, "b")
    if color:
        _sub(rPr, "color", val=color)
    if size_pt:
        _sub(rPr, "sz", val=str(int(size_pt * 2)))
    return st


def _add_list_style(styles, style_id, name, num_id):
    """A list paragraph style that references a w:num via w:numPr."""
    st = _sub(styles, "style", type="paragraph", styleId=style_id)
    st.set(_w("customStyle"), "1")
    _sub(st, "name", val=name)
    _sub(st, "basedOn", val="ListParagraph")
    _sub(st, "qFormat")
    pPr = _sub(st, "pPr")
    numPr = _sub(pPr, "numPr")
    _sub(numPr, "numId", val=str(num_id))
    return st


def _add_table_style(styles):
    """A custom ``w:type='table'`` style: header-row shading + row banding."""
    st = _sub(styles, "style", type="table", styleId="AcmeTable")
    st.set(_w("customStyle"), "1")
    _sub(st, "name", val="Acme Table")
    _sub(st, "basedOn", val="TableNormal")
    _sub(st, "uiPriority", val="99")
    # Base table: thin navy grid + banded-row size hint.
    tblPr = _sub(st, "tblPr")
    _sub(tblPr, "tblStyleRowBandSize", val="1")
    borders = _sub(tblPr, "tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _sub(borders, side, val="single", sz="4", space="0", color=ACME_NAVY)
    # Default cell run color.
    rPr = _sub(st, "rPr")
    _sub(rPr, "color", val=ACME_NAVY)
    # First-row (header) conditional formatting: navy fill + white bold text.
    fr = _sub(st, "tblStylePr", type="firstRow")
    fr_rpr = _sub(fr, "rPr")
    _sub(fr_rpr, "b")
    _sub(fr_rpr, "color", val=WHITE)
    fr_tcpr = _sub(fr, "tcPr")
    _sub(fr_tcpr, "shd", val="clear", color="auto", fill=ACME_NAVY)
    # Banded rows: light-blue fill on every other row.
    band = _sub(st, "tblStylePr", type="band1Horz")
    band_tcpr = _sub(band, "tcPr")
    _sub(band_tcpr, "shd", val="clear", color="auto", fill=ACME_BAND)
    return st


def _ensure_list_paragraph_style(styles):
    """python-docx's default styles.xml has no 'List Paragraph'; add a minimal one."""
    for st in styles.findall(_w("style")):
        if st.get(_w("styleId")) == "ListParagraph":
            return
    st = _sub(styles, "style", type="paragraph", styleId="ListParagraph")
    _sub(st, "name", val="List Paragraph")
    _sub(st, "basedOn", val="Normal")
    _sub(st, "uiPriority", val="34")
    _sub(st, "qFormat")
    pPr = _sub(st, "pPr")
    _sub(pPr, "ind", left="720")


def _ensure_toc_styles(styles):
    """Add TOC entry styles + TOCHeading + Caption + Footnote styles if absent."""
    have = {st.get(_w("styleId")) for st in styles.findall(_w("style"))}

    def add_simple(style_id, name, *, based_on="Normal", ui="39"):
        if style_id in have:
            return
        st = _sub(styles, "style", type="paragraph", styleId=style_id)
        _sub(st, "name", val=name)
        _sub(st, "basedOn", val=based_on)
        _sub(st, "uiPriority", val=ui)
        if style_id.startswith("TOC"):
            pPr = _sub(st, "pPr")
            _sub(pPr, "tabs")  # placeholder; real entries carry their own tabs

    for lvl in (1, 2, 3):
        add_simple(f"TOC{lvl}", f"TOC {lvl}")
    add_simple("TOCHeading", "TOC Heading", based_on="Heading1")
    add_simple("TableofFigures", "Table of Figures")
    if "Caption" not in have:
        st = _sub(styles, "style", type="paragraph", styleId="Caption")
        _sub(st, "name", val="Caption")
        _sub(st, "basedOn", val="Normal")
        _sub(st, "uiPriority", val="35")
        _sub(st, "qFormat")
        rPr = _sub(st, "rPr")
        _sub(rPr, "i")
        _sub(rPr, "color", val=ACME_TEAL)
        _sub(rPr, "sz", val="18")
    if "FootnoteText" not in have:
        st = _sub(styles, "style", type="paragraph", styleId="FootnoteText")
        _sub(st, "name", val="Footnote Text")
        _sub(st, "basedOn", val="Normal")
        rPr = _sub(st, "rPr")
        _sub(rPr, "sz", val="20")
    if "FootnoteReference" not in have:
        st = _sub(styles, "style", type="character", styleId="FootnoteReference")
        _sub(st, "name", val="Footnote Reference")
        rPr = _sub(st, "rPr")
        _sub(rPr, "vertAlign", val="superscript")


def _build_styles(doc):
    styles = doc.styles.element
    _ensure_list_paragraph_style(styles)
    _ensure_toc_styles(styles)
    # Branded paragraph styles.
    _add_paragraph_style(
        styles,
        "AcmeCoverTitle",
        "Acme Cover Title",
        color=ACME_NAVY,
        bold=True,
        size_pt=28,
    )
    _add_paragraph_style(
        styles, "AcmeCoverSubtitle", "Acme Cover Subtitle", color=ACME_TEAL, size_pt=14
    )
    _add_paragraph_style(
        styles,
        "AcmeCallout",
        "Acme Callout",
        color=ACME_NAVY,
        shading=ACME_LIGHT,
        box_border=ACME_TEAL,
    )
    # List styles -> reference w:num 1 (bullet L1), 2 (bullet L2), 3 (number L1).
    _add_list_style(styles, "AcmeBulletL1", "Acme Bullet L1", num_id=1)
    _add_list_style(styles, "AcmeBulletL2", "Acme Bullet L2", num_id=2)
    _add_list_style(styles, "AcmeNumberL1", "Acme Number L1", num_id=3)
    # Branded table style.
    _add_table_style(styles)


# ---------------------------------------------------------------------------
# NUMBERING - a real ``word/numbering.xml`` with abstractNum + num.
# python-docx exposes no numbering authoring API for a fresh document, so the
# part is built with lxml and attached to the package.
# ---------------------------------------------------------------------------
def _populate_numbering(root) -> None:
    """Fill an existing ``w:numbering`` element with Acme abstractNum/num defs.

    The python-docx default template already ships a (empty) ``word/numbering.xml``
    part, already related from ``document.xml``. We MUST reuse that part - adding a
    second part of the same name corrupts the zip - so this mutates the existing
    ``CT_Numbering`` element in place rather than authoring a fresh part.
    """
    for child in list(root):
        root.remove(child)

    def abstract(aid, levels):
        an = _sub(root, "abstractNum", abstractNumId=str(aid))
        _sub(an, "multiLevelType", val="hybridMultilevel")
        for lvl, (fmt, text, indent) in enumerate(levels):
            lvl_el = _sub(an, "lvl", ilvl=str(lvl))
            _sub(lvl_el, "start", val="1")
            _sub(lvl_el, "numFmt", val=fmt)
            _sub(lvl_el, "lvlText", val=text)
            _sub(lvl_el, "lvlJc", val="left")
            pPr = _sub(lvl_el, "pPr")
            _sub(pPr, "ind", left=str(indent), hanging="360")
            if fmt == "bullet":
                rPr = _sub(lvl_el, "rPr")
                rfonts = _sub(rPr, "rFonts")
                rfonts.set(_w("ascii"), "Symbol")
                rfonts.set(_w("hAnsi"), "Symbol")
                rfonts.set(_w("hint"), "default")
        return an

    # abstractNum 0: two-level bullet (filled then hollow square).
    abstract(0, [("bullet", "", 720), ("bullet", "", 1440)])
    # abstractNum 1: decimal numbered list.
    abstract(1, [("decimal", "%1.", 720)])

    # num bindings: 1->bullet(ilvl0 entry), 2->bullet, 3->decimal.
    for num_id, aid in ((1, 0), (2, 0), (3, 1)):
        n = _sub(root, "num", numId=str(num_id))
        _sub(n, "abstractNumId", val=str(aid))


# ---------------------------------------------------------------------------
# FOOTNOTES - a real ``word/footnotes.xml`` with the two reserved separators
# plus one authored footnote (id 2).
# ---------------------------------------------------------------------------
def _build_footnotes_xml() -> bytes:
    nsmap = {"w": W}
    root = etree.Element(_w("footnotes"), nsmap=nsmap)

    def sep(fid, kind):
        fn = _sub(root, "footnote", type=kind, id=str(fid))
        p = _sub(fn, "p")
        r = _sub(p, "r")
        _sub(r, kind if kind == "separator" else "continuationSeparator")

    s = _sub(root, "footnote", type="separator", id="-1")
    p = _sub(s, "p")
    r = _sub(p, "r")
    _sub(r, "separator")
    c = _sub(root, "footnote", type="continuationSeparator", id="0")
    p = _sub(c, "p")
    r = _sub(p, "r")
    _sub(r, "continuationSeparator")

    # The authored footnote (id 2).
    fn = _sub(root, "footnote", id="2")
    p = _sub(fn, "p")
    pPr = _sub(p, "pPr")
    _sub(pPr, "pStyle", val="FootnoteText")
    ref_run = _sub(p, "r")
    ref_rpr = _sub(ref_run, "rPr")
    _sub(ref_rpr, "rStyle", val="FootnoteReference")
    _sub(ref_run, "footnoteRef")
    txt = _sub(p, "r")
    t = _sub(txt, "t")
    t.text = " Acme Corp is a fictional company used only for testing."
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _attach_part(doc, partname, content_type, xml_bytes, rel_type):
    """Register a new package part + content-type override + document rel.

    Returns the relationship id assigned to ``document.xml -> partname``.
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    package = doc.part.package
    puri = PackURI("/" + partname)
    part = Part(puri, content_type, xml_bytes, package)
    rid = doc.part.relate_to(part, rel_type)
    return rid


# ---------------------------------------------------------------------------
# COVER - a block-level w:sdt title content control + placeholder paragraphs.
# ---------------------------------------------------------------------------
def _cover_title_sdt():
    """A block-level ``w:sdt`` cover-title content control (lxml; python-docx
    cannot author block-level SDTs)."""
    sdt = _el("sdt")
    sdtPr = _sub(sdt, "sdtPr")
    rpr = _sub(sdtPr, "rPr")
    _sub(rpr, "color", val=ACME_NAVY)
    _sub(rpr, "sz", val="56")
    _sub(rpr, "b")
    _sub(sdtPr, "alias", val="Title")
    _sub(sdtPr, "tag", val="acme_title")
    _sub(sdtPr, "id", val="101")
    _sub(sdtPr, "showingPlcHdr")
    placeholder = _sub(sdtPr, "placeholder")
    _sub(placeholder, "docPart", val="DefaultPlaceholder_Title")
    _sub(sdtPr, "text")
    _sub(sdt, "sdtEndPr")
    sdtContent = _sub(sdt, "sdtContent")
    p = _sub(sdtContent, "p")
    pPr = _sub(p, "pPr")
    _sub(pPr, "pStyle", val="AcmeCoverTitle")
    r = _sub(p, "r")
    rpr2 = _sub(r, "rPr")
    _sub(rpr2, "color", val=ACME_NAVY)
    _sub(rpr2, "sz", val="56")
    _sub(rpr2, "b")
    t = _sub(r, "t")
    t.text = "Insert title here"
    return sdt


def _build_cover(doc):
    body = doc.element.body
    # 1) Block-level SDT title (inserted as the very first body child).
    sdt = _cover_title_sdt()
    body.insert(0, sdt)

    # 2..n) Placeholder paragraphs for subtitle / doc-id / date, each in the
    # cover region (before the TOC), each a short single-line slot.
    sub_p = _p(doc, "{{subtitle}} - an internal Acme Corp brief", "AcmeCoverSubtitle")
    docid_p = doc.add_paragraph("Document ID: {{doc_id}}")
    date_p = doc.add_paragraph("{{date}}")
    # Move them right after the SDT (they were appended at the end of the body).
    for p in (date_p._p, docid_p._p, sub_p._p):
        body.remove(p)
        sdt.addnext(p)


# ---------------------------------------------------------------------------
# INDEX FRONT MATTER - three real complex fields with cached demo entries.
# ---------------------------------------------------------------------------
def _toc_heading(doc, text):
    return _p(doc, text, "TOCHeading")


def _toc_entry(doc, label, page, *, style):
    """A cached TOC/index entry paragraph: a nested PAGEREF field + tab + page."""
    p = _p(doc, "", style)
    pp = p._p
    # The entry hyperlink text run.
    r = _sub(pp, "r")
    t = _sub(r, "t")
    t.text = label
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    # A right tab + a cached PAGEREF (begin / instr / separate / page / end).
    tab = _sub(pp, "r")
    _sub(tab, "tab")
    pp.append(_fldchar("begin"))
    pp.append(_run(f" PAGEREF _Toc{page:04d} \\h ", instr=True))
    pp.append(_fldchar("separate"))
    pp.append(_run(str(page)))
    pp.append(_fldchar("end"))
    return p


def _complex_field(doc, instr, cached_paragraph_builder):
    """Wrap a complex field around cached result paragraphs.

    Layout (Word's pattern for a multi-paragraph field):
      p0: [begin(dirty)] [instrText]
      p1..pn: cached entry paragraphs (the "result"), each a normal paragraph
      pE: [end]
    A ``separate`` fldChar precedes the cached result; the closing ``end`` sits
    in its own trailing paragraph.
    """
    # Field begin + instruction (its own paragraph).
    begin_p = doc.add_paragraph()
    bp = begin_p._p
    bp.append(_fldchar("begin", dirty=True))
    bp.append(_run(instr, instr=True))
    bp.append(_fldchar("separate"))
    # Cached result entries (real styled paragraphs in between).
    cached_paragraph_builder(doc)
    # Field end (its own paragraph).
    end_p = doc.add_paragraph()
    end_p._p.append(_fldchar("end"))


def _build_index_front_matter(doc):
    # --- Table of Contents (outline) ---
    _toc_heading(doc, "Table of Contents")

    def _toc_entries(d):
        _toc_entry(d, "1  Overview", 3, style="TOC1")
        _toc_entry(d, "1.1  Scope", 3, style="TOC2")
        _toc_entry(d, "1.2  Audience", 4, style="TOC2")
        _toc_entry(d, "2  Methodology", 5, style="TOC1")
        _toc_entry(d, "2.1  Data sources", 5, style="TOC2")
        _toc_entry(d, "3  Results", 6, style="TOC1")

    _complex_field(doc, 'TOC \\o "1-3" \\h \\z \\u ', _toc_entries)

    # --- Table of Tables ---
    _toc_heading(doc, "Table of Tables")

    def _tot_entries(d):
        _toc_entry(d, "Table 1. Acme quarterly revenue", 5, style="TableofFigures")
        _toc_entry(d, "Table 2. Acme regional split", 6, style="TableofFigures")

    _complex_field(doc, 'TOC \\h \\z \\c "Table" ', _tot_entries)

    # --- Table of Figures ---
    _toc_heading(doc, "Table of Figures")

    def _tof_entries(d):
        _toc_entry(d, "Figure 1. Acme Corp logo mark", 2, style="TableofFigures")
        _toc_entry(d, "Figure 2. Acme growth curve", 6, style="TableofFigures")

    _complex_field(doc, 'TOC \\h \\z \\c "Figure" ', _tof_entries)


# ---------------------------------------------------------------------------
# DEMO BODY - lists, table+caption, figure+caption, callout, footnote, demo
# heading content. Everything below the index front matter is the freeform body
# a generation would clear.
# ---------------------------------------------------------------------------
def _seq_caption(doc, prefix, seq_name, tail, *, style="Caption"):
    """A real ``SEQ`` caption paragraph: 'Prefix N. tail' with a live SEQ field."""
    p = _p(doc, "", style)
    pp = p._p
    pp.append(_run(f"{prefix} "))
    pp.append(_fldchar("begin"))
    pp.append(_run(f" SEQ {seq_name} \\* ARABIC ", instr=True))
    pp.append(_fldchar("separate"))
    pp.append(_run("1"))
    pp.append(_fldchar("end"))
    pp.append(_run(f". {tail}"))
    return p


def _build_lists(doc):
    _p(doc, "Key Acme principles", "Heading1")
    _p(doc, "First-level bullet about Acme widgets", "AcmeBulletL1")
    _p(doc, "Second-level bullet detail", "AcmeBulletL2")
    _p(doc, "Another second-level detail", "AcmeBulletL2")
    _p(doc, "Another first-level bullet", "AcmeBulletL1")
    _p(doc, "Acme rollout steps", "Heading2")
    _p(doc, "Define the Acme brand profile", "AcmeNumberL1")
    _p(doc, "Extract the template surface", "AcmeNumberL1")
    _p(doc, "Generate the branded document", "AcmeNumberL1")


def _build_table(doc):
    _p(doc, "Acme quarterly revenue", "Heading2")
    table = doc.add_table(rows=3, cols=4)
    table.style = "Acme Table"
    # Tell Word which conditional formats to apply (first row + banding).
    tblPr = table._tbl.tblPr
    _sub(
        tblPr,
        "tblLook",
        firstRow="1",
        lastRow="0",
        firstColumn="0",
        lastColumn="0",
        noHBand="0",
        noVBand="1",
    )
    hdr = ("Quarter", "Revenue", "Growth", "Region")
    for c, label in zip(table.rows[0].cells, hdr):
        c.text = label
    data = [
        ("Q1", "$3.2M", "+8%", "North"),
        ("Q2", "$3.5M", "+9%", "South"),
    ]
    for r, row in enumerate(data, start=1):
        for c, val in zip(table.rows[r].cells, row):
            c.text = val
    _seq_caption(doc, "Table", "Table", "Acme quarterly revenue (demo data).")


def _build_figure(doc, logo_rid, logo_cx, logo_cy):
    _p(doc, "Acme brand mark", "Heading2")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run._r.append(_inline_drawing(logo_rid, logo_cx, logo_cy, "AcmeLogoFigure", 100))
    _seq_caption(doc, "Figure", "Figure", "Acme Corp logo mark (synthetic).")


def _build_callout(doc):
    _p(
        doc,
        "Note: this is a synthetic Acme Corp callout box. Replace the body "
        "content during generation; the cover and the index front matter are "
        "preserved.",
        "AcmeCallout",
    )


def _build_footnote_paragraph(doc, footnote_id):
    _p(doc, "Acme footnote demo", "Heading2")
    body = doc.add_paragraph("Acme Corp")
    body.add_run(" is a registered placeholder brand")
    # Append a footnoteReference run (id -> footnotes.xml).
    fr = _sub(body._p, "r")
    frpr = _sub(fr, "rPr")
    _sub(frpr, "rStyle", val="FootnoteReference")
    _sub(fr, "footnoteReference", id=str(footnote_id))
    body.add_run(" used throughout this template.")


def _build_demo_body(doc):
    _p(doc, "Example heading", "Heading1")
    doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. This demo "
        "paragraph is instruction content that a generation run should clear "
        "and replace with the user's real Acme content."
    )
    _p(doc, "Example subheading", "Heading2")
    doc.add_paragraph(
        "Praesent commodo cursus magna. Curabitur blandit tempus porttitor. "
        "Acme Corp placeholder body text continues here."
    )


# ---------------------------------------------------------------------------
# DRAWINGS - an inline picture (figure) and a header logo, both referencing the
# same image part by relationship id.
# ---------------------------------------------------------------------------
def _inline_drawing(rid, cx, cy, name, doc_pr_id):
    drawing = _el("drawing")
    inline = etree.SubElement(drawing, f"{{{WP}}}inline")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")
    ext = etree.SubElement(inline, f"{{{WP}}}extent")
    ext.set("cx", str(cx))
    ext.set("cy", str(cy))
    eff = etree.SubElement(inline, f"{{{WP}}}effectExtent")
    for k in ("l", "t", "r", "b"):
        eff.set(k, "0")
    docpr = etree.SubElement(inline, f"{{{WP}}}docPr")
    docpr.set("id", str(doc_pr_id))
    docpr.set("name", name)
    cnv = etree.SubElement(inline, f"{{{WP}}}cNvGraphicFramePr")
    locks = etree.SubElement(cnv, f"{{{A}}}graphicFrameLocks")
    locks.set("noChangeAspect", "1")
    graphic = etree.SubElement(inline, f"{{{A}}}graphic")
    gdata = etree.SubElement(graphic, f"{{{A}}}graphicData")
    gdata.set("uri", PIC)
    pic = etree.SubElement(gdata, f"{{{PIC}}}pic")
    nvpic = etree.SubElement(pic, f"{{{PIC}}}nvPicPr")
    cnvpr = etree.SubElement(nvpic, f"{{{PIC}}}cNvPr")
    cnvpr.set("id", "0")
    cnvpr.set("name", name)
    etree.SubElement(nvpic, f"{{{PIC}}}cNvPicPr")
    blipfill = etree.SubElement(pic, f"{{{PIC}}}blipFill")
    blip = etree.SubElement(blipfill, f"{{{A}}}blip")
    blip.set(f"{{{R}}}embed", rid)
    stretch = etree.SubElement(blipfill, f"{{{A}}}stretch")
    etree.SubElement(stretch, f"{{{A}}}fillRect")
    sppr = etree.SubElement(pic, f"{{{PIC}}}spPr")
    xfrm = etree.SubElement(sppr, f"{{{A}}}xfrm")
    off = etree.SubElement(xfrm, f"{{{A}}}off")
    off.set("x", "0")
    off.set("y", "0")
    extb = etree.SubElement(xfrm, f"{{{A}}}ext")
    extb.set("cx", str(cx))
    extb.set("cy", str(cy))
    geom = etree.SubElement(sppr, f"{{{A}}}prstGeom")
    geom.set("prst", "rect")
    etree.SubElement(geom, f"{{{A}}}avLst")
    return drawing


def _build_header_footer(doc, logo_rid, cx, cy):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run()
    # The header logo references the header part's OWN relationship to the image.
    run._r.append(_inline_drawing(logo_rid, cx, cy, "AcmeHeaderLogo", 200))

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = "Page "
    fpp = fp._p
    fpp.append(_fldchar("begin"))
    fpp.append(_run(" PAGE ", instr=True))
    fpp.append(_fldchar("separate"))
    fpp.append(_run("1"))
    fpp.append(_fldchar("end"))


def _relate_image_to(part, image_blob, partname="media/image1.png"):
    """Ensure ``part`` (document or header) relates to a shared image part.

    The image part is added once to the package; each consuming part gets its own
    relationship id back.
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    package = part.package
    puri = PackURI("/word/" + partname)
    image_part = None
    for p in package.iter_parts():
        if p.partname == puri:
            image_part = p
            break
    if image_part is None:
        image_part = Part(puri, "image/png", image_blob, package)
    rid = part.relate_to(image_part, f"{R}/image")
    return rid


# ---------------------------------------------------------------------------
# SECTIONS - convert the document into a portrait section followed by a
# landscape section.
# ---------------------------------------------------------------------------
def _add_landscape_section(doc):
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    # Swap page width/height for landscape (python-docx does not auto-swap).
    new_section.page_width = Twips(15840)  # 11"
    new_section.page_height = Twips(12240)  # 8.5"
    _p(doc, "Acme landscape appendix", "Heading1")
    doc.add_paragraph(
        "This appendix sits in a landscape section. Wide Acme tables and figures "
        "live here. Demo content the generator may clear."
    )
    return new_section


# ---------------------------------------------------------------------------
# settings.xml - footnotePr + a docId-free, deterministic settings part already
# exists; we just request field update on open so cached TOC/SEQ recompute.
# ---------------------------------------------------------------------------
def _request_update_fields(doc):
    settings = doc.settings.element
    if settings.find(_w("updateFields")) is None:
        uf = _el("updateFields", val="true")
        settings.insert(0, uf)


def build(out: Path = OUT) -> Path:
    doc = Document()  # python-docx default template (single portrait section)

    # 1) Styles (paragraph + list + table) authored into styles.xml.
    _build_styles(doc)

    # 2) Numbering: the default template already ships an (empty) numbering part,
    # already related from document.xml. Populate it in place (a duplicate part
    # of the same name would corrupt the zip).
    _populate_numbering(doc.part.numbering_part.element)

    # 3) Footnotes part attached; footnote id 2 is the authored note.
    _attach_part(
        doc,
        "word/footnotes.xml",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        _build_footnotes_xml(),
        f"{R}/footnotes",
    )

    # 4) Shared logo image; relate it to the document part (for the figure) and
    # to the header part (for the header logo) independently.
    logo_blob = _synthetic_logo_png()
    doc_logo_rid = _relate_image_to(doc.part, logo_blob)

    # 5) Cover (block-level SDT title + placeholder slots), then the three index
    # fields, then the demo body. add_* appends in document order, so build the
    # front matter first, then the body.
    _build_cover(doc)
    _build_index_front_matter(doc)

    # Body content (freeform; a generation clears this region).
    _build_lists(doc)
    _build_table(doc)
    # Figure uses the document-part image relationship.
    fig_cx, fig_cy = 914400, 304800  # 1.0in x ~0.33in @ EMU
    _build_figure(doc, doc_logo_rid, fig_cx, fig_cy)
    _build_callout(doc)
    _build_footnote_paragraph(doc, footnote_id=2)
    _build_demo_body(doc)

    # 6) Header (logo) + footer (PAGE field). The header part needs its OWN
    # relationship to the shared image part.
    section = doc.sections[0]
    header_logo_rid = _relate_image_to(section.header.part, logo_blob)
    _build_header_footer(doc, header_logo_rid, 762000, 254000)

    # 7) A second, landscape section after the portrait body.
    _add_landscape_section(doc)

    # 8) Ask Word to refresh the cached fields on open.
    _request_update_fields(doc)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"built {path}")
