# SPDX-License-Identifier: MIT
"""Regression tests for brand appearance capture across the six INDEPENDENT axes
(font / size / color / geometry / table / numbering).

Covers the four layers, per axis:
  - capture: the dominant direct value is recorded into the document defaults
    (theme.fonts.body for font/size, theme.text.body for color, theme.geometry.body
    for paragraph geometry, theme.table.body for table facts) and per-role
    appearance, each axis independently; a no-dominant document captures nothing;
  - resolver: a role's own captured value wins per axis; otherwise the document body
    value fills in (font for every role, but size/color ONLY for the paragraph/body
    family, never a heading; geometry/table with NO family gate), including for a
    missing-role stub;
  - apply: generated runs/paragraphs get the captured value as direct formatting via
    independent per-axis set-only-when-unset guards; a profile with NO captured
    appearance leaves everything inherited (no regression); re-runs stay
    byte-identical;
  - verify: a value the shell does not prove it contains is an ERROR, a shell-backed
    value is accepted, and an empty-appearance profile produces no finding.

The run-typography trio (font / size / color) applies on all formats; geometry
(Cluster D1), table conditional formats (Cluster D2), and numbering (Cluster D3)
are docx-only. The D3 numbering edge cases live in test_numbering_fidelity.py;
this file exercises the other five axes plus the shared orchestration.
"""

from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.enum.dml import MSO_COLOR_TYPE, MSO_THEME_COLOR
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from brandkit.common import text as textutil
from brandkit.formats.docx import extract as docx_extract
from brandkit.formats.docx import generate as docx_generate
from brandkit.formats.docx import typography
from brandkit.ir import model as ir
from brandkit.common import appearance as common_appearance
from brandkit.profile import schema
from brandkit.profile.resolver import ProfileResolver
from brandkit.qa import checks_deterministic

_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _profile(theme=None, roles=None):
    prof = schema.build_envelope("docx", {"name": "t"})
    prof["surface"] = {"docx": {}}
    prof["roles"] = roles or {"_index": []}
    if theme is not None:
        prof["theme"] = theme
    return prof


def _shell(tmp_path, *, heading=True, size_hp=None, hex_color=None):
    shell = tmp_path / "shell.docx"
    d = Document()
    if heading:
        d.add_paragraph("x", style="Heading 1")
    # Seed run-level provenance the size/color verifier reads from document.xml: an
    # explicit w:sz and/or an explicit w:color on a real run.
    if size_hp is not None or hex_color is not None:
        run = d.add_paragraph().add_run("provenance")
        if size_hp is not None:
            run.font.size = Pt(size_hp / 2)
        if hex_color is not None:
            run.font.color.rgb = RGBColor.from_string(hex_color)
    d.save(shell)
    return shell


# ---------------------------------------------------------------------------
# capture
# ---------------------------------------------------------------------------
class CaptureTest(unittest.TestCase):
    def test_dominant_body_font_captured_into_theme(self):
        doc = Document()
        for _ in range(5):
            run = doc.add_paragraph().add_run("body text")
            run.font.name = "Roboto"
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertEqual(theme["fonts"]["body"]["latin"], "Roboto")
        self.assertGreaterEqual(theme["fonts"]["body"]["confidence"], 0.6)

    def test_per_role_font_captured_from_role_style(self):
        doc = Document()
        for _ in range(4):
            run = doc.add_paragraph(style="Heading 1").add_run("H")
            run.font.name = "Montserrat Black"
        # more body runs in a different font: the heading capture must stay
        # style-scoped (Montserrat) while the document body font is Roboto.
        for _ in range(8):
            doc.add_paragraph().add_run("b").font.name = "Roboto"
        roles = {
            "_index": ["heading.1"],
            "heading.1": {
                "resolver": {
                    "type": "named_style",
                    "style_id": "Heading1",
                    "style_name": "Heading 1",
                },
                "appearance": {},
            },
        }
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, roles, theme)
        self.assertEqual(
            roles["heading.1"]["appearance"]["font"]["latin"], "Montserrat Black"
        )
        self.assertEqual(theme["fonts"]["body"]["latin"], "Roboto")

    def test_no_explicit_font_captures_nothing(self):
        doc = Document()
        for _ in range(5):
            doc.add_paragraph().add_run("inherits the style font")  # no run font set
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertNotIn("body", theme["fonts"])

    def test_below_dominance_threshold_captures_nothing(self):
        doc = Document()
        fonts = ["Roboto", "Arial", "Times New Roman", "Courier"]  # 1/4 each, no winner
        for f in fonts:
            doc.add_paragraph().add_run("x").font.name = f
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertNotIn("body", theme["fonts"])


# ---------------------------------------------------------------------------
# Cluster E2: faked-heading-in-body-style detection (pure statistic, docx-first)
# ---------------------------------------------------------------------------
class PseudoHeadingDetectTest(unittest.TestCase):
    """The detector surfaces a body-style size/color OUTLIER as a pseudo_heading fact
    (a pure statistic vs the captured dominant body appearance), and a uniform body
    surfaces nothing - so the comprehend bundle stays byte-identical."""

    def _uniform_body_doc(self, *, size_pt=12.0, n=8):
        doc = Document()
        for _ in range(n):
            run = doc.add_paragraph().add_run("ordinary body text")
            run.font.size = Pt(size_pt)
            run.font.color.theme_color = MSO_THEME_COLOR.TEXT_1
        return doc

    def test_size_outlier_is_surfaced(self):
        # A body-style run at 22pt (44hp) when the body is 12pt (24hp) is a clear size
        # outlier (1.83x) - surfaced with the CAPTURED size, not a synthesized one.
        # (size_hp is round(pt*2): 12pt->24hp, 22pt->44hp.)
        doc = self._uniform_body_doc()
        big = doc.add_paragraph().add_run("a faked heading line")
        big.font.size = Pt(22)
        big.font.color.theme_color = MSO_THEME_COLOR.TEXT_1
        theme = {"colors": {}, "fonts": {}, "text": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        typography.capture_pseudo_headings(doc, {"_index": []}, theme)
        facts = theme["pseudo_headings"]
        self.assertEqual(len(facts), 1, facts)
        self.assertEqual(facts[0]["size_hp"], 44)
        self.assertNotIn("color", facts[0])  # color matches body -> not an outlier
        # Evidence is coarse + brand-text-free (no run text leaks in).
        self.assertIn("44hp", facts[0]["evidence"])
        self.assertNotIn("faked heading", facts[0]["evidence"])

    def test_color_outlier_is_surfaced(self):
        # A body-style run in accent1 when the body is text1 is a color outlier.
        doc = self._uniform_body_doc()
        accent = doc.add_paragraph().add_run("a faked colored heading")
        accent.font.size = Pt(12)  # same size as body -> only color is the outlier
        accent.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        theme = {"colors": {}, "fonts": {}, "text": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        typography.capture_pseudo_headings(doc, {"_index": []}, theme)
        facts = theme["pseudo_headings"]
        self.assertEqual(len(facts), 1, facts)
        self.assertNotIn("size_hp", facts[0])  # size matches body -> not an outlier
        self.assertEqual(facts[0]["color"], {"kind": "theme", "theme": "accent1"})
        self.assertIn("off-body color", facts[0]["evidence"])

    def test_uniform_body_surfaces_nothing(self):
        doc = self._uniform_body_doc()
        theme = {"colors": {}, "fonts": {}, "text": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        typography.capture_pseudo_headings(doc, {"_index": []}, theme)
        # NO key written, so the comprehend bundle stays byte-identical.
        self.assertNotIn("pseudo_headings", theme)

    def test_multiple_distinct_outliers_each_surfaced(self):
        doc = self._uniform_body_doc()
        for pt in (22, 28):
            run = doc.add_paragraph().add_run("a faked heading")
            run.font.size = Pt(pt)
            run.font.color.theme_color = MSO_THEME_COLOR.TEXT_1
        theme = {"colors": {}, "fonts": {}, "text": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        typography.capture_pseudo_headings(doc, {"_index": []}, theme)
        sizes = sorted(f.get("size_hp") for f in theme["pseudo_headings"])
        self.assertEqual(sizes, [44, 56])  # 22pt->44hp, 28pt->56hp

    def test_no_body_dominant_surfaces_nothing(self):
        # A doc whose body has no captured dominant (every run a different size) has no
        # dominant to call an outlier against - the no-capture path, byte-identical.
        doc = Document()
        for pt in (10, 12, 14, 16):
            doc.add_paragraph().add_run("x").font.size = Pt(pt)
        theme = {"colors": {}, "fonts": {}, "text": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        typography.capture_pseudo_headings(doc, {"_index": []}, theme)
        self.assertNotIn("pseudo_headings", theme)

    def test_named_heading_run_is_not_a_candidate(self):
        # A run under a REAL heading style already carries its heading role; it is not
        # a "faked" heading and must never be surfaced as a body-style outlier.
        doc = self._uniform_body_doc()
        head = doc.add_paragraph(style="Heading 1").add_run("a real heading")
        head.font.size = Pt(22)
        theme = {"colors": {}, "fonts": {}, "text": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        typography.capture_pseudo_headings(doc, {"_index": []}, theme)
        self.assertNotIn("pseudo_headings", theme)

    def test_detector_is_deterministic(self):
        # Same template -> identical facts (stable refs, stable order).
        def run():
            doc = self._uniform_body_doc()
            big = doc.add_paragraph().add_run("faked")
            big.font.size = Pt(22)
            big.font.color.theme_color = MSO_THEME_COLOR.TEXT_1
            theme = {"colors": {}, "fonts": {}, "text": {}}
            typography.capture_fonts(doc, {"_index": []}, theme)
            typography.capture_pseudo_headings(doc, {"_index": []}, theme)
            return theme["pseudo_headings"]

        self.assertEqual(run(), run())


# ---------------------------------------------------------------------------
# resolver: role-specific font wins; body font is the fallback (incl. stub)
# ---------------------------------------------------------------------------
class ResolverAppearanceTest(unittest.TestCase):
    def _prof(self):
        return {
            "kind": "docx",
            "theme": {"colors": {}, "fonts": {"body": {"latin": "Roboto"}}},
            "roles": {
                "_index": ["heading.1"],
                "heading.1": {
                    "resolver": {"type": "named_style", "style_id": "Heading1"},
                    "appearance": {"font": {"latin": "Montserrat Black"}},
                    "status": "robust",
                    "confidence": 1.0,
                },
            },
        }

    def test_role_font_wins_over_body(self):
        op = ProfileResolver(self._prof()).resolve_role("heading.1")
        self.assertEqual(op.appearance["font"]["latin"], "Montserrat Black")

    def test_missing_role_stub_gets_body_font(self):
        op = ProfileResolver(self._prof()).resolve_role(
            "paragraph", fallback="paragraph"
        )
        self.assertEqual(op.appearance["font"]["latin"], "Roboto")

    def test_no_body_font_yields_empty_appearance(self):
        prof = self._prof()
        prof["theme"]["fonts"] = {}
        prof["roles"]["heading.1"]["appearance"] = {}
        op = ProfileResolver(prof).resolve_role("paragraph", fallback="paragraph")
        self.assertEqual(op.appearance, {})


# ---------------------------------------------------------------------------
# apply at generate time
# ---------------------------------------------------------------------------
class ApplyTest(unittest.TestCase):
    def test_body_font_applied_to_generated_runs(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={"colors": {}, "fonts": {"body": {"latin": "Roboto"}}}
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "hello world"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            fonts = {r.font.name for p in Document(out).paragraphs for r in p.runs}
            self.assertIn("Roboto", fonts)

    def test_role_font_wins_at_apply_time(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={"colors": {}, "fonts": {"body": {"latin": "Roboto"}}},
                roles={
                    "_index": ["heading.1"],
                    "heading.1": {
                        "resolver": {
                            "type": "named_style",
                            "style_id": "Heading1",
                            "style_name": "Heading 1",
                        },
                        "appearance": {"font": {"latin": "Montserrat Black"}},
                    },
                },
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Heading(level=1, runs=[{"t": "Title"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            fonts = {r.font.name for p in Document(out).paragraphs for r in p.runs}
            self.assertIn("Montserrat Black", fonts)

    def test_no_captured_typography_leaves_runs_unfonted(self):
        # Regression: a profile with no theme.fonts.body and no role appearance must
        # produce runs with no direct font (exactly the pre-feature behavior).
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(theme={"colors": {}, "fonts": {}})
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "plain"}])])
            docx_generate.generate(prof, shell, idoc, out)
            fonts = {r.font.name for p in Document(out).paragraphs for r in p.runs}
            self.assertEqual(fonts, {None})

    def test_body_font_applied_to_table_cell_runs(self):
        # Dimension 4: a table cell paragraph carries no python-docx style, so its
        # runs are branded by the table writer from the resolved table.default op
        # (which falls back to the document body font).
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={"colors": {}, "fonts": {"body": {"latin": "Roboto"}}}
            )
            table = ir.Table(
                columns=[],
                rows=[[ir.TableCell(runs=[{"t": "cell text"}])]],
                role="default",
            )
            idoc = ir.IntermediateDocument(blocks=[table])
            docx_generate.generate(prof, shell, idoc, out)
            cell_fonts = {
                r.font.name
                for t in Document(out).tables
                for row in t.rows
                for c in row.cells
                for p in c.paragraphs
                for r in p.runs
            }
            self.assertIn("Roboto", cell_fonts)

    def test_no_typography_leaves_table_cell_runs_unfonted(self):
        # Regression: with no captured typography, table cell runs stay unfonted.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(theme={"colors": {}, "fonts": {}})
            table = ir.Table(
                columns=[],
                rows=[[ir.TableCell(runs=[{"t": "cell text"}])]],
                role="default",
            )
            idoc = ir.IntermediateDocument(blocks=[table])
            docx_generate.generate(prof, shell, idoc, out)
            cell_fonts = {
                r.font.name
                for t in Document(out).tables
                for row in t.rows
                for c in row.cells
                for p in c.paragraphs
                for r in p.runs
            }
            self.assertEqual(cell_fonts, {None})

    def test_table_cell_hyperlink_run_carries_brand_font(self):
        # A hyperlink run is raw w:r XML (not a python-docx Run), so _apply_appearance
        # cannot reach it; the writer injects w:rFonts (ascii+hAnsi only, no cs) from
        # the resolved op so even link text in a cell carries the brand typeface.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={"colors": {}, "fonts": {"body": {"latin": "Roboto"}}}
            )
            cell = ir.TableCell(runs=[{"t": "site", "link": "https://example.com"}])
            table = ir.Table(columns=[], rows=[[cell]], role="default")
            idoc = ir.IntermediateDocument(blocks=[table])
            docx_generate.generate(prof, shell, idoc, out)
            doc = Document(out)
            cell_xml = doc.tables[0].rows[0].cells[0]._tc.xml
            self.assertIn("w:hyperlink", cell_xml)
            # The rFonts on the link run names the brand font on ascii + hAnsi only.
            rfonts = (
                doc.tables[0]
                .rows[0]
                .cells[0]
                ._tc.findall(
                    f".//{{{_W_NS}}}hyperlink/{{{_W_NS}}}r/"
                    f"{{{_W_NS}}}rPr/{{{_W_NS}}}rFonts"
                )
            )
            self.assertEqual(len(rfonts), 1)
            self.assertEqual(rfonts[0].get(f"{{{_W_NS}}}ascii"), "Roboto")
            self.assertEqual(rfonts[0].get(f"{{{_W_NS}}}hAnsi"), "Roboto")
            self.assertIsNone(rfonts[0].get(f"{{{_W_NS}}}cs"))


# ---------------------------------------------------------------------------
# verify: fail-closed against the shell's available fonts
# ---------------------------------------------------------------------------
class AppearanceTargetsCheckTest(unittest.TestCase):
    def test_font_absent_from_shell_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(
                theme={"colors": {}, "fonts": {"body": {"latin": "ZZZ Bogus Font"}}}
            )
            findings = checks_deterministic.check_appearance_targets(shell, prof)
            errs = [f for f in findings if f.check == "appearance_targets_exist"]
            self.assertTrue(errs)
            self.assertTrue(
                all(f.severity == schema.Severity.ERROR.value for f in errs)
            )

    def test_shell_font_is_accepted(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            # Arial is always in the docDefaults baseline / fontTable.
            prof = _profile(theme={"colors": {}, "fonts": {"body": {"latin": "Arial"}}})
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )

    def test_empty_appearance_profile_has_no_finding(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(theme={"colors": {}, "fonts": {}})
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )


class PromotedAppearanceShellBackedTest(unittest.TestCase):
    """E2: a promoted size/color on a heading role is re-validated SHELL-BACKED by the
    EXISTING ``check_appearance_targets`` (it walks ``roles[*].appearance``), so a
    promoted value the shell does not carry is an ERROR - the model can never inject a
    value the template lacks."""

    def _profile_with_promoted_size(self, size_hp):
        # A heading role carrying a role-SPECIFIC promoted size (as _derive_promote_
        # appearance would write it). Arial keeps the font axis shell-backed so only the
        # promoted size is under test.
        return _profile(
            theme={"colors": {}, "fonts": {"body": {"latin": "Arial"}}},
            roles={
                "_index": ["heading.1"],
                "heading.1": {
                    "resolver": {"type": "named_style", "style_id": "Heading1"},
                    "appearance": {"size_hp": size_hp},
                },
            },
        )

    def test_promoted_size_present_in_shell_passes(self):
        with tempfile.TemporaryDirectory() as td:
            # The shell carries a 22pt (44hp) run, so a promoted 44hp is shell-backed.
            shell = _shell(Path(td), size_hp=44)
            prof = self._profile_with_promoted_size(44)
            errs = [
                f
                for f in checks_deterministic.check_appearance_targets(shell, prof)
                if f.severity == schema.Severity.ERROR.value
            ]
            self.assertEqual(errs, [])

    def test_promoted_size_absent_from_shell_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            # The shell carries only a 44hp run; a promoted 99hp is NOT shell-backed.
            shell = _shell(Path(td), size_hp=44)
            prof = self._profile_with_promoted_size(99)
            errs = [
                f
                for f in checks_deterministic.check_appearance_targets(shell, prof)
                if f.severity == schema.Severity.ERROR.value
            ]
            self.assertTrue(errs)
            self.assertTrue(any("99" in f.message for f in errs), errs)


# ---------------------------------------------------------------------------
# extract: theme major/minor fonts are TRUTHFUL (read from the package)
# ---------------------------------------------------------------------------
def _docx_with_parts(tmp_path, *, theme: str | None, styles: str | None) -> Path:
    """A minimal package carrying only the parts the font reader touches."""
    path = tmp_path / "shell.docx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("[Content_Types].xml", "<x/>")
        if theme is not None:
            z.writestr("word/theme/theme1.xml", theme)
        if styles is not None:
            z.writestr("word/styles.xml", styles)
    return path


def _theme_xml(major: str | None, minor: str | None) -> str:
    def font(tag: str, face: str | None) -> str:
        if face is None:
            return f"<a:{tag}/>"
        return f'<a:{tag}><a:latin typeface="{face}"/></a:{tag}>'

    return (
        f'<a:theme xmlns:a="{_A_NS}"><a:themeElements>'
        f'<a:fontScheme name="X">'
        f"{font('majorFont', major)}{font('minorFont', minor)}"
        f"</a:fontScheme></a:themeElements></a:theme>"
    )


def _styles_xml(ascii_face: str | None) -> str:
    rfonts = "" if ascii_face is None else f'<w:rFonts w:ascii="{ascii_face}"/>'
    return (
        f'<w:styles xmlns:w="{_W_NS}"><w:docDefaults><w:rPrDefault>'
        f"<w:rPr>{rfonts}</w:rPr></w:rPrDefault></w:docDefaults></w:styles>"
    )


class ThemeFontsExtractTest(unittest.TestCase):
    def test_default_shell_reads_theme_latin_and_arial_baseline(self):
        # python-docx's stock template: theme major=Calibri, minor=Cambria, and a
        # docDefaults rFonts that carries only a THEME reference (no literal ascii),
        # so the major fallback keeps Word's Arial baseline and the minor fallback
        # is None (no real ascii to record).
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            fonts = docx_extract._extract_theme(shell)["fonts"]
            self.assertEqual(fonts["major"], {"latin": "Calibri", "fallback": "Arial"})
            self.assertEqual(fonts["minor"], {"latin": "Cambria", "fallback": None})

    def test_major_minor_typeface_is_read_from_theme1(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _docx_with_parts(
                Path(td),
                theme=_theme_xml("Montserrat", "Inter"),
                styles=_styles_xml(None),
            )
            fonts = docx_extract._extract_theme(shell)["fonts"]
            self.assertEqual(fonts["major"]["latin"], "Montserrat")
            self.assertEqual(fonts["minor"]["latin"], "Inter")

    def test_empty_typeface_is_treated_as_none(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _docx_with_parts(
                Path(td),
                theme=_theme_xml("", "Inter"),
                styles=_styles_xml(None),
            )
            fonts = docx_extract._extract_theme(shell)["fonts"]
            self.assertIsNone(fonts["major"]["latin"])
            self.assertEqual(fonts["minor"]["latin"], "Inter")

    def test_doc_default_ascii_drives_the_fallbacks(self):
        # An explicit docDefaults ascii is the document's real baseline: the minor
        # fallback IS that ascii, and the major fallback uses it instead of Arial.
        with tempfile.TemporaryDirectory() as td:
            shell = _docx_with_parts(
                Path(td),
                theme=_theme_xml("Montserrat", "Inter"),
                styles=_styles_xml("Georgia"),
            )
            fonts = docx_extract._extract_theme(shell)["fonts"]
            self.assertEqual(fonts["major"]["fallback"], "Georgia")
            self.assertEqual(fonts["minor"]["fallback"], "Georgia")

    def test_missing_theme_part_yields_none_without_crashing(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _docx_with_parts(Path(td), theme=None, styles=None)
            fonts = docx_extract._extract_theme(shell)["fonts"]
            self.assertEqual(fonts["major"], {"latin": None, "fallback": "Arial"})
            self.assertEqual(fonts["minor"], {"latin": None, "fallback": None})

    def test_theme_latin_widens_the_appearance_allow_set(self):
        # The truthful major/minor latin faces are exactly what widens the
        # allow-set check_appearance_targets validates an applied body font against.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            theme = docx_extract._extract_theme(shell)
            theme["fonts"]["body"] = {"latin": "Cambria"}  # the minor latin face
            prof = _profile(theme=theme)
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )


# ---------------------------------------------------------------------------
# capture: size and color as INDEPENDENT axes
# ---------------------------------------------------------------------------
class CaptureSizeColorTest(unittest.TestCase):
    def test_dominant_body_size_captured_into_theme_fonts(self):
        doc = Document()
        for _ in range(5):
            doc.add_paragraph().add_run("body").font.size = Pt(11)  # 22 half-points
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertEqual(theme["fonts"]["body"]["size_hp"], 22)
        self.assertGreaterEqual(theme["fonts"]["body"]["size_confidence"], 0.6)

    def test_dominant_body_color_hex_captured_into_theme_text(self):
        doc = Document()
        for _ in range(5):
            doc.add_paragraph().add_run("body").font.color.rgb = RGBColor.from_string(
                "1F4E79"
            )
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        # The additive theme.text key (theme.fonts untouched).
        self.assertEqual(
            theme["text"]["body"]["color"], {"kind": "hex", "hex": "1F4E79"}
        )
        self.assertGreaterEqual(theme["text"]["body"]["color_confidence"], 0.6)
        self.assertNotIn("color", theme["fonts"].get("body", {}))

    def test_dominant_body_color_theme_token_captured(self):
        doc = Document()
        for _ in range(5):
            doc.add_paragraph().add_run(
                "body"
            ).font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertEqual(
            theme["text"]["body"]["color"], {"kind": "theme", "theme": "accent1"}
        )

    def test_size_and_color_are_independent_of_font(self):
        # A run with an explicit SIZE but no explicit FONT must capture size only,
        # leaving the font axis untouched (independence).
        doc = Document()
        for _ in range(5):
            doc.add_paragraph().add_run("body").font.size = Pt(12)  # 24 half-points
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertEqual(theme["fonts"]["body"]["size_hp"], 24)
        self.assertNotIn("latin", theme["fonts"]["body"])

    def test_per_role_size_color_captured_independently(self):
        doc = Document()
        for _ in range(4):
            run = doc.add_paragraph(style="Heading 1").add_run("H")
            run.font.size = Pt(18)  # 36 half-points
            run.font.color.rgb = RGBColor.from_string("C00000")
        roles = {
            "_index": ["heading.1"],
            "heading.1": {
                "resolver": {
                    "type": "named_style",
                    "style_id": "Heading1",
                    "style_name": "Heading 1",
                },
                "appearance": {},
            },
        }
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, roles, theme)
        appearance = roles["heading.1"]["appearance"]
        self.assertEqual(appearance["size_hp"], 36)
        self.assertEqual(appearance["color"], {"kind": "hex", "hex": "C00000"})

    def test_no_explicit_size_or_color_captures_nothing(self):
        doc = Document()
        for _ in range(5):
            doc.add_paragraph().add_run("inherits everything")
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertNotIn("body", theme["fonts"])
        self.assertNotIn("text", theme)

    def test_below_threshold_size_captures_nothing(self):
        doc = Document()
        for hp_pt in (10, 11, 12, 14):  # 1/4 each, no winner
            doc.add_paragraph().add_run("x").font.size = Pt(hp_pt)
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertNotIn("size_hp", theme["fonts"].get("body", {}))


# ---------------------------------------------------------------------------
# resolver: size/color independence + the heading-no-body-default rule
# ---------------------------------------------------------------------------
class ResolverSizeColorTest(unittest.TestCase):
    def _prof(self, heading_appearance=None):
        return {
            "kind": "docx",
            "theme": {
                "colors": {},
                "fonts": {"body": {"latin": "Roboto", "size_hp": 22}},
                "text": {"body": {"color": {"kind": "hex", "hex": "1F4E79"}}},
            },
            "roles": {
                "_index": ["heading.1", "paragraph"],
                "heading.1": {
                    "resolver": {"type": "named_style", "style_id": "Heading1"},
                    "appearance": heading_appearance or {},
                    "status": "robust",
                    "confidence": 1.0,
                },
                "paragraph": {
                    "resolver": {"type": "named_style", "style_id": "Normal"},
                    "appearance": {},
                    "status": "robust",
                    "confidence": 1.0,
                },
            },
        }

    def test_body_size_color_flow_to_paragraph(self):
        op = ProfileResolver(self._prof()).resolve_role("paragraph")
        self.assertEqual(op.appearance["size_hp"], 22)
        self.assertEqual(op.appearance["color"], {"kind": "hex", "hex": "1F4E79"})
        # Font body-default flows too (v1 behavior).
        self.assertEqual(op.appearance["font"]["latin"], "Roboto")

    def test_body_size_color_do_not_flow_to_heading(self):
        # The CRITICAL rule: body size/color must NEVER override a heading's intrinsic
        # style size/color. The font body-default still flows.
        op = ProfileResolver(self._prof()).resolve_role("heading.1")
        self.assertNotIn("size_hp", op.appearance)
        self.assertNotIn("color", op.appearance)
        self.assertEqual(op.appearance["font"]["latin"], "Roboto")

    def test_role_specific_size_color_apply_to_heading(self):
        # A role-SPECIFIC captured size/color still applies to a heading (only the
        # BODY default is gated off it).
        prof = self._prof(
            heading_appearance={
                "size_hp": 48,
                "color": {"kind": "theme", "theme": "accent1"},
            }
        )
        op = ProfileResolver(prof).resolve_role("heading.1")
        self.assertEqual(op.appearance["size_hp"], 48)
        self.assertEqual(op.appearance["color"], {"kind": "theme", "theme": "accent1"})

    def test_missing_role_stub_gets_body_size_color(self):
        # The body/paragraph stub keeps the body size/color defaults.
        op = ProfileResolver(self._prof()).resolve_role(
            "paragraph", fallback="paragraph"
        )
        self.assertEqual(op.appearance["size_hp"], 22)
        self.assertEqual(op.appearance["color"], {"kind": "hex", "hex": "1F4E79"})

    def test_pre_feature_profile_resolves_empty_appearance(self):
        # Backward-compat: a profile with no appearance/theme.text/size_hp resolves
        # to an empty appearance, exactly as before this feature.
        prof = {
            "kind": "docx",
            "theme": {"colors": {}, "fonts": {}},
            "roles": {
                "_index": ["paragraph"],
                "paragraph": {
                    "resolver": {"type": "named_style", "style_id": "Normal"},
                    "appearance": {},
                },
            },
        }
        op = ProfileResolver(prof).resolve_role("paragraph")
        self.assertEqual(op.appearance, {})


# ---------------------------------------------------------------------------
# apply: per-axis guards, independence, idempotency, backward-compat
# ---------------------------------------------------------------------------
class ApplySizeColorTest(unittest.TestCase):
    def _para_runs(self, out):
        return [r for p in Document(out).paragraphs for r in p.runs]

    def test_body_size_applied_to_generated_runs(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(theme={"colors": {}, "fonts": {"body": {"size_hp": 22}}})
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "sized body"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            sizes = {r.font.size for r in self._para_runs(out) if r.text}
            self.assertIn(Pt(11), sizes)  # 22 half-points -> 11pt

    def test_body_hex_color_applied_to_generated_runs(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "hex", "hex": "1F4E79"}}},
                }
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "colored body"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            colored = [
                r
                for r in self._para_runs(out)
                if r.text and r.font.color.type == MSO_COLOR_TYPE.RGB
            ]
            self.assertTrue(colored)
            self.assertEqual(str(colored[0].font.color.rgb), "1F4E79")

    def test_role_theme_color_applied_to_generated_runs(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={"colors": {}, "fonts": {}},
                roles={
                    "_index": ["heading.1"],
                    "heading.1": {
                        "resolver": {
                            "type": "named_style",
                            "style_id": "Heading1",
                            "style_name": "Heading 1",
                        },
                        "appearance": {"color": {"kind": "theme", "theme": "accent1"}},
                    },
                },
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Heading(level=1, runs=[{"t": "Title"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            themed = [
                r
                for r in self._para_runs(out)
                if r.text and r.font.color.type == MSO_COLOR_TYPE.THEME
            ]
            self.assertTrue(themed)
            self.assertEqual(themed[0].font.color.theme_color, MSO_THEME_COLOR.ACCENT_1)

    def test_size_without_font_still_applies(self):
        # Per-axis-guard independence: a role with size-but-no-font applies the size.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(theme={"colors": {}, "fonts": {"body": {"size_hp": 28}}})
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "sized but unfonted"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            runs = [r for r in self._para_runs(out) if r.text]
            self.assertTrue(runs)
            self.assertIn(Pt(14), {r.font.size for r in runs})  # 28 hp -> 14pt
            self.assertEqual({r.font.name for r in runs}, {None})  # font untouched

    def test_unmapped_theme_token_skips_and_records_finding(self):
        # A theme token outside the closed table is SKIPPED (color left inherited)
        # with an INFO finding, never a raise.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "theme", "theme": "bogus"}}},
                }
            )
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "x"}])])
            findings: list = []
            docx_generate.generate(prof, shell, idoc, out, findings=findings)
            runs = [r for r in self._para_runs(out) if r.text]
            self.assertEqual({r.font.color.type for r in runs}, {None})
            self.assertTrue(
                any(f.check == "appearance_color_skipped" for f in findings)
            )

    def test_no_typography_leaves_runs_unsized_and_uncolored(self):
        # Backward-compat: a pre-feature profile (no appearance/theme.text/size_hp)
        # leaves runs with no direct size or color.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(theme={"colors": {}, "fonts": {}})
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "plain"}])])
            docx_generate.generate(prof, shell, idoc, out)
            runs = [r for r in self._para_runs(out) if r.text]
            self.assertEqual({r.font.size for r in runs}, {None})
            self.assertEqual({r.font.color.type for r in runs}, {None})

    def test_size_color_paths_are_byte_idempotent(self):
        # Generating twice with size + hex color must be byte-identical.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out1 = Path(td) / "out1.docx"
            out2 = Path(td) / "out2.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {"body": {"size_hp": 22}},
                    "text": {"body": {"color": {"kind": "hex", "hex": "1F4E79"}}},
                }
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "idempotent body"}])]
            )
            docx_generate.generate(prof, shell, idoc, out1)
            docx_generate.generate(prof, shell, idoc, out2)
            self.assertEqual(out1.read_bytes(), out2.read_bytes())

    def test_theme_color_path_is_byte_idempotent(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out1 = Path(td) / "out1.docx"
            out2 = Path(td) / "out2.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "theme", "theme": "accent1"}}},
                }
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "themed body"}])]
            )
            docx_generate.generate(prof, shell, idoc, out1)
            docx_generate.generate(prof, shell, idoc, out2)
            self.assertEqual(out1.read_bytes(), out2.read_bytes())


# ---------------------------------------------------------------------------
# verify: re-validate size/color against the shell's true provenance
# ---------------------------------------------------------------------------
class AppearanceSizeColorCheckTest(unittest.TestCase):
    def test_off_template_size_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            # The shell carries a run at 22 half-points; an applied 99 is off-template.
            shell = _shell(Path(td), size_hp=22)
            prof = _profile(theme={"colors": {}, "fonts": {"body": {"size_hp": 99}}})
            findings = checks_deterministic.check_appearance_targets(shell, prof)
            errs = [f for f in findings if f.check == "appearance_targets_exist"]
            self.assertTrue(errs)
            self.assertTrue(
                all(f.severity == schema.Severity.ERROR.value for f in errs)
            )

    def test_on_template_size_is_accepted(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td), size_hp=22)
            prof = _profile(theme={"colors": {}, "fonts": {"body": {"size_hp": 22}}})
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )

    def test_off_palette_hex_color_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "hex", "hex": "ABCDEF"}}},
                }
            )
            findings = checks_deterministic.check_appearance_targets(shell, prof)
            errs = [f for f in findings if f.check == "appearance_targets_exist"]
            self.assertTrue(errs)
            self.assertTrue(
                all(f.severity == schema.Severity.ERROR.value for f in errs)
            )

    def test_run_provenance_hex_color_is_accepted(self):
        with tempfile.TemporaryDirectory() as td:
            # The hex is on a run in the shell, even though not a palette slot.
            shell = _shell(Path(td), hex_color="123456")
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "hex", "hex": "123456"}}},
                }
            )
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )

    def test_palette_theme_token_color_is_accepted(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            # accent1 is present in python-docx's stock theme palette.
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "theme", "theme": "accent1"}}},
                }
            )
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )

    def test_text1_token_maps_to_dk1_and_is_accepted(self):
        # The closed alias table must map text1 -> dk1 (present in the palette),
        # which resolve_theme_color's tx1/bg1 subset would NOT cover.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "theme", "theme": "text1"}}},
                }
            )
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )

    def test_empty_appearance_profile_has_no_size_color_finding(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(theme={"colors": {}, "fonts": {}})
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )


# ---------------------------------------------------------------------------
# verify: fold theme.palette refs into check_appearance_targets (model-driven color)
# ---------------------------------------------------------------------------
class PaletteRefAppearanceCheckTest(unittest.TestCase):
    """Every theme.palette[*].ref is an applicable run color (a token resolves to
    it), so check_appearance_targets must re-validate each against the shell, the
    same loop as a role/body color - fail-closed on a ref the shell cannot back."""

    def test_palette_theme_slot_ref_present_in_shell_accepted(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            # accent1 is in python-docx's stock clrScheme palette.
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "palette": {
                        "accent1": {"ref": {"kind": "theme", "theme": "accent1"}}
                    },
                }
            )
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )

    def test_palette_clrscheme_slot_ref_dk1_accepted(self):
        # A SEEDED palette keys theme refs by the clrScheme slot name directly
        # (dk1/lt1/hlink...), which has no WML themeColor alias; the check must
        # accept it because it is itself a parsed-palette slot.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "palette": {"dk1": {"ref": {"kind": "theme", "theme": "dk1"}}},
                }
            )
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )

    def test_palette_ref_absent_from_shell_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            # A hex ref the shell neither declares in its palette nor uses on a run.
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "palette": {
                        "hex:ABCDEF": {"ref": {"kind": "hex", "hex": "ABCDEF"}}
                    },
                }
            )
            findings = checks_deterministic.check_appearance_targets(shell, prof)
            errs = [
                f
                for f in findings
                if f.check == "appearance_targets_exist"
                and f.severity == schema.Severity.ERROR.value
            ]
            self.assertTrue(errs, [f.message for f in findings])
            self.assertTrue(any("theme.palette.hex:ABCDEF" in f.message for f in errs))

    def test_palette_theme_token_ref_absent_slot_is_error(self):
        # A theme ref naming a slot the shell's clrScheme does not carry fails closed.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "palette": {
                        "bogus": {"ref": {"kind": "theme", "theme": "notaslot"}}
                    },
                }
            )
            findings = checks_deterministic.check_appearance_targets(shell, prof)
            self.assertTrue(
                any(
                    f.check == "appearance_targets_exist"
                    and f.severity == schema.Severity.ERROR.value
                    for f in findings
                )
            )


# ---------------------------------------------------------------------------
# verify: check_color_token_targets (color-token membership against theme.palette)
# ---------------------------------------------------------------------------
class ColorTokenTargetsCheckTest(unittest.TestCase):
    def _present_comp_profile(self, palette, annotations):
        prof = _profile(theme={"colors": {}, "fonts": {}, "palette": palette})
        prof["comprehension"] = schema.empty_comprehension()
        prof["comprehension"]["status"] = schema.ComprehensionStatus.PRESENT.value
        prof["comprehension"]["palette_annotations"] = annotations
        return prof

    def test_annotation_key_in_palette_is_ok(self):
        prof = self._present_comp_profile(
            {"accent1": {"ref": {"kind": "theme", "theme": "accent1"}}},
            {"accent1": {"name": "primary"}},
        )
        self.assertEqual(checks_deterministic.check_color_token_targets(prof), [])

    def test_annotation_key_absent_from_palette_is_error(self):
        prof = self._present_comp_profile(
            {"accent1": {"ref": {"kind": "theme", "theme": "accent1"}}},
            {"accent9": {"name": "ghost"}},
        )
        findings = checks_deterministic.check_color_token_targets(prof)
        self.assertEqual(len(findings), 1, [f.message for f in findings])
        self.assertEqual(findings[0].check, "color_token_targets_exist")
        self.assertEqual(findings[0].severity, schema.Severity.ERROR.value)
        self.assertIn("accent9", findings[0].message)

    def test_annotation_key_into_empty_palette_is_error(self):
        # Fail-closed on empty: a key into an EMPTY palette is itself an ERROR.
        prof = self._present_comp_profile({}, {"accent1": {"name": "x"}})
        findings = checks_deterministic.check_color_token_targets(prof)
        self.assertTrue(
            any(f.severity == schema.Severity.ERROR.value for f in findings)
        )

    def test_absent_comprehension_is_a_noop(self):
        prof = _profile(
            theme={
                "colors": {},
                "fonts": {},
                "palette": {"accent1": {"ref": {"kind": "theme", "theme": "accent1"}}},
            }
        )
        # default comprehension is 'absent'
        self.assertEqual(checks_deterministic.check_color_token_targets(prof), [])

    def test_wired_into_run_qa(self):
        from brandkit.qa.gate import run_qa

        prof = self._present_comp_profile(
            {"accent1": {"ref": {"kind": "theme", "theme": "accent1"}}},
            {"accent9": {"name": "ghost"}},
        )
        report = run_qa(None, prof, qa="fast", shell=None)
        self.assertTrue(
            any(
                f.check == "color_token_targets_exist"
                and f.severity == schema.Severity.ERROR.value
                for f in report.findings
            ),
            [f.message for f in report.findings],
        )


def _run_with_theme_color(token: str):
    """A docx run carrying an explicit ``w:themeColor`` token, written as raw XML so
    even a spec-valid-but-unmappable token ('none'/'phClr') or the 'UNMAPPED'
    sentinel can be driven through the capture reader."""
    run = Document().add_paragraph().add_run("x")
    rpr = run._r.get_or_add_rPr()
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "auto")
    color.set(qn("w:themeColor"), token)
    rpr.append(color)
    return run


# ---------------------------------------------------------------------------
# Regression tests for the code+quality review fixes (apply/verify symmetry,
# capture crash-safety, heading-fallback gating, verify alias/ERROR coverage).
# ---------------------------------------------------------------------------
class ReviewFixRegressionTest(unittest.TestCase):
    def test_capture_ignores_unmappable_theme_color_without_crashing(self):
        # A spec-valid w:themeColor python-docx cannot map (e.g. 'none'/'phClr')
        # makes .theme_color raise; capture must swallow it, not crash extraction.
        for token in ("none", "phClr"):
            with self.subTest(token=token):
                self.assertIsNone(typography._run_color(_run_with_theme_color(token)))

    def test_capture_drops_unmapped_theme_sentinel(self):
        token = MSO_THEME_COLOR.NOT_THEME_COLOR.xml_value  # "UNMAPPED"
        self.assertIsNone(typography._run_color(_run_with_theme_color(token)))

    def test_minority_accent_color_is_not_captured_as_body(self):
        # 10 runs, only 2 carry an explicit accent color while 8 inherit: the accent
        # must NOT become the body color (it does not dominate ALL runs). This is the
        # real-template inversion (a 2% blue accent over mostly-inherited body text).
        doc = Document()
        for i in range(10):
            run = doc.add_paragraph().add_run("body")
            if i < 2:
                run.font.color.rgb = RGBColor.from_string("3D85C6")
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertNotIn("text", theme)  # no body color captured

    def test_explicit_color_on_all_runs_is_still_captured(self):
        # Sanity peer: when the color IS on (almost) every run it dominates and is
        # captured - the dominance gate is over ALL runs, not just explicit ones.
        doc = Document()
        for _ in range(10):
            doc.add_paragraph().add_run("b").font.color.rgb = RGBColor.from_string(
                "333333"
            )
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertEqual(
            theme["text"]["body"]["color"], {"kind": "hex", "hex": "333333"}
        )

    def test_apply_unmapped_theme_token_skips_with_finding(self):
        run = Document().add_paragraph().add_run("x")
        findings = []
        docx_generate._brand_run_color(
            run,
            {"kind": "theme", "theme": MSO_THEME_COLOR.NOT_THEME_COLOR.xml_value},
            findings,
        )
        self.assertIsNone(run.font.color.type)  # left inherited, nothing written
        self.assertTrue(any(f.check == "appearance_color_skipped" for f in findings))

    def test_apply_normalizes_hash_and_short_hex(self):
        for spelling, expected in (("#1F4E79", "1F4E79"), ("#fff", "FFFFFF")):
            with self.subTest(spelling=spelling):
                run = Document().add_paragraph().add_run("x")
                docx_generate._brand_run_color(
                    run, {"kind": "hex", "hex": spelling}, []
                )
                self.assertEqual(str(run.font.color.rgb), expected)

    def test_apply_malformed_hex_skips_with_finding_not_crash(self):
        run = Document().add_paragraph().add_run("x")
        findings = []
        docx_generate._brand_run_color(run, {"kind": "hex", "hex": "zzzzzz"}, findings)
        self.assertIsNone(run.font.color.type)
        self.assertTrue(any(f.check == "appearance_color_skipped" for f in findings))

    def test_heading_fallback_does_not_inherit_body_size_color(self):
        # A heading.* that falls back to the paragraph style must keep the body FONT
        # (family-agnostic) but NOT the body size/color default (gated on the
        # originally requested role, so the Heading style's intrinsic size survives).
        prof = {
            "kind": "docx",
            "theme": {
                "colors": {},
                "fonts": {"body": {"latin": "Roboto", "size_hp": 32}},
                "text": {"body": {"color": {"kind": "hex", "hex": "1F4E79"}}},
            },
            "roles": {
                "_index": ["paragraph"],
                "paragraph": {
                    "resolver": {"type": "named_style", "style_id": "Normal"},
                },
            },
        }
        op = ProfileResolver(prof).resolve_role("heading.1", fallback="paragraph")
        self.assertEqual((op.appearance.get("font") or {}).get("latin"), "Roboto")
        self.assertIsNone(op.appearance.get("size_hp"))
        self.assertIsNone(op.appearance.get("color"))

    def test_verify_accepts_non_identity_theme_aliases(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            for token in (
                "background1",
                "background2",
                "light1",
                "light2",
                "hyperlink",
                "followedHyperlink",
                "text2",
            ):
                prof = _profile(
                    theme={
                        "colors": {},
                        "fonts": {},
                        "text": {"body": {"color": {"kind": "theme", "theme": token}}},
                    }
                )
                with self.subTest(token=token):
                    self.assertEqual(
                        checks_deterministic.check_appearance_targets(shell, prof), []
                    )

    def test_verify_errors_on_unmapped_theme_token(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "theme", "theme": "bogus"}}},
                }
            )
            findings = checks_deterministic.check_appearance_targets(shell, prof)
            errs = [f for f in findings if f.check == "appearance_targets_exist"]
            self.assertTrue(errs)
            self.assertTrue(
                all(f.severity == schema.Severity.ERROR.value for f in errs)
            )

    def test_color_only_apply_leaves_font_and_size_untouched(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td), hex_color="1F4E79")
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "text": {"body": {"color": {"kind": "hex", "hex": "1F4E79"}}},
                }
            )
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "x"}])])
            docx_generate.generate(prof, shell, idoc, out)
            colored = [
                r
                for p in Document(out).paragraphs
                for r in p.runs
                if r.font.color.type is not None
            ]
            self.assertTrue(colored)
            self.assertEqual({str(r.font.color.rgb) for r in colored}, {"1F4E79"})
            self.assertEqual({r.font.name for r in colored}, {None})  # font untouched
            self.assertEqual({r.font.size for r in colored}, {None})  # size untouched


# ---------------------------------------------------------------------------
# theme.palette capture (the UNDERSTAND half of model-driven color)
# ---------------------------------------------------------------------------
def _link_run(para, text, *, hex_color=None):
    """Append a run physically nested under a ``w:hyperlink`` element (a real
    link run), optionally carrying an explicit ``w:color`` hex, and return it."""
    hyperlink = OxmlElement("w:hyperlink")
    run = para.add_run(text)
    if hex_color is not None:
        run.font.color.rgb = RGBColor.from_string(hex_color)
    para._p.remove(run._r)
    hyperlink.append(run._r)
    para._p.append(hyperlink)
    return run


class CapturePaletteTest(unittest.TestCase):
    def _theme(self, colors=None, palette_roles=None):
        return {
            "colors": colors if colors is not None else {},
            "palette_roles": palette_roles or {},
            "fonts": {},
        }

    def test_dominant_accent_and_role_color_captured(self):
        # A dominant body color (12 runs), a SPARSE accent (3 runs, below the
        # dominance gate but at the accent floor), and a per-role theme color.
        doc = Document()
        for _ in range(12):
            doc.add_paragraph().add_run("body").font.color.rgb = RGBColor.from_string(
                "333333"
            )
        for _ in range(3):
            doc.add_paragraph().add_run("accent").font.color.rgb = RGBColor.from_string(
                "C00000"
            )
        for _ in range(4):
            run = doc.add_paragraph(style="Heading 1").add_run("H")
            run.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        roles = {
            "_index": ["heading.1"],
            "heading.1": {
                "resolver": {
                    "type": "named_style",
                    "style_id": "Heading1",
                    "style_name": "Heading 1",
                },
                "appearance": {},
            },
        }
        theme = self._theme(
            colors={"accent1": {"hex": "4472C4"}, "dk1": {"hex": "000000"}},
            palette_roles={"primary": {"theme": "accent1"}, "text": {"theme": "dk1"}},
        )
        typography.capture_fonts(doc, roles, theme)
        typography.capture_palette(doc, roles, theme)
        palette = theme["palette"]
        # Dominant body color -> dominant; sparse accent -> accent (no dominance gate).
        self.assertEqual(palette["hex:333333"]["frequency"], "dominant")
        self.assertEqual(palette["hex:333333"]["ref"], {"kind": "hex", "hex": "333333"})
        self.assertEqual(palette["hex:C00000"]["frequency"], "accent")
        # The per-role theme color folds in as a role.appearance + run.color fact and
        # the palette_role map names it (palette_role) - byte-identical ref to
        # typography._color_obj.
        self.assertEqual(
            palette["accent1"]["ref"], {"kind": "theme", "theme": "accent1"}
        )
        wheres = {p["where"] for p in palette["accent1"]["provenance"]}
        self.assertIn("role.appearance", wheres)
        self.assertIn("run.color", wheres)
        self.assertIn("palette_role", wheres)
        # name/purpose/use_when are null in the deterministic path (model fills them).
        for entry in palette.values():
            self.assertIsNone(entry["name"])
            self.assertIsNone(entry["purpose"])
            self.assertIsNone(entry["use_when"])

    def test_provenance_is_sorted_and_deduped(self):
        doc = Document()
        for _ in range(4):
            run = doc.add_paragraph(style="Heading 1").add_run("H")
            run.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        roles = {
            "_index": ["heading.1"],
            "heading.1": {
                "resolver": {
                    "type": "named_style",
                    "style_id": "Heading1",
                    "style_name": "Heading 1",
                },
                "appearance": {},
            },
        }
        theme = self._theme(
            colors={"accent1": {"hex": "4472C4"}},
            palette_roles={"primary": {"theme": "accent1"}},
        )
        typography.capture_fonts(doc, roles, theme)
        typography.capture_palette(doc, roles, theme)
        provenance = theme["palette"]["accent1"]["provenance"]
        keys = [(p["where"], p["detail"]) for p in provenance]
        self.assertEqual(keys, sorted(keys))
        self.assertEqual(len(keys), len(set(keys)))

    def test_sparse_accent_does_not_require_dominance(self):
        # 3 runs of an accent on a body of 20 inheriting runs: the accent NEVER
        # dominates, yet the low-floor aggregation records it as an accent entry.
        doc = Document()
        for i in range(20):
            run = doc.add_paragraph().add_run("body")
            if i < 3:
                run.font.color.rgb = RGBColor.from_string("3D85C6")
        theme = self._theme()
        typography.capture_palette(doc, {"_index": []}, theme)
        self.assertIn("hex:3D85C6", theme["palette"])
        self.assertEqual(theme["palette"]["hex:3D85C6"]["frequency"], "accent")

    def test_below_accent_floor_is_rare(self):
        # 2 runs (below the floor of 3) of an off-theme color: still recorded (it is
        # observed) but as ``rare``, never ``accent``.
        doc = Document()
        for i in range(20):
            run = doc.add_paragraph().add_run("body")
            if i < 2:
                run.font.color.rgb = RGBColor.from_string("AB12CD")
        theme = self._theme()
        typography.capture_palette(doc, {"_index": []}, theme)
        self.assertEqual(theme["palette"]["hex:AB12CD"]["frequency"], "rare")

    def test_empty_palette_when_nothing_observed(self):
        # No theme colors, no palette_roles, no observed run/link/role colors -> {}.
        doc = Document()
        for _ in range(5):
            doc.add_paragraph().add_run("inherits everything")
        theme = self._theme()
        typography.capture_palette(doc, {"_index": []}, theme)
        self.assertEqual(theme["palette"], {})

    def test_link_color_falls_back_to_theme_hlink_slot(self):
        # No explicit link-colored run, but the theme declares hlink/folHlink: those
        # slots carry a link.color where-fact (the link palette is non-empty).
        doc = Document()
        theme = self._theme(
            colors={"hlink": {"hex": "0563C1"}, "folHlink": {"hex": "954F72"}}
        )
        typography.capture_palette(doc, {"_index": []}, theme)
        for slot in ("hlink", "folHlink"):
            wheres = {p["where"] for p in theme["palette"][slot]["provenance"]}
            self.assertIn("link.color", wheres)

    def test_explicit_link_run_color_recorded_as_link_color(self):
        # A run physically under a w:hyperlink with an explicit off-theme hex is an
        # observed link color (link.color where-fact on its hex entry).
        doc = Document()
        para = doc.add_paragraph()
        for _ in range(1):
            _link_run(para, "site", hex_color="0000EE")
        theme = self._theme()
        typography.capture_palette(doc, {"_index": []}, theme)
        wheres = {p["where"] for p in theme["palette"]["hex:0000EE"]["provenance"]}
        self.assertIn("link.color", wheres)
        self.assertIn("run.color", wheres)

    def test_re_extract_is_byte_identical(self):
        doc = Document()
        for _ in range(6):
            doc.add_paragraph().add_run("b").font.color.rgb = RGBColor.from_string(
                "112233"
            )
        for _ in range(3):
            doc.add_paragraph().add_run("a").font.color.rgb = RGBColor.from_string(
                "445566"
            )
        theme1 = self._theme(colors={"accent1": {"hex": "4472C4"}})
        theme2 = self._theme(colors={"accent1": {"hex": "4472C4"}})
        typography.capture_palette(doc, {"_index": []}, theme1)
        rt = _round_trip(doc)
        self.addCleanup(lambda: Path(rt).unlink(missing_ok=True))
        typography.capture_palette(Document(rt), {"_index": []}, theme2)
        import json

        self.assertEqual(
            json.dumps(theme1["palette"], sort_keys=True),
            json.dumps(theme2["palette"], sort_keys=True),
        )

    def test_capture_palette_is_idempotent_on_rerun(self):
        # Calling capture_palette twice on the SAME theme dict must not double up
        # provenance or change anything (de-dup + stable keys).
        doc = Document()
        for _ in range(5):
            run = doc.add_paragraph(style="Heading 1").add_run("H")
            run.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        roles = {
            "_index": ["heading.1"],
            "heading.1": {
                "resolver": {
                    "type": "named_style",
                    "style_id": "Heading1",
                    "style_name": "Heading 1",
                },
                "appearance": {},
            },
        }
        theme = self._theme(
            colors={"accent1": {"hex": "4472C4"}},
            palette_roles={"primary": {"theme": "accent1"}},
        )
        typography.capture_fonts(doc, roles, theme)
        import copy

        typography.capture_palette(doc, roles, theme)
        once = copy.deepcopy(theme["palette"])
        typography.capture_palette(doc, roles, theme)
        self.assertEqual(theme["palette"], once)


def _round_trip(doc):
    """Save a python-docx Document to a temp path and return that path (so a second
    capture pass reads from a fresh package, proving capture is input-stable)."""
    fd = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    fd.close()
    doc.save(fd.name)
    return fd.name


# ---------------------------------------------------------------------------
# extract: theme.palette is captured + saved + schema-valid (end-to-end)
# ---------------------------------------------------------------------------
class PaletteExtractEndToEndTest(unittest.TestCase):
    def test_empty_theme_carries_palette(self):
        self.assertEqual(schema._empty_theme()["palette"], {})

    def test_build_envelope_theme_has_palette(self):
        prof = schema.build_envelope("docx", {"name": "t"})
        self.assertEqual(prof["theme"]["palette"], {})

    def test_extract_saves_schema_valid_profile_with_palette(self):
        import json

        # Build a tiny in-memory shell with a dominant body color so the palette is
        # non-empty, then run the full extract pipeline and validate the profile.
        with tempfile.TemporaryDirectory() as td:
            tmpl = Path(td) / "tmpl.docx"
            d = Document()
            for _ in range(8):
                d.add_paragraph().add_run("body").font.color.rgb = RGBColor.from_string(
                    "204060"
                )
            d.save(tmpl)
            saved = docx_extract.extract(str(tmpl), "pal", scope="project", cwd=td)
            prof = json.loads(Path(saved).read_text())
            self.assertEqual(schema.validate(prof), [])
            self.assertIn("palette", prof["theme"])
            self.assertEqual(
                prof["theme"]["palette"]["hex:204060"]["frequency"], "dominant"
            )


# ---------------------------------------------------------------------------
# schema._validate_palette: shape-only structural validation
# ---------------------------------------------------------------------------
class ValidatePaletteTest(unittest.TestCase):
    def test_none_and_empty_are_ok(self):
        self.assertEqual(schema._validate_palette(None), [])
        self.assertEqual(schema._validate_palette({}), [])

    def test_well_formed_entry_is_ok(self):
        palette = {
            "accent1": {
                "ref": {"kind": "theme", "theme": "accent1"},
                "provenance": [{"where": "run.color", "detail": "accent1"}],
                "frequency": "accent",
                "name": None,
                "purpose": None,
                "use_when": None,
            },
            "hex:112233": {
                "ref": {"kind": "hex", "hex": "112233"},
                "provenance": [],
                "frequency": "rare",
            },
        }
        self.assertEqual(schema._validate_palette(palette), [])

    def test_non_object_is_flagged(self):
        self.assertTrue(schema._validate_palette([1, 2]))

    def test_bad_ref_kind_is_flagged(self):
        palette = {"x": {"ref": {"kind": "bogus"}}}
        problems = schema._validate_palette(palette)
        self.assertTrue(any("ref.kind" in p for p in problems))

    def test_illegal_where_is_flagged(self):
        palette = {
            "accent1": {
                "ref": {"kind": "theme", "theme": "accent1"},
                "provenance": [{"where": "made_up", "detail": "x"}],
            }
        }
        problems = schema._validate_palette(palette)
        self.assertTrue(any("where" in p for p in problems))

    def test_illegal_frequency_is_flagged(self):
        palette = {
            "accent1": {
                "ref": {"kind": "theme", "theme": "accent1"},
                "frequency": "sometimes",
            }
        }
        problems = schema._validate_palette(palette)
        self.assertTrue(any("frequency" in p for p in problems))

    def test_captured_palette_passes_validation(self):
        # The real capture output (theme + accent) must be structurally valid.
        doc = Document()
        for _ in range(6):
            doc.add_paragraph().add_run("b").font.color.rgb = RGBColor.from_string(
                "334455"
            )
        theme = {
            "colors": {"accent1": {"hex": "4472C4"}},
            "palette_roles": {"primary": {"theme": "accent1"}},
            "fonts": {},
        }
        typography.capture_palette(doc, {"_index": []}, theme)
        self.assertEqual(schema._validate_palette(theme["palette"]), [])


# ---------------------------------------------------------------------------
# APPLY: model-driven run COLOR token (resolve off theme.palette, apply via the
# existing _brand_run_color; structural hex rejection in normalize_runs).
# ---------------------------------------------------------------------------
def _palette_entry(kind: str, value: str) -> dict:
    """A minimal valid theme.palette entry whose ref is {kind, ...}."""
    ref = (
        {"kind": "theme", "theme": value}
        if kind == "theme"
        else {"kind": "hex", "hex": value}
    )
    return {
        "ref": ref,
        "provenance": [],
        "frequency": "accent",
        "name": None,
        "purpose": None,
        "use_when": None,
    }


def _palette_profile(palette: dict):
    """A docx profile carrying ``theme.palette`` (and no body color default)."""
    return _profile(theme={"colors": {}, "fonts": {}, "palette": palette})


class RunColorTokenNormalizeTest(unittest.TestCase):
    """normalize_runs preserves a valid color TOKEN and STRUCTURALLY drops a hex."""

    def test_valid_theme_slot_token_is_preserved(self):
        runs = textutil.normalize_runs([{"t": "x", "color": "accent1"}])
        self.assertEqual(runs, [{"t": "x", "color": "accent1"}])

    def test_dotted_named_token_is_preserved(self):
        runs = textutil.normalize_runs([{"t": "x", "color": "brand.primary"}])
        self.assertEqual(runs, [{"t": "x", "color": "brand.primary"}])

    def test_hex_shaped_token_is_dropped(self):
        # A lowercase 6-char hex passes the role-id regex but is rejected as hex.
        runs = textutil.normalize_runs([{"t": "x", "color": "abcdef"}])
        self.assertEqual(runs, [{"t": "x"}])

    def test_uppercase_hex_token_is_dropped(self):
        runs = textutil.normalize_runs([{"t": "x", "color": "1F4E79"}])
        self.assertEqual(runs, [{"t": "x"}])

    def test_hash_bearing_token_is_dropped(self):
        runs = textutil.normalize_runs([{"t": "x", "color": "#fff"}])
        self.assertEqual(runs, [{"t": "x"}])

    def test_hex_prefixed_palette_key_is_dropped(self):
        # The off-theme palette key shape `hex:RRGGBB` is NOT a legal run token.
        runs = textutil.normalize_runs([{"t": "x", "color": "hex:1F4E79"}])
        self.assertEqual(runs, [{"t": "x"}])

    def test_ir_run_carries_token_through_parse(self):
        # The IR round-trips the token; a hex-shaped value never enters the IDoc.
        doc = ir.parse_idoc(
            {
                "blocks": [
                    {"type": "paragraph", "runs": [{"t": "ok", "color": "accent1"}]},
                    {"type": "paragraph", "runs": [{"t": "no", "color": "abcdef"}]},
                ]
            }
        )
        self.assertEqual(doc.blocks[0].runs, [{"t": "ok", "color": "accent1"}])
        self.assertEqual(doc.blocks[1].runs, [{"t": "no"}])


class ResolveColorTest(unittest.TestCase):
    def test_known_token_resolves_to_ref(self):
        prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
        ref = ProfileResolver(prof).resolve_color("accent1")
        self.assertEqual(ref, {"kind": "theme", "theme": "accent1"})

    def test_unknown_token_resolves_to_none(self):
        prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
        self.assertIsNone(ProfileResolver(prof).resolve_color("accent9"))

    def test_empty_palette_resolves_to_none(self):
        self.assertIsNone(ProfileResolver(_profile()).resolve_color("accent1"))

    def test_falsy_token_resolves_to_none(self):
        prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
        self.assertIsNone(ProfileResolver(prof).resolve_color(None))


class ApplyRunColorTokenTest(unittest.TestCase):
    def _para_runs(self, out):
        return [r for p in Document(out).paragraphs for r in p.runs]

    def test_theme_token_resolves_and_applies(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "tinted", "color": "accent1"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            themed = [
                r
                for r in self._para_runs(out)
                if r.text and r.font.color.type == MSO_COLOR_TYPE.THEME
            ]
            self.assertTrue(themed)
            self.assertEqual(themed[0].font.color.theme_color, MSO_THEME_COLOR.ACCENT_1)

    def test_hex_token_resolves_and_applies(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _palette_profile({"brand.primary": _palette_entry("hex", "1F4E79")})
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "hexed", "color": "brand.primary"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            colored = [
                r
                for r in self._para_runs(out)
                if r.text and r.font.color.type == MSO_COLOR_TYPE.RGB
            ]
            self.assertTrue(colored)
            self.assertEqual(str(colored[0].font.color.rgb), "1F4E79")

    def test_minted_palette_alias_resolves_and_applies(self):
        # Cluster E1 cross-format leg (docx): an off-theme hex:RRGGBB accent the model
        # NAMED an ALIAS for (minted via the real merge path, ref byte-copied) is
        # addressable as a clean dotted run-color token on docx and applies as the
        # captured RGB (zero resolver change - the alias is just another palette key).
        from brandkit.profile import comprehension as comp_mod

        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _palette_profile({"hex:1F4E79": _palette_entry("hex", "1F4E79")})
            res = comp_mod.merge(
                prof,
                {"palette_annotations": {"hex:1F4E79": {"alias": "accent.brandblue"}}},
            )
            self.assertTrue(res.ok, res.problems)
            self.assertIn("accent.brandblue", prof["theme"]["palette"])
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(runs=[{"t": "aliased", "color": "accent.brandblue"}])
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            colored = [
                r
                for r in self._para_runs(out)
                if r.text and r.font.color.type == MSO_COLOR_TYPE.RGB
            ]
            self.assertTrue(colored)
            self.assertEqual(str(colored[0].font.color.rgb), "1F4E79")

    def test_unknown_token_leaves_inherited_with_info_finding(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "x", "color": "accent9"}])]
            )
            findings: list = []
            docx_generate.generate(prof, shell, idoc, out, findings=findings)
            runs = [r for r in self._para_runs(out) if r.text]
            self.assertEqual({r.font.color.type for r in runs}, {None})  # inherited
            infos = [f for f in findings if f.check == "color_token_unresolved"]
            self.assertTrue(infos)
            self.assertEqual(infos[0].severity, schema.Severity.INFO.value)

    def test_explicit_token_wins_over_body_default(self):
        # Precedence: an explicit run token (accent1) is the first writer; the body
        # color default (a hex) must NOT overwrite it.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "palette": {"accent1": _palette_entry("theme", "accent1")},
                    "text": {"body": {"color": {"kind": "hex", "hex": "1F4E79"}}},
                }
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "wins", "color": "accent1"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            runs = [r for r in self._para_runs(out) if r.text]
            self.assertTrue(runs)
            # The explicit token wins: the run is a THEME color, not the body hex.
            self.assertEqual(runs[0].font.color.type, MSO_COLOR_TYPE.THEME)
            self.assertEqual(runs[0].font.color.theme_color, MSO_THEME_COLOR.ACCENT_1)

    def test_run_with_no_color_key_is_unchanged(self):
        # Backward-compat: a run carrying no color token is left uncolored even when
        # the profile HAS a palette (no token -> no per-run color).
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "plain"}])])
            docx_generate.generate(prof, shell, idoc, out)
            runs = [r for r in self._para_runs(out) if r.text]
            self.assertEqual({r.font.color.type for r in runs}, {None})

    def test_colored_hyperlink_run_carries_the_color(self):
        # A safe-url hyperlink run is raw XML; the resolved token color is injected
        # as w:color on its rPr (theme token -> w:themeColor).
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(
                        runs=[
                            {
                                "t": "link",
                                "link": "https://example.com",
                                "color": "accent1",
                            }
                        ]
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            doc = Document(out)
            colors = doc.element.findall(f".//{{{_W_NS}}}hyperlink//{{{_W_NS}}}color")
            self.assertTrue(colors)
            self.assertEqual(colors[0].get(f"{{{_W_NS}}}themeColor"), "accent1")

    def test_colored_hex_hyperlink_run_carries_the_hex(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _palette_profile({"brand.primary": _palette_entry("hex", "1F4E79")})
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(
                        runs=[
                            {
                                "t": "link",
                                "link": "https://example.com",
                                "color": "brand.primary",
                            }
                        ]
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            doc = Document(out)
            colors = doc.element.findall(f".//{{{_W_NS}}}hyperlink//{{{_W_NS}}}color")
            self.assertTrue(colors)
            self.assertEqual(colors[0].get(f"{{{_W_NS}}}val"), "1F4E79")

    def test_colored_underlined_hyperlink_color_precedes_u_in_rpr(self):
        # Regression: w:color must be inserted at the schema-correct CT_RPr position
        # (before w:u / w:sz / w:vertAlign), NOT appended last - else an underlined,
        # colored safe-url link emits non-conformant OOXML run-property child order.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(
                        runs=[
                            {
                                "t": "link",
                                "link": "https://example.com",
                                "color": "accent1",
                                "u": True,
                            }
                        ]
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            doc = Document(out)
            rpr = doc.element.find(
                f".//{{{_W_NS}}}hyperlink/{{{_W_NS}}}r/{{{_W_NS}}}rPr"
            )
            self.assertIsNotNone(rpr)
            order = [c.tag.split("}")[1] for c in rpr]
            self.assertIn("color", order)
            self.assertIn("u", order)
            self.assertLess(order.index("color"), order.index("u"))

    def test_run_color_token_is_byte_idempotent(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out1 = Path(td) / "out1.docx"
            out2 = Path(td) / "out2.docx"
            prof = _palette_profile({"accent1": _palette_entry("theme", "accent1")})
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(runs=[{"t": "a", "color": "accent1"}]),
                    ir.Paragraph(
                        runs=[
                            {
                                "t": "b",
                                "link": "https://example.com",
                                "color": "accent1",
                            }
                        ]
                    ),
                ]
            )
            docx_generate.generate(prof, shell, idoc, out1)
            docx_generate.generate(prof, shell, idoc, out2)
            self.assertEqual(out1.read_bytes(), out2.read_bytes())


class ThemeColorHexFallbackTest(unittest.TestCase):
    """A theme-token color carries the resolved hex in w:color@w:val ALONGSIDE the
    themeColor, so renderers that ignore a run themeColor (headless LibreOffice) still
    show the brand color; Word still uses the live themeColor."""

    def _prof(self, *, with_colors=True):
        theme = {
            "colors": {"accent1": {"hex": "4F81BD"}} if with_colors else {},
            "fonts": {},
            "palette": {"accent1": {"ref": {"kind": "theme", "theme": "accent1"}}},
        }
        return _profile(theme=theme)

    def test_resolve_color_enriches_theme_token_with_hex(self):
        op = ProfileResolver(self._prof()).resolve_color("accent1")
        self.assertEqual(op, {"kind": "theme", "theme": "accent1", "hex": "4F81BD"})

    def test_resolve_color_theme_without_theme_colors_has_no_hex(self):
        op = ProfileResolver(self._prof(with_colors=False)).resolve_color("accent1")
        self.assertEqual(op, {"kind": "theme", "theme": "accent1"})

    def test_applied_theme_color_run_carries_val_and_themecolor(self):
        import re
        import zipfile

        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "branded", "color": "accent1"}])]
            )
            docx_generate.generate(self._prof(), shell, idoc, out)
            xml = (
                zipfile.ZipFile(out)
                .read("word/document.xml")
                .decode("utf-8", "replace")
            )
            m = re.search(r'<w:color\b[^>]*w:themeColor="accent1"[^>]*/>', xml)
            self.assertIsNotNone(m, "expected a themeColor=accent1 run color")
            self.assertIn('w:val="4F81BD"', m.group(0))


class HyperlinkRunAppearanceTest(unittest.TestCase):
    """A hyperlink run (raw w:r under w:hyperlink, not in para.runs) receives the
    SAME captured appearance as the surrounding body runs, so a link does not render
    smaller / unfonted than the text around it."""

    def test_hyperlink_run_gets_body_font_and_size(self):
        import re
        import zipfile

        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {"body": {"latin": "Roboto", "size_hp": 32}},
                }
            )
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(
                        runs=[
                            {"t": "see "},
                            {"t": "the site", "link": "https://example.com"},
                        ]
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            xml = (
                zipfile.ZipFile(out)
                .read("word/document.xml")
                .decode("utf-8", "replace")
            )
            link = re.search(r"<w:hyperlink\b.*?</w:hyperlink>", xml, re.S)
            self.assertIsNotNone(link, "expected a hyperlink in the output")
            block = link.group(0)
            self.assertIn('<w:sz w:val="32"/>', block)  # body 16pt reached the link
            self.assertIn('w:ascii="Roboto"', block)  # body font reached the link


def _doc_xml(out) -> str:
    import zipfile

    return zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")


class SlotTokenAndCompositionTest(unittest.TestCase):
    """The review-confirmed apply/verify asymmetry + three-axis composition.

    A clrScheme-slot color token (dk1/lt1/hlink) has no WordprocessingML themeColor
    member, but the resolver enriches it with the concrete hex and apply realizes it
    via that hex (RGB) instead of dropping it to inherited. Plus: a single run can
    carry token-color + body-size + body-font together (links included)."""

    def _prof(self, token, hexv):
        return _profile(
            theme={
                "colors": {token: {"hex": hexv}},
                "fonts": {"body": {"latin": "Roboto", "size_hp": 22}},
                "palette": {token: {"ref": {"kind": "theme", "theme": token}}},
            }
        )

    def test_clrscheme_slot_token_applies_via_hex(self):
        for token, hexv in (("dk1", "123456"), ("hlink", "0563C1")):
            with self.subTest(token=token), tempfile.TemporaryDirectory() as td:
                shell = _shell(Path(td))
                out = Path(td) / "o.docx"
                idoc = ir.IntermediateDocument(
                    blocks=[ir.Paragraph(runs=[{"t": "x", "color": token}])]
                )
                docx_generate.generate(self._prof(token, hexv), shell, idoc, out)
                colored = [
                    r
                    for p in Document(out).paragraphs
                    for r in p.runs
                    if r.font.color.type is not None
                ]
                self.assertTrue(colored, f"{token} should apply, not inherit")
                self.assertEqual({str(r.font.color.rgb) for r in colored}, {hexv})

    def test_three_axes_compose_on_one_run(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "o.docx"
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "x", "color": "accent1"}])]
            )
            docx_generate.generate(self._prof("accent1", "4F81BD"), shell, idoc, out)
            runs = [
                r
                for p in Document(out).paragraphs
                for r in p.runs
                if (r.text or "").strip()
            ]
            self.assertTrue(runs)
            r0 = runs[0]
            self.assertEqual(r0.font.name, "Roboto")  # body font
            self.assertEqual(int(r0.font.size.pt * 2), 22)  # body size
            self.assertEqual(str(r0.font.color.rgb), "4F81BD")  # token color

    def test_colored_link_gets_font_size_color_together(self):
        import re

        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "o.docx"
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(
                        runs=[
                            {
                                "t": "link",
                                "link": "https://example.com",
                                "color": "accent1",
                            }
                        ]
                    )
                ]
            )
            docx_generate.generate(self._prof("accent1", "4F81BD"), shell, idoc, out)
            block = re.search(
                r"<w:hyperlink\b.*?</w:hyperlink>", _doc_xml(out), re.S
            ).group(0)
            self.assertIn('w:ascii="Roboto"', block)  # font
            self.assertIn('<w:sz w:val="22"/>', block)  # size
            self.assertIn('w:themeColor="accent1"', block)  # theme color
            self.assertIn('w:val="4F81BD"', block)  # + hex fallback for LibreOffice

    def test_link_theme_color_without_theme_colors_has_no_val(self):
        import re

        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "o.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "palette": {
                        "accent1": {"ref": {"kind": "theme", "theme": "accent1"}}
                    },
                }
            )
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Paragraph(
                        runs=[
                            {
                                "t": "link",
                                "link": "https://example.com",
                                "color": "accent1",
                            }
                        ]
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            block = re.search(
                r"<w:hyperlink\b.*?</w:hyperlink>", _doc_xml(out), re.S
            ).group(0)
            color_el = re.search(r"<w:color\b[^>]*/>", block).group(0)
            self.assertIn('w:themeColor="accent1"', color_el)
            self.assertNotIn(
                "w:val=", color_el
            )  # no theme.colors hex -> themeColor only


# ===========================================================================
# Cluster D1: paragraph GEOMETRY (spacing / indent / borders / shading), DOCX-ONLY.
# A NEW appearance axis under the SAME _dominant floor: capture, resolver merge (no
# family gate), apply (set-only-when-unset), byte-identity (no-geometry path), the
# honest fail-closed check (well-formed + observed-floor), and a docx end-to-end.
# ===========================================================================
def _set_spacing(para, *, before=None, after=None, line=None, line_rule=None):
    """Set explicit ``w:pPr/w:spacing`` attributes on a python-docx paragraph."""
    ppr = para._p.get_or_add_pPr()
    sp = ppr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        ppr.append(sp)
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
    if line_rule is not None:
        sp.set(qn("w:lineRule"), line_rule)


def _set_indent(para, *, left=None, right=None, first_line=None, hanging=None):
    """Set explicit ``w:pPr/w:ind`` attributes on a python-docx paragraph."""
    ppr = para._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if left is not None:
        ind.set(qn("w:left"), str(left))
    if right is not None:
        ind.set(qn("w:right"), str(right))
    if first_line is not None:
        ind.set(qn("w:firstLine"), str(first_line))
    if hanging is not None:
        ind.set(qn("w:hanging"), str(hanging))


def _set_top_border(para, *, val="single", sz="4", space="1", color="auto"):
    """Set an explicit ``w:pPr/w:pBdr/w:top`` border element on a paragraph."""
    ppr = para._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), val)
    top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), space)
    top.set(qn("w:color"), color)
    pbdr.append(top)


def _set_shading(para, fill):
    """Set an explicit ``w:pPr/w:shd@w:fill`` on a paragraph."""
    ppr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def _heading_role(appearance=None):
    entry = {
        "resolver": {
            "type": "named_style",
            "style_id": "Heading1",
            "style_name": "Heading 1",
        }
    }
    if appearance is not None:
        entry["appearance"] = appearance
    return {"_index": ["heading.1"], "heading.1": entry}


def _out_ppr_xml(out):
    """Every ``w:pPr`` of the generated doc as serialized strings (for raw assertions)."""
    from lxml import etree

    doc = Document(out)
    return [
        etree.tostring(p._p.find(qn("w:pPr")), encoding="unicode")
        for p in doc.paragraphs
        if p._p.find(qn("w:pPr")) is not None
    ]


# ---------------------------------------------------------------------------
# capture: dominant geometry floor, independent fields, absent when no dominance
# ---------------------------------------------------------------------------
class CaptureGeometryTest(unittest.TestCase):
    def test_dominant_spacing_before_captured_into_role(self):
        doc = Document()
        for _ in range(5):
            _set_spacing(doc.add_paragraph(style="Heading 1"), before=240)
        roles = _heading_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        geom = roles["heading.1"]["appearance"]["geometry"]
        self.assertEqual(geom["spacing"]["before_twips"], 240)
        self.assertGreaterEqual(geom["confidence"]["spacing.before_twips"], 0.6)

    def test_dominant_body_geometry_captured_into_theme(self):
        doc = Document()
        for _ in range(6):
            _set_indent(doc.add_paragraph(), left=720)
        roles = {"_index": []}
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        self.assertEqual(theme["geometry"]["body"]["indentation"]["left_twips"], 720)

    def test_independent_geometry_fields(self):
        # A role captures indent but NOT spacing: only the indent axis is recorded.
        doc = Document()
        for _ in range(4):
            _set_indent(doc.add_paragraph(style="Heading 1"), left=360)
        roles = _heading_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        geom = roles["heading.1"]["appearance"]["geometry"]
        self.assertEqual(geom["indentation"]["left_twips"], 360)
        self.assertNotIn("spacing", geom)

    def test_below_dominance_geometry_floor_captures_nothing(self):
        # Only 2 paragraphs carry spacing (< MIN_RUNS=3): nothing is captured.
        doc = Document()
        for _ in range(2):
            _set_spacing(doc.add_paragraph(style="Heading 1"), before=240)
        roles = _heading_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        self.assertNotIn("geometry", roles["heading.1"].get("appearance", {}))

    def test_no_explicit_geometry_captures_nothing(self):
        doc = Document()
        for _ in range(5):
            doc.add_paragraph("inherits geometry")  # no explicit w:pPr geometry
        roles = {"_index": []}
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        self.assertNotIn("geometry", theme)

    def test_minority_geometry_value_is_not_captured(self):
        # 5 paragraphs: 2 with 240, 3 inherit. 240 is 2/5 = 40% < 60% -> not captured.
        doc = Document()
        for _ in range(2):
            _set_spacing(doc.add_paragraph(), before=240)
        for _ in range(3):
            doc.add_paragraph("inherit")
        roles = {"_index": []}
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        self.assertNotIn("geometry", theme)

    def test_border_capture_as_serialized_element(self):
        doc = Document()
        for _ in range(4):
            _set_top_border(doc.add_paragraph(style="Heading 1"))
        roles = _heading_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        borders = roles["heading.1"]["appearance"]["geometry"]["borders"]
        self.assertIn("top", borders)
        self.assertIn("w:top", borders["top"])
        self.assertIn('w:val="single"', borders["top"])

    def test_shading_capture_normalizes_hex(self):
        doc = Document()
        for _ in range(4):
            _set_shading(doc.add_paragraph(style="Heading 1"), "ffEE00")
        roles = _heading_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        self.assertEqual(
            roles["heading.1"]["appearance"]["geometry"]["shading"]["fill_hex"],
            "FFEE00",
        )

    def test_geometry_capture_is_idempotent_on_rerun(self):
        doc = Document()
        for _ in range(5):
            _set_spacing(doc.add_paragraph(style="Heading 1"), before=240, after=120)
        roles = _heading_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_geometry(doc, roles, theme)
        first = dict(roles["heading.1"]["appearance"]["geometry"])
        typography.capture_geometry(doc, roles, theme)
        self.assertEqual(roles["heading.1"]["appearance"]["geometry"], first)


# ---------------------------------------------------------------------------
# resolver: role-specific geometry wins; body geometry fills in for EVERY role
# (NO family gate, unlike size/color)
# ---------------------------------------------------------------------------
class ResolverGeometryTest(unittest.TestCase):
    def _prof(self, *, body=None, heading_geom=None):
        theme = {"colors": {}, "fonts": {}}
        if body is not None:
            theme["geometry"] = {"body": body}
        appearance = {"geometry": heading_geom} if heading_geom is not None else {}
        return {
            "kind": "docx",
            "theme": theme,
            "roles": {
                "_index": ["heading.1"],
                "heading.1": {
                    "resolver": {"type": "named_style", "style_id": "Heading1"},
                    "appearance": appearance,
                    "status": "robust",
                    "confidence": 1.0,
                },
            },
        }

    def test_role_geometry_wins_over_body(self):
        prof = self._prof(
            body={"indentation": {"left_twips": 720}},
            heading_geom={"indentation": {"left_twips": 360}},
        )
        op = ProfileResolver(prof).resolve_role("heading.1")
        self.assertEqual(op.appearance["geometry"]["indentation"]["left_twips"], 360)

    def test_body_geometry_flows_to_heading_no_family_gate(self):
        # CRITICAL: geometry has NO family gate (unlike body size/color). A heading
        # with no captured geometry DOES inherit the body geometry default.
        prof = self._prof(body={"spacing": {"before_twips": 240}})
        op = ProfileResolver(prof).resolve_role("heading.1")
        self.assertEqual(op.appearance["geometry"]["spacing"]["before_twips"], 240)

    def test_heading_with_no_geometry_and_no_body_stays_clean(self):
        prof = self._prof()
        op = ProfileResolver(prof).resolve_role("heading.1")
        self.assertNotIn("geometry", op.appearance)

    def test_pre_d1_profile_resolves_without_geometry(self):
        prof = self._prof()
        op = ProfileResolver(prof).resolve_role("paragraph", fallback="paragraph")
        self.assertNotIn("geometry", op.appearance)


# ---------------------------------------------------------------------------
# apply: geometry written to w:pPr; set-only-when-unset; borders copied verbatim
# ---------------------------------------------------------------------------
class ApplyGeometryTest(unittest.TestCase):
    def _prof_with_body_geometry(self, geometry):
        return _profile(
            theme={"colors": {}, "fonts": {}, "geometry": {"body": geometry}}
        )

    def test_spacing_applied_when_unset(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = self._prof_with_body_geometry(
                {"spacing": {"before_twips": 240, "after_twips": 120}}
            )
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "x"}])])
            docx_generate.generate(prof, shell, idoc, out)
            ppr_xml = "\n".join(_out_ppr_xml(out))
            self.assertIn('w:before="240"', ppr_xml)
            self.assertIn('w:after="120"', ppr_xml)

    def test_indent_applied_independent_from_spacing(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = self._prof_with_body_geometry({"indentation": {"left_twips": 720}})
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "x"}])])
            docx_generate.generate(prof, shell, idoc, out)
            ppr_xml = "\n".join(_out_ppr_xml(out))
            self.assertIn('w:left="720"', ppr_xml)
            self.assertNotIn("w:spacing", ppr_xml)

    def test_geometry_not_written_when_set(self):
        # Set-only-when-unset: a paragraph whose style already carries explicit
        # spacing-before is NOT clobbered. We assert by writing geometry onto a
        # paragraph that already has an authored value via the IR is not possible
        # directly; instead we apply the backend twice and confirm the SECOND value
        # (a different one) does not overwrite the first.
        from brandkit.formats.docx import generate as gen

        doc = Document()
        para = doc.add_paragraph("x")
        _set_spacing(para, before=360)  # authored
        gen._apply_paragraph_geometry(para, {"spacing": {"before_twips": 240}})
        sp = para._p.find(qn("w:pPr")).find(qn("w:spacing"))
        self.assertEqual(sp.get(qn("w:before")), "360")  # authored value preserved

    def test_borders_applied_as_element_copy(self):
        # The captured serialized border is re-emitted byte-identically.
        from brandkit.formats.docx import generate as gen
        from lxml import etree

        src = Document()
        p = src.add_paragraph("x")
        _set_top_border(p, val="double", sz="12", color="FF0000")
        captured = etree.tostring(
            p._p.find(qn("w:pPr")).find(qn("w:pBdr")).find(qn("w:top")),
            encoding="unicode",
        )
        doc = Document()
        target = doc.add_paragraph("y")
        gen._apply_paragraph_geometry(target, {"borders": {"top": captured}})
        applied = etree.tostring(
            target._p.find(qn("w:pPr")).find(qn("w:pBdr")).find(qn("w:top")),
            encoding="unicode",
        )
        self.assertEqual(applied, captured)

    def test_shading_applied_set_only_when_unset(self):
        from brandkit.formats.docx import generate as gen

        doc = Document()
        para = doc.add_paragraph("x")
        gen._apply_paragraph_geometry(para, {"shading": {"fill_hex": "ABCDEF"}})
        shd = para._p.find(qn("w:pPr")).find(qn("w:shd"))
        self.assertEqual(shd.get(qn("w:fill")), "ABCDEF")
        # second apply with a different fill must NOT clobber the now-set fill
        gen._apply_paragraph_geometry(para, {"shading": {"fill_hex": "123456"}})
        shd = para._p.find(qn("w:pPr")).find(qn("w:shd"))
        self.assertEqual(shd.get(qn("w:fill")), "ABCDEF")

    def test_empty_geometry_is_noop(self):
        from brandkit.formats.docx import generate as gen

        doc = Document()
        para = doc.add_paragraph("x")
        had_ppr = para._p.find(qn("w:pPr")) is not None
        gen._apply_paragraph_geometry(para, {})
        # no geometry -> no injected pPr beyond what already existed
        self.assertEqual(para._p.find(qn("w:pPr")) is not None, had_ppr)

    def test_geometry_and_font_size_compose(self):
        # Geometry composes with the run typography axes, independently.
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td), size_hp=22)
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {"body": {"latin": "Roboto", "size_hp": 22}},
                    "geometry": {"body": {"spacing": {"before_twips": 240}}},
                }
            )
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "x"}])])
            docx_generate.generate(prof, shell, idoc, out)
            runs = [r for p in Document(out).paragraphs for r in p.runs if r.text]
            self.assertIn("Roboto", {r.font.name for r in runs})
            self.assertIn(Pt(11), {r.font.size for r in runs})
            self.assertIn('w:before="240"', "\n".join(_out_ppr_xml(out)))


# ---------------------------------------------------------------------------
# byte-identity: a no-geometry profile takes ZERO new branches
# ---------------------------------------------------------------------------
class GeometryByteIdentityTest(unittest.TestCase):
    def _gen_hash(self, prof):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "hi"}])])
            docx_generate.generate(prof, shell, idoc, out)
            import hashlib

            return hashlib.sha256(out.read_bytes()).hexdigest()

    def test_no_geometry_profile_is_unchanged_from_empty(self):
        # A profile with an empty/absent geometry produces the same bytes as a profile
        # that never had a geometry key at all (zero-branch on no-capture).
        plain = _profile(theme={"colors": {}, "fonts": {}})
        with_empty_geom = _profile(theme={"colors": {}, "fonts": {}, "geometry": {}})
        self.assertEqual(self._gen_hash(plain), self._gen_hash(with_empty_geom))

    def test_geometry_apply_is_byte_idempotent(self):
        prof = _profile(
            theme={
                "colors": {},
                "fonts": {},
                "geometry": {"body": {"spacing": {"before_twips": 240}}},
            }
        )
        self.assertEqual(self._gen_hash(prof), self._gen_hash(prof))


# ---------------------------------------------------------------------------
# check: well-formed + observed-floor passes; malformed/synthesized rejected
# ---------------------------------------------------------------------------
class GeometryTargetsCheckTest(unittest.TestCase):
    def _shell_with_geometry(self, tmp_path, geom_fn):
        shell = tmp_path / "shell.docx"
        d = Document()
        for _ in range(2):
            geom_fn(d.add_paragraph("provenance"))
        d.save(shell)
        return shell

    def _prof(self, geometry):
        prof = schema.build_envelope("docx", {"name": "g"})
        prof["surface"] = {"docx": {}}
        prof["theme"] = {"colors": {}, "fonts": {}, "geometry": {"body": geometry}}
        prof["roles"] = {"_index": []}
        return prof

    def test_observed_spacing_passes(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_geometry(
                Path(td), lambda p: _set_spacing(p, before=240)
            )
            prof = self._prof({"spacing": {"before_twips": 240}})
            findings = checks_deterministic.check_geometry_targets(shell, prof)
            self.assertEqual(findings, [])

    def test_synthesized_spacing_not_observed_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_geometry(
                Path(td), lambda p: _set_spacing(p, before=240)
            )
            prof = self._prof({"spacing": {"before_twips": 333}})  # not on the shell
            findings = checks_deterministic.check_geometry_targets(shell, prof)
            self.assertTrue(
                any(
                    f.check == "appearance_geometry_targets"
                    and f.severity == schema.Severity.ERROR.value
                    for f in findings
                )
            )

    def test_out_of_range_twips_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_geometry(
                Path(td), lambda p: _set_spacing(p, before=240)
            )
            prof = self._prof({"spacing": {"before_twips": 999999}})
            findings = checks_deterministic.check_geometry_targets(shell, prof)
            self.assertTrue(any("out of the sane" in f.message for f in findings))

    def test_non_integer_twips_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_geometry(
                Path(td), lambda p: _set_spacing(p, before=240)
            )
            prof = self._prof({"spacing": {"before_twips": "not_an_int"}})
            findings = checks_deterministic.check_geometry_targets(shell, prof)
            self.assertTrue(any("not an integer" in f.message for f in findings))

    def test_observed_border_passes_and_synthesized_fails(self):
        from lxml import etree

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_geometry(Path(td), lambda p: _set_top_border(p))
            # read the shell's own serialized top border to use as the captured value
            facts = checks_deterministic._docx_collect_geometry_facts(shell)
            observed_top = next(iter(facts.borders["top"]))
            ok = self._prof({"borders": {"top": observed_top}})
            self.assertEqual(checks_deterministic.check_geometry_targets(shell, ok), [])
            # a synthesized (different) border is rejected
            synth = OxmlElement("w:top")
            synth.set(qn("w:val"), "wave")
            synth.set(qn("w:sz"), "48")
            synth.set(qn("w:color"), "00FF00")
            bad = self._prof(
                {"borders": {"top": etree.tostring(synth, encoding="unicode")}}
            )
            findings = checks_deterministic.check_geometry_targets(shell, bad)
            self.assertTrue(any("border" in f.message for f in findings))

    def test_malformed_border_missing_val_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_geometry(Path(td), lambda p: _set_top_border(p))
            # a border element missing the required w:val attribute
            bad_xml = (
                '<w:top xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main" w:sz="4"/>'
            )
            prof = self._prof({"borders": {"top": bad_xml}})
            findings = checks_deterministic.check_geometry_targets(shell, prof)
            self.assertTrue(
                any("valid WordprocessingML border" in f.message for f in findings)
            )

    def test_malformed_shading_hex_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_geometry(
                Path(td), lambda p: _set_shading(p, "ABCDEF")
            )
            prof = self._prof({"shading": {"fill_hex": "nothex"}})
            findings = checks_deterministic.check_geometry_targets(shell, prof)
            self.assertTrue(any("not a valid #RRGGBB" in f.message for f in findings))

    def test_no_geometry_profile_has_no_finding(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = schema.build_envelope("docx", {"name": "g"})
            prof["theme"] = {"colors": {}, "fonts": {}}
            prof["roles"] = {"_index": []}
            self.assertEqual(
                checks_deterministic.check_geometry_targets(shell, prof), []
            )

    def test_non_docx_kind_is_noop(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = self._prof({"spacing": {"before_twips": 240}})
            prof["kind"] = "pptx"
            self.assertEqual(
                checks_deterministic.check_geometry_targets(shell, prof), []
            )

    def test_wired_into_run_qa(self):
        # The geometry check runs inside run_qa: a synthesized value surfaces as a
        # finding from the gate, not only from the standalone function.
        from brandkit.qa import gate

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_geometry(
                Path(td), lambda p: _set_spacing(p, before=240)
            )
            out = Path(td) / "out.docx"
            prof = self._prof({"spacing": {"before_twips": 333}})
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "x"}])])
            docx_generate.generate(prof, shell, idoc, out)
            # Injected visual seam: the asserted finding is deterministic L0,
            # assembled before the (here degraded, soffice-free) visual branch.
            report = gate.run_qa(out, prof, shell=shell, visual=(False, []))
            self.assertTrue(
                any(f.check == "appearance_geometry_targets" for f in report.findings)
            )


# ---------------------------------------------------------------------------
# end-to-end: extract a template with dominant geometry -> profile -> generate
# -> geometry applied -> check passes
# ---------------------------------------------------------------------------
class GeometryEndToEndTest(unittest.TestCase):
    def test_extract_apply_verify_geometry(self):
        with tempfile.TemporaryDirectory() as td:
            tmp = Path(td)
            template = tmp / "template.docx"
            d = Document()
            # A dominant body indentation across many paragraphs.
            for _ in range(6):
                _set_indent(d.add_paragraph("body content"), left=720)
            d.save(template)

            roles = {"_index": []}
            theme = {"colors": {}, "fonts": {}}
            typography.capture_geometry(Document(template), roles, theme)
            self.assertEqual(
                theme["geometry"]["body"]["indentation"]["left_twips"], 720
            )

            prof = _profile(
                theme={"colors": {}, "fonts": {}, "geometry": theme["geometry"]}
            )
            out = tmp / "out.docx"
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "generated"}])]
            )
            docx_generate.generate(prof, template, idoc, out)
            self.assertIn('w:left="720"', "\n".join(_out_ppr_xml(out)))
            # the captured value is observed on the template -> check passes
            self.assertEqual(
                checks_deterministic.check_geometry_targets(template, prof), []
            )

    def test_extract_full_profile_carries_geometry(self):
        with tempfile.TemporaryDirectory() as td:
            tmp = Path(td)
            template = tmp / "template.docx"
            d = Document()
            for _ in range(6):
                _set_spacing(d.add_paragraph("body"), before=200, after=200)
            d.save(template)
            saved = docx_extract.extract(template, "geomprof", cwd=tmp)
            import json

            profile = json.loads(Path(saved).read_text())
            body_geom = profile["theme"]["geometry"]["body"]
            self.assertEqual(body_geom["spacing"]["before_twips"], 200)
            self.assertEqual(body_geom["spacing"]["after_twips"], 200)


# ===========================================================================
# Cluster D2: TABLE conditional-format fidelity (tblLook / table style / cell margins),
# DOCX-ONLY. A NEW appearance axis under the SAME _dominant floor: capture, resolver
# merge (no family gate), apply (set-only-when-unset; the synthetic python-docx default
# tblLook is replaced), byte-identity (no-table path), the honest fail-closed check
# (tblLook shape + style-ref membership + cell-margin observed-floor), and a docx
# end-to-end. The band FILLS stay in the shell's style part - the engine only toggles.
# ===========================================================================
def _define_table_style(doc, style_id="AcmeTable", name="Acme Table"):
    """Add a custom ``w:type='table'`` style so a captured/applied style id is a member
    of the shell's table-style inventory (the check's name-membership floor)."""
    styles = doc.styles.element
    st = OxmlElement("w:style")
    st.set(qn("w:type"), "table")
    st.set(qn("w:styleId"), style_id)
    st.set(qn("w:customStyle"), "1")
    nm = OxmlElement("w:name")
    nm.set(qn("w:val"), name)
    st.append(nm)
    bo = OxmlElement("w:basedOn")
    bo.set(qn("w:val"), "TableNormal")
    st.append(bo)
    styles.append(st)


def _brand_table(
    doc,
    *,
    tbllook="01E0",
    margins=None,
    style_id="AcmeTable",
    style_name="Acme Table",
):
    """Add a 2x2 table carrying an explicit ``w:tblLook@w:val``, a ``w:tblStyle`` ref,
    and optional ``w:tblCellMar`` margins (a ``{side: twips}`` dict)."""
    t = doc.add_table(rows=2, cols=2)
    if style_name is not None:
        t.style = style_name
    tblpr = t._tbl.tblPr
    if style_id is not None and tblpr.find(qn("w:tblStyle")) is None:
        st = OxmlElement("w:tblStyle")
        st.set(qn("w:val"), style_id)
        tblpr.insert(0, st)
    if tbllook is not None:
        look = tblpr.find(qn("w:tblLook"))
        if look is None:
            look = OxmlElement("w:tblLook")
            tblpr.append(look)
        look.set(qn("w:val"), tbllook)
    if margins:
        cm = OxmlElement("w:tblCellMar")
        for side, w in margins.items():
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(w))
            el.set(qn("w:type"), "dxa")
            cm.append(el)
        tblpr.append(cm)
    return t


def _table_role(appearance=None):
    entry = {
        "resolver": {
            "type": "named_style",
            "style_id": "AcmeTable",
            "style_name": "Acme Table",
        }
    }
    if appearance is not None:
        entry["appearance"] = appearance
    return {"_index": ["table.default"], "table.default": entry}


def _out_tblpr_xml(out):
    """Every ``w:tblPr`` of the generated doc as serialized strings (raw assertions)."""
    from lxml import etree

    doc = Document(out)
    return [
        etree.tostring(t._tbl.tblPr, encoding="unicode")
        for t in doc.tables
        if t._tbl.tblPr is not None
    ]


# ---------------------------------------------------------------------------
# capture: dominant tblLook / style / cell-margins; absent when the template declares
# none; independent fields; below-dominance captures nothing
# ---------------------------------------------------------------------------
class CaptureTableAppearanceTest(unittest.TestCase):
    def test_dominant_table_tbllook_captured(self):
        doc = Document()
        _define_table_style(doc)
        for _ in range(3):
            _brand_table(doc, tbllook="01E0", margins=None)
        roles = _table_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_table_appearance(doc, roles, theme)
        # 0x01E0 = firstRow|lastRow|firstColumn|lastColumn.
        self.assertEqual(theme["table"]["body"]["tblLook"], 0x01E0)
        self.assertEqual(
            roles["table.default"]["appearance"]["table"]["tblLook"], 0x01E0
        )

    def test_table_style_reference_captured(self):
        doc = Document()
        _define_table_style(doc)
        for _ in range(3):
            _brand_table(doc)
        roles = _table_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_table_appearance(doc, roles, theme)
        self.assertEqual(theme["table"]["body"]["style_id"], "AcmeTable")

    def test_table_cell_margins_captured(self):
        doc = Document()
        _define_table_style(doc)
        for _ in range(3):
            _brand_table(doc, margins={"top": 120, "bottom": 120, "left": 80})
        roles = _table_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_table_appearance(doc, roles, theme)
        margins = theme["table"]["body"]["cell_margins"]
        self.assertEqual(margins["top_twips"], 120)
        self.assertEqual(margins["bottom_twips"], 120)
        self.assertEqual(margins["left_twips"], 80)
        # right was never declared -> absent (independent fields)
        self.assertNotIn("right_twips", margins)

    def test_no_table_declares_nothing_is_absent(self):
        # A doc with NO tables leaves theme.table absent (no-op).
        doc = Document()
        for _ in range(5):
            doc.add_paragraph("body")
        roles = _table_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_table_appearance(doc, roles, theme)
        self.assertNotIn("table", theme)
        self.assertNotIn("table", roles["table.default"].get("appearance", {}))

    def test_below_dominance_table_floor_captures_nothing(self):
        # Only 2 tables (< MIN_RUNS=3): nothing is captured.
        doc = Document()
        _define_table_style(doc)
        for _ in range(2):
            _brand_table(doc, tbllook="01E0")
        roles = _table_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_table_appearance(doc, roles, theme)
        self.assertNotIn("table", theme)

    def test_minority_tbllook_value_is_not_captured(self):
        # 5 tables: 2 carry a custom 01E0, 3 carry the python-docx default 04A0. The
        # default (3/5) dominates instead, so a minority custom look does not win.
        doc = Document()
        _define_table_style(doc)
        for _ in range(2):
            _brand_table(doc, tbllook="01E0")
        for _ in range(3):
            _brand_table(doc, tbllook=None)  # keeps python-docx default 04A0
        roles = _table_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_table_appearance(doc, roles, theme)
        self.assertEqual(theme["table"]["body"]["tblLook"], 0x04A0)

    def test_table_capture_is_idempotent_on_rerun(self):
        doc = Document()
        _define_table_style(doc)
        for _ in range(3):
            _brand_table(doc, margins={"left": 80, "right": 80})
        roles = _table_role(appearance={})
        theme = {"colors": {}, "fonts": {}}
        typography.capture_table_appearance(doc, roles, theme)
        first = dict(roles["table.default"]["appearance"]["table"])
        typography.capture_table_appearance(doc, roles, theme)
        self.assertEqual(roles["table.default"]["appearance"]["table"], first)


# ---------------------------------------------------------------------------
# resolver: role-specific table appearance wins; body fills in for EVERY table role
# (NO family gate, like geometry)
# ---------------------------------------------------------------------------
class ResolverTableAppearanceTest(unittest.TestCase):
    def _prof(self, *, body=None, role_table=None):
        theme = {"colors": {}, "fonts": {}}
        if body is not None:
            theme["table"] = {"body": body}
        appearance = {"table": role_table} if role_table is not None else {}
        return {
            "kind": "docx",
            "theme": theme,
            "roles": {
                "_index": ["table.default"],
                "table.default": {
                    "resolver": {"type": "named_style", "style_id": "AcmeTable"},
                    "appearance": appearance,
                    "status": "robust",
                    "confidence": 1.0,
                },
            },
        }

    def test_role_table_wins_over_body(self):
        prof = self._prof(body={"tblLook": 0x04A0}, role_table={"tblLook": 0x01E0})
        op = ProfileResolver(prof).resolve_role("table.default", fallback=None)
        self.assertEqual(op.appearance["table"]["tblLook"], 0x01E0)

    def test_body_table_flows_to_table_role_no_family_gate(self):
        # Table appearance has NO family gate: a table role with no captured table
        # appearance DOES inherit the body table default.
        prof = self._prof(body={"style_id": "AcmeTable"})
        op = ProfileResolver(prof).resolve_role("table.default", fallback=None)
        self.assertEqual(op.appearance["table"]["style_id"], "AcmeTable")

    def test_pre_d2_profile_resolves_without_table(self):
        prof = self._prof()
        op = ProfileResolver(prof).resolve_role("table.default", fallback=None)
        self.assertNotIn("table", op.appearance)


# ---------------------------------------------------------------------------
# apply: tblLook / style / margins written to w:tblPr; set-only-when-unset; the
# synthetic python-docx default tblLook is replaced; KPI-as-table inherits
# ---------------------------------------------------------------------------
class ApplyTableAppearanceTest(unittest.TestCase):
    def _shell_with_table_style(self, tmp_path):
        shell = tmp_path / "shell.docx"
        d = Document()
        _define_table_style(d)
        d.save(shell)
        return shell

    def _prof_with_body_table(self, table):
        return _profile(theme={"colors": {}, "fonts": {}, "table": {"body": table}})

    def test_tbllook_applied_replacing_synthetic_default(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_table_style(Path(td))
            out = Path(td) / "out.docx"
            prof = self._prof_with_body_table({"tblLook": 0x01E0})
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Table(
                        columns=[{"t": "A"}],
                        rows=[[ir.TableCell(runs=[{"t": "1"}])]],
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            tblpr_xml = "\n".join(_out_tblpr_xml(out))
            self.assertIn('w:val="01E0"', tblpr_xml)
            self.assertNotIn('w:val="04A0"', tblpr_xml)  # synthetic default replaced

    def test_table_style_applied(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_table_style(Path(td))
            out = Path(td) / "out.docx"
            prof = self._prof_with_body_table({"style_id": "AcmeTable"})
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Table(
                        columns=[{"t": "A"}],
                        rows=[[ir.TableCell(runs=[{"t": "1"}])]],
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            self.assertIn("AcmeTable", "\n".join(_out_tblpr_xml(out)))

    def test_cell_margins_applied(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_table_style(Path(td))
            out = Path(td) / "out.docx"
            prof = self._prof_with_body_table(
                {"cell_margins": {"top_twips": 120, "left_twips": 80}}
            )
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Table(
                        columns=[{"t": "A"}],
                        rows=[[ir.TableCell(runs=[{"t": "1"}])]],
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            tblpr_xml = "\n".join(_out_tblpr_xml(out))
            self.assertIn('w:w="120"', tblpr_xml)
            self.assertIn('w:w="80"', tblpr_xml)

    def test_authored_tbllook_not_clobbered(self):
        # Set-only-when-unset: a tblPr that already carries an AUTHORED (non-default)
        # tblLook is never overwritten by the captured value.
        from brandkit.formats.docx import generate as gen

        doc = Document()
        t = doc.add_table(rows=1, cols=1)
        t._tbl.tblPr.find(qn("w:tblLook")).set(qn("w:val"), "00A0")  # authored
        op = ProfileResolver(
            {
                "kind": "docx",
                "theme": {},
                "roles": {
                    "_index": ["table.default"],
                    "table.default": {
                        "resolver": {"type": "named_style", "style_id": "AcmeTable"},
                        "appearance": {"table": {"tblLook": 0x01E0}},
                    },
                },
            }
        ).resolve_role("table.default", fallback=None)
        gen._apply_table_appearance(t, op)
        self.assertEqual(t._tbl.tblPr.find(qn("w:tblLook")).get(qn("w:val")), "00A0")

    def test_authored_margin_side_not_clobbered(self):
        from brandkit.formats.docx import generate as gen

        doc = Document()
        t = doc.add_table(rows=1, cols=1)
        tblpr = t._tbl.tblPr
        cm = OxmlElement("w:tblCellMar")
        left = OxmlElement("w:left")
        left.set(qn("w:w"), "999")
        left.set(qn("w:type"), "dxa")
        cm.append(left)
        tblpr.append(cm)
        op = ProfileResolver(
            {
                "kind": "docx",
                "theme": {"table": {"body": {"cell_margins": {"left_twips": 80}}}},
                "roles": {
                    "_index": ["table.default"],
                    "table.default": {
                        "resolver": {"type": "named_style", "style_id": "AcmeTable"},
                        "appearance": {},
                    },
                },
            }
        ).resolve_role("table.default", fallback=None)
        gen._apply_table_appearance(t, op)
        left_el = t._tbl.tblPr.find(qn("w:tblCellMar")).find(qn("w:left"))
        self.assertEqual(left_el.get(qn("w:w")), "999")  # authored value preserved

    def test_kpi_as_table_inherits_table_appearance(self):
        # KPI / synthetic tables route through the SAME _write_table path, so they get
        # the table.default appearance.table for free (no separate KPI styling).
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell_with_table_style(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={
                    "colors": {},
                    "fonts": {},
                    "table": {"body": {"tblLook": 0x01E0}},
                },
                roles=_table_role(appearance={}),
            )
            # A KPI block synthesizes a table internally via _write_table.
            idoc = ir.IntermediateDocument(
                blocks=[ir.Kpi(items=[ir.KpiItem(label="Revenue", value="$3.2M")])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            self.assertIn('w:val="01E0"', "\n".join(_out_tblpr_xml(out)))

    def test_empty_table_appearance_is_noop(self):
        from brandkit.formats.docx import generate as gen

        doc = Document()
        t = doc.add_table(rows=1, cols=1)
        op = ProfileResolver(
            {
                "kind": "docx",
                "theme": {},
                "roles": {
                    "_index": ["table.default"],
                    "table.default": {
                        "resolver": {"type": "named_style", "style_id": "AcmeTable"},
                        "appearance": {},
                    },
                },
            }
        ).resolve_role("table.default", fallback=None)
        # No captured table appearance -> op_table is None -> no mutation.
        look_before = t._tbl.tblPr.find(qn("w:tblLook")).get(qn("w:val"))
        gen._apply_table_appearance(t, op)
        self.assertEqual(
            t._tbl.tblPr.find(qn("w:tblLook")).get(qn("w:val")), look_before
        )


# ---------------------------------------------------------------------------
# byte-identity: a no-table profile takes ZERO new branches
# ---------------------------------------------------------------------------
class TableByteIdentityTest(unittest.TestCase):
    def _gen_hash(self, prof):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Table(
                        columns=[{"t": "A"}],
                        rows=[[ir.TableCell(runs=[{"t": "1"}])]],
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            import hashlib

            return hashlib.sha256(out.read_bytes()).hexdigest()

    def test_no_table_profile_is_unchanged_from_empty(self):
        plain = _profile(theme={"colors": {}, "fonts": {}})
        with_empty_table = _profile(theme={"colors": {}, "fonts": {}, "table": {}})
        self.assertEqual(self._gen_hash(plain), self._gen_hash(with_empty_table))

    def test_table_apply_is_byte_idempotent(self):
        prof = _profile(
            theme={"colors": {}, "fonts": {}, "table": {"body": {"tblLook": 0x01E0}}}
        )
        self.assertEqual(self._gen_hash(prof), self._gen_hash(prof))


# ---------------------------------------------------------------------------
# check: tblLook shape + style-ref membership + cell-margin observed-floor
# ---------------------------------------------------------------------------
class TableTargetsCheckTest(unittest.TestCase):
    def _shell(self, tmp_path, *, style=True, margins=None):
        shell = tmp_path / "shell.docx"
        d = Document()
        if style:
            _define_table_style(d)
        if margins is not None:
            _brand_table(d, tbllook=None, margins=margins, style_name=None)
        d.save(shell)
        return shell

    def _prof(self, table):
        prof = schema.build_envelope("docx", {"name": "t"})
        prof["surface"] = {"docx": {}}
        prof["theme"] = {"colors": {}, "fonts": {}, "table": {"body": table}}
        prof["roles"] = {"_index": []}
        return prof

    def test_check_table_targets_style_membership(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            ok = self._prof({"style_id": "AcmeTable"})
            self.assertEqual(checks_deterministic.check_table_targets(shell, ok), [])
            bad = self._prof({"style_id": "BrandTable"})  # not in the shell
            findings = checks_deterministic.check_table_targets(shell, bad)
            self.assertTrue(
                any(
                    f.check == "appearance_table_targets"
                    and f.severity == schema.Severity.ERROR.value
                    and "not a table style the shell defines" in f.message
                    for f in findings
                )
            )

    def test_check_table_targets_tbllook_wellformed(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            ok = self._prof({"tblLook": 0x01E0})
            self.assertEqual(checks_deterministic.check_table_targets(shell, ok), [])
            # a bit outside the valid flag set is malformed
            bad = self._prof({"tblLook": 0x8000})
            findings = checks_deterministic.check_table_targets(shell, bad)
            self.assertTrue(
                any("outside the valid flags" in f.message for f in findings)
            )
            # out of the 16-bit range
            oor = self._prof({"tblLook": 99999})
            findings = checks_deterministic.check_table_targets(shell, oor)
            self.assertTrue(any("16-bit range" in f.message for f in findings))
            # a non-integer bitmask
            nonint = self._prof({"tblLook": "01E0"})
            findings = checks_deterministic.check_table_targets(shell, nonint)
            self.assertTrue(
                any("is not an integer bitmask" in f.message for f in findings)
            )

    def test_check_table_targets_margins_observed_floor(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td), margins={"left": 80, "right": 80})
            ok = self._prof({"cell_margins": {"left_twips": 80}})
            self.assertEqual(checks_deterministic.check_table_targets(shell, ok), [])
            # an un-observed margin is rejected (observed-floor)
            synth = self._prof({"cell_margins": {"left_twips": 999}})
            findings = checks_deterministic.check_table_targets(shell, synth)
            self.assertTrue(any("not observed" in f.message for f in findings))
            # out-of-range margin is rejected (shape)
            oor = self._prof({"cell_margins": {"left_twips": -50000}})
            findings = checks_deterministic.check_table_targets(shell, oor)
            self.assertTrue(any("out of the sane" in f.message for f in findings))

    def test_no_table_profile_has_no_finding(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            prof = schema.build_envelope("docx", {"name": "t"})
            prof["theme"] = {"colors": {}, "fonts": {}}
            prof["roles"] = {"_index": []}
            self.assertEqual(checks_deterministic.check_table_targets(shell, prof), [])

    def test_non_docx_kind_is_noop(self):
        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            prof = self._prof({"tblLook": 0x8000})
            prof["kind"] = "pptx"
            self.assertEqual(checks_deterministic.check_table_targets(shell, prof), [])

    def test_wired_into_run_qa(self):
        from brandkit.qa import gate

        with tempfile.TemporaryDirectory() as td:
            shell = self._shell(Path(td))
            out = Path(td) / "out.docx"
            prof = self._prof({"style_id": "BrandTable"})  # undefined -> ERROR
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Table(
                        columns=[{"t": "A"}],
                        rows=[[ir.TableCell(runs=[{"t": "1"}])]],
                    )
                ]
            )
            docx_generate.generate(prof, shell, idoc, out)
            # Injected visual seam: the asserted finding is deterministic L0,
            # assembled before the (here degraded, soffice-free) visual branch.
            report = gate.run_qa(out, prof, shell=shell, visual=(False, []))
            self.assertTrue(
                any(f.check == "appearance_table_targets" for f in report.findings)
            )


# ---------------------------------------------------------------------------
# end-to-end: extract a template with dominant table facts -> profile -> generate
# -> table appearance applied -> check passes
# ---------------------------------------------------------------------------
class TableEndToEndTest(unittest.TestCase):
    def test_extract_apply_verify_table_appearance(self):
        with tempfile.TemporaryDirectory() as td:
            tmp = Path(td)
            template = tmp / "template.docx"
            d = Document()
            _define_table_style(d)
            for _ in range(3):
                _brand_table(d, tbllook="01E0", margins={"left": 80, "right": 80})
            d.save(template)

            saved = docx_extract.extract(template, "tblprof", cwd=tmp)
            import json

            profile = json.loads(Path(saved).read_text())
            body = profile["theme"]["table"]["body"]
            self.assertEqual(body["tblLook"], 0x01E0)
            self.assertEqual(body["style_id"], "AcmeTable")
            self.assertEqual(body["cell_margins"]["left_twips"], 80)

            out = tmp / "out.docx"
            idoc = ir.IntermediateDocument(
                blocks=[
                    ir.Table(
                        columns=[{"t": "A"}],
                        rows=[[ir.TableCell(runs=[{"t": "1"}])]],
                    )
                ]
            )
            docx_generate.generate(profile, template, idoc, out)
            tblpr_xml = "\n".join(_out_tblpr_xml(out))
            self.assertIn('w:val="01E0"', tblpr_xml)
            self.assertIn("AcmeTable", tblpr_xml)
            self.assertIn('w:w="80"', tblpr_xml)
            # the captured facts are shell-backed -> check passes
            self.assertEqual(
                checks_deterministic.check_table_targets(template, profile), []
            )


class AppearanceApplyDegradedTest(unittest.TestCase):
    """E3: the uniform parity ledger. A captured appearance axis the format backend
    does not realize surfaces as ONE INFO ``appearance_apply_degraded`` finding per
    (role, axis), naming only role id + axis; a fully-realized op emits nothing
    (byte-identical existing paths)."""

    class _TrioBackend:
        """A pptx/xlsx-like backend: realizes only the run-typography trio."""

        realized_axes = frozenset({"font", "size_hp", "color"})

        def runs_of(self, target):
            return []

    class _FullBackend(_TrioBackend):
        realized_axes = frozenset(
            {"font", "size_hp", "color", "geometry", "table", "numbering"}
        )

    class _LegacyBackend:
        """A pre-E3 backend with NO declaration and NO geometry hook."""

        def runs_of(self, target):
            return []

    @staticmethod
    def _op(role_id: str, appearance: dict):
        from types import SimpleNamespace

        return SimpleNamespace(role_id=role_id, appearance=appearance)

    def test_unrealizable_axis_emits_one_info_finding(self):
        findings: list = []
        op = self._op("table.default", {"table": {"tblLook": 32}})
        common_appearance.apply_role_appearance(
            self._TrioBackend(), object(), op, findings
        )
        self.assertEqual(len(findings), 1)
        f = findings[0]
        self.assertEqual(f.check, "appearance_apply_degraded")
        self.assertEqual(f.severity, schema.Severity.INFO.value)
        self.assertEqual(f.location, "table.default:table")
        # The message names only role id + axis, never a brand value.
        self.assertNotIn("32", f.message)

    def test_deduplicated_per_role_axis(self):
        findings: list = []
        op = self._op("table.default", {"table": {"tblLook": 32}})
        backend = self._TrioBackend()
        common_appearance.apply_role_appearance(backend, object(), op, findings)
        common_appearance.apply_role_appearance(backend, object(), op, findings)
        self.assertEqual(len(findings), 1)

    def test_fully_realized_backend_emits_nothing(self):
        findings: list = []
        op = self._op(
            "table.default",
            {"table": {"tblLook": 32}, "numbering": {"num_id": "2"}},
        )
        common_appearance.apply_role_appearance(
            self._FullBackend(), object(), op, findings
        )
        self.assertEqual(findings, [])

    def test_empty_appearance_emits_nothing(self):
        findings: list = []
        op = self._op("paragraph", {})
        common_appearance.apply_role_appearance(
            self._TrioBackend(), object(), op, findings
        )
        self.assertEqual(findings, [])

    def test_legacy_backend_inference(self):
        # No declaration + no paragraphs_of hook: geometry is NOT realized -> ledger
        # entry; the trio stays realized -> no entry for a color-only op.
        findings: list = []
        op = self._op("paragraph", {"geometry": {"spacing": {"before_twips": 240}}})
        common_appearance.apply_role_appearance(
            self._LegacyBackend(), object(), op, findings
        )
        self.assertEqual([f.location for f in findings], ["paragraph:geometry"])
        findings2: list = []
        op2 = self._op("paragraph", {"color": {"kind": "hex", "hex": "FF0000"}})
        common_appearance.apply_role_appearance(
            self._LegacyBackend(), object(), op2, findings2
        )
        self.assertEqual(findings2, [])

    def test_real_backends_declare_expected_axes(self):
        from brandkit.formats.pptx import generate as pptx_generate
        from brandkit.formats.xlsx import generate as xlsx_generate

        self.assertEqual(
            docx_generate.DOCX_BACKEND.realized_axes,
            frozenset({"font", "size_hp", "color", "geometry", "table", "numbering"}),
        )
        for backend in (pptx_generate.PPTX_BACKEND, xlsx_generate.XLSX_BACKEND):
            self.assertEqual(
                backend.realized_axes, frozenset({"font", "size_hp", "color"})
            )

    def test_not_an_l0_invariant_and_not_learnable(self):
        self.assertNotIn("appearance_apply_degraded", schema.DEFAULT_L0_INVARIANTS)
        self.assertNotIn("appearance_apply_degraded", schema.LEARNABLE_CHECKS)


if __name__ == "__main__":
    unittest.main()
